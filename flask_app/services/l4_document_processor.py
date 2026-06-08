#!/usr/bin/env python3
"""
Level 4 Document Processor
Extracts text from Word documents and splits into sections for device design extraction
"""

import logging
import re
from typing import Dict, List, Optional
from pathlib import Path
import io

logger = logging.getLogger(__name__)


class L4DocumentProcessor:
    """Processes Level 4 Word documents and extracts relevant sections"""
    
    # Section headings to split on
    SECTION_HEADINGS = [
        "Device Design Data Considerations",
        "Mandibular Advancement Device (Nighttime) – Design Data Considerations",
        "Lower TMJ Appliance (Daytime) – Design Data Considerations",
        "Oral Appliance Options for Consideration",
        "Sleep Study Data",
        "Clinical Background, Complaints & Goals",
        "Structural Observations from Imaging Data",
        "Possible Treatment Considerations",
        "Observations"
    ]
    
    def __init__(self):
        pass
    
    def extract_text_from_docx(self, docx_path: str) -> str:
        """
        Extract text from Word document
        
        Args:
            docx_path: Path to DOCX file
            
        Returns:
            Extracted text content
        """
        try:
            from docx import Document
            
            doc = Document(docx_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)
            
            # Also extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_text.append(" | ".join(row_text))
                if table_text:
                    text_parts.append("\n".join(table_text))
            
            return "\n\n".join(text_parts)
            
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            raise
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {docx_path}: {e}", exc_info=True)
            raise
    
    def extract_text_from_docx_bytes(self, docx_bytes: bytes) -> str:
        """
        Extract text from Word document bytes
        
        Args:
            docx_bytes: DOCX file bytes
            
        Returns:
            Extracted text content
        """
        try:
            from docx import Document
            
            doc = Document(io.BytesIO(docx_bytes))
            text_parts = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)
            
            # Also extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_text.append(" | ".join(row_text))
                if table_text:
                    text_parts.append("\n".join(table_text))
            
            return "\n\n".join(text_parts)
            
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            raise
        except Exception as e:
            logger.error(f"Error extracting text from DOCX bytes: {e}", exc_info=True)
            raise
    
    def split_into_sections(self, text: str) -> Dict[str, str]:
        """
        Split document text into sections based on headings
        
        Args:
            text: Full document text
            
        Returns:
            Dictionary mapping section names to section text
        """
        sections = {}
        
        # Normalize text for matching
        text_lower = text.lower()
        
        # Find all section headings and their positions
        section_positions = []
        for heading in self.SECTION_HEADINGS:
            # Try exact match first
            pattern = re.escape(heading)
            matches = list(re.finditer(pattern, text, re.IGNORECASE))
            for match in matches:
                section_positions.append((match.start(), heading))
            
            # Also try variations (with/without dashes, different spacing)
            variations = [
                heading.replace("–", "-"),
                heading.replace("-", "–"),
                heading.replace("  ", " "),
            ]
            for variation in variations:
                if variation != heading:
                    pattern = re.escape(variation)
                    matches = list(re.finditer(pattern, text, re.IGNORECASE))
                    for match in matches:
                        # Avoid duplicates
                        if not any(abs(match.start() - pos[0]) < 10 for pos in section_positions):
                            section_positions.append((match.start(), heading))
        
        # Sort by position
        section_positions.sort(key=lambda x: x[0])
        
        # Extract sections
        for i, (start_pos, heading) in enumerate(section_positions):
            # Find end position (next section or end of document)
            if i + 1 < len(section_positions):
                end_pos = section_positions[i + 1][0]
            else:
                end_pos = len(text)
            
            section_text = text[start_pos:end_pos].strip()
            sections[heading] = section_text
        
        # If no sections found, return full text under a default key
        if not sections:
            logger.warning("No section headings found, returning full text")
            sections["full_document"] = text
        
        return sections
    
    def extract_patient_id(self, text: str, filename: str) -> Optional[str]:
        """
        Extract patient ID / case ID from text or filename
        
        Args:
            text: Document text
            filename: Source filename
            
        Returns:
            Patient ID if found, None otherwise
        """
        # Try to extract from filename first (e.g., "Example 1 (case YS 1982).docx" -> "YS 1982")
        case_match = re.search(r'case\s+([A-Za-z0-9\s]+)', filename, re.IGNORECASE)
        if case_match:
            return case_match.group(1).strip()
        
        # Try to extract from text (look for "case" or "patient" patterns)
        case_patterns = [
            r'case\s+([A-Za-z0-9\s]+)',
            r'patient\s+id[:\s]+([A-Za-z0-9\s]+)',
            r'case\s+id[:\s]+([A-Za-z0-9\s]+)',
        ]
        
        for pattern in case_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        
        return None
    
    def extract_age_sex(self, text: str) -> Dict[str, Optional[str]]:
        """
        Extract age and sex from report text
        
        Args:
            text: Document text
            
        Returns:
            Dictionary with 'age' and 'sex' keys
        """
        result = {"age": None, "sex": None}
        
        # Look for "Personal details" or "Gender" and "Age" patterns
        # Pattern: "Gender: | M | Age: | 42"
        personal_pattern = r'Gender[:\s|]+\s*([MF])[:\s|]+\s*Age[:\s|]+\s*(\d+)'
        match = re.search(personal_pattern, text, re.IGNORECASE)
        if match:
            result["sex"] = match.group(1).upper()
            result["age"] = match.group(2)
            return result
        
        # Try separate patterns
        sex_patterns = [
            r'Gender[:\s]+\s*([MF])',
            r'Sex[:\s]+\s*([MF])',
        ]
        for pattern in sex_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                result["sex"] = match.group(1).upper()
                break
        
        age_patterns = [
            r'Age[:\s]+\s*(\d+)',
        ]
        for pattern in age_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                result["age"] = match.group(1)
                break
        
        return result
    
    def process_document(self, docx_path: str) -> Dict[str, any]:
        """
        Process a Level 4 document and extract structured information
        
        Args:
            docx_path: Path to DOCX file
            
        Returns:
            Dictionary with extracted text, sections, and metadata
        """
        filename = Path(docx_path).name
        
        # Extract text
        text = self.extract_text_from_docx(docx_path)
        
        # Split into sections
        sections = self.split_into_sections(text)
        
        # Extract patient ID
        patient_id = self.extract_patient_id(text, filename)
        
        # Extract age and sex
        demographics = self.extract_age_sex(text)
        
        return {
            'filename': filename,
            'full_text': text,
            'sections': sections,
            'patient_id': patient_id,
            'age': demographics.get('age'),
            'sex': demographics.get('sex')
        }
    
    def process_document_bytes(self, docx_bytes: bytes, filename: str) -> Dict[str, any]:
        """
        Process a Level 4 document from bytes and extract structured information
        
        Args:
            docx_bytes: DOCX file bytes
            filename: Source filename
            
        Returns:
            Dictionary with extracted text, sections, and metadata
        """
        # Extract text
        text = self.extract_text_from_docx_bytes(docx_bytes)
        
        # Split into sections
        sections = self.split_into_sections(text)
        
        # Extract patient ID
        patient_id = self.extract_patient_id(text, filename)
        
        # Extract age and sex
        demographics = self.extract_age_sex(text)
        
        return {
            'filename': filename,
            'full_text': text,
            'sections': sections,
            'patient_id': patient_id,
            'age': demographics.get('age'),
            'sex': demographics.get('sex')
        }
