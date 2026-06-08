#!/usr/bin/env python3
"""
Batch normalize OSA clinical reports into standardized Level-4 format.

This script processes all DOCX/PDF/TXT files in the input directory, normalizes them
using an LLM, and saves the normalized reports to the output directory.

Usage:
    python3 scripts/batch_normalize_reports.py \
        --input-dir "/home/ec2-user/patient_data/Report Examples" \
        --output-dir "/home/ec2-user/patient_data/Report Examples/normalized" \
        --provider bedrock

Requires:
    - python-docx (for DOCX files)
    - PyPDF2 or pdfplumber (for PDF files)
    - LLM API keys configured in environment variables
"""

from __future__ import annotations

import argparse
import json
import os
import sys
import time
from pathlib import Path
from typing import Iterable, List

# Add parent directory to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent))

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import openai
except ImportError:
    openai = None

try:
    from anthropic import Anthropic
except ImportError:
    Anthropic = None

# Import normalization system prompt from routes
# We'll define it here to avoid Flask dependency
_NORMALIZATION_SYSTEM_PROMPT = """You are a clinical report normalization specialist. Your job is to normalize ANY raw OSA clinical report into a standardized Level-4 OSA Report format suitable for knowledge base ingestion.

CRITICAL RULES:

1. DO NOT FABRICATE CLINICAL VALUES
   - If a field is missing, write "Not provided"
   - NEVER infer or assume values (no O2 nadir, positional AHI, BMI, etc. unless explicitly stated)

2. PRESERVE CLINICAL VOICE
   - Use anatomical correctness
   - Use medical-grade terminology
   - Maintain consistent phrasing

3. NEVER INFER DATA NOT EXPLICITLY STATED
   - No assumptions about any clinical values
   - If not in the source, mark as "Not provided"

4. ALWAYS INCLUDE ALL MANDATORY SECTIONS
   - Even if empty, include every section
   - Follow the exact section order

5. ALWAYS INCLUDE DISCLAIMERS VERBATIM
   - Top disclaimer and final disclaimer must be preserved exactly

OUTPUT FORMAT (MANDATORY - FOLLOW EXACTLY):

# OSA Data Assessment Report

## Personal Details
| Personal details | Gender: | <value or "Not provided"> | Age: | <value or "Not provided"> | BMI: | <value or "Not provided"> |

## Clinical Background, Complaints & Goals
| Clinical background: | <value or "Not provided"> |
| Patient complaints: | <value or "Not provided"> |
| Patient goals: | <value or "Not provided"> |

## ENT Findings
<value or "Not provided.">

## Sleep Study Data
| AHI | <value or "Not provided"> | REM AHI | <value or "Not provided"> |
| RDI | <value or "Not provided"> | REM RDI | <value or "Not provided"> |
| ODI | <value or "Not provided"> | REM ODI | <value or "Not provided"> |
| Supine AHI | <value or "Not provided"> | Supine RDI | <value or "Not provided"> |
| Supine ODI | <value or "Not provided"> | Non-Supine AHI | <value or "Not provided"> |
| Snoring % | <value or "Not provided"> | O2 Nadir | <value or "Not provided"> |
| Sleep Efficiency | <value or "Not provided"> | Total Sleep Time | <value or "Not provided"> |

## Observations
• <bullet-point summary of OSA severity & patterns>
• <bullet of positional dependence>
• <bullet of desaturation severity>
• <bullet of snoring patterns>
• <bullet of apnea/hypopnea breakdown>
• <bullet of any clinically relevant pattern>

If data missing → write: "Sleep observations not provided."

## Structural Observations from Imaging Data
**Important Note:** This section presents observations based on imaging data and does not constitute an official radiological interpretation. Any imaging findings must be reviewed by a certified radiologist or physician before making clinical decisions.

| Key Observations | Details |
| Obstruction Sites | <value or "Not provided"> |
| Bite & Jaw Structure | <value or "Not provided"> |
| Soft Palate & Uvula | <value or "Not provided"> |
| Tongue Position | <value or "Not provided"> |
| Hyoid Bone | <value or "Not provided"> |
| Nasal & Sinus | <value or "Not provided"> |

**Conclusion:**
<high-level imaging interpretation or "Not provided.">

## Possible Treatment Considerations
• <bullet of airway stabilization>
• <bullet of tongue positioning>
• <bullet of nasal airflow optimization>
• <bullet of positional therapy>
• <bullet of weight management if BMI > 30>
• <bullet of CPAP vs OAT logic>

If data missing → write general, non-diagnostic considerations.

## Device Design Data Considerations
| Parameter | Data-Based Consideration |
| Mandibular Advancement | <value or "Not provided"> |
| Vertical Opening | <value or "Not provided"> |
| Anterior Window | <value or "Not provided"> |
| Retention Features | <value or "Not provided"> |
| Material | <value or "Not provided"> |
| Pre-set | <value or "Not provided"> |
| Anterior Acrylic | <value or "Not provided"> |
| Coverage | <value or "Not provided"> |
| Clinical Notes | <value or "Not provided"> |

## Recommendations for Further Evaluation
• <ENT evaluation if nasal/sinus issues>
• <Follow-up sleep test after 90 days>
• <Weight management if BMI > 30>
• <DISE if airway unclear>

If nothing available → write: "No further evaluation recommendations provided."

## Oral Appliance Options for Consideration
| Device | | Key Features |
| Emerald Herbst | | Strong, durable, high-density acrylic |
| Respire Herbst Pink AT | | Metal mesh embedded, high-density acrylic |
| Daynaflex Herbst | | Enhanced tongue space, stain-resistant PMMA |

**Disclaimer**
This AI-generated report assists in analyzing medical imaging and clinical data. It does not constitute a medical diagnosis or treatment recommendation. All clinical decisions must be made by qualified healthcare professionals. This report is for informational purposes only and should not replace professional medical judgment.

FORMATTING RULES:
- Single Markdown document
- Clean tables
- No bullet formatting drift
- No duplicate headings
- No extra commentary
- No images, footnotes, HTML, physician names, or PHI

Your output must contain ALL sections in this exact order, with exact section headings, even if data is missing."""


def extract_text_from_docx(docx_path: Path) -> str:
    """Extract text from a DOCX file."""
    if Document is None:
        raise SystemExit("python-docx is required. Install via: pip install python-docx")
    
    document = Document(docx_path)
    parts: List[str] = []
    
    # Extract paragraphs
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        parts.append(text if text else "")
    
    # Extract tables
    for table in document.tables:
        parts.append("")
        for row in table.rows:
            cell_text = [cell.text.strip() for cell in row.cells]
            if any(cell_text):
                parts.append(" | ".join(cell_text))
        parts.append("")
    
    # Clean up multiple blank lines
    text_lines: List[str] = []
    previous_blank = False
    for line in parts:
        if line.strip():
            text_lines.append(line)
            previous_blank = False
        else:
            if not previous_blank:
                text_lines.append("")
            previous_blank = True
    
    result = "\n".join(text_lines).strip()
    return result + "\n" if result else ""


def extract_text_from_pdf(pdf_path: Path) -> str:
    """Extract text from a PDF file."""
    text_parts: List[str] = []
    
    # Try pdfplumber first (better quality)
    if pdfplumber is not None:
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
            return "\n\n".join(text_parts)
        except Exception as e:
            print(f"  Warning: pdfplumber failed: {e}, trying PyPDF2...")
    
    # Fallback to PyPDF2
    if PyPDF2 is not None:
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
            return "\n\n".join(text_parts)
        except Exception as e:
            raise SystemExit(f"Failed to extract PDF text: {e}")
    
    raise SystemExit("No PDF library available. Install via: pip install pdfplumber or pip install PyPDF2")


def extract_text_from_file(file_path: Path) -> str:
    """Extract text from a file based on its extension."""
    if file_path.suffix.lower() == '.txt':
        return file_path.read_text(encoding='utf-8', errors='ignore')
    elif file_path.suffix.lower() == '.docx':
        return extract_text_from_docx(file_path)
    elif file_path.suffix.lower() == '.pdf':
        return extract_text_from_pdf(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_path.suffix}")


def normalize_with_bedrock(raw_text: str) -> str:
    """Normalize report using Bedrock (requires AWS credentials)."""
    try:
        import boto3
        from botocore.exceptions import ClientError
        
        bedrock_runtime = boto3.client('bedrock-runtime', region_name='us-east-1')
        
        # Prepare messages
        messages = [
            {"role": "user", "content": f"""Normalize the following raw patient report into the Standard Level-4 OSA Report Format defined in the "OSA REPORT NORMALIZATION SPECIFICATION v1.0".

Preserve all real values, insert "Not provided" for missing values, and follow all formatting and wording rules strictly.

RAW REPORT TO NORMALIZE:

{raw_text}

TASK:
Generate a fully normalized Level-4 OSA Report following the exact template structure. Include all mandatory sections even if data is missing."""}
        ]
        
        # Use Claude model via Bedrock
        body = {
            "anthropic_version": "bedrock-2023-05-31",
            "max_tokens": 4000,
            "temperature": 0.1,
            "system": _NORMALIZATION_SYSTEM_PROMPT,
            "messages": messages
        }
        
        # Use Claude 4 Sonnet (default model) - use inference profile format
        response = bedrock_runtime.invoke_model(
            modelId="us.anthropic.claude-sonnet-4-20250514-v1:0",
            body=json.dumps(body)
        )
        
        response_body = json.loads(response['body'].read())
        return response_body['content'][0]['text']
    except Exception as e:
        raise SystemExit(f"Bedrock normalization failed: {e}")


def normalize_with_openai(raw_text: str, api_key: str) -> str:
    """Normalize report using OpenAI."""
    if openai is None:
        raise SystemExit("openai package not installed. Install via: pip install openai")
    
    try:
        openai.api_key = api_key
        completion = openai.chat.completions.create(
            model=os.getenv('NORMALIZE_OPENAI_MODEL', 'gpt-4o'),
            messages=[
                {'role': 'system', 'content': _NORMALIZATION_SYSTEM_PROMPT},
                {'role': 'user', 'content': f"""Normalize the following raw patient report into the Standard Level-4 OSA Report Format defined in the "OSA REPORT NORMALIZATION SPECIFICATION v1.0".

Preserve all real values, insert "Not provided" for missing values, and follow all formatting and wording rules strictly.

RAW REPORT TO NORMALIZE:

{raw_text}

TASK:
Generate a fully normalized Level-4 OSA Report following the exact template structure. Include all mandatory sections even if data is missing."""}
            ],
            temperature=0.1,
            max_tokens=4000,
        )
        return completion.choices[0].message.content
    except Exception as e:
        raise SystemExit(f"OpenAI normalization failed: {e}")


def normalize_with_claude(raw_text: str, api_key: str) -> str:
    """Normalize report using Claude API."""
    if Anthropic is None:
        raise SystemExit("anthropic package not installed. Install via: pip install anthropic")
    
    try:
        client = Anthropic(api_key=api_key)
        resp = client.messages.create(
            model=os.getenv('NORMALIZE_CLAUDE_MODEL', 'claude-3-5-sonnet-20241022-v2:0'),
            max_tokens=4000,
            temperature=0.1,
            system=_NORMALIZATION_SYSTEM_PROMPT,
            messages=[{
                'role': 'user',
                'content': f"""Normalize the following raw patient report into the Standard Level-4 OSA Report Format defined in the "OSA REPORT NORMALIZATION SPECIFICATION v1.0".

Preserve all real values, insert "Not provided" for missing values, and follow all formatting and wording rules strictly.

RAW REPORT TO NORMALIZE:

{raw_text}

TASK:
Generate a fully normalized Level-4 OSA Report following the exact template structure. Include all mandatory sections even if data is missing."""}
            ],
        )
        text_blocks = [block.text for block in resp.content if getattr(block, 'type', '') == 'text']
        return '\n'.join(text_blocks)
    except Exception as e:
        raise SystemExit(f"Claude normalization failed: {e}")


def add_metadata_tags(normalized_report: str, report_level: str = "4", role: str = "style_reference", use: str = "formatting_only", structure_authoritative: bool = False, style_reference: bool = False, include_structure_authoritative: bool = True) -> str:
    """Add metadata tags to normalized report"""
    metadata = {
        "report_level": report_level,
        "role": role
    }
    # Add structure_authoritative field (True or False) if include_structure_authoritative is True
    if include_structure_authoritative:
        metadata["structure_authoritative"] = structure_authoritative
    # Add style_reference if True
    if style_reference:
        metadata["style_reference"] = True
    # Note: "use" field is not included in metadata - only report_level, role, and boolean flags
    metadata_json = json.dumps(metadata, indent=2)
    
    # Add metadata at the top of the report
    tagged_report = f"""<!--METADATA_START-->
{metadata_json}
<!--METADATA_END-->

{normalized_report}"""
    
    return tagged_report


def normalize_report(raw_text: str, provider: str, api_key: str | None = None) -> str:
    """Normalize a raw report using the specified provider."""
    if provider == 'bedrock':
        return normalize_with_bedrock(raw_text)
    elif provider == 'openai':
        if not api_key:
            api_key = os.getenv('LEVEL4_OPENAI_API_KEY') or os.getenv('OPENAI_API_KEY')
        if not api_key:
            raise SystemExit("OpenAI API key required. Set LEVEL4_OPENAI_API_KEY or OPENAI_API_KEY environment variable.")
        return normalize_with_openai(raw_text, api_key)
    elif provider == 'claude':
        if not api_key:
            api_key = os.getenv('LEVEL4_ANTHROPIC_API_KEY') or os.getenv('ANTHROPIC_API_KEY')
        if not api_key:
            raise SystemExit("Anthropic API key required. Set LEVEL4_ANTHROPIC_API_KEY or ANTHROPIC_API_KEY environment variable.")
        return normalize_with_claude(raw_text, api_key)
    else:
        raise SystemExit(f"Unknown provider: {provider}. Use 'bedrock', 'openai', or 'claude'.")


def iter_report_files(input_dir: Path) -> Iterable[Path]:
    """Iterate over all report files (DOCX, PDF, TXT) in the input directory."""
    for ext in ['.docx', '.pdf', '.txt']:
        yield from sorted(input_dir.glob(f"*{ext}"))


def batch_normalize(input_dir: Path, output_dir: Path, provider: str, api_key: str | None = None, skip_existing: bool = True, add_tags: bool = False, report_level: str = "4", role: str = "style_reference", use: str = "formatting_only", structure_authoritative: bool = False, style_reference: bool = False, include_structure_authoritative: bool = True) -> None:
    """Normalize all reports in the input directory."""
    files = list(iter_report_files(input_dir))
    if not files:
        print(f"No report files found in {input_dir}")
        return
    
    output_dir.mkdir(parents=True, exist_ok=True)
    
    total = len(files)
    success_count = 0
    error_count = 0
    
    print(f"\nFound {total} report files to normalize")
    print(f"Provider: {provider}")
    print(f"Output directory: {output_dir}\n")
    
    for idx, file_path in enumerate(files, 1):
        output_path = output_dir / f"{file_path.stem}_normalized.txt"
        
        # Skip if already exists
        if skip_existing and output_path.exists():
            print(f"[{idx}/{total}] SKIP: {file_path.name} (already normalized)")
            continue
        
        print(f"[{idx}/{total}] Processing: {file_path.name}...")
        
        try:
            # Extract text
            raw_text = extract_text_from_file(file_path)
            if not raw_text.strip():
                print(f"  Warning: No text extracted from {file_path.name}")
                continue
            
            # Normalize
            print(f"  Normalizing with {provider}...")
            normalized = normalize_report(raw_text, provider, api_key)
            
            # Add metadata tags if requested
            if add_tags:
                normalized = add_metadata_tags(normalized, report_level, role, use, structure_authoritative, style_reference, include_structure_authoritative)
                tag_info = f"level={report_level}, role={role}"
                if include_structure_authoritative:
                    tag_info += f", structure_authoritative={str(structure_authoritative).lower()}"
                if style_reference:
                    tag_info += ", style_reference=true"
                print(f"  Added metadata tags ({tag_info})")
            
            # Save
            output_path.write_text(normalized, encoding='utf-8')
            print(f"  ✓ Saved: {output_path.name}")
            success_count += 1
            
            # Rate limiting - be nice to APIs
            if idx < total:
                time.sleep(1)
                
        except KeyboardInterrupt:
            print("\n\nInterrupted by user. Exiting...")
            break
        except Exception as e:
            print(f"  ✗ Error: {e}")
            error_count += 1
            continue
    
    print(f"\n{'='*60}")
    print(f"Batch normalization complete!")
    print(f"  Total files: {total}")
    print(f"  Successful: {success_count}")
    print(f"  Errors: {error_count}")
    print(f"  Skipped: {total - success_count - error_count}")
    print(f"{'='*60}\n")


def parse_args(argv: List[str]) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Batch normalize OSA clinical reports into standardized Level-4 format."
    )
    parser.add_argument(
        "--input-dir",
        type=Path,
        required=True,
        help="Directory containing report files (DOCX, PDF, TXT)."
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        required=True,
        help="Directory to write normalized reports."
    )
    parser.add_argument(
        "--provider",
        choices=['bedrock', 'openai', 'claude'],
        default='bedrock',
        help="LLM provider to use (default: bedrock)."
    )
    parser.add_argument(
        "--api-key",
        type=str,
        default=None,
        help="API key (if not set, uses environment variables)."
    )
    parser.add_argument(
        "--no-skip-existing",
        action='store_true',
        help="Re-normalize files even if output already exists."
    )
    parser.add_argument(
        "--add-tags",
        action='store_true',
        help="Add metadata tags to normalized reports."
    )
    parser.add_argument(
        "--report-level",
        type=str,
        default="4",
        help="Report level for metadata tags (default: 4)."
    )
    parser.add_argument(
        "--role",
        type=str,
        default="style_reference",
        help="Role for metadata tags (default: style_reference)."
    )
    parser.add_argument(
        "--use",
        type=str,
        default="formatting_only",
        help="Use field for metadata tags (default: formatting_only)."
    )
    parser.add_argument(
        "--structure-authoritative",
        action='store_true',
        help="Mark structure as authoritative in metadata tags."
    )
    parser.add_argument(
        "--style-reference",
        action='store_true',
        help="Mark as style reference in metadata tags."
    )
    return parser.parse_args(argv)


def main(argv: List[str] | None = None) -> None:
    args = parse_args(argv or sys.argv[1:])
    input_dir: Path = args.input_dir.expanduser().resolve()
    output_dir: Path = args.output_dir.expanduser().resolve()
    
    if not input_dir.exists() or not input_dir.is_dir():
        raise SystemExit(f"Input directory is invalid: {input_dir}")
    
    batch_normalize(
        input_dir=input_dir,
        output_dir=output_dir,
        provider=args.provider,
        api_key=args.api_key,
        skip_existing=not args.no_skip_existing,
        add_tags=args.add_tags,
        report_level=args.report_level,
        role=args.role,
        use=args.use,
        structure_authoritative=args.structure_authoritative,
        style_reference=args.style_reference
    )


if __name__ == "__main__":
    main()

