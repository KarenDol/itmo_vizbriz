#!/usr/bin/env python3
"""
Document Parser Service
Converts PDF/Word documents to structured JSON for audio script generation
"""

import json
import logging
import io
import tempfile
import os
import re
from typing import Dict, Any, Optional
from flask_app.services.bedrock_service import BedrockService

logger = logging.getLogger(__name__)


def clean_json_text(text: str) -> str:
    """
    Clean JSON text by removing control characters and extracting JSON from markdown
    
    Args:
        text: Raw text that may contain JSON
        
    Returns:
        Cleaned JSON string
    """
    if not text:
        return ""
    
    # First, try to extract JSON from markdown code blocks
    json_text = text.strip()
    
    # Remove markdown code blocks
    if "```json" in json_text:
        json_text = json_text.split("```json")[1].split("```")[0].strip()
    elif "```" in json_text:
        # Find first ``` and last ```
        parts = json_text.split("```")
        if len(parts) >= 3:
            # Take the middle part (between first and last ```)
            json_text = "```".join(parts[1:-1]).strip()
    
    # More aggressive control character removal
    # JSON only allows: space (32+), \n (10), \r (13), \t (9)
    # Remove ALL other control characters (0x00-0x1F except 9, 10, 13)
    cleaned = ""
    for char in json_text:
        code = ord(char)
        # Keep only: printable ASCII (32-126), newline (10), tab (9), carriage return (13)
        # Also allow extended ASCII and Unicode for content, but escape properly
        if code == 9 or code == 10 or code == 13:  # \t, \n, \r
            cleaned += char
        elif code >= 32:  # All printable characters
            # For string content, we need to be careful with quotes and backslashes
            # But we'll let json.loads handle proper escaping
            cleaned += char
        # All other control chars (0x00-0x1F except 9, 10, 13) are removed
    
    # Remove any remaining problematic characters that might cause issues
    # Remove zero-width spaces and other invisible Unicode characters
    cleaned = re.sub(r'[\u200B-\u200D\uFEFF]', '', cleaned)  # Zero-width spaces
    
    # Remove trailing commas before closing braces/brackets (common JSON error)
    cleaned = re.sub(r',\s*}', '}', cleaned)
    cleaned = re.sub(r',\s*]', ']', cleaned)
    
    # Try to find JSON object boundaries if text is mixed
    # Look for first { and last }
    first_brace = cleaned.find('{')
    last_brace = cleaned.rfind('}')
    if first_brace >= 0 and last_brace > first_brace:
        cleaned = cleaned[first_brace:last_brace + 1]
    
    # Final cleanup: remove any leading/trailing whitespace and ensure it starts with {
    cleaned = cleaned.strip()
    if not cleaned.startswith('{'):
        first_brace = cleaned.find('{')
        if first_brace >= 0:
            cleaned = cleaned[first_brace:]
    
    return cleaned


class DocumentParserService:
    """Service for parsing PDF/Word documents and converting to structured JSON"""
    
    def __init__(self):
        self.bedrock = BedrockService()
    
    def extract_text_from_pdf(self, pdf_bytes: bytes) -> str:
        """
        Extract text from PDF file
        
        Args:
            pdf_bytes: PDF file bytes
            
        Returns:
            Extracted text content
        """
        try:
            import pdfplumber
            
            text_parts = []
            with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                    
                    # Also try to extract tables
                    tables = page.extract_tables()
                    for table in tables:
                        if table:
                            table_text = "\n".join(["\t".join([str(cell) if cell else "" for cell in row]) for row in table])
                            text_parts.append(f"\n[TABLE]\n{table_text}\n[/TABLE]\n")
            
            return "\n\n".join(text_parts)
            
        except ImportError:
            # Fallback to PyPDF2
            try:
                import PyPDF2
                text_parts = []
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                return "\n\n".join(text_parts)
            except Exception as e:
                logger.error(f"Failed to extract text from PDF: {e}")
                raise
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}", exc_info=True)
            raise
    
    def extract_text_from_docx(self, docx_bytes: bytes) -> str:
        """
        Extract text from Word document
        
        Args:
            docx_bytes: DOCX file bytes
            
        Returns:
            Extracted text content
        """
        try:
            from docx import Document
            
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                temp_file.write(docx_bytes)
                temp_file_path = temp_file.name
            
            try:
                doc = Document(temp_file_path)
                text_parts = []
                
                # Extract paragraphs
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text)
                
                # Extract tables
                for table in doc.tables:
                    table_rows = []
                    for row in table.rows:
                        row_cells = [cell.text.strip() for cell in row.cells]
                        table_rows.append("\t".join(row_cells))
                    if table_rows:
                        text_parts.append(f"\n[TABLE]\n" + "\n".join(table_rows) + "\n[/TABLE]\n")
                
                return "\n\n".join(text_parts)
            finally:
                # Clean up temp file
                try:
                    os.unlink(temp_file_path)
                except:
                    pass
                    
        except ImportError:
            raise RuntimeError("python-docx library not installed. Install with: pip install python-docx")
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}", exc_info=True)
            raise
    
    def structure_text_to_json(self, text: str, document_type: str = "sleep_study_report") -> Dict[str, Any]:
        """
        Use Bedrock to structure extracted text into JSON format
        
        Args:
            text: Extracted text from document
            document_type: Type of document (e.g., "sleep_study_report", "level4_report")
            
        Returns:
            Structured JSON data
        """
        try:
            # Truncate text if too long (Bedrock has token limits)
            max_chars = 10000  # Reasonable limit for prompt
            if len(text) > max_chars:
                text = text[:max_chars] + "\n\n[Document truncated...]"
            
            system_prompt = """You are a medical document parser. Extract structured data from sleep study reports.
Output a clean JSON structure that can be used for generating patient-friendly audio scripts.
Focus on key metrics, findings, and recommendations."""

            user_prompt = f"""Extract structured data from this {document_type}:

{text}

Output a JSON object with this structure (use null for missing values):
{{
  "study_type": "baseline" or "follow-up" or null,
  "patient_info": {{
    "age": <number or null>,
    "gender": "male" or "female" or null
  }},
  "metrics": {{
    "AHI": <number or null>,
    "ODI": <number or null>,
    "SpO2_nadir": <number or null>,
    "SpO2_mean": <number or null>,
    "total_sleep_time": <number in hours or null>,
    "sleep_efficiency": <percentage or null>
  }},
  "severity": "mild" or "moderate" or "severe" or null,
  "key_findings": ["finding1", "finding2", ...],
  "symptoms": ["symptom1", "symptom2", ...],
  "recommendations": ["recommendation1", "recommendation2", ...],
  "notes": "any additional relevant notes or null"
}}

Rules:
- Only extract information that is explicitly stated in the document
- Use null for missing values, not 0 or empty strings
- Keep findings and recommendations in plain language
- Do not invent or infer values that aren't in the document"""

            messages = [
                {"role": "user", "content": system_prompt + "\n\n" + user_prompt}
            ]
            
            result = self.bedrock.invoke_model(
                messages=messages,
                model="claude_35_sonnet_v2",
                max_tokens=2000,
                temperature=0.1,
                endpoint="document_parser_structure"
            )
            
            if not result.get("success"):
                logger.error(f"Failed to structure document: {result.get('error')}")
                return {"error": result.get("error", "Unknown error")}
            
            response_text = result.get("response", "")
            
            # Clean and extract JSON
            json_text = clean_json_text(response_text)
            
            # Helper function to fix common JSON issues
            def fix_json_common_issues(text):
                """Fix common JSON formatting issues"""
                fixed = text
                # Remove trailing commas
                fixed = re.sub(r',\s*}', '}', fixed)
                fixed = re.sub(r',\s*]', ']', fixed)
                # Fix unquoted keys (basic attempt)
                fixed = re.sub(r'(\w+):', r'"\1":', fixed)
                # Remove comments (JSON doesn't allow comments)
                fixed = re.sub(r'//.*?$', '', fixed, flags=re.MULTILINE)
                fixed = re.sub(r'/\*.*?\*/', '', fixed, flags=re.DOTALL)
                return fixed
            
            # Try multiple parsing strategies
            parsing_strategies = [
                ("standard", lambda: json.loads(json_text)),
                ("extract_braces", lambda: json.loads(json_text[json_text.find('{'):json_text.rfind('}')+1])),
                ("remove_escaped_newlines", lambda: json.loads(re.sub(r'\\n', ' ', json_text))),
                ("fix_common_issues", lambda: json.loads(fix_json_common_issues(json_text))),
                ("remove_all_newlines", lambda: json.loads(re.sub(r'\n', ' ', json_text))),
            ]
            
            for strategy_name, parse_func in parsing_strategies:
                try:
                    structured_data = parse_func()
                    logger.info(f"Successfully structured document into JSON using {strategy_name} strategy")
                    return {"success": True, "data": structured_data}
                except json.JSONDecodeError as e:
                    # Log detailed error information
                    error_msg = str(e)
                    if hasattr(e, 'pos'):
                        pos = e.pos
                        start = max(0, pos - 50)
                        end = min(len(json_text), pos + 50)
                        context = json_text[start:end]
                        problem_char = json_text[pos] if pos < len(json_text) else 'EOF'
                        logger.debug(f"Strategy {strategy_name} failed at position {pos}: {error_msg}")
                        logger.debug(f"Problem character: {repr(problem_char)} (code: {ord(problem_char) if pos < len(json_text) else 'N/A'})")
                        logger.debug(f"Context: ...{context}...")
                    else:
                        logger.debug(f"Strategy {strategy_name} failed: {error_msg}")
                    continue
                except (UnicodeDecodeError, ValueError) as e:
                    logger.debug(f"Strategy {strategy_name} failed: {e}")
                    continue
                except Exception as e:
                    logger.debug(f"Strategy {strategy_name} failed with unexpected error: {e}")
                    continue
            
            # If all strategies fail, try manual extraction of key fields
            logger.error(f"All JSON parsing strategies failed")
            logger.error(f"Cleaned JSON text (first 2000 chars): {json_text[:2000]}")
            logger.error(f"Cleaned JSON text (last 500 chars): {json_text[-500:]}")
            logger.error(f"Original response (first 1000 chars): {response_text[:1000]}")
            
            # Last resort: try to extract just the JSON portion more carefully
            try:
                # Find the JSON object more carefully
                start_idx = json_text.find('{')
                if start_idx >= 0:
                    # Count braces to find the matching closing brace
                    brace_count = 0
                    end_idx = start_idx
                    for i in range(start_idx, len(json_text)):
                        if json_text[i] == '{':
                            brace_count += 1
                        elif json_text[i] == '}':
                            brace_count -= 1
                            if brace_count == 0:
                                end_idx = i + 1
                                break
                    
                    if end_idx > start_idx:
                        json_snippet = json_text[start_idx:end_idx]
                        # Try to fix and parse
                        for fix_strategy in [json_snippet, fix_json_common_issues(json_snippet)]:
                            try:
                                structured_data = json.loads(fix_strategy)
                                logger.info("Successfully parsed JSON using brace counting with fixes")
                                return {"success": True, "data": structured_data}
                            except:
                                continue
            except Exception as e:
                logger.error(f"Brace counting strategy also failed: {e}")
            
            # Final fallback: try to use Bedrock again with stricter prompt
            logger.warning("Attempting to regenerate JSON with stricter formatting requirements")
            try:
                strict_prompt = f"""Extract structured data from this {document_type}:

{text[:5000]}

Output ONLY valid JSON (no markdown, no code blocks, no explanations). Use this exact structure:
{{
  "study_type": null,
  "patient_info": {{"age": null, "gender": null}},
  "metrics": {{"AHI": null, "ODI": null, "SpO2_nadir": null, "SpO2_mean": null, "total_sleep_time": null, "sleep_efficiency": null}},
  "severity": null,
  "key_findings": [],
  "symptoms": [],
  "recommendations": [],
  "notes": null
}}

Output ONLY the JSON object, nothing else."""

                strict_messages = [
                    {"role": "user", "content": strict_prompt}
                ]
                
                strict_result = self.bedrock.invoke_model(
                    messages=strict_messages,
                    model="claude_35_sonnet_v2",
                    max_tokens=2000,
                    temperature=0.0,  # Lower temperature for more consistent output
                    endpoint="document_parser_strict_retry"
                )
                
                if strict_result.get("success"):
                    strict_response = strict_result.get("response", "")
                    strict_json = clean_json_text(strict_response)
                    try:
                        structured_data = json.loads(strict_json)
                        logger.info("Successfully parsed JSON from strict retry")
                        return {"success": True, "data": structured_data}
                    except:
                        pass
            except Exception as e:
                logger.error(f"Strict retry also failed: {e}")
            
            return {"error": f"Failed to parse JSON after all strategies. Response length: {len(response_text)} chars. First 200 chars: {response_text[:200]}"}
                
        except Exception as e:
            logger.error(f"Error structuring document: {e}", exc_info=True)
            return {"error": str(e)}
    
    def parse_document(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse a document (PDF or Word) and convert to structured JSON
        
        Args:
            file_bytes: Document file bytes
            filename: Original filename (for determining file type)
            
        Returns:
            Dict with success status and structured JSON data or error
        """
        try:
            filename_lower = filename.lower()
            
            # Extract text based on file type
            if filename_lower.endswith('.pdf'):
                text = self.extract_text_from_pdf(file_bytes)
            elif filename_lower.endswith('.docx') or filename_lower.endswith('.doc'):
                text = self.extract_text_from_docx(file_bytes)
            else:
                return {"error": f"Unsupported file type: {filename}. Supported: PDF, DOCX"}
            
            if not text or len(text.strip()) < 50:
                return {"error": "Document appears to be empty or could not extract text"}
            
            logger.info(f"Extracted {len(text)} characters from {filename}")
            
            # Structure the text into JSON
            structure_result = self.structure_text_to_json(text)
            
            if "error" in structure_result:
                return structure_result
            
            return {
                "success": True,
                "raw_text": text[:500] + "..." if len(text) > 500 else text,  # Include preview
                "structured_data": structure_result.get("data", {})
            }
            
        except Exception as e:
            logger.error(f"Error parsing document: {e}", exc_info=True)
            return {"error": str(e)}
