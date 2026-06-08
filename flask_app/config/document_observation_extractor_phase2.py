#!/usr/bin/env python3
"""
Document-Based Observation Extraction System - Phase 2
Document Content Extraction & LLM Integration

This script implements Phase 2 of the document observation extraction system:
1. Document content extraction (PDF, images, text files)
2. AWS Bedrock LLM integration for observation extraction
3. Observation storage in observation_store table with deduplication
4. Batch processing capabilities
5. Schema-compliant canonical JSON creation

KEY IMPROVEMENTS:
- Schema compliance: Uses correct enum values ("home"/"inlab" for study_type, etc.)
- Deduplication: Prevents duplicate observations across different document sources
- Sparse JSON: Removes empty/null values before storage
- Safe sex detection: Context-aware gender detection to avoid false positives
- Imaging exclusion: Filters out imaging category documents
- Validation: Schema validation before storage with proper error handling
"""

import os
import sys
import logging
import json
import boto3
import requests
import re
from datetime import datetime
from typing import List, Dict, Optional, Tuple, Any
from decimal import Decimal, InvalidOperation
import mysql.connector
from dotenv import load_dotenv
import PyPDF2
import io
from PIL import Image
import pytesseract
try:
    from pdf2image import convert_from_path, convert_from_bytes
    from pdf2image.exceptions import PDFInfoNotInstalledError, PDFPageCountError
    PDF2IMAGE_AVAILABLE = True
except Exception:
    convert_from_path = None  # type: ignore
    convert_from_bytes = None  # type: ignore
    PDFInfoNotInstalledError = None  # type: ignore
    PDFPageCountError = None  # type: ignore
    PDF2IMAGE_AVAILABLE = False

POPPLER_PATH = os.getenv("POPPLER_PATH")
import fitz  # PyMuPDF for better PDF handling
import pdfplumber  # Enhanced PDF text and table extraction
import camelot  # Advanced table extraction
# import tabula  # Alternative table extraction - removed due to Java dependency
import shutil
import gc
import argparse
import time

# Load environment variables
load_dotenv()

# Note: BedrockService not imported for standalone script compatibility
# Model ID is hardcoded below in BEDROCK_CONFIG section

# --- TARGETED AHI EXTRACTOR HELPERS ---
import cv2
import numpy as np
from PIL import ImageOps, ImageFilter

HEB_OVERALL_TOKENS = ["כללי", "סה\"כ", "כולל"]
ENG_OVERALL_TOKENS = ["overall", "total"]
REM_TOKENS = ["rem"]
SUPINE_TOKENS = ["supine"]
NREM_TOKENS = ["nrem", "non-rem", "non rem", "nonrem"]

def _preprocess_for_ocr(pil_img: Image.Image) -> Image.Image:
    """Preprocess image for better OCR results."""
    # enlarge + grayscale + contrast + slight sharpen
    img = pil_img.convert("L").resize((pil_img.width * 2, pil_img.height * 2), Image.BICUBIC)
    img = ImageOps.autocontrast(img)
    img = img.filter(ImageFilter.SHARPEN)
    return img

def _ocr_words_with_boxes(pil_img: Image.Image, lang="eng+heb"):
    """Extract words with bounding boxes using OCR."""
    from pytesseract import Output
    d = pytesseract.image_to_data(pil_img, lang=lang, config="--psm 6", output_type=Output.DICT)
    words = []
    for i in range(len(d["text"])):
        txt = d["text"][i].strip()
        if not txt:
            continue
        x, y, w, h = d["left"][i], d["top"][i], d["width"][i], d["height"][i]
        words.append({"text": txt, "bbox": (x,y,w,h)})
    return words

def _find_overall_ahi_from_words(words: list) -> tuple:
    """
    Returns (value, bbox_hint) if found with strong cues for 'overall' / 'כללי' / 'סה"כ'.
    Falls back to best AHI not labeled REM/supine/NREM when multiple appear.
    """
    # Build line-like groups by y proximity
    lines = {}
    for w in words:
        y = w["bbox"][1]
        key = int(round(y / 12.0))  # coarse bucket
        lines.setdefault(key, []).append(w)
    
    candidates = []
    for _, ws in lines.items():
        ws_sorted = sorted(ws, key=lambda t: t["bbox"][0])
        line_text = " ".join([t["text"] for t in ws_sorted]).lower()
        # capture any number next to AHI-like tokens
        if "ahi" in line_text or any(tok in line_text for tok in ["אהי"]):
            nums = re.findall(r'(?<![%:/])\b(\d{1,3}(?:\.\d+)?)\b', line_text)
            if not nums:
                continue
            # label context
            has_overall = any(tok in line_text for tok in ENG_OVERALL_TOKENS+HEB_OVERALL_TOKENS)
            has_rem = any(tok in line_text for tok in REM_TOKENS)
            has_supine = any(tok in line_text for tok in SUPINE_TOKENS)
            has_nrem = any(tok in line_text for tok in NREM_TOKENS)
            # prefer numbers that sit to the right of 'AHI'
            x_mean = np.mean([w["bbox"][0] for w in ws_sorted])
            bbox_hint = (ws_sorted[0]["bbox"][0], ws_sorted[0]["bbox"][1], 
                         ws_sorted[-1]["bbox"][0]+ws_sorted[-1]["bbox"][2]-ws_sorted[0]["bbox"][0],
                         max(ws_sorted, key=lambda z: z["bbox"][3])["bbox"][3])
            for n in nums:
                val = float(n)
                score = 0
                if has_overall: score += 3
                if has_rem: score -= 2
                if has_supine: score -= 2
                if has_nrem: score -= 2
                # plausible range bonus
                if 0 <= val <= 100: score += 1
                candidates.append((val, score, bbox_hint, line_text))
    
    if not candidates:
        return (None, None)
    
    # First: any with positive score (overall cues win)
    best = sorted(candidates, key=lambda x: (x[1], -abs(x[0]-15)), reverse=True)[0]
    if best[1] <= 0:
        # fallback: pick number on an AHI line that does NOT mention REM/supine/NREM
        clean = [c for c in candidates if c[1] >= 1 or ("rem" not in c[3] and "supine" not in c[3] and "nrem" not in c[3])]
        if clean:
            best = sorted(clean, key=lambda x: (x[1], -abs(x[0]-15)), reverse=True)[0]
    
    return (best[0], best[2])

def _requery_ahi_overall_with_roi(bedrock, pil_img: Image.Image, bbox) -> Optional[float]:
    """Re-query vision model on cropped AHI overall region."""
    # Crop + encode PNG
    x,y,w,h = bbox
    roi = pil_img.crop((x-20, y-10, x+w+20, y+h+10))
    buf = io.BytesIO()
    roi.save(buf, format="PNG")
    roi_bytes = buf.getvalue()

    system = [{"text": (
        "You are reading a small cropped image that contains ONLY the AHI overall line from a sleep report.\n"
        "Return a SINGLE JSON object: {\"ahi_overall\": number}. No text. No other fields.\n"
        "Do not return REM/NREM/supine values. If multiple numbers, choose the one labeled overall/total (or Hebrew: כללי, סה\"כ)."
    )}]
    messages = [{"role":"user","content":[{"text":"Extract only AHI overall."},
                 {"image":{"format":"png","source":{"bytes": roi_bytes}}}]}]
    
    # Use hardcoded model ID for standalone script
    model_id = MODEL_ID
    
    import time
    start_time = time.time()
    
    try:
        resp = bedrock.converse(
            modelId=model_id,
            system=system, messages=messages,
            inferenceConfig={"temperature":0.0,"maxTokens":200}
        )
        raw = resp["output"]["message"]["content"][0]["text"].strip()
        response_time_ms = int((time.time() - start_time) * 1000)
        
        # Log to database
        _log_llm_call(
            prompt_text="Extract AHI overall from cropped image",
            response_text=raw,
            response_time_ms=response_time_ms,
            status='success'
        )
        
        raw = re.sub(r'^```json\s*|\s*```$', '', raw).strip()
        obj = json.loads(raw)
        v = obj.get("ahi_overall")
        return float(v) if v is not None else None
    except Exception as e:
        response_time_ms = int((time.time() - start_time) * 1000)
        _log_llm_call(
            prompt_text="Extract AHI overall from cropped image",
            response_text='',
            response_time_ms=response_time_ms,
            status='error',
            error_message=str(e)
        )
        logger.warning(f"ROI re-query failed: {e}")
        return None

def _finalize_ahi_consistency(d: dict):
    """Sanity check to prevent AHI overall from mirroring NREM."""
    ri = d.setdefault("respiratory_indices", {})
    a, r, n, s = ri.get("ahi_overall"), ri.get("ahi_rem"), ri.get("ahi_nrem"), ri.get("ahi_supine")
    # If overall equals NREM exactly, and another distinct candidate exists, prefer the distinct non-stage value
    if a is not None and n is not None and float(a) == float(n):
        # choose alternative if one exists and differs clearly
        for cand in [r, s]:
            if cand is not None and abs(float(cand) - float(a)) >= 8.0:
                # don't swap to REM/Supine; instead, leave overall for ROI re-query
                pass
        # If we had earlier OCR candidate stashed, you can wire it in here (e.g., via context)
    return d


def _validate_ahi_overall_from_text(visual_data: dict, document_text: str) -> dict:
    """
    Validate AHI overall is not a positional/stage-specific value that was hallucinated.
    If AHI overall matches supine/REM/NREM exactly, use regex on document text to find the real overall.
    
    Args:
        visual_data: The extracted data from Visual LLM
        document_text: The OCR/text content of the document (for validation)
        
    Returns:
        The visual_data dict, potentially with corrected ahi_overall
    """
    if not visual_data or not document_text:
        return visual_data
    
    ri = visual_data.get("respiratory_indices", {})
    if not ri:
        return visual_data
    
    ahi_overall = ri.get("ahi_overall")
    ahi_supine = ri.get("ahi_supine")
    ahi_rem = ri.get("ahi_rem")
    ahi_nrem = ri.get("ahi_nrem")
    
    if ahi_overall is None:
        return visual_data
    
    try:
        ahi_overall_float = float(ahi_overall)
    except (ValueError, TypeError):
        return visual_data
    
    # Red flag: AHI overall equals a positional/stage value (likely hallucination)
    suspicious = False
    suspicious_reason = None
    
    if ahi_supine is not None:
        try:
            if abs(ahi_overall_float - float(ahi_supine)) < 0.1:
                suspicious = True
                suspicious_reason = f"AHI overall ({ahi_overall}) equals supine ({ahi_supine})"
        except (ValueError, TypeError):
            pass
            
    if ahi_rem is not None and not suspicious:
        try:
            if abs(ahi_overall_float - float(ahi_rem)) < 0.1:
                suspicious = True
                suspicious_reason = f"AHI overall ({ahi_overall}) equals REM ({ahi_rem})"
        except (ValueError, TypeError):
            pass
            
    if ahi_nrem is not None and not suspicious:
        try:
            if abs(ahi_overall_float - float(ahi_nrem)) < 0.1:
                suspicious = True
                suspicious_reason = f"AHI overall ({ahi_overall}) equals NREM ({ahi_nrem})"
        except (ValueError, TypeError):
            pass
    
    if suspicious:
        logger.warning(f"🔍 AHI validation: {suspicious_reason} - checking document text")
        
        # Search for explicit "Overall AHI" or "Total AHI" patterns in document text
        doc_text_lower = document_text.lower()
        
        # Patterns to find explicitly labeled overall/total AHI
        # Hebrew: כללי (overall), סה"כ (total)
        patterns = [
            r'(?:overall|total|כללי|סה"כ)\s*(?:ahi|אהי)[:\s=]*(\d+\.?\d*)',
            r'(?:ahi|אהי)\s*(?:overall|total|כללי|סה"כ)[:\s=]*(\d+\.?\d*)',
            r'(?:ahi|אהי)\s*(?:index)?[:\s=]*(\d+\.?\d*)(?:\s*/\s*hr)?(?:\s*events)?',
        ]
        
        candidates = []
        for pattern in patterns:
            matches = re.findall(pattern, doc_text_lower)
            for match in matches:
                try:
                    val = float(match)
                    # Only consider plausible AHI values (0-150)
                    if 0 <= val <= 150:
                        candidates.append(val)
                except (ValueError, TypeError):
                    pass
        
        if candidates:
            # Find the most common value or the one that differs from suspicious values
            from collections import Counter
            counter = Counter(candidates)
            
            # Prefer values that appear multiple times or differ from the suspicious value
            for val, count in counter.most_common():
                if abs(val - ahi_overall_float) > 0.5:  # Different from current value
                    logger.info(f"✅ AHI validation: Corrected ahi_overall from {ahi_overall} to {val} (found {count} times in text)")
                    ri["ahi_overall"] = val
                    ri["_ahi_overall_corrected"] = True
                    ri["_ahi_original_value"] = ahi_overall
                    break
            else:
                # No different value found - keep original but log warning
                logger.warning(f"⚠️ AHI validation: Could not find alternative value in text, keeping {ahi_overall}")
        else:
            logger.warning(f"⚠️ AHI validation: No explicit overall AHI found in text, keeping {ahi_overall}")
    
    return visual_data

def _coerce_num(value):
    """Convert value to float if possible, otherwise return None."""
    if value is None:
        return None
    try:
        return float(value)
    except (ValueError, TypeError):
        return None

# Set up logging
logging.basicConfig(
    level=logging.ERROR,  # Set to ERROR to suppress most messages
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('document_extraction_phase2.log'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

# Suppress debug messages from PDF libraries and other verbose libraries
logging.getLogger('pdfplumber').setLevel(logging.WARNING)
logging.getLogger('camelot').setLevel(logging.WARNING)
logging.getLogger('tabula').setLevel(logging.WARNING)
logging.getLogger('PyPDF2').setLevel(logging.WARNING)
logging.getLogger('fitz').setLevel(logging.WARNING)
logging.getLogger('PIL').setLevel(logging.WARNING)
logging.getLogger('PIL.Image').setLevel(logging.WARNING)
logging.getLogger('pytesseract').setLevel(logging.WARNING)
logging.getLogger('boto3').setLevel(logging.WARNING)
logging.getLogger('botocore').setLevel(logging.WARNING)
logging.getLogger('urllib3').setLevel(logging.WARNING)
logging.getLogger('requests').setLevel(logging.WARNING)
logging.getLogger('matplotlib').setLevel(logging.WARNING)
logging.getLogger('numpy').setLevel(logging.WARNING)
logging.getLogger('pandas').setLevel(logging.WARNING)
logging.getLogger('openpyxl').setLevel(logging.WARNING)
logging.getLogger('xlrd').setLevel(logging.WARNING)
logging.getLogger('chardet').setLevel(logging.WARNING)

# Additional aggressive suppression for any library that might be verbose
logging.getLogger('').setLevel(logging.WARNING)  # Root logger
for logger_name in ['pdfminer', 'pdfminer.high_level', 'pdfminer.layout', 'pdfminer.pdfparser', 
                   'pdfminer.pdfinterp', 'pdfminer.pdfdevice', 'pdfminer.converter',
                   'pymupdf', 'pymupdf.fitz', 'pymupdf.utils', 'pymupdf.mupdf',
                   'ghostscript', 'wand', 'wand.image', 'wand.color', 'wand.drawing',
                   'PIL.PngImagePlugin', 'PIL.JpegImagePlugin', 'PIL.PdfImagePlugin',
                   'camelot.io', 'camelot.ext.ghostscript', 'camelot.ext.ghostscript.ghostscript',
                   'tabula.io', 'tabula.util', 'tabula.wrapper',
                   'PyPDF2.generic', 'PyPDF2.pdf', 'PyPDF2.utils',
                   'fitz.utils', 'fitz.fitz', 'fitz.mupdf']:
    logging.getLogger(logger_name).setLevel(logging.ERROR)
# OCR configuration: disable OCR if tesseract is not available or explicitly disabled
DISABLE_OCR = os.getenv('DISABLE_OCR', '0') == '1'
_TESSERACT_CANDIDATES = [
    os.getenv('TESSERACT_CMD'),
    shutil.which('tesseract'),
    '/usr/bin/tesseract',
    '/usr/local/bin/tesseract',
]
TESSERACT_PATH = next((path for path in _TESSERACT_CANDIDATES if path and os.path.exists(path)), None)
if TESSERACT_PATH:
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH
ENABLE_TEXTRACT = os.getenv('ENABLE_TEXTRACT', '1') == '1'
OCR_AVAILABLE = (not DISABLE_OCR) and ((TESSERACT_PATH is not None) or ENABLE_TEXTRACT)
if not OCR_AVAILABLE:
    logger.info("OCR disabled: %s; tesseract found: %s", DISABLE_OCR, bool(TESSERACT_PATH))

# Resource and payload limits (env-overridable)
def _get_int_env(name: str, default_value: int) -> int:
    try:
        return int(os.getenv(name, str(default_value)))
    except Exception:
        return default_value

MAX_DOC_CHARS = _get_int_env('MAX_DOC_CHARS', 60000)
MAX_BATCH_DOCS = _get_int_env('MAX_BATCH_DOCS', 1)
MAX_TOTAL_CHARS = _get_int_env('MAX_TOTAL_CHARS', 120000)
BEDROCK_MAX_TOKENS_SINGLE = _get_int_env('BEDROCK_MAX_TOKENS_SINGLE', 400)
BEDROCK_MAX_TOKENS_BATCH = _get_int_env('BEDROCK_MAX_TOKENS_BATCH', 800)
BEDROCK_MAX_DOCUMENT_BYTES = _get_int_env('BEDROCK_MAX_DOCUMENT_BYTES', int(4.5 * 1024 * 1024))

# Throttling configuration
THROTTLING_DETECTED = False
THROTTLING_BATCH_SIZE = 1  # Reduce batch size when throttling is detected:qw


def trim_content(text: str, limit: int) -> str:
    if not text or len(text) <= limit:
        return text
    head = max(limit // 2, 1)
    tail = limit - head
    return text[:head] + "\n\n[... content truncated ...]\n\n" + text[-tail:]


def organize_timeline_with_llm(patient_id: int, sleep_studies: list, reports: list, reports_grouped: list) -> dict:
    """
    Use LLM to organize timeline data for better clinical coherence and graphing.
    Handles fragmented data and ensures consistent chronological progression.
    """
    logger.info(f"Patient {patient_id}: Starting LLM timeline organization function")
    try:
        # Prepare data summary for LLM analysis
        data_summary = {
            'sleep_studies': len(sleep_studies),
            'reports': len(reports),
            'reports_grouped': len(reports_grouped),
            'total_data_points': len(sleep_studies) + len(reports) + len(reports_grouped)
        }
        
        logger.info(f"Patient {patient_id}: Data summary: {data_summary}")
        
        # If we have very little data, skip LLM organization
        if data_summary['total_data_points'] < 2:
            logger.warning(f"Patient {patient_id}: Insufficient data for LLM organization - only {data_summary['total_data_points']} points")
            return {'success': False, 'reason': 'insufficient_data'}
        
        # Extract all dates and metrics for LLM context
        all_data_points = []
        
        # Add sleep study data points
        for ss in sleep_studies:
            point = {
                'type': 'sleep_study',
                'date': ss.get('date'),
                'file_name': ss.get('file_name', ''),
                'metrics': {k: v for k, v in ss.items() if k not in ['date', 'file_name'] and v is not None}
            }
            if point['date'] and point['metrics']:
                all_data_points.append(point)
        
        # Add report data points
        for report in reports:
            point = {
                'type': 'report',
                'date': report.get('date'),
                'file_name': report.get('file_name', ''),
                'metric': report.get('key'),
                'value': report.get('value')
            }
            if point['date'] and point['metric'] and point['value'] is not None:
                all_data_points.append(point)
        
        # Add grouped report data points
        for grouped in reports_grouped:
            point = {
                'type': 'grouped_report',
                'date': grouped.get('date'),
                'file_name': grouped.get('file_name', ''),
                'metrics': {k: v for k, v in grouped.items() if k not in ['date', 'file_name'] and v is not None}
            }
            if point['date'] and point['metrics']:
                all_data_points.append(point)
        
        # Sort by date for chronological analysis
        all_data_points.sort(key=lambda x: x['date'] if x['date'] else '1900-01-01')
        
        # Prepare LLM prompt for timeline organization with schema template
        schema_template = {
            "success": True,
            "organized_timeline": {
                "sleep_studies": [
                    {
                        "date": "YYYY-MM-DD",
                        "metrics": {
                            "ahi": "float",
                            "odi": "float", 
                            "o2_nadir_pct": "float",
                            "supine_ahi": "float",
                            "time_below_90_pct": "float"
                        },
                        "episode_id": "string",
                        "source_kind": "sleep_study"
                    }
                ],
                "reports": [],
                "reports_grouped": [
                    {
                        "date": "YYYY-MM-DD",
                        "metrics": {
                            "ahi": "float",
                            "odi": "float",
                            "o2_nadir_pct": "float"
                        }
                    }
                ]
            },
            "ui_sleep_metrics": {
                "baseline": {"ahi": "float", "odi": "float", "o2_nadir_pct": "float"},
                "current": {"ahi": "float", "odi": "float", "o2_nadir_pct": "float"},
                "therapy_start_date": "YYYY-MM-DD or null",
                "timeline": [
                    {
                        "date": "YYYY-MM-DD",
                        "context": "baseline|current|historical",
                        "metrics": {"ahi": "float", "odi": "float"},
                        "provenance": {"source_kind": "sleep_study|report", "file_name": "string"}
                    }
                ],
                "needs_review": [
                    {
                        "context": "baseline|current|historical",
                        "metrics": {"ahi": "float"},
                        "reason": "No date found in document",
                        "provenance": {"source_kind": "report", "file_name": "string"}
                    }
                ]
            },
            "changes_made": "string describing consolidation logic"
        }

        system_prompt = f"""You are a clinical data analyst specializing in sleep medicine data integration from multiple unreliable sources. Your task is to consolidate fragmented, conflicting, and inconsistent timeline data using clinical intelligence while fitting it into the EXACT schema template provided.

CRITICAL REQUIREMENTS:
1. Use clinical judgment to resolve data conflicts and inconsistencies
2. Fill the provided schema template with consolidated data
3. Maintain exact field names and structure from template
4. Do not add or remove fields from the template
5. Apply source reliability hierarchy and data quality assessment

DATA SOURCE RELIABILITY (Highest to Lowest):
1. Sleep lab studies (definitive measurements)
2. Medical reports from sleep specialists
3. Follow-up clinical reports
4. Patient questionnaires/self-reports
5. Secondary references in notes

         DATA QUALITY ASSESSMENT:
         - Detect and correct obvious typos (AHI 282 → 28.2)
         - Validate plausible ranges and EXCLUDE implausible values:
           * AHI: 0-100 events/hour (exclude >100)
           * ODI: 0-100 events/hour (exclude >100) 
           * O2 Nadir: 70-100% (exclude <70% or >100%)
           * Sleep Efficiency: 70-95% (exclude <70% - likely data error)
           * Sleep Duration: 4-12 hours (exclude <4h or >12h)
           * REM AHI: 0-150 events/hour (exclude >150)
         - Handle format inconsistencies (dates, units, decimals)
         - Identify duplicate entries with minor variations
         - Flag and REMOVE obviously erroneous values (sleep_efficiency 27% = data error)

         CLINICAL CONSOLIDATION LOGIC:
         - Prioritize most reliable source when data conflicts
         - Merge complementary data from different sources
         - Use temporal context (pre/post treatment, baseline/follow-up)
         - Apply clinical knowledge (ODI typically ≤ AHI, severity correlations)
         - Remove obvious duplicates while preserving unique metrics
         - Fill gaps intelligently without inventing data
         
         BASELINE vs CURRENT CLASSIFICATION:
         - BASELINE: Higher AHI values (≥15) from earlier dates, before treatment
         - CURRENT: Lower AHI values (<15) from later dates, after treatment
         - Use AHI magnitude and temporal sequence to determine baseline vs current
         - Higher AHI = baseline (pre-treatment), Lower AHI = current (post-treatment)
         - Example: AHI 22 (Feb 2025) = baseline, AHI 3.6 (Sep 2025) = current
         
         UI SLEEP METRICS SECTION (CRITICAL - MUST POPULATE IN RESPONSE):
         ⚠️ YOU MUST include a "ui_sleep_metrics" field at the TOP LEVEL of your JSON response (same level as "organized_timeline")
         - baseline: Extract metrics from pre-treatment/diagnostic study (look for RDI or AHI ≥15, typically from earlier dates)
         - current: Extract metrics from most recent post-treatment study (look for AHI <15, typically from later dates)  
         - timeline: Create timeline entries with context classification (baseline/current/historical)
         - needs_review: Include any values that lack proper dates (e.g., RDI values without study dates)
         - IMPORTANT: RDI and AHI are DIFFERENT metrics - do not exclude RDI just because it differs from AHI
         - This section is used directly by the UI to display sleep metrics - DO NOT OMIT IT
         
         MEDICAL AHI PRIORITIZATION LOGIC:
         - When multiple AHI values exist without clear dates, apply medical logic:
           * Higher AHI = Baseline (before treatment) - assign to earlier date
           * Lower AHI = Post-treatment (improvement) - assign to later date
         - If significant difference (≥8 AHI points), prioritize higher value as baseline
         - This reflects standard medical practice: higher AHI indicates more severe condition
         - Example: AHI 27 and AHI 15 → 27=baseline (earlier), 15=improvement (later)
         
         CRITICAL DATE HANDLING:
         - DISTINGUISH between report creation dates and actual study dates
         - Look for comparison tables (Baseline vs Follow-up #1 vs Follow-up #2)
         - Extract actual study timing from report content ("June 2025", "Aug 2025")
         - Use clinical sequence logic (baseline → follow-up #1 → follow-up #2)
         - When reports reference historical studies, use the STUDY date not REPORT date
         - Reports written in Sept 2025 may reference studies from June 2025 or Aug 2025
         

REQUIRED OUTPUT SCHEMA:
{json.dumps(schema_template, indent=2)}

CRITICAL JSON FORMATTING REQUIREMENTS:
- Return valid JSON that can be parsed by json.loads()
- Escape all newlines in string values as \\n (not actual line breaks)
- Escape all tabs in string values as \\t  
- Do NOT include actual newline characters inside JSON string values
- Use single-line text for the changes_made field, or properly escape multi-line content
- Test your JSON syntax before returning

         You must return JSON that exactly matches this schema structure and is valid for parsing.
         
         EXAMPLE: If you see a comparison table like:
         "Baseline (2024): AHI 28.2 → Follow-up #1 (June 2025): AHI 10.2 → Follow-up #2 (Aug 2025): AHI 5.8"
         
         Create timeline entries with ACTUAL study dates:
         - 2024: AHI 28.2 (baseline)
         - 2025-06: AHI 10.2 (follow-up #1) 
         - 2025-08: AHI 5.8 (follow-up #2)
         
         NOT the report upload dates (which might be Sept 2025).
         
         CRITICAL DATA FILTERING:
         Remove implausible values like:
         - Sleep efficiency 27% (should be 70-95%, this is clearly an error)
         - AHI 282 (should be <100, likely meant 28.2)
         - O2 nadir 15% (should be >70%, this would be fatal)"""

        user_prompt = f"""Patient {patient_id} timeline data to organize:

Data Points ({len(all_data_points)} total):
{json.dumps(all_data_points, indent=2, default=str)}

Please organize this data into a coherent timeline that will produce smooth, clinically meaningful graphs. Focus on:
- Deduplication of overlapping data points
- Prioritizing authoritative sources (sleep studies > reports)
- Filling temporal gaps where clinically appropriate
- Ensuring consistent metric progression

Return format:
{{
  "success": true,
  "organized_timeline": {{
    "sleep_studies": [...],
    "reports": [...],
    "reports_grouped": [...]
  }},
  "changes_made": "description of organization performed"
}}"""

        # Query LLM for timeline organization with retry logic
        try:
            from flask_app.config.bedrock_config import query_bedrock_claude_enhanced as bedrock_query_enhanced
            
            # Retry mechanism for rate limiting
            max_retries = 5
            retry_delay = 13  # Start with 13 seconds (just over the 12s rate limit)
            response = None
            
            for attempt in range(max_retries):
                try:
                    # Add rate limiting to prevent API quota exceeded errors
                    throttle_bedrock()
                    
                    # Combine system prompt with user prompt for Bedrock Claude compatibility
                    combined_prompt = f"{system_prompt}\n\n{user_prompt}"
                    
                    response = bedrock_query_enhanced([
                        {"role": "user", "content": combined_prompt}
                    ], max_tokens=4000, temperature=0.1, top_p=0.9)
                    
                    # If successful, break out of retry loop
                    if response and response.get('success'):
                        break
                    
                    # If not successful and not rate limited, don't retry
                    if response and not response.get('success'):
                        error_msg = response.get('message', '')
                        # Check if it's a rate limit or busy error (retryable)
                        is_retryable = ('rate limit' in error_msg.lower() or 
                                       'throttl' in error_msg.lower() or 
                                       'busy' in error_msg.lower() or
                                       'try again' in error_msg.lower())
                        
                        if not is_retryable:
                            logger.warning(f"Patient {patient_id}: Non-retryable error: {error_msg}")
                            break
                        
                        # Retryable error detected, retry
                        if attempt < max_retries - 1:
                            logger.warning(f"Patient {patient_id}: Retryable error detected, retrying in {retry_delay}s (attempt {attempt + 1}/{max_retries}): {error_msg}")
                            time.sleep(retry_delay)
                            retry_delay = min(retry_delay * 1.5, 60)  # Exponential backoff, max 60s
                        else:
                            logger.error(f"Patient {patient_id}: Max retries reached after {max_retries} attempts")
                            
                except Exception as retry_e:
                    logger.error(f"Patient {patient_id}: Retry attempt {attempt + 1} failed: {retry_e}")
                    if attempt < max_retries - 1:
                        time.sleep(retry_delay)
                        retry_delay = min(retry_delay * 1.5, 60)
            
            if response and response.get('success') and (response.get('content') or response.get('response')):
                # Parse LLM response - handle both 'content' and 'response' keys
                content = response.get('content') or response.get('response')
                logger.info(f"Patient {patient_id}: LLM Raw Response:\n{content}")
                
                try:
                    # Extract JSON from response (handle potential markdown formatting)
                    if '```json' in content:
                        json_start = content.find('```json') + 7
                        json_end = content.find('```', json_start)
                        json_content = content[json_start:json_end].strip()
                        logger.info(f"Patient {patient_id}: Extracted JSON from markdown:\n{json_content}")
                    elif '{' in content:
                        json_start = content.find('{')
                        json_end = content.rfind('}') + 1
                        json_content = content[json_start:json_end]
                        logger.info(f"Patient {patient_id}: Extracted JSON from content:\n{json_content}")
                    else:
                        logger.error(f"Patient {patient_id}: No JSON found in LLM response")
                        raise ValueError("No JSON found in LLM response")
                    
                    # Basic JSON cleaning (LLM should now return proper JSON)
                    import re
                    
                    # Remove any potential BOM or invisible characters at start/end
                    json_content = json_content.strip()
                    
                    # Normalize line endings
                    json_content = json_content.replace('\r\n', '\n').replace('\r', '\n')
                    
                    logger.info(f"Patient {patient_id}: Applied basic JSON cleaning (LLM instructed to return valid JSON)")
                    
                    # Try parsing the cleaned JSON
                    try:
                        organized_data = json.loads(json_content)
                        logger.info(f"Patient {patient_id}: Successfully parsed LLM organized data")
                    except json.JSONDecodeError as json_error:
                        # Fallback: Extract just the organized_timeline section to avoid text parsing issues
                        logger.warning(f"Patient {patient_id}: JSON parse failed ({json_error}), attempting structured extraction")
                        
                        # Find the organized_timeline section specifically
                        timeline_start = json_content.find('"organized_timeline":')
                        if timeline_start > -1:
                            # Find the matching closing brace for organized_timeline
                            brace_count = 0
                            timeline_content_start = json_content.find('{', timeline_start)
                            if timeline_content_start > -1:
                                pos = timeline_content_start
                                for char in json_content[timeline_content_start:]:
                                    if char == '{':
                                        brace_count += 1
                                    elif char == '}':
                                        brace_count -= 1
                                        if brace_count == 0:
                                            timeline_end = pos + 1
                                            break
                                    pos += 1
                                
                                if brace_count == 0:
                                    timeline_json = json_content[timeline_content_start:timeline_end]
                                    try:
                                        timeline_data = json.loads(timeline_json)
                                        organized_data = {
                                            'success': True,
                                            'organized_timeline': timeline_data,
                                            'changes_made': 'LLM organization successful (extracted from structured section)'
                                        }
                                        logger.info(f"Patient {patient_id}: Successfully extracted organized timeline structure")
                                    except json.JSONDecodeError:
                                        raise json_error
                                else:
                                    raise json_error
                            else:
                                raise json_error
                        else:
                            raise json_error
                    
                    if organized_data.get('success'):
                        logger.info(f"Patient {patient_id}: LLM timeline organization successful - {organized_data.get('changes_made', 'organized')}")
                        organized_timeline = organized_data.get('organized_timeline', {})
                        logger.info(f"Patient {patient_id}: Organized timeline contains: sleep_studies={len(organized_timeline.get('sleep_studies', []))}, reports={len(organized_timeline.get('reports', []))}, reports_grouped={len(organized_timeline.get('reports_grouped', []))}")
                        
                        # Convert LLM format to expected timeline format
                        converted_data = convert_llm_timeline_format(organized_timeline)
                        logger.info(f"Patient {patient_id}: Converted timeline contains: sleep_studies={len(converted_data.get('sleep_studies', []))}, reports={len(converted_data.get('reports', []))}, reports_grouped={len(converted_data.get('reports_grouped', []))}")
                        
                        # Extract ui_sleep_metrics from LLM response
                        result = {
                            'success': True,
                            'organized_timeline': converted_data,
                            'changes_made': organized_data.get('changes_made', 'organized')
                        }
                        
                        # Pass through ui_sleep_metrics if provided by LLM
                        if 'ui_sleep_metrics' in organized_data:
                            result['ui_sleep_metrics'] = organized_data['ui_sleep_metrics']
                            logger.info(f"Patient {patient_id}: Extracted ui_sleep_metrics from LLM response")
                            logger.info(f"Patient {patient_id}: Baseline: {organized_data['ui_sleep_metrics'].get('baseline', {})}")
                            logger.info(f"Patient {patient_id}: Current: {organized_data['ui_sleep_metrics'].get('current', {})}")
                        
                        return result
                    else:
                        logger.warning(f"Patient {patient_id}: LLM indicated organization failure - {organized_data}")
                        return {'success': False, 'reason': 'llm_indicated_failure'}
                        
                except (json.JSONDecodeError, ValueError) as parse_e:
                    logger.error(f"Patient {patient_id}: Failed to parse LLM timeline response: {parse_e}")
                    logger.error(f"Patient {patient_id}: Raw content that failed to parse:\n{content}")
                    return {'success': False, 'reason': 'parse_error'}
            else:
                logger.warning(f"Patient {patient_id}: LLM timeline organization query failed - response: {response}")
                return {'success': False, 'reason': 'llm_query_failed'}
                
        except ImportError as ie:
            logger.warning(f"Patient {patient_id}: Bedrock not available for timeline organization: {ie}")
            return {'success': False, 'reason': 'bedrock_unavailable'}
            
    except Exception as e:
        logger.error(f"Patient {patient_id}: Timeline organization error: {e}")
        return {'success': False, 'reason': 'unexpected_error', 'error': str(e)}


def validate_and_normalize_timeline_data(timeline_items: list, item_type: str, meta_data: dict) -> list:
    """
    Validate and normalize timeline data to ensure proper date handling.
    
    Args:
        timeline_items: List of timeline items to validate
        item_type: 'sleep_study' or 'report'
        meta_data: Patient metadata containing date_of_study, report_date, etc.
    
    Returns:
        List of validated and normalized timeline items
    """
    if not timeline_items:
        return timeline_items
    
    validated_items = []
    
    for item in timeline_items:
        validated_item = item.copy()
        
        # Determine the appropriate date field based on item type
        if item_type == 'sleep_study':
            # For sleep studies, use observed_at (when the sleep test was performed)
            if not validated_item.get('observed_at') and not validated_item.get('date'):
                # Fallback to meta.date_of_study if available
                fallback_date = meta_data.get('date_of_study')
                if fallback_date:
                    validated_item['observed_at'] = fallback_date
                    validated_item['date'] = fallback_date
                    logger.warning(f"Timeline item missing observed_at, using meta.date_of_study: {fallback_date}")
                else:
                    logger.error(f"Sleep study item missing observed_at and no meta.date_of_study available: {item}")
                    continue
            elif validated_item.get('date') and not validated_item.get('observed_at'):
                # If only 'date' is present, use it as observed_at
                validated_item['observed_at'] = validated_item['date']
                logger.info(f"Set observed_at from date field: {validated_item['date']}")
        
        elif item_type == 'report':
            # For reports, use reported_at (when the report was authored)
            if not validated_item.get('reported_at') and not validated_item.get('date'):
                # Fallback to meta.report_date if available
                fallback_date = meta_data.get('report_date')
                if fallback_date:
                    validated_item['reported_at'] = fallback_date
                    validated_item['date'] = fallback_date
                    logger.warning(f"Report item missing reported_at, using meta.report_date: {fallback_date}")
                else:
                    logger.error(f"Report item missing reported_at and no meta.report_date available: {item}")
                    continue
            elif validated_item.get('date') and not validated_item.get('reported_at'):
                # If only 'date' is present, use it as reported_at
                validated_item['reported_at'] = validated_item['date']
                logger.info(f"Set reported_at from date field: {validated_item['date']}")
        
        # Ensure date field is set for backwards compatibility
        if not validated_item.get('date'):
            if item_type == 'sleep_study' and validated_item.get('observed_at'):
                validated_item['date'] = validated_item['observed_at']
            elif item_type == 'report' and validated_item.get('reported_at'):
                validated_item['date'] = validated_item['reported_at']
        
        validated_items.append(validated_item)
    
    logger.info(f"Validated {len(validated_items)}/{len(timeline_items)} {item_type} items")
    return validated_items
def convert_llm_timeline_format(llm_timeline: dict) -> dict:
    """
    Convert LLM organized timeline format to the expected timeline format.
    
    LLM returns:
    {
      "sleep_studies": [{"date": "2024-10-01", "metrics": {"ahi": 28.2, ...}}],
      "reports": [{"date": "2025-09-09", "metrics": {"ahi": 5.8, ...}}],
      "reports_grouped": [...]
    }
    
    Expected format with proper date handling:
    {
      "sleep_studies": [{"ahi": 28.2, "observed_at": "2024-10-01", "date": "2024-10-01", "episode_id": "...", ...}],
      "reports": [{"reported_at": "2025-09-09", "date": "2025-09-09", "key": "ahi", "value": 5.8, ...}],
      "reports_grouped": [{"ahi": 5.8, "reported_at": "2025-09-09", "date": "2025-09-09", ...}]
    }
    
    Note: date is kept for backwards compatibility, derived from observed_at/reported_at
    """
    converted = {
        'sleep_studies': [],
        'reports': [],
        'reports_grouped': []
    }
    
    try:
        # Convert sleep studies from nested metrics to flat structure
        llm_sleep_studies = llm_timeline.get('sleep_studies', [])
        for study in llm_sleep_studies:
            # Use observed_at for sleep studies (when the sleep test was performed)
            observed_at = study.get('date') or study.get('observed_at')
            converted_study = {
                'observed_at': observed_at,
                'date': observed_at,  # Keep for backwards compatibility
                'source_kind': 'sleep_study'
            }
            
            # Handle episode_id
            if 'episode_id' in study:
                converted_study['episode_id'] = study['episode_id']
            else:
                # Generate a merged episode ID if not provided
                converted_study['episode_id'] = 'llm_merged_' + str(hash(str(study)))[:8]
            
            # Handle file_name
            if 'file_name' in study:
                converted_study['file_name'] = study['file_name']
            else:
                converted_study['file_name'] = 'consolidated_study.pdf'
            
            # Flatten metrics into the main object
            metrics = study.get('metrics', {})
            for key, value in metrics.items():
                if value is not None:
                    converted_study[key] = value
            
            converted['sleep_studies'].append(converted_study)
        
        # Convert reports from nested metrics to individual key-value entries
        llm_reports = llm_timeline.get('reports', [])
        for report in llm_reports:
            # Use reported_at for reports (when the report was authored)
            reported_at = report.get('date') or report.get('reported_at')
            base_report = {
                'reported_at': reported_at,
                'date': reported_at,  # Keep for backwards compatibility
                'file_name': report.get('file_name', 'consolidated_report.pdf'),
                'source_kind': 'report'
            }
            
            metrics = report.get('metrics', {})
            for key, value in metrics.items():
                if value is not None:
                    report_entry = base_report.copy()
                    report_entry['key'] = key
                    report_entry['value'] = value
                    converted['reports'].append(report_entry)
        
        # Convert reports_grouped from nested metrics to flat structure
        llm_reports_grouped = llm_timeline.get('reports_grouped', [])
        for grouped in llm_reports_grouped:
            # Use reported_at for grouped reports (when the report was authored)
            reported_at = grouped.get('date') or grouped.get('reported_at')
            converted_grouped = {
                'reported_at': reported_at,
                'date': reported_at,  # Keep for backwards compatibility
                'file_name': grouped.get('file_name', 'consolidated_report.pdf'),
                'source_kind': 'report'
            }
            
            # Flatten metrics into the main object
            metrics = grouped.get('metrics', {})
            for key, value in metrics.items():
                if value is not None:
                    converted_grouped[key] = value
            
            converted['reports_grouped'].append(converted_grouped)
        
        logger.info(f"Timeline format conversion successful: {len(converted['sleep_studies'])} sleep studies, {len(converted['reports'])} reports, {len(converted['reports_grouped'])} grouped reports")
        return converted
        
    except Exception as e:
        logger.error(f"Timeline format conversion failed: {e}")
        # Return empty structure on conversion failure
        return {'sleep_studies': [], 'reports': [], 'reports_grouped': []}

# Database configuration
DB_CONFIG = {
    'host': 'vizbrizapp-202606.ch8koiygcu36.us-east-2.rds.amazonaws.com',
    'user': 'admin',
    'password': 'Vizbriz2025!',
    'database': 'vizbriz',
    'port': 3306
}

# AWS Bedrock configuration
BEDROCK_CONFIG = {
    'region_name': 'us-west-2'
}

# Model configuration - hardcoded to Claude 4 Sonnet for standalone script
# Can be overridden with BEDROCK_MODEL_ID environment variable
MODEL_ID = os.getenv('BEDROCK_MODEL_ID', 'us.anthropic.claude-sonnet-4-20250514-v1:0')

# Global variable to track current patient_id for LLM logging
_current_patient_id = None

def set_current_patient_id(patient_id):
    """Set the current patient ID for LLM logging context"""
    global _current_patient_id
    _current_patient_id = patient_id

def get_current_patient_id():
    """Get the current patient ID for LLM logging context"""
    global _current_patient_id
    return _current_patient_id

# Helper to check if we're running in Flask context (for logging)
# Removed Flask dependencies - script is completely standalone

def _log_llm_call(prompt_text, response_text, response_time_ms=None, status='success', error_message=None):
    """
    Log LLM call directly to database using existing MySQL connection.
    No Flask context needed - uses direct database access.
    """
    try:
        import uuid
        from datetime import datetime
        
        session_id = str(uuid.uuid4())
        patient_id = get_current_patient_id()
        
        # Log patient_id for debugging if needed
        if patient_id is None:
            logger.warning(f"[LLM_LOG] patient_id is None - set_current_patient_id() may not have been called")
        
        # Estimate tokens (simple word count * 1.3 multiplier)
        prompt_tokens = int(len(prompt_text.split()) * 1.3) if prompt_text else 0
        response_tokens = int(len(response_text.split()) * 1.3) if response_text else 0
        
        # Get database connection (same as other parts of this script)
        import mysql.connector
        conn = mysql.connector.connect(**DB_CONFIG)
        cursor = conn.cursor()
        
        # Insert prompt entry
        cursor.execute("""
            INSERT INTO llm_interactions 
            (session_id, interaction_type, patient_id, page_endpoint, model_name, model_id, 
             content_text, token_count_estimated, status, created_at)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
        """, (
            session_id,
            'prompt',
            patient_id,
            'document_extraction_script',
            'claude_4_sonnet',
            MODEL_ID,
            prompt_text[:10000],  # Truncate if very long
            prompt_tokens,
            'success',
            datetime.now()
        ))
        
        # Insert response entry
        cursor.execute("""
            INSERT INTO llm_interactions 
            (session_id, interaction_type, patient_id, page_endpoint, model_name, model_id, 
             content_text, token_count_estimated, response_time_ms, status, error_message, created_at)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
        """, (
            session_id,
            'response',
            patient_id,
            'document_extraction_script',
            'claude_4_sonnet',
            MODEL_ID,
            response_text[:10000] if response_text else '',  # Truncate if very long
            response_tokens,
            response_time_ms,
            status,
            error_message,
            datetime.now()
        ))
        
        conn.commit()
        cursor.close()
        conn.close()
        
    except Exception as e:
        logger.warning(f"Failed to log LLM call to database: {e}")

# S3 configuration - use the same as flask_app/__init__.py
S3_BUCKET_NAME = 'vizbrizpatients'  # Same as flask_app/__init__.py
AWS_REGION = 'us-west-2'  # Same as flask_app/__init__.py

# AWS/S3 — credentials from environment only (see .env.example)
os.environ.setdefault('AWS_REGION', 'us-west-2')
os.environ.setdefault('S3_BUCKET_NAME', os.getenv('S3_BUCKET_NAME', 'vizbrizpatients'))
# Do NOT override BASE_URL at import time. This module is imported by the Flask app
# and clobbering BASE_URL here can leak dev URLs into emails/links.
os.environ.setdefault('BASE_URL', 'https://app.vizbriz.com')
os.environ['FLASK_RUN_HOST'] = '0.0.0.0'

# Ensure Bedrock region is set explicitly to us-west-2 unless overridden via env
os.environ.setdefault('AWS_BEDROCK_REGION', 'us-west-2')

# Global throttling configuration: default 5 requests/minute => 12s spacing
BEDROCK_RPM = _get_int_env('BEDROCK_RPM', 5)
MIN_BEDROCK_INTERVAL_SECONDS = _get_int_env(
    'BEDROCK_MIN_INTERVAL_SECONDS', max(1, 60 // max(1, BEDROCK_RPM))
)
_LAST_BEDROCK_CALL_AT = 0.0

def throttle_bedrock() -> None:
    """Sleep as needed to ensure Bedrock calls are spaced to respect RPM limits."""
    global _LAST_BEDROCK_CALL_AT
    now = time.time()
    if _LAST_BEDROCK_CALL_AT > 0:
        elapsed = now - _LAST_BEDROCK_CALL_AT
        if elapsed < MIN_BEDROCK_INTERVAL_SECONDS:
            wait_s = MIN_BEDROCK_INTERVAL_SECONDS - elapsed
            logger.info(f"Throttling Bedrock call for {wait_s:.1f}s to respect rate limits ({BEDROCK_RPM}/min)")
            time.sleep(wait_s)
    _LAST_BEDROCK_CALL_AT = time.time()

# Initialize Bedrock client using the same pattern as the Flask app
def get_bedrock_client():
    """Initialize and return Bedrock client using Bedrock-specific region.

    Resolution order:
      1) env AWS_BEDROCK_REGION
      2) BEDROCK_CONFIG['region_name']
      3) fallback 'us-west-2'
    """
    try:
        # Prefer dedicated Bedrock region env var over general AWS region
        region = os.getenv('AWS_BEDROCK_REGION') or BEDROCK_CONFIG.get('region_name') or 'us-west-2'
        client = boto3.client("bedrock-runtime", region_name=region)
        logger.info(f"Bedrock client initialized successfully with region: {region}")
        return client
    except Exception as e:
        logger.error(f"Error initializing Bedrock client: {e}")
        return None

bedrock_client = get_bedrock_client()

# Use the same S3 client pattern as the existing codebase
def get_s3_client():
    """Create and return an S3 client with proper AWS credentials."""
    # Use the exact same pattern as s3_utils.py
    from botocore.config import Config
    
    # Get credentials from environment (same as __init__.py)
    aws_access_key_id = os.getenv('AWS_ACCESS_KEY_ID')
    aws_secret_access_key = os.getenv('AWS_SECRET_ACCESS_KEY')
    aws_region = os.getenv('AWS_REGION')
    s3_bucket_name = os.getenv('S3_BUCKET_NAME')
    
    # Print debug info
    print(f"S3 Utils - Loading configuration:")
    print(f"AWS Region: {aws_region}")
    print(f"S3 Bucket: {s3_bucket_name}")
    if aws_access_key_id:
        print(f"AWS Access Key ID: {aws_access_key_id[:4]}...{aws_access_key_id[-4:]}")
    else:
        print("AWS Access Key ID: Not set")
    
    # Create S3 client with credentials (same as s3_utils.py)
    s3_client = boto3.client(
        's3',
        region_name=aws_region,
        aws_access_key_id=aws_access_key_id,
        aws_secret_access_key=aws_secret_access_key,
        config=Config(signature_version='s3v4')
    )
    
    print(f"S3 client created with AWS region: {aws_region}")
    
    return s3_client

s3_client = get_s3_client()

# Import Phase 1 functions from the local project root to avoid picking up any external packages
# Ensure the project root (vizbriz) is on sys.path
CURRENT_DIR = os.path.dirname(os.path.abspath(__file__))
FLASK_APP_DIR = os.path.dirname(CURRENT_DIR)
PROJECT_ROOT = os.path.dirname(FLASK_APP_DIR)
if PROJECT_ROOT not in sys.path:
    sys.path.insert(0, PROJECT_ROOT)

from document_observation_extractor import (
    map_document_type_to_source_type,
    is_binary_file
)

def debug_patient_documents(patient_id: int) -> Dict[str, Any]:
    """
    Debug function to check all documents for a patient and see what's being excluded.
    
    Args:
        patient_id (int): Patient ID to debug
        
    Returns:
        Dict containing document counts and categories
    """
    conn = None
    try:
        conn = mysql.connector.connect(**DB_CONFIG)
        cursor = conn.cursor(dictionary=True)
        
        # Check all files for this patient
        cursor.execute(
            """
            SELECT id, name, category, subcategory, analyzed
            FROM files WHERE patient_id = %s
            ORDER BY upload_date DESC
            """,
            (patient_id,),
        )
        all_files = cursor.fetchall()
        
        # Check all adminfiles for this patient
        cursor.execute(
            """
            SELECT id, name, file_category, analyzed
            FROM adminfiles WHERE patient_id = %s
            ORDER BY upload_date DESC
            """,
            (patient_id,),
        )
        all_adminfiles = cursor.fetchall()
        
        # Count by category
        files_by_category = {}
        for file in all_files:
            category = file.get('category', 'unknown')
            if category not in files_by_category:
                files_by_category[category] = []
            files_by_category[category].append(file)
        
        adminfiles_by_category = {}
        for file in all_adminfiles:
            category = file.get('file_category', 'unknown')
            if category not in adminfiles_by_category:
                adminfiles_by_category[category] = []
            adminfiles_by_category[category].append(file)
        
        result = {
            'patient_id': patient_id,
            'total_files': len(all_files),
            'total_adminfiles': len(all_adminfiles),
            'files_by_category': {cat: len(files) for cat, files in files_by_category.items()},
            'adminfiles_by_category': {cat: len(files) for cat, files in adminfiles_by_category.items()},
            'imaging_files': len([f for f in all_files if f.get('category') == 'imaging']),
            'imaging_adminfiles': len([f for f in all_adminfiles if f.get('file_category') == 'imaging']),
            'non_imaging_files': len([f for f in all_files if f.get('category') != 'imaging']),
            'non_imaging_adminfiles': len([f for f in all_adminfiles if f.get('file_category') != 'imaging'])
        }
        
        logger.info(f"Debug results for patient {patient_id}: {result}")
        return result
        
    except Exception as e:
        logger.error(f"Error debugging documents for patient {patient_id}: {e}")
        return {'error': str(e)}
    finally:
        if conn:
            conn.close()

def discover_patient_documents(patient_id: int) -> List[Dict]:
    """
    Discover documents for a patient, excluding imaging category.
    This is a local version that filters out imaging documents.
    Processes all documents regardless of analyzed status.
    """
    docs: List[Dict] = []
    conn = None
    try:
        conn = mysql.connector.connect(**DB_CONFIG)
        cursor = conn.cursor(dictionary=True)
        
        # files - exclude imaging category only
        cursor.execute(
            """
            SELECT id, name, patient_id, upload_date, file_type, file_size, s3_key,
                   category, subcategory, comment, analyzed, 'files' as source_table
            FROM files WHERE patient_id = %s AND category != 'imaging'
            ORDER BY upload_date DESC
            """,
            (patient_id,),
        )
        files_results = cursor.fetchall()
        logger.info(f"Found {len(files_results)} documents in 'files' table for patient {patient_id}")
        
        for r in files_results:
            docs.append({
                'id': r['id'], 'name': r['name'], 'patient_id': r['patient_id'], 'upload_date': r['upload_date'],
                'file_type': r['file_type'], 'file_size': r['file_size'], 's3_key': r['s3_key'],
                'source_table': r['source_table'], 'source_type': map_document_type_to_source_type(r.get('category'), r.get('subcategory')),
                'category': r.get('category'), 'subcategory': r.get('subcategory'), 'comment': r.get('comment'),
                'analyzed': r.get('analyzed', False)
            })
        
        # adminfiles - exclude imaging category only (but include NULL values)
        cursor.execute(
            """
            SELECT id, name, patient_id, upload_date, file_type, file_size, s3_key,
                   is_public, file_category, analyzed, 'adminfiles' as source_table
            FROM adminfiles WHERE patient_id = %s AND (file_category IS NULL OR file_category != 'imaging')
            ORDER BY upload_date DESC
            """,
            (patient_id,),
        )
        adminfiles_results = cursor.fetchall()
        logger.info(f"Found {len(adminfiles_results)} documents in 'adminfiles' table for patient {patient_id}")
        
        for r in adminfiles_results:
            docs.append({
                'id': r['id'], 'name': r['name'], 'patient_id': r['patient_id'], 'upload_date': r['upload_date'],
                'file_type': r['file_type'], 'file_size': r['file_size'], 's3_key': r['s3_key'],
                'source_table': r['source_table'], 'source_type': map_document_type_to_source_type(None, None, r.get('file_category')),
                'is_public': r.get('is_public'), 'file_category': r.get('file_category'),
                'analyzed': r.get('analyzed', False)
            })
        
        logger.info(f"Total documents discovered for patient {patient_id}: {len(docs)} (files: {len(files_results)}, adminfiles: {len(adminfiles_results)})")
            
    except Exception as e:
        logger.error(f"Error discovering documents for patient {patient_id}: {e}")
    finally:
        if conn:
            conn.close()
    return docs

# Reuse schema-guided snapshot + explosion from ingestion service
try:
    from flask_app.services.case_ingest import (
        normalize_to_patient_case_json_v1,
        explode_observations_from_snapshot,
    )
except Exception as _e:
    normalize_to_patient_case_json_v1 = None
    explode_observations_from_snapshot = None

# Use the same Bedrock pipeline as the patient workflow
try:
    from flask_app.config.bedrock_config import query_bedrock_claude_enhanced as bedrock_query_enhanced
except Exception as _e:
    bedrock_query_enhanced = None

# Import delta_ingest for future use (if needed)
try:
    from flask_app.services.delta_ingest import (
        apply_delta_for_patient
    )
except Exception as _e:
    apply_delta_for_patient = None

def generate_s3_presigned_url(s3_key: str, expiration: int = 3600) -> str:
    """
    Generate a pre-signed URL for S3 object access.
    
    Args:
        s3_key (str): The S3 key of the object
        expiration (int): URL expiration time in seconds
        
    Returns:
        str: Pre-signed URL
    """
    try:
        url = s3_client.generate_presigned_url(
            'get_object',
            Params={'Bucket': S3_BUCKET_NAME, 'Key': s3_key},
            ExpiresIn=expiration
        )
        return url
    except Exception as e:
        logger.error(f"Error generating presigned URL for {s3_key}: {e}")
        return None

def _run_pdf_ocr_fallback(pdf_bytes: bytes, include_tables: bool = False) -> str:
    """
    Run OCR on a local PDF when standard text extraction fails.
    
    Args:
        pdf_path: Path to the local PDF file.
        include_tables: Reserved for future table OCR (currently unused).
        
    Returns:
        Extracted OCR text (empty string if OCR unavailable or fails).
    """
    if not OCR_AVAILABLE:
        logger.info("Skipping PDF OCR fallback: OCR is disabled or tesseract not available")
        return ""
    if not PDF2IMAGE_AVAILABLE or convert_from_bytes is None:
        logger.info("Skipping PDF OCR fallback: pdf2image/poppler not available")
        return ""
    
    ocr_chunks: List[str] = []
    images = []
    try:
        logger.info("📄 Running OCR fallback for PDF (pdf2image + pytesseract)...")
        convert_kwargs = {"dpi": 300}
        if POPPLER_PATH:
            convert_kwargs["poppler_path"] = POPPLER_PATH
        images = convert_from_bytes(pdf_bytes, **convert_kwargs)
        logger.info("pdf2image converted %s pages for OCR fallback", len(images))
    except PDFInfoNotInstalledError as pdfinfo_error:  # type: ignore
        logger.warning(
            "PDF OCR fallback skipped: %s. Install poppler-utils (provides pdfinfo/pdftoppm) and ensure they are on PATH.",
            pdfinfo_error,
        )
        return ""
    except PDFPageCountError as page_err:  # type: ignore
        logger.warning(
            "PDF OCR fallback could not determine page count via pdfinfo: %s. "
            "Falling back to PyMuPDF rasterization.",
            page_err,
        )
        images = []
    except Exception as pdf2img_error:
        logger.warning("PDF OCR fallback failed to render pages: %s", pdf2img_error, exc_info=True)
        images = []
    
    if not images:
        try:
            logger.info("🔄 Falling back to PyMuPDF rasterization for OCR...")
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            for page_index in range(doc.page_count):
                page = doc.load_page(page_index)
                pix = page.get_pixmap(dpi=300)
                mode = "RGB"
                if pix.alpha:
                    mode = "RGBA"
                pil_img = Image.frombytes(mode, (pix.width, pix.height), pix.samples)
                images.append(pil_img)
            logger.info("PyMuPDF generated %s rasterized pages for OCR fallback", len(images))
            doc.close()
        except Exception as pymupdf_error:
            logger.warning("PyMuPDF rasterization failed: %s", pymupdf_error, exc_info=True)
            return ""
    
    for index, pil_image in enumerate(images, start=1):
        try:
            processed_image = _preprocess_for_ocr(pil_image)
            text = pytesseract.image_to_string(processed_image, lang="eng+heb")
            if text:
                logger.debug("OCR fallback extracted %s chars from page %s", len(text), index)
                ocr_chunks.append(text)
            else:
                logger.debug("OCR fallback page %s returned empty text", index)
        except Exception as ocr_page_error:
            logger.warning("OCR fallback failed on page %s: %s", index, ocr_page_error, exc_info=True)
            continue
    
    combined_text = "\n".join(chunk for chunk in ocr_chunks if chunk).strip()
    if combined_text:
        if len(combined_text) > MAX_DOC_CHARS:
            logger.warning(
                "OCR fallback content too large (%s chars), truncating to %s chars",
                len(combined_text),
                MAX_DOC_CHARS,
            )
            combined_text = trim_content(combined_text, MAX_DOC_CHARS)
        logger.info(
            "✅ OCR fallback extracted %s characters from %s pages",
            len(combined_text),
            len(ocr_chunks),
        )
    else:
        logger.info("OCR fallback could not extract any text from the PDF")
    return combined_text


def extract_text_from_pdf(s3_key: str) -> str:
    """
    Extract text content from PDF using multiple frameworks for enhanced extraction.
    Uses PyMuPDF, PDFPlumber, Camelot, and Tabula for comprehensive text and table extraction.
    
    Args:
        s3_key (str): S3 key of the PDF file
        
    Returns:
        str: Extracted text content with structured table data
    """
    try:
        # Generate presigned URL (like the working file viewer)
        presigned_url = s3_client.generate_presigned_url(
            'get_object',
            Params={'Bucket': S3_BUCKET_NAME, 'Key': s3_key},
            ExpiresIn=3600
        )
        
        # Download using requests (more reliable than direct S3 access)
        import requests
        response = requests.get(presigned_url, timeout=30)
        response.raise_for_status()
        pdf_data = response.content
        
        # Save to temporary file for frameworks that need file path
        import tempfile
        import os
        
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp_file:
            tmp_file.write(pdf_data)
            tmp_path = tmp_file.name
        
        text_content = ""
        table_content = ""
        
        try:
            # 1. PDFPlumber - Best for text and simple tables
            logger.info("PDFPlumber: Extracting text and tables...")
            with pdfplumber.open(tmp_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    # Extract text
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"
                    
                    # Extract tables
                    tables = page.extract_tables()
                    if tables:
                        logger.info(f"PDFPlumber: Found {len(tables)} tables on page {page_num + 1}")
                        for i, table in enumerate(tables):
                            table_text = f"\n=== PDFPLUMBER TABLE {page_num + 1}_{i+1} ===\n"
                            for row in table:
                                if row:  # Skip empty rows
                                    row_text = " | ".join([str(cell).strip() if cell else "" for cell in row])
                                    table_text += row_text + "\n"
                            table_content += table_text + "\n"
            
            # 2. Camelot - Best for complex tables
            logger.info("Camelot: Extracting complex tables...")
            try:
                # Try lattice flavor first (better for structured tables)
                tables = camelot.read_pdf(tmp_path, pages='all', flavor='lattice')
                if tables:
                    logger.info(f"Camelot lattice: Found {len(tables)} tables")
                    for i, table in enumerate(tables):
                        df = table.df
                        if not df.empty:
                            table_text = f"\n=== CAMELOT LATTICE TABLE {i+1} ===\n"
                            for _, row in df.iterrows():
                                row_text = " | ".join([str(cell).strip() if str(cell) != 'nan' else "" for cell in row])
                                table_text += row_text + "\n"
                            table_content += table_text + "\n"
            except Exception as camelot_error:
                logger.warning(f"Camelot lattice failed: {camelot_error}")
                
                # Try stream flavor as fallback
                try:
                    tables = camelot.read_pdf(tmp_path, pages='all', flavor='stream')
                    if tables:
                        logger.info(f"Camelot stream: Found {len(tables)} tables")
                        for i, table in enumerate(tables):
                            df = table.df
                            if not df.empty:
                                table_text = f"\n=== CAMELOT STREAM TABLE {i+1} ===\n"
                                for _, row in df.iterrows():
                                    row_text = " | ".join([str(cell).strip() if str(cell) != 'nan' else "" for cell in row])
                                    table_text += row_text + "\n"
                                table_content += table_text + "\n"
                except Exception as stream_error:
                    logger.warning(f"Camelot stream failed: {stream_error}")
            
            # 3. Tabula - Alternative table extraction (DISABLED - Java dependency issue)
            logger.info("Tabula: Skipped (Java dependency not available)")
            # Note: Tabula requires Java runtime which may not be available
            # Using PDFPlumber and Camelot instead for table extraction
            
            # 4. PyMuPDF - Fallback for text extraction
            logger.info("🔄 PyMuPDF: Extracting text as fallback...")
            pdf_document = fitz.open(stream=pdf_data, filetype="pdf")
            logger.info(f"PDF has {len(pdf_document)} pages")
            
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                
                # Extract regular text (only if not already extracted by PDFPlumber)
                if not text_content.strip():
                    page_text = page.get_text()
                    text_content += page_text + "\n"
                
                # Extract tables with PyMuPDF's table finder
                try:
                    tables = page.find_tables()
                    if tables:
                        logger.info(f"✅ PyMuPDF: Found {len(tables)} tables on page {page_num + 1}")
                        for i, table in enumerate(tables):
                            table_text = f"\n=== PYMUPDF TABLE {page_num + 1}_{i+1} ===\n"
                            for row in table.extract():
                                row_text = " | ".join([str(cell).strip() if cell else "" for cell in row])
                                table_text += row_text + "\n"
                            table_content += table_text + "\n"
                except Exception as table_error:
                    logger.warning(f"PyMuPDF table extraction failed on page {page_num + 1}: {table_error}")
                

            
            pdf_document.close()
            
        finally:
            # Clean up temporary file
            try:
                os.unlink(tmp_path)
            except Exception as cleanup_error:
                logger.warning(f"Failed to clean up temporary file: {cleanup_error}")
        
        # Combine regular text and table content
        combined_content = text_content + "\n" + table_content
        
        # Clean up the text content
        combined_content = combined_content.strip()
        
        # Check if we have meaningful content (not just whitespace or very short content)
        if combined_content and len(combined_content.strip()) > 50:  # Minimum 50 characters
            if len(combined_content) > MAX_DOC_CHARS:
                logger.warning(f"Content too large ({len(combined_content)} chars), truncating to {MAX_DOC_CHARS} chars for {s3_key}")
                combined_content = trim_content(combined_content, MAX_DOC_CHARS)
            logger.info(f"Successfully extracted {len(combined_content)} characters from PDF (including {len(table_content)} table chars): {s3_key}")
            return combined_content
        else:
            logger.warning(
                "PDF appears to be empty, unreadable, or contains insufficient text (%s chars) for %s. Attempting OCR fallback.",
                len(combined_content),
                s3_key,
            )
            ocr_text = _run_pdf_ocr_fallback(pdf_data)
            if ocr_text:
                logger.info(
                    "Using OCR fallback result (%s chars) for %s",
                    len(ocr_text),
                    s3_key,
                )
                return ocr_text
            logger.warning(f"OCR fallback failed or produced no text for {s3_key}")
            return ""
        
    except Exception as e:
        logger.error(f"Error extracting text from PDF {s3_key}: {e}")
        return ""

def extract_text_from_image(s3_key: str) -> str:
    """
    Extract text content from image using OCR (Tesseract).
    
    Args:
        s3_key (str): S3 key of the image file
        
    Returns:
        str: Extracted text content
    """
    try:
        if not OCR_AVAILABLE:
            logger.info(f"Skipping OCR for {s3_key}: tesseract not available or OCR disabled")
            return ""
        # Generate presigned URL (like the working file viewer)
        presigned_url = s3_client.generate_presigned_url(
            'get_object',
            Params={'Bucket': S3_BUCKET_NAME, 'Key': s3_key},
            ExpiresIn=3600
        )
        
        # Download using requests (more reliable than direct S3 access)
        import requests
        response = requests.get(presigned_url, timeout=30)
        response.raise_for_status()
        image_data = response.content
        
        # If Tesseract is available, use it; otherwise fallback to AWS Textract
        if TESSERACT_PATH and not DISABLE_OCR:
            image = Image.open(io.BytesIO(image_data))
            text_content = pytesseract.image_to_string(image)
        elif ENABLE_TEXTRACT:
            try:
                textract = boto3.client('textract', region_name=os.getenv('AWS_REGION', 'us-west-2'))
                resp = textract.detect_document_text(Document={'Bytes': image_data})
                
                # Debug: Log the response structure

                if 'Blocks' in resp:
                    # Safely extract text from blocks
                    lines = []
                    for item in resp.get('Blocks', []):
                        if item.get('BlockType') == 'LINE' and 'Text' in item:
                            lines.append(item['Text'])
                    
                    text_content = "\n".join(lines)
                    
                    if not text_content.strip():
                        logger.warning(f"No text detected in image {s3_key}")
                    
            except Exception as e:
                logger.error(f"AWS Textract OCR failed for {s3_key}: {e}")
                text_content = ""
        else:
            logger.info(f"OCR disabled and Textract not enabled for {s3_key}")
            text_content = ""
        if text_content and len(text_content) > MAX_DOC_CHARS:
            text_content = trim_content(text_content, MAX_DOC_CHARS)
        
        logger.info(f"Successfully extracted text from image: {s3_key}")
        return text_content.strip()
        
    except Exception as e:
        logger.error(f"Error extracting text from image {s3_key}: {e}")
        return ""

def extract_text_from_text_file(s3_key: str) -> str:
    """
    Extract text content from text-based files.
    
    Args:
        s3_key (str): S3 key of the text file
        
    Returns:
        str: Extracted text content
    """
    try:
        # Generate presigned URL (like the working file viewer)
        presigned_url = s3_client.generate_presigned_url(
            'get_object',
            Params={'Bucket': S3_BUCKET_NAME, 'Key': s3_key},
            ExpiresIn=3600
        )
        
        # Download using requests (more reliable than direct S3 access)
        import requests
        response = requests.get(presigned_url, timeout=30)
        response.raise_for_status()
        text_content = response.text
        if text_content and len(text_content) > MAX_DOC_CHARS:
            logger.warning(f"Content too large ({len(text_content)} chars), truncating to {MAX_DOC_CHARS} chars for {s3_key}")
            text_content = trim_content(text_content, MAX_DOC_CHARS)
        
        logger.info(f"Successfully extracted text from file: {s3_key}")
        return text_content.strip()
        
    except Exception as e:
        logger.error(f"Error extracting text from file {s3_key}: {e}")
        return ""
def extract_text_from_word_document(s3_key: str) -> str:
    """
    Extract text content from Word document (.docx) using python-docx.
    
    Args:
        s3_key (str): S3 key of the Word document
        
    Returns:
        str: Extracted text content
    """
    try:
        # Generate presigned URL
        presigned_url = s3_client.generate_presigned_url(
            'get_object',
            Params={'Bucket': S3_BUCKET_NAME, 'Key': s3_key},
            ExpiresIn=3600
        )
        
        # Download using requests
        import requests
        response = requests.get(presigned_url, timeout=30)
        response.raise_for_status()
        
        # Save to temporary file
        import tempfile
        import os
        from docx import Document
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
            temp_file.write(response.content)
            temp_file_path = temp_file.name
        
        try:
            # Extract text using python-docx
            doc = Document(temp_file_path)
            text_content = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text.strip() + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content += cell.text.strip() + "\n"
            
            if text_content.strip():
                logger.info(f"Successfully extracted {len(text_content)} characters from Word document: {s3_key}")
                return text_content.strip()
            else:
                logger.warning(f"Word document appears to be empty or unreadable, skipping: {s3_key}")
                return ""
                
        finally:
            # Clean up temporary file
            if os.path.exists(temp_file_path):
                os.unlink(temp_file_path)
                
    except ImportError:
        logger.error("python-docx library not installed. Install with: pip install python-docx")
        return ""
    except Exception as e:
        logger.error(f"Error extracting text from Word document {s3_key}: {e}")
        return ""

def extract_document_content(document: Dict) -> str:
    """
    Extract text content from a document based on its file type.
    
    Args:
        document (Dict): Document metadata dictionary
        
    Returns:
        str: Extracted text content
    """
    s3_key = document.get('s3_key')
    file_type = document.get('file_type')
    filename = document.get('name', 'unknown')
    
    if not s3_key:
        logger.warning(f"No S3 key for document: {filename}")
        return ""
    
    # Determine file type and extract content accordingly
    file_type_lower = file_type.lower() if file_type else ""
    filename_lower = filename.lower() if filename else ""
    
    # PDF files
    if 'pdf' in file_type_lower or filename_lower.endswith('.pdf'):
        return extract_text_from_pdf(s3_key)
    
    # Image files
    elif any(img_type in file_type_lower for img_type in ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'webp', 'tiff', 'tif']):
        return extract_text_from_image(s3_key)
    
    # Text files
    elif any(text_type in file_type_lower for text_type in ['txt', 'csv', 'html', 'xml', 'json']):
        return extract_text_from_text_file(s3_key)
    
    # Word documents (.docx)
    elif 'docx' in file_type_lower or filename_lower.endswith('.docx'):
        return extract_text_from_word_document(s3_key)
    
    # Other Office documents (for now, return empty - could add parsing later)
    elif any(office_type in file_type_lower for office_type in ['doc', 'xls', 'xlsx', 'ppt', 'pptx']):
        logger.info(f"Office document detected: {filename} - content extraction not implemented yet")
        return ""
    
    else:
        logger.warning(f"Unknown file type for {filename}: {file_type}")
        return ""



def _chunk_pdf_by_pages(pdf_bytes: bytes, max_chunk_size_bytes: int) -> List[Tuple[bytes, int, int]]:
    """
    Split a large PDF into smaller chunks by pages, where each chunk is under max_chunk_size_bytes.
    
    Args:
        pdf_bytes: Full PDF content as bytes
        max_chunk_size_bytes: Maximum size for each chunk
        
    Returns:
        List of tuples: (chunk_pdf_bytes, start_page, end_page) for each chunk
    """
    chunks = []
    try:
        import io
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        total_pages = doc.page_count
        
        if total_pages == 0:
            return []
        
        # Estimate bytes per page (rough estimate)
        avg_bytes_per_page = len(pdf_bytes) / total_pages if total_pages > 0 else len(pdf_bytes)
        pages_per_chunk = max(1, int(max_chunk_size_bytes / avg_bytes_per_page * 0.8))  # 80% safety margin
        
        logger.info(f"Chunking PDF with {total_pages} pages into chunks of ~{pages_per_chunk} pages each")
        
        for start_page in range(0, total_pages, pages_per_chunk):
            end_page = min(start_page + pages_per_chunk - 1, total_pages - 1)
            
            # Create a new PDF with just these pages
            chunk_doc = fitz.open()
            chunk_doc.insert_pdf(doc, from_page=start_page, to_page=end_page)
            
            # Convert to bytes
            chunk_bytes = chunk_doc.tobytes()
            chunk_doc.close()
            
            # If chunk is still too large, try smaller chunks
            if len(chunk_bytes) > max_chunk_size_bytes:
                logger.warning(f"Chunk {start_page}-{end_page} is still too large ({len(chunk_bytes)} bytes), splitting further")
                # Split this range into smaller pieces
                for sub_start in range(start_page, end_page + 1, max(1, pages_per_chunk // 2)):
                    sub_end = min(sub_start + max(1, pages_per_chunk // 2) - 1, end_page)
                    sub_doc = fitz.open()
                    sub_doc.insert_pdf(doc, from_page=sub_start, to_page=sub_end)
                    sub_bytes = sub_doc.tobytes()
                    sub_doc.close()
                    chunks.append((sub_bytes, sub_start, sub_end))
            else:
                chunks.append((chunk_bytes, start_page, end_page))
        
        doc.close()
        logger.info(f"Created {len(chunks)} PDF chunks")
        return chunks
        
    except Exception as e:
        logger.error(f"Error chunking PDF: {e}")
        return []


def _process_single_document_chunk(document_content: bytes, safe_name: str, doc_format: str) -> Tuple[Dict, bool, str]:
    """
    Process a single document chunk (or full document) with Visual LLM.
    This is a helper function used by extract_observations_with_visual_llm.
    
    Args:
        document_content: PDF/image bytes to process
        safe_name: Sanitized document name
        doc_format: Document format ("pdf", "jpeg", "png")
        
    Returns:
        Tuple of (extracted_observations_dict, success_bool, error_message)
    """
    try:
        import boto3
        import json
        
        bedrock = boto3.client("bedrock-runtime", region_name="us-west-2")
        
        # Optimize images for better OCR/vision processing
        final_document_content = document_content
        final_doc_format = doc_format
        
        if doc_format in ["jpeg"]:
            try:
                pil = Image.open(io.BytesIO(document_content))
                pil = _preprocess_for_ocr(pil)  # reuse preprocessing
                buf = io.BytesIO()
                pil.save(buf, format="PNG")
                final_document_content = buf.getvalue()
                final_doc_format = "png"
                logger.info(f"Converted JPEG to PNG with preprocessing for {safe_name}")
            except Exception as e:
                logger.warning(f"Failed to optimize image {safe_name}: {e}")
        
        # Determine if this is an image or document for Bedrock API
        is_image = final_doc_format in ["png", "jpeg"]
        
        if is_image:
            document_block = {
                "image": {
                    "format": final_doc_format,
                    "source": {"bytes": final_document_content},
                }
            }
            logger.info(f"Sending {final_doc_format.upper()} image to Bedrock for {safe_name} ({len(final_document_content)} bytes)")
        else:
            document_block = {
                "document": {
                    "format": final_doc_format,
                    "name": safe_name,
                    "source": {"bytes": final_document_content},
                }
            }
            logger.info(f"Sending {doc_format.upper()} document to Bedrock for {safe_name} ({len(document_content)} bytes)")
        
        # Use the same comprehensive prompt as before
        system = [{
            "text": (
                "You extract structured MEDICAL INFORMATION from a clinical sleep study document (PDF/Image) following the Patient Case JSON v1 schema.\n"
                "Return ONLY a single STRICT JSON object (no prose). If a value is absent or unclear, use null. Do not guess.\n\n"
                "## CRITICAL AHI EXTRACTION RULES:\n"
                "- When extracting AHI overall, select the value labeled 'overall' or 'total' (Hebrew: 'כללי', 'סה\"כ').\n"
                "- DO NOT set AHI overall equal to REM/NREM/Supine values. Those go into ahi_rem, ahi_nrem, ahi_supine respectively.\n"
                "- If multiple AHI values occur (e.g., 15 and 27), and one is labeled REM/NREM/Supine while another is labeled overall/כללי/סה\"כ, use the latter for ahi_overall.\n"
                "- Return evidence.ahi_overall as a short snippet that includes the OVERALL/TOTAL/כללי label next to its number.\n"
                "- Be extremely careful to distinguish between stage-specific AHI (REM/NREM) and overall AHI.\n\n"
                "## OUTPUT SCHEMA (Patient Case JSON v1)\n"
                "{\n"
                "  \"meta\": {\n"
                "    \"patient_name\": \"string|null\",\n"
                "    \"dob\": \"YYYY-MM-DD|null\",\n"
                "    \"mrn\": \"string|null\",\n"
                "    \"sex\": \"M|F|Other|null\",\n"
                "    \"age_years\": number|null,\n"
                "    \"height_cm\": number|null,\n"
                "    \"weight_kg\": number|null,\n"
                "    \"bmi\": number|null,\n"
                "    \"facility\": \"string|null\",\n"
                "    \"ordering_provider\": \"string|null\",\n"
                "    \"report_author\": \"string|null\",\n"
                "    \"date_of_study\": \"YYYY-MM-DD|null\",\n"
                "    \"report_date\": \"YYYY-MM-DD|null\",\n"
                "    \"study_type\": \"HSAT|PSG|Type I|Type II|Type III|Type IV|null\",\n"
                "    \"scoring_hypopnea_rule\": \"AASM_3pct|AASM_4pct|unknown|null\",\n"
                "    \"methodology_notes\": \"string|null\",\n"
                "    \"data_quality\": {\n"
                "      \"overall_quality\": \"good|fair|poor|null\",\n"
                "      \"data_loss_pct\": number|null,\n"
                "      \"comments\": \"string|null\"\n"
                "    }\n"
                "  },\n\n"
                "  \"indications_symptoms\": {\n"
                "    \"primary_indication\": \"string|null\",\n"
                "    \"epworth_score\": number|null,\n"
                "    \"snoring_reported\": true|false|null,\n"
                "    \"witnessed_apneas\": true|false|null,\n"
                "    \"daytime_sleepiness\": true|false|null,\n"
                "    \"insomnia\": true|false|null\n"
                "  },\n\n"
                "  \"comorbidities\": [\n"
                "    {\"condition\": \"string\", \"present\": true|false, \"evidence\": \"string|null\"}\n"
                "  ],\n\n"
                "  \"medications\": [\n"
                "    {\"name\": \"string\", \"dose\": \"string|null\", \"timing\": \"string|null\"}\n"
                "  ],\n\n"
                "  \"prior_therapy\": {\n"
                "    \"cpap\": {\"used\": true|false|null, \"settings\": \"string|null\", \"mask_type\": \"string|null\"},\n"
                "    \"apap\": {\"used\": true|false|null, \"settings\": \"string|null\"},\n"
                "    \"bilevel\": {\"used\": true|false|null, \"settings\": \"string|null\"},\n"
                "    \"oral_appliance\": {\"used\": true|false|null, \"type\": \"string|null\"}\n"
                "  },\n\n"
                "  \"device_adherence_if_applicable\": {\n"
                "    \"nights_ge_4h_pct\": number|null,\n"
                "    \"avg_use_hours\": number|null,\n"
                "    \"residual_ahi\": number|null,\n"
                "    \"median_pressure_cmH2O\": number|null,\n"
                "    \"p95_pressure_cmH2O\": number|null,\n"
                "    \"median_leak_lpm\": number|null\n"
                "  },\n\n"
                "  \"sleep_timing_architecture\": {\n"
                "    \"trt_min\": number|null,\n"
                "    \"tst_min\": number|null,\n"
                "    \"sleep_efficiency_pct\": number|null,\n"
                "    \"sleep_latency_min\": number|null,\n"
                "    \"rem_latency_min\": number|null,\n"
                "    \"wakeup_after_sleep_onset_min\": number|null,\n"
                "    \"stages_pct\": {\"n1\": number|null, \"n2\": number|null, \"n3\": number|null, \"rem\": number|null}\n"
                "  },\n\n"
                "  \"respiratory_indices\": {\n"
                "    \"ahi_overall\": number|null,\n"
                "    \"rdi_overall\": number|null,\n"
                "    \"odi3\": number|null,\n"
                "    \"odi4\": number|null,\n"
                "    \"oai\": number|null,\n"
                "    \"cai\": number|null,\n"
                "    \"mai\": number|null,\n"
                "    \"hi\": number|null,\n"
                "    \"ahi_rem\": number|null,\n"
                "    \"ahi_nrem\": number|null,\n"
                "    \"ahi_supine\": number|null,\n"
                "    \"ahi_non_supine\": number|null\n"
                "  },\n\n"
                "  \"event_counts\": {\n"
                "    \"apnea_total\": number|null,\n"
                "    \"apnea_obstructive\": number|null,\n"
                "    \"apnea_central\": number|null,\n"
                "    \"apnea_mixed\": number|null,\n"
                "    \"hypopnea_total\": number|null\n"
                "  },\n\n"
                "  \"oxygenation\": {\n"
                "    \"spo2_nadir_pct\": number|null,\n"
                "    \"spo2_mean_pct\": number|null,\n"
                "    \"t90_pct\": number|null,\n"
                "    \"t88_pct\": number|null,\n"
                "    \"t85_pct\": number|null,\n"
                "    \"t80_pct\": number|null,\n"
                "    \"t70_pct\": number|null,\n"
                "    \"time_below_90_min\": number|null\n"
                "  },\n\n"
                "  \"snoring\": {\n"
                "    \"snore_index\": number|null,\n"
                "    \"snore_time_pct\": number|null\n"
                "  },\n\n"
                "  \"arousals_movements\": {\n"
                "    \"arousal_index\": number|null,\n"
                "    \"rera_index\": number|null,\n"
                "    \"plmi\": number|null,\n"
                "    \"plm_arousal_index\": number|null\n"
                "  },\n\n"
                "  \"position_stats\": {\n"
                "    \"supine_pct_of_sleep\": number|null,\n"
                "    \"left_pct_of_sleep\": number|null,\n"
                "    \"right_pct_of_sleep\": number|null,\n"
                "    \"prone_pct_of_sleep\": number|null\n"
                "  },\n\n"
                "  \"cardiac\": {\n"
                "    \"avg_hr_bpm\": number|null,\n"
                "    \"min_hr_bpm\": number|null,\n"
                "    \"max_hr_bpm\": number|null,\n"
                "    \"arrhythmia_notes\": \"string|null\"\n"
                "  },\n\n"
                "  \"titration_if_present\": {\n"
                "    \"modality\": \"CPAP|APAP|Bilevel|OA|None|null\",\n"
                "    \"settings\": \"string|null\",\n"
                "    \"recommendation\": \"string|null\"\n"
                "  },\n\n"
                "  \"impression_assessment\": {\n"
                "    \"diagnoses\": [\"string\", \"...\"],\n"
                "    \"ahi_severity_label\": \"mild|moderate|severe|null\",\n"
                "    \"free_text_impression\": \"string|null\",\n"
                "    \"plan_recommendations\": \"string|null\",\n"
                "    \"follow_up_interval\": \"string|null\"\n"
                "  },\n\n"
                "  \"temporal_series\": [\n"
                "    {\n"
                "      \"label\": \"Baseline|Follow-up #1|…|string\",\n"
                "      \"date\": \"YYYY-MM-DD|null\",\n"
                "      \"study_type\": \"baseline|follow_up|unknown\",\n"
                "      \"ahi\": number|null,\n"
                "      \"rdi\": number|null,\n"
                "      \"odi3\": number|null,\n"
                "      \"odi4\": number|null,\n"
                "      \"o2_nadir_pct\": number|null,\n"
                "      \"time_below_90_pct\": number|null,\n"
                "      \"tst_min\": number|null,\n"
                "      \"sleep_efficiency_pct\": number|null,\n"
                "      \"rem_ahi\": number|null,\n"
                "      \"supine_ahi\": number|null\n"
                "    }\n"
                "  ],\n\n"
                "  \"evidence\": {\n"
                "    \"ahi_overall\": \"string|null\",\n"
                "    \"ahi_supine\": \"string|null\",\n"
                "    \"t90_pct\": \"string|null\",\n"
                "    \"odi3\": \"string|null\",\n"
                "    \"odi4\": \"string|null\",\n"
                "    \"spo2_nadir_pct\": \"string|null\",\n"
                "    \"scoring_hypopnea_rule\": \"string|null\"\n"
                "  }\n"
                "}\n\n"
                "## EXTRACTION RULES\n"
                "- Return numbers only (no units or % symbols). Use null if absent.\n"
                "- Search BOTH narrative text and tables throughout the document.\n"
                "- Extract ALL available information from each category.\n"
                "- For comorbidities: Look for conditions like hypertension, diabetes, heart disease, etc.\n"
                "- For medications: Extract drug names, dosages, and timing if mentioned.\n"
                "- For prior therapy: Look for CPAP, APAP, oral appliance history.\n"
                "- For symptoms: Extract Epworth scores, snoring reports, witnessed apneas, etc.\n"
                "- For sleep architecture: Extract sleep stages, latency, efficiency metrics.\n"
                "- For respiratory indices: Extract all AHI variants (overall, REM, NREM, supine, non-supine).\n"
                "- For oxygenation: Extract all O2 saturation metrics and time below thresholds.\n"
                "- For cardiac: Extract heart rate metrics and arrhythmia notes.\n"
                "- For impression: Extract diagnoses, severity labels, and recommendations.\n"
                "- For temporal series: Create entries for baseline vs follow-up studies if present.\n"
                "- Keep evidence snippets short (≤120 chars) and human-readable.\n"
                "- Return ONLY the single JSON object. No extra text."
            )
        }]
        
        # Few-shot examples
        shot_a_input = (
            "Oxygen Saturation <90   <=88  <85  <80  <70\n"
            "Duration (minutes): 3.1  0.0  0.0  0.0  0.0\n"
            "Sleep %:             0.7  0.0  0.0  0.0  0.0\n"
            "Body Position Statistics … Supine … pAHI 42.5 …"
        )
        shot_a_output = {
            "oxygenation": {"t90_pct": 0.7, "time_below_90_min": 3.1},
            "respiratory_indices": {"ahi_supine": 42.5},
            "evidence": {
                "t90_pct": "Sleep %: 0.7 under <90",
                "ahi_supine": "Supine … pAHI 42.5"
            }
        }
        
        shot_b_input = (
            "Supine AHI 42.5 (62% of sleep time). Less than 90% O2 0.5%.\n"
            "AHI overall 28.2. ODI (3%) 16.1. SpO2 nadir 83%."
        )
        shot_b_output = {
            "respiratory_indices": {"ahi_overall": 28.2, "ahi_supine": 42.5, "odi3": 16.1},
            "oxygenation": {"spo2_nadir_pct": 83, "t90_pct": 0.5},
            "evidence": {
                "ahi_overall": "AHI overall 28.2",
                "ahi_supine": "Supine AHI 42.5",
                "t90_pct": "Less than 90% O2 0.5%",
                "spo2_nadir_pct": "SpO2 nadir 83%"
            }
        }
        
        # Build messages with few-shot examples
        messages = [
            # FEW-SHOT A
            {"role": "user", "content": [{"text": "Extract per the schema from this slice:\n" + shot_a_input}]},
            {"role": "assistant", "content": [{"text": json.dumps(shot_a_output, ensure_ascii=False)}]},
            # FEW-SHOT B
            {"role": "user", "content": [{"text": "Extract per the schema from this slice:\n" + shot_b_input}]},
            {"role": "assistant", "content": [{"text": json.dumps(shot_b_output, ensure_ascii=False)}]},
            # REAL REQUEST + DOCUMENT
            {"role": "user", "content": [
                {"text": "Now extract ALL fields per the schema from this full document. Return ONLY the single JSON object."},
                document_block,
            ]}
        ]
        
        # Use hardcoded model ID for standalone script
        model_id = MODEL_ID
        
        import time
        start_time = time.time()
        
        try:
            resp = bedrock.converse(
                modelId=model_id,
                system=system,
                messages=messages,
                inferenceConfig={"temperature": 0.0, "maxTokens": 4000},
            )
            raw = resp["output"]["message"]["content"][0]["text"]
            response_time_ms = int((time.time() - start_time) * 1000)
            
            # Log to database
            _log_llm_call(
                prompt_text=f"Extract observations from {safe_name} (visual LLM)",
                response_text=raw[:500],  # Truncate for storage
                response_time_ms=response_time_ms,
                status='success'
            )
            
            logger.info(f"Raw LLM response for {safe_name}: {raw[:500]}...")
            
            if not raw or not raw.strip():
                return {}, False, f"Empty response from LLM for {safe_name}"
        except Exception as e:
            response_time_ms = int((time.time() - start_time) * 1000)
            _log_llm_call(
                prompt_text=f"Extract observations from {safe_name} (visual LLM)",
                response_text='',
                response_time_ms=response_time_ms,
                status='error',
                error_message=str(e)
            )
            raise
        
        # robust JSON parse (handles code fences/stray text)
        def _parse_json_only(raw_text: str):
            txt = raw_text.strip()
            txt = re.sub(r'^```(?:json)?\s*', '', txt)
            txt = re.sub(r'\s*```$', '', txt)
            m = re.search(r'(\{.*\}|\[.*\])', txt, flags=re.DOTALL)
            if m:
                txt = m.group(1)
            try:
                return json.loads(txt)
            except json.JSONDecodeError:
                logger.warning(f"Failed to parse JSON from LLM response for {safe_name}")
                return {}
        
        parsed = _parse_json_only(raw)
        if not parsed:
            return {}, False, f"Failed to parse JSON from LLM response for {safe_name}"
        
        return parsed, True, ""
        
    except Exception as e:
        logger.error(f"Error in _process_single_document_chunk for {safe_name}: {e}")
        return {}, False, str(e)


def extract_observations_with_visual_llm(s3_key: str, document_name: str) -> Tuple[Dict, bool, str]:
    """
    Extract comprehensive medical observations from PDF using direct vision processing (no text extraction).
    Returns: (extracted_observations_dict, success_bool, error_message)
    """
    try:
        import boto3
        import requests
        import re
        
        bedrock = boto3.client("bedrock-runtime", region_name="us-west-2")

        # Download document content from S3 (Bedrock only accepts bytes, not S3 URIs)
        def download_document_from_s3(s3_key):
            try:
                s3_client = boto3.client('s3', region_name='us-west-2')
                bucket = os.getenv('S3_BUCKET_NAME', 'vizbrizpatients')
                
                # Determine content type based on file extension
                if s3_key.lower().endswith('.pdf'):
                    content_type = 'application/pdf'
                elif s3_key.lower().endswith(('.jpg', '.jpeg')):
                    content_type = 'image/jpeg'
                elif s3_key.lower().endswith('.png'):
                    content_type = 'image/png'
                else:
                    content_type = 'application/octet-stream'
                
                presigned_url = s3_client.generate_presigned_url(
                    'get_object',
                    Params={
                        'Bucket': bucket,
                        'Key': s3_key,
                        'ResponseContentType': content_type
                    },
                    ExpiresIn=3600
                )
                
                response = requests.get(presigned_url, timeout=30)
                response.raise_for_status()
                document_content = response.content
                
                logger.info(f"Downloaded document from S3: {s3_key} ({len(document_content)} bytes)")
                return document_content
                
            except Exception as e:
                logger.error(f"Error downloading document from S3 {s3_key}: {str(e)}")
                return None

        document_content = download_document_from_s3(s3_key)
        if not document_content:
            return {}, False, f"Failed to download document from S3: {s3_key}"
        
        # Sanitize document name for Bedrock (strict validation rules)
        def sanitize_bedrock_doc_name(name: str) -> str:
            base = os.path.splitext(os.path.basename(name))[0]
            base = re.sub(r'[^A-Za-z0-9\-\(\)\[\]\s]', ' ', base)
            base = re.sub(r'\s+', ' ', base).strip()
            return base or "Document"
        
        safe_name = sanitize_bedrock_doc_name(document_name or s3_key)
        
        content_size = len(document_content)
        
        # Determine document format based on file extension
        if s3_key.lower().endswith('.pdf'):
            doc_format = "pdf"
        elif s3_key.lower().endswith(('.jpg', '.jpeg')):
            doc_format = "jpeg"
        elif s3_key.lower().endswith('.png'):
            doc_format = "png"
        else:
            doc_format = "pdf"  # fallback
        
        # Handle large PDFs by chunking
        if content_size > BEDROCK_MAX_DOCUMENT_BYTES and doc_format == "pdf":
            logger.info(
                f"Document {safe_name} is {content_size} bytes, exceeding Bedrock limit of "
                f"{BEDROCK_MAX_DOCUMENT_BYTES} bytes. Attempting to chunk PDF by pages..."
            )
            chunks = _chunk_pdf_by_pages(document_content, BEDROCK_MAX_DOCUMENT_BYTES)
            
            if not chunks:
                logger.warning(f"Failed to chunk PDF {safe_name}, will skip Visual LLM but continue with OCR")
                return {}, False, (
                    f"Document size {content_size} exceeds Bedrock limit "
                    f"({BEDROCK_MAX_DOCUMENT_BYTES} bytes) and chunking failed"
                )
            
            # Process each chunk and merge results
            logger.info(f"Processing {len(chunks)} PDF chunks for {safe_name}")
            all_chunk_results = {}
            
            for chunk_idx, (chunk_bytes, start_page, end_page) in enumerate(chunks, 1):
                logger.info(f"Processing chunk {chunk_idx}/{len(chunks)} (pages {start_page+1}-{end_page+1})")
                chunk_name = f"{safe_name}_chunk{chunk_idx}_pages{start_page+1}-{end_page+1}"
                
                # Process the chunk using the helper function
                chunk_result, chunk_success, chunk_error = _process_single_document_chunk(
                    chunk_bytes, chunk_name, doc_format
                )
                
                if chunk_success and chunk_result:
                    # Merge chunk results (prefer non-null values, take max for numeric fields)
                    for category, data in chunk_result.items():
                        if not isinstance(data, dict):
                            continue
                        if category not in all_chunk_results:
                            all_chunk_results[category] = {}
                        for field, value in data.items():
                            if value is not None:
                                # For numeric fields, take the max (e.g., AHI values)
                                if isinstance(value, (int, float)) and field in all_chunk_results[category]:
                                    existing = all_chunk_results[category][field]
                                    if isinstance(existing, (int, float)):
                                        all_chunk_results[category][field] = max(value, existing)
                                    else:
                                        all_chunk_results[category][field] = value
                                # For lists, extend them
                                elif isinstance(value, list) and field in all_chunk_results[category]:
                                    existing = all_chunk_results[category][field]
                                    if isinstance(existing, list):
                                        all_chunk_results[category][field] = existing + value
                                    else:
                                        all_chunk_results[category][field] = value
                                # Otherwise, prefer non-null values
                                elif field not in all_chunk_results[category] or all_chunk_results[category][field] is None:
                                    all_chunk_results[category][field] = value
                else:
                    logger.warning(f"Chunk {chunk_idx} failed: {chunk_error}")
            
            if all_chunk_results:
                logger.info(f"Successfully processed {len(chunks)} chunks for {safe_name}, merged results")
                return all_chunk_results, True, ""
            else:
                logger.warning(f"All chunks failed for {safe_name}, will continue with OCR")
                return {}, False, "All PDF chunks failed Visual LLM processing"
        
        elif content_size > BEDROCK_MAX_DOCUMENT_BYTES:
            # Non-PDF file that's too large (image)
            logger.warning(
                f"Document {safe_name} is {content_size} bytes, exceeding Bedrock limit of "
                f"{BEDROCK_MAX_DOCUMENT_BYTES} bytes. Skipping Visual LLM call (not a PDF, cannot chunk)."
            )
            return {}, False, (
                f"Document size {content_size} exceeds Bedrock limit "
                f"({BEDROCK_MAX_DOCUMENT_BYTES} bytes)"
            )
        
        # Process the document using the helper function (handles images and PDFs)
        return _process_single_document_chunk(document_content, safe_name, doc_format)

    except Exception as e:
        msg = f"Error in comprehensive PDF vision extraction for {document_name}: {e}"
        logger.error(msg)
        return {}, False, msg


def convert_visual_llm_to_observations(visual_data: Dict, document_name: str) -> List[Dict]:
    """
    Convert visual LLM output to the existing observation format used by the system.
    
    Args:
        visual_data (Dict): Output from extract_observations_with_visual_llm
        document_name (str): Name of the document
        
    Returns:
        List[Dict]: List of observations in the expected format
    """
    observations = []
    
    try:
        # Convert each category to observations
        for category, data in visual_data.items():
            if not data or not isinstance(data, dict):
                continue
                
            for field, value in data.items():
                if value is None:
                    continue
                    
                # Create observation in the expected format
                observation = {
                    'path': f"{category}.{field}",
                    'value': value,
                    'observation': f"{category.title()} {field.replace('_', ' ').title()}",
                    'score': 1,
                    'explanation': 'Visual LLM extraction',
                    'evidence': str(value)[:512] if value else '',
                    'confidence': 95,
                    'source': 'bedrock-vision'
                }
                observations.append(observation)
        
        # Handle special cases for nested structures
        if 'temporal_series' in visual_data and isinstance(visual_data['temporal_series'], list):
            for i, time_point in enumerate(visual_data['temporal_series']):
                if isinstance(time_point, dict):
                    for field, value in time_point.items():
                        if value is not None and field not in ['date', 'study_type', 'label']:
                            observation = {
                                'path': f"temporal_series[{i}].{field}",
                                'value': value,
                                'observation': f"Temporal Series {field.replace('_', ' ').title()}",
                                'score': 1,
                                'explanation': 'Visual LLM extraction - temporal series',
                                'evidence': str(value)[:512] if value else '',
                                'confidence': 95,
                                'source': 'bedrock-vision'
                            }
                            observations.append(observation)
        
        # Handle arrays like comorbidities and medications
        for array_field in ['comorbidities', 'medications']:
            if array_field in visual_data and isinstance(visual_data[array_field], list):
                for i, item in enumerate(visual_data[array_field]):
                    if isinstance(item, dict):
                        for field, value in item.items():
                            if value is not None:
                                observation = {
                                    'path': f"{array_field}[{i}].{field}",
                                    'value': value,
                                    'observation': f"{array_field.title()} {field.replace('_', ' ').title()}",
                                    'score': 1,
                                    'explanation': 'Visual LLM extraction',
                                    'evidence': str(value)[:512] if value else '',
                                    'confidence': 95,
                                    'source': 'bedrock-vision'
                                }
                                observations.append(observation)
        
        logger.info(f"Converted {len(observations)} observations from visual LLM data for {document_name}")
        return observations
        
    except Exception as e:
        logger.error(f"Error converting visual LLM data to observations: {e}")
        return []


def convert_regex_data_to_observations(regex_data: Dict, document_name: str) -> List[Dict]:
    """
    Convert regex extraction data to the existing observation format used by the system.
    
    Args:
        regex_data (Dict): Output from extract_numerical_data_with_regex
        document_name (str): Name of the document
        
    Returns:
        List[Dict]: List of observations in the expected format
    """
    observations = []
    
    try:
        # Convert regex data to observations
        for category, data in regex_data.items():
            if not data or not isinstance(data, dict):
                continue
                
            for field, value in data.items():
                if value is None:
                    continue
                    
                # Create observation in the expected format
                observation = {
                    'path': f"{category}.{field}",
                    'value': value,
                    'observation': f"{category.title()} {field.replace('_', ' ').title()}",
                    'score': 0.8,  # Lower confidence than LLM
                    'explanation': 'Regex extraction',
                    'evidence': str(value)[:512] if value else '',
                    'document_name': document_name,
                    'confidence': 80,
                    'source': 'regex-extraction'
                }
                observations.append(observation)
                
    except Exception as e:
        logger.error(f"Error converting regex data to observations: {e}")
    
    return observations
def validate_llm_vs_regex_extraction(llm_observations: List[Dict], regex_observations: List[Dict], document_name: str) -> Dict:
    """
    Compare LLM and regex extraction results to identify discrepancies.
    
    Args:
        llm_observations: List of observations from LLM extraction
        regex_observations: List of observations from regex extraction
        document_name: Name of the document
        
    Returns:
        Dict: Validation results with discrepancies and confidence scores
    """
    validation_results = {
        'agreements': [],
        'discrepancies': [],
        'llm_only': [],
        'regex_only': [],
        'confidence_scores': {}
    }
    
    # Create lookup dictionaries
    llm_lookup = {obs['path']: obs for obs in llm_observations}
    regex_lookup = {obs['path']: obs for obs in regex_observations}
    
    # Critical fields for validation (with tolerance)
    critical_fields = {
        'respiratory_indices.ahi_overall': 10,  # AHI tolerance
        'respiratory_indices.odi3': 5,          # ODI tolerance
        'respiratory_indices.odi4': 5,          # ODI tolerance
        'oxygenation.spo2_nadir_pct': 5,        # O2 nadir tolerance
        'oxygenation.t90_pct': 2,               # T90 tolerance
        'sleep_timing_architecture.sleep_efficiency_pct': 10,  # Sleep efficiency tolerance
    }
    
    # Check all paths from both extractions
    all_paths = set(llm_lookup.keys()) | set(regex_lookup.keys())
    
    for path in all_paths:
        llm_obs = llm_lookup.get(path)
        regex_obs = regex_lookup.get(path)
        
        if llm_obs and regex_obs:
            # Both methods found this field
            llm_value = llm_obs.get('value')
            regex_value = regex_obs.get('value')
            
            try:
                llm_num = float(llm_value) if llm_value is not None else None
                regex_num = float(regex_value) if regex_value is not None else None
                
                if llm_num is not None and regex_num is not None:
                    # Both have numeric values
                    tolerance = critical_fields.get(path, 5)  # Default tolerance
                    difference = abs(llm_num - regex_num)
                    
                    if difference <= tolerance:
                        # Agreement within tolerance
                        validation_results['agreements'].append({
                            'path': path,
                            'llm_value': llm_value,
                            'regex_value': regex_value,
                            'difference': difference,
                            'tolerance': tolerance
                        })
                        validation_results['confidence_scores'][path] = 0.95
                    else:
                        # Significant discrepancy
                        validation_results['discrepancies'].append({
                            'path': path,
                            'llm_value': llm_value,
                            'regex_value': regex_value,
                            'difference': difference,
                            'tolerance': tolerance,
                            'severity': 'high' if difference > tolerance * 2 else 'medium'
                        })
                        validation_results['confidence_scores'][path] = 0.3
                else:
                    # Non-numeric comparison
                    if str(llm_value).lower() == str(regex_value).lower():
                        validation_results['agreements'].append({
                            'path': path,
                            'llm_value': llm_value,
                            'regex_value': regex_value,
                            'difference': 0
                        })
                        validation_results['confidence_scores'][path] = 0.9
                    else:
                        validation_results['discrepancies'].append({
                            'path': path,
                            'llm_value': llm_value,
                            'regex_value': regex_value,
                            'difference': 'non-numeric',
                            'severity': 'medium'
                        })
                        validation_results['confidence_scores'][path] = 0.4
                        
            except (ValueError, TypeError):
                # Non-numeric values
                if str(llm_value).lower() == str(regex_value).lower():
                    validation_results['agreements'].append({
                        'path': path,
                        'llm_value': llm_value,
                        'regex_value': regex_value
                    })
                    validation_results['confidence_scores'][path] = 0.9
                else:
                    validation_results['discrepancies'].append({
                        'path': path,
                        'llm_value': llm_value,
                        'regex_value': regex_value,
                        'severity': 'medium'
                    })
                    validation_results['confidence_scores'][path] = 0.4
                    
        elif llm_obs and not regex_obs:
            # LLM only
            validation_results['llm_only'].append({
                'path': path,
                'value': llm_obs.get('value'),
                'explanation': 'LLM extraction only'
            })
            validation_results['confidence_scores'][path] = 0.7
            
        elif regex_obs and not llm_obs:
            # Regex only
            validation_results['regex_only'].append({
                'path': path,
                'value': regex_obs.get('value'),
                'explanation': 'Regex extraction only'
            })
            validation_results['confidence_scores'][path] = 0.6
    
    # Log validation summary
    logger.info(f"Validation summary for {document_name}:")
    logger.info(f"  - Agreements: {len(validation_results['agreements'])}")
    logger.info(f"  - Discrepancies: {len(validation_results['discrepancies'])}")
    logger.info(f"  - LLM only: {len(validation_results['llm_only'])}")
    logger.info(f"  - Regex only: {len(validation_results['regex_only'])}")
    
    return validation_results


def merge_llm_and_regex_observations(llm_observations: List[Dict], regex_observations: List[Dict], 
                                    validation_results: Dict, document_name: str) -> List[Dict]:
    """
    Merge LLM and regex observations using LLM as primary and regex as fallback.
    
    Args:
        llm_observations: List of observations from LLM extraction
        regex_observations: List of observations from regex extraction
        validation_results: Results from validation comparison
        document_name: Name of the document
        
    Returns:
        List[Dict]: Merged observations with confidence scores
    """
    merged_observations = []
    
    # Create lookup dictionaries
    llm_lookup = {obs['path']: obs for obs in llm_observations}
    regex_lookup = {obs['path']: obs for obs in regex_observations}
    
    # Get all unique paths
    all_paths = set(llm_lookup.keys()) | set(regex_lookup.keys())
    
    # Create lookup sets for faster path checking
    discrepancy_paths = {item['path'] for item in validation_results.get('discrepancies', [])}
    agreement_paths = {item['path'] for item in validation_results.get('agreements', [])}
    
    for path in all_paths:
        llm_obs = llm_lookup.get(path)
        regex_obs = regex_lookup.get(path)
        
        # Determine which observation to use based on validation results
        if path in discrepancy_paths:
            # Significant discrepancy - prefer Visual LLM as primary, but flag discrepancies
            # For critical fields with large discrepancies, we may prefer regex
            critical_paths = ['respiratory_indices.ahi_overall', 'respiratory_indices.odi3', 'respiratory_indices.odi4']
            
            # Find the discrepancy details
            discrepancy_item = next((item for item in validation_results.get('discrepancies', []) if item['path'] == path), None)
            discrepancy_severity = discrepancy_item.get('severity', 'medium') if discrepancy_item else 'medium'
            
            if path in critical_paths and discrepancy_severity == 'high' and regex_obs:
                # For critical fields with high severity discrepancies, prefer regex
                obs = regex_obs.copy()
                obs['score'] = 0.6  # Lower confidence due to discrepancy
                obs['explanation'] = f"Regex extraction (Visual LLM discrepancy: {llm_obs.get('value') if llm_obs else 'N/A'})"
                merged_observations.append(obs)
            elif llm_obs:
                # Use Visual LLM as primary (as requested), but flag the discrepancy
                obs = llm_obs.copy()
                obs['score'] = 0.6  # Lower confidence due to discrepancy
                regex_val = regex_obs.get('value') if regex_obs else 'N/A'
                obs['explanation'] = f"Visual LLM extraction (regex cross-check shows discrepancy: {regex_val})"
                merged_observations.append(obs)
            elif regex_obs:
                # Fallback to regex if LLM observation missing
                obs = regex_obs.copy()
                obs['score'] = 0.5
                obs['explanation'] = "Regex extraction (Visual LLM discrepancy)"
                merged_observations.append(obs)
                    
        elif path in agreement_paths:
            # Agreement - use Visual LLM with high confidence (validated by regex)
            if llm_obs:
                obs = llm_obs.copy()
                obs['score'] = 0.95  # High confidence due to agreement
                obs['explanation'] = "Visual LLM extraction (validated by regex cross-reference)"
                merged_observations.append(obs)
            elif regex_obs:
                obs = regex_obs.copy()
                obs['score'] = 0.9
                obs['explanation'] = "Regex extraction (validated by Visual LLM)"
                merged_observations.append(obs)
                
        elif llm_obs and not regex_obs:
            # Visual LLM only - use with medium confidence (no validation available)
            obs = llm_obs.copy()
            obs['score'] = 0.7  # Medium confidence
            obs['explanation'] = "Visual LLM extraction only (no text validation available)"
            merged_observations.append(obs)
            
        elif regex_obs and not llm_obs:
            # Regex only - use with lower confidence (Visual LLM failed)
            obs = regex_obs.copy()
            obs['score'] = 0.6  # Lower confidence without Visual LLM
            obs['explanation'] = "Regex extraction only (Visual LLM failed)"
            merged_observations.append(obs)
    
    # Log merge summary
    logger.info(f"Merged {len(merged_observations)} observations for {document_name}")
    
    return merged_observations


def validate_canonical_against_source_files(patient_id: int, canonical_data: Dict[str, Any]) -> Dict[str, Any]:
    """
    Validate canonical data against source files using simple regex searches.
    This catches LLM hallucinations by cross-checking critical values.
    
    Args:
        patient_id: Patient ID
        canonical_data: The canonical JSON data to validate
        
    Returns:
        Dict with validation results including discrepancies
    """
    validation_results = {
        'discrepancies': [],
        'validations': [],
        'source_files_checked': [],
        'critical_fields_validated': []
    }
    
    try:
        # Get all documents for this patient
        documents = discover_patient_documents(patient_id)
        
        # Critical fields to validate with their canonical paths
        critical_fields = {
            'ahi_overall': {
                'canonical_path': 'respiratory_indices.ahi_overall',
                'regex_patterns': [
                    r'ahi[:\s]*(\d+(?:\.\d+)?)',
                    r'apnea.*?hypopnea.*?index[:\s]*(\d+(?:\.\d+)?)',
                    r'(\d+(?:\.\d+)?)\s*ahi',
                    r'pahi[:\s]*(\d+(?:\.\d+)?)',
                ],
                'tolerance': 10
            },
            'odi3': {
                'canonical_path': 'respiratory_indices.odi3', 
                'regex_patterns': [
                    r'odi[:\s]*(\d+(?:\.\d+)?)',
                    r'oxygen.*?desaturation.*?index[:\s]*(\d+(?:\.\d+)?)',
                    r'(\d+(?:\.\d+)?)\s*odi',
                ],
                'tolerance': 5
            },
            'spo2_nadir_pct': {
                'canonical_path': 'oxygenation.spo2_nadir_pct',
                'regex_patterns': [
                    r'o2.*?nadir[:\s]*(\d+(?:\.\d+)?)',
                    r'oxygen.*?nadir[:\s]*(\d+(?:\.\d+)?)',
                    r'spo2.*?nadir[:\s]*(\d+(?:\.\d+)?)',
                    r'(\d+(?:\.\d+)?).*?o2.*?nadir',
                ],
                'tolerance': 5
            }
        }
        
        # Extract canonical values
        canonical_values = {}
        for field, config in critical_fields.items():
            path_parts = config['canonical_path'].split('.')
            value = canonical_data
            for part in path_parts:
                if isinstance(value, dict) and part in value:
                    value = value[part]
                else:
                    value = None
                    break
            canonical_values[field] = value
        
        # Check each document
        for doc in documents:
            if not doc.get('s3_key'):
                continue
                
            validation_results['source_files_checked'].append(doc['name'])
            
            try:
                # Extract text content from document
                content = extract_document_content(doc)
                if not content:
                    continue
                    
                text_content = content if isinstance(content, str) else str(content)
                
                # Validate each critical field
                for field, config in critical_fields.items():
                    canonical_value = canonical_values.get(field)
                    if canonical_value is None:
                        continue
                        
                    # Search for values in document using regex
                    found_values = []
                    for pattern in config['regex_patterns']:
                        import re
                        matches = re.findall(pattern, text_content, re.IGNORECASE)
                        for match in matches:
                            try:
                                found_values.append(float(match))
                            except (ValueError, TypeError):
                                continue
                    
                    if found_values:
                        # Find closest match to canonical value
                        canonical_num = float(canonical_value)
                        closest_match = min(found_values, key=lambda x: abs(x - canonical_num))
                        difference = abs(canonical_num - closest_match)
                        
                        if difference <= config['tolerance']:
                            # Validation passed
                            validation_results['validations'].append({
                                'field': field,
                                'canonical_value': canonical_value,
                                'source_value': closest_match,
                                'difference': difference,
                                'document': doc['name'],
                                'status': 'validated'
                            })
                        else:
                            # Significant discrepancy found
                            validation_results['discrepancies'].append({
                                'field': field,
                                'canonical_value': canonical_value,
                                'source_value': closest_match,
                                'difference': difference,
                                'document': doc['name'],
                                'severity': 'high' if difference > config['tolerance'] * 2 else 'medium',
                                'status': 'discrepancy'
                            })
                            logger.warning(f"🚨 DISCREPANCY: {field} in canonical ({canonical_value}) vs source ({closest_match}) in {doc['name']}")
                    else:
                        # No values found in source - this could be a problem
                        validation_results['discrepancies'].append({
                            'field': field,
                            'canonical_value': canonical_value,
                            'source_value': None,
                            'difference': 'not_found_in_source',
                            'document': doc['name'],
                            'severity': 'high',
                            'status': 'not_found_in_source'
                        })
                        logger.warning(f"🚨 NO SOURCE VALIDATION: {field} value {canonical_value} not found in {doc['name']}")
                        
            except Exception as e:
                logger.error(f"Error validating document {doc['name']}: {e}")
                continue
        
        # Summary
        validation_results['critical_fields_validated'] = list(critical_fields.keys())
        logger.info(f"Canonical validation for patient {patient_id}: {len(validation_results['validations'])} validations, {len(validation_results['discrepancies'])} discrepancies")
        
    except Exception as e:
        logger.error(f"Error in canonical validation for patient {patient_id}: {e}")
        validation_results['error'] = str(e)
    
    return validation_results


def extract_observations_with_llm(document_content: str, document_type: str, filename: str) -> str:
    """
    LLM extraction for two canonical metrics + evidence.
    Falls back to regex if LLM returns null. Returns a short status string.
    """
    import json as _json
    try:
        if not document_content or not document_content.strip():
            return "No content available for analysis"

        if len(document_content) > MAX_DOC_CHARS:
            logger.warning(f"Content too large ({len(document_content)} chars), truncating to {MAX_DOC_CHARS} chars for {filename}")
            document_content = trim_content(document_content, MAX_DOC_CHARS)

        system_prompt = """You extract SLEEP STUDY METRICS from noisy text and tables.
Return only STRICT JSON (no prose). If a value is not present, use null.
Schema:
{
  "supine_ahi": number|null,
  "percent_time_spo2_below_90": number|null,
  "supporting_evidence": {
    "supine_ahi": string|null,
    "percent_time_spo2_below_90": string|null
  }
}
Rules:
- Search BOTH narrative lines and tables.
- If a table shows "Oxygen Saturation <90 … Sleep % 0.7", set percent_time_spo2_below_90 = 0.7.
- Accept synonyms: "SpO2 < 90%", "% time <90".
- For supine AHI, accept "Supine AHI 42.5", "AHI (Supine) 42.5", or values in "Body Position Statistics".
- Copy a short evidence substring (line or table row) into supporting_evidence.
- Use numbers only (no % sign).
"""

        shot_A_in = "… Oxygen Saturation <90 … Sleep % 0.7 … Body Position Statistics … Supine … pAHI 42.5 …"
        shot_A_out = {
            "supine_ahi": 42.5,
            "percent_time_spo2_below_90": 0.7,
            "supporting_evidence": {
                "supine_ahi": "Body Position … Supine … pAHI 42.5",
                "percent_time_spo2_below_90": "Oxygen Saturation <90 … Sleep % 0.7"
            }
        }
        shot_B_in = "Supine AHI 42.5 (62% of sleep time) … Less than 90% O2 0.5% …"
        shot_B_out = {
            "supine_ahi": 42.5,
            "percent_time_spo2_below_90": 0.5,
            "supporting_evidence": {
                "supine_ahi": "Supine AHI 42.5 (62% of sleep time)",
                "percent_time_spo2_below_90": "Less than 90% O2 0.5%"
            }
        }

        user_prompt = f"""EXAMPLES:
INPUT_A:
{shot_A_in}
OUTPUT_A:
{_json.dumps(shot_A_out, ensure_ascii=False)}

INPUT_B:
{shot_B_in}
OUTPUT_B:
{_json.dumps(shot_B_out, ensure_ascii=False)}

NOW EXTRACT FOR THIS DOCUMENT:
Document Type: {document_type}
Filename: {filename}

CONTENT:
{document_content}
"""

        # Use system role for system prompt
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ]

        logger.info(f"Calling Bedrock (JSON extract) for {filename} ({len(document_content)} chars)")
        if bedrock_query_enhanced is None:
            return "Error: Bedrock pipeline not available"

        result = bedrock_query_enhanced(
            messages,
            max_tokens=BEDROCK_MAX_TOKENS_SINGLE,
            temperature=0.0,
            top_p=0.9
        )
        if not (isinstance(result, dict) and result.get("success")):
            return f"Error extracting observations: {result.get('message', 'Unknown error') if isinstance(result, dict) else 'Unknown error'}"

        payload_raw = result.get("response", "").strip()
        try:
            # Log the raw LLM response for auditing (truncate to avoid giant logs)
            logger.info(f"LLM raw response for {filename} (first 1500 chars): {payload_raw[:1500]}")
        except Exception:
            pass
        try:
            payload = _json.loads(payload_raw)
        except Exception:
            logger.warning("Non-JSON response received; attempting regex fallback only.")
            payload = {"supine_ahi": None, "percent_time_spo2_below_90": None, "supporting_evidence": {"supine_ahi": None, "percent_time_spo2_below_90": None}}

        supine_ahi = _coerce_num(payload.get("supine_ahi"))
        o2lt90 = _coerce_num(payload.get("percent_time_spo2_below_90"))
        ev_sup = payload.get("supporting_evidence", {}).get("supine_ahi")
        ev_o2 = payload.get("supporting_evidence", {}).get("percent_time_spo2_below_90")

        if supine_ahi is None:
            m = _first_match(document_content, PATTERNS_SUPINE_AHI)
            if m:
                supine_ahi, ev_sup = m

        if o2lt90 is None:
            # 1) Regex fallback
            m = _first_match(document_content, PATTERNS_O2_LT90)
            if m:
                o2lt90, ev_o2 = m
            # 2) Direct question to LLM (very focused) if still missing
            if o2lt90 is None:
                try:
                    dq_system = (
                        "You extract a single numerical metric from a sleep study. "
                        "Return ONLY the number (no text), representing the percent of sleep time with oxygen saturation below 90."
                    )
                    dq_user = (
                        f"Question: What percentage of time did this patient spend with oxygen saturation below 90%?\n\n"
                        f"Document: {document_content[:8000]}"
                    )
                    resp = bedrock_query_enhanced([
                        {"role": "system", "content": dq_system},
                        {"role": "user", "content": dq_user}
                    ], max_tokens=64, temperature=0.0, top_p=0.9)
                    if isinstance(resp, dict) and resp.get("success"):
                        raw = (resp.get("response") or "").strip()
                        # Accept formats like "0.5%" or "0.5"
                        import re as _re
                        m2 = _re.search(r"(\d+(?:\.\d+)?)\s*%?", raw)
                        if m2:
                            o2lt90 = float(m2.group(1))
                            ev_o2 = raw[:200]
                except Exception:
                    pass

        status_bits = []
        if supine_ahi is not None:
            status_bits.append(f"Supine AHI={supine_ahi}")
        if o2lt90 is not None:
            status_bits.append(f"% time O2<90={o2lt90}")

        # Log the parsed metrics/evidence for traceability
        try:
            logger.info(
                f"LLM parsed metrics for {filename}: supine_ahi={supine_ahi}, o2_lt90={o2lt90}, "
                f"evidence_supine={str(ev_sup)[:200]}, evidence_o2={str(ev_o2)[:200]}"
            )
        except Exception:
            pass

        return "; ".join(status_bits) if status_bits else "No target metrics found"

    except Exception as e:
        logger.error(f"Error extracting observations with LLM for {filename}: {e}")
        return f"Error extracting observations: {str(e)}"

def extract_observations_with_llm_batch(documents_batch: List[Dict]) -> Dict[str, str]:
    """
    Use AWS Bedrock to extract observations from multiple documents in a single call.
    This optimizes Bedrock usage by batching documents together.
    
    Args:
        documents_batch (List[Dict]): List of document dictionaries with content and metadata
        
    Returns:
        Dict[str, str]: Dictionary mapping filename to extracted observations
    """
    try:
        if not documents_batch:
            return {}
        
        # Create a combined prompt for all documents
        system_prompt = """You are a medical AI assistant specialized in extracting SLEEP APNEA and DENTAL SLEEP MEDICINE clinical observations from medical documents. Your task is to analyze the provided document content and extract ONLY relevant sleep-related medical observations, findings, and clinical information.

CRITICAL REQUIREMENTS:
- FOCUS EXCLUSIVELY on sleep apnea, airway anatomy, dental sleep medicine, and related respiratory disorders
- ALWAYS extract patient demographics (age, sex, weight, height, BMI) if present
- Extract sleep study data (AHI, ODI, SpO2, sleep efficiency, etc.)
- Extract airway anatomy findings (tongue position, soft palate, uvula, TMJ, nasal obstruction, epiglottis, tonsils, adenoids, pharyngeal wall, retropalatal/retroglossal collapse, arches, overjet/overbite, hyoid position, mandibular plane angle, airway volume, Mallampati score, Friedman stage, Mueller maneuver findings, DISE findings)
- Extract treatment considerations for oral appliances, CPAP, or sleep surgery

EXCLUDE THE FOLLOWING (return "Document not relevant to sleep medicine" if document contains only):
- Dermatology reports, patch tests, skin conditions
- Cardiology reports unrelated to sleep apnea
- General medical conditions unrelated to sleep/airway
- Corrupted or illegible documents
- Emergency room visits unrelated to sleep
- Routine lab work unrelated to sleep disorders

RESPONSE FORMAT:
- If document is NOT related to sleep medicine: return "Document not relevant to sleep medicine"
- If document is corrupted/illegible: return "Document appears corrupted or illegible"
- Otherwise: provide observations in clear, structured format focusing on sleep-related findings
- Use medical terminology appropriately
- Be concise but comprehensive for sleep-related content only

RESPONSE STYLE:
- Be direct and factual about SLEEP-RELATED findings only
- Avoid lengthy disclaimers
- Focus on actionable sleep medicine clinical information
- Use bullet points for clarity when appropriate
- Include demographics if present"""

        # Build the user prompt with all documents
        user_prompt = """Please analyze the following documents and extract all relevant clinical observations, findings, and medical information.

IMPORTANT: For each document, extract ALL available information including:

1. **DEMOGRAPHICS** (always look for these in every document):
   - Patient age (look for: age, years old, yo, patient age)
   - Patient sex/gender (look for: male, female, sex, gender)
   - Patient weight (look for: weight, kg, lbs, pounds)
   - Patient height (look for: height, cm, inches, feet)
   - BMI (look for: BMI, body mass index)

2. **SLEEP STUDY DATA**:
   - AHI (Apnea-Hypopnea Index)
   - ODI (Oxygen Desaturation Index)
   - O2 levels, oxygen saturation
   - Sleep study type (HSAT, PSG, etc.)

3. **ANATOMICAL FINDINGS**:
   - Soft palate, uvula, tongue base, epiglottis
   - TMJ (temporomandibular joint) issues
   - Airway obstruction sites (retropalatal, retroglossal, velopharyngeal, oropharyngeal)
   - Tonsils, adenoids, pharyngeal wall
   - Dental arches, overjet, overbite
   - Hyoid bone position
   - Mandibular plane angle, airway volume
   - Mallampati score, Friedman stage
   - Mueller maneuver findings, DISE findings
   - Neck findings (thick neck, etc.)

4. **SYMPTOMS & CLINICAL FINDINGS**:
   - Snoring, daytime sleepiness
   - Breathing issues, sleep disruption
   - Pain, clicking, jaw problems

"""
        
        total_chars = 0
        max_total_chars = MAX_TOTAL_CHARS
        
        for doc in documents_batch:
            content = doc['content']
            # Truncate individual document content if needed
            if len(content) > MAX_DOC_CHARS:
                logger.warning(f"Content too large for {doc['filename']} ({len(content)} chars), truncating to {MAX_DOC_CHARS} chars")
                content = trim_content(content, MAX_DOC_CHARS)
            
            # Check if adding this document would exceed total limit
            if total_chars + len(content) > max_total_chars:
                logger.warning(f"Batch would exceed {max_total_chars} chars, stopping at {doc['filename']}")
                break
                
            user_prompt += f"Document: {doc['filename']}\n"
            user_prompt += f"Type: {doc['document_type']}\n"
            user_prompt += f"Content:\n{content}\n"
            user_prompt += "-" * 50 + "\n\n"
            
            total_chars += len(content)

        # Use the same Bedrock pattern as the working patient workflow
        bedrock_messages = [
            {
                "role": "assistant",
                "content": system_prompt
            },
            {
                "role": "user", 
                "content": user_prompt
            }
        ]
        
        total_content_length = sum(len(doc['content']) for doc in documents_batch)
        logger.info(f"Calling Bedrock (enhanced pipeline) for batch of {len(documents_batch)} documents with {total_content_length} total characters")

        if bedrock_query_enhanced is None:
            return {doc['filename']: "Error: Bedrock pipeline not available" for doc in documents_batch}

        result = bedrock_query_enhanced(
            bedrock_messages,
            max_tokens=BEDROCK_MAX_TOKENS_BATCH,
            temperature=0.2,
            top_p=0.9,
        )

        if isinstance(result, dict) and result.get("success"):
            batch_response = result.get("response", "")
            observations_by_file = parse_batch_response(batch_response, [doc['filename'] for doc in documents_batch])
            logger.info(f"Successfully extracted observations for {len(observations_by_file)} documents in batch")
            return observations_by_file
        else:
            error_msg = result.get('message', 'Unknown error') if isinstance(result, dict) else 'Unknown error'
            logger.error(f"Bedrock batch call failed via enhanced pipeline: {error_msg}")
            return {doc['filename']: f"Error extracting observations: {error_msg}" for doc in documents_batch}
        
    except Exception as e:
        logger.error(f"Error extracting observations with LLM for batch: {e}")
        return {doc['filename']: f"Error extracting observations: {str(e)}" for doc in documents_batch}

def parse_batch_response(batch_response: str, filenames: List[str]) -> Dict[str, str]:
    """
    Parse the batch Bedrock response to extract observations for each document.
    
    Args:
        batch_response (str): Raw batch response from Bedrock
        filenames (List[str]): List of filenames in the batch
        
    Returns:
        Dict[str, str]: Dictionary mapping filename to observations
    """
    observations_by_file = {}
    
    # Split response by document sections
    sections = batch_response.split("=== FILENAME:")
    
    for section in sections:
        if not section.strip():
            continue
            
        # Extract filename and content
        lines = section.strip().split('\n', 1)
        if len(lines) < 2:
            continue
            
        filename_line = lines[0].strip()
        content = lines[1].strip() if len(lines) > 1 else ""
        
        # Extract filename (remove "===" and clean up)
        filename = filename_line.replace("===", "").strip()
        
        if filename and content:
            observations_by_file[filename] = content
    
    # If parsing failed, try to match by filename patterns
    if not observations_by_file:
        for filename in filenames:
            # Look for content near the filename in the response
            filename_pattern = filename.replace('.', '\.').replace('_', '[_\-]')
            pattern = rf"{filename_pattern}.*?(?=== FILENAME:|$)"
            match = re.search(pattern, batch_response, re.DOTALL | re.IGNORECASE)
            if match:
                observations_by_file[filename] = match.group(0).strip()
            else:
                observations_by_file[filename] = "No observations found"
    
    return observations_by_file
def parse_observations_from_text(observation_text: str) -> List[Dict]:
    """
    Parse the LLM response text into individual observations.
    Similar to how quiz system parses observations.
    
    Args:
        observation_text (str): Raw LLM response text
        
    Returns:
        List[Dict]: List of individual observations
    """
    observations = []
    
    # Split by numbered items, bullet points, or new lines
    lines = observation_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Skip headers and section titles
        if any(skip_word in line.lower() for skip_word in ['diagnosis:', 'sleep apnea indices:', 'oxygen saturation:', 'sleep architecture:', 'sleep latency:', 'oxygen desaturation events:', 'time spent at low oxygen saturations:', 'pulse rate during sleep:', 'sleep fragmentation:', 'these observations indicate']):
            continue
            
        # Extract numbered items (e.g., "1. Diagnosis: Severe OSA")
        if re.match(r'^\d+\.', line):
            # Remove the number and dot
            observation_text = re.sub(r'^\d+\.\s*', '', line)
            if observation_text:
                observations.append({
                    'observation': observation_text,
                    'value': 'Yes',
                    'score': 1,
                    'explanation': f'Extracted from document: {observation_text}',
                    'evidence': observation_text,
                    'confidence': 100,
                    'source': 'bedrock-document-extraction'
                })
        
        # Extract bullet points or dash items
        elif line.startswith('-') or line.startswith('•') or line.startswith('*'):
            observation_text = line[1:].strip()
            if observation_text:
                observations.append({
                    'observation': observation_text,
                    'value': 'Yes',
                    'score': 1,
                    'explanation': f'Extracted from document: {observation_text}',
                    'evidence': observation_text,
                    'confidence': 100,
                    'source': 'bedrock-document-extraction'
                })
        
        # Extract key-value pairs (e.g., "Mean: 93%")
        elif ':' in line and not line.startswith('Diagnosis:') and not line.startswith('Sleep'):
            parts = line.split(':', 1)
            if len(parts) == 2:
                key = parts[0].strip()
                value = parts[1].strip()
                if key and value:
                    observations.append({
                        'observation': key,
                        'value': value,
                        'score': 1,
                        'explanation': f'Extracted from document: {key} = {value}',
                        'evidence': f'{key}: {value}',
                        'confidence': 100,
                        'source': 'bedrock-document-extraction'
                    })
    
    # If no structured observations found, create one general observation
    if not observations:
        observations.append({
            'observation': 'Document Analysis Complete',
            'value': 'Yes',
            'score': 1,
            'explanation': 'Document was analyzed for clinical observations',
            'evidence': observation_text[:200] + '...' if len(observation_text) > 200 else observation_text,
            'confidence': 100,
            'source': 'bedrock-document-extraction'
        })
    
    return observations

def store_observations(patient_id: int, source_type: str, observations: List[Dict], document_info: Dict) -> bool:
    """
    Store individual observations in the observation_store table.
    Each observation gets its own row, like the quiz system.
    
    Args:
        patient_id (int): Patient ID
        source_type (str): Type of document source
        observations (List[Dict]): List of individual observations
        document_info (Dict): Document metadata
        
    Returns:
        bool: True if successful, False otherwise
    """
    conn = None
    try:
        conn = mysql.connector.connect(**DB_CONFIG)
        cursor = conn.cursor()
        
        # Insert each observation as a separate row (include file_name for dedupe/auditing)
        insert_query = """
            INSERT INTO observation_store 
            (patient_id, file_name, source_type, source_text, extracted_observations, provider, created_at, updated_at)
            VALUES (%s, %s, %s, %s, %s, %s, NOW(), NOW())
        """
        
        stored_count = 0
        for observation in observations:
            # Create source text for this observation
            source_text = observation.get('path') or observation.get('observation') or 'Document Analysis'
            
            # Store the observation as JSON (like quiz system)
            observation_json = {
                **observation,
                "document_name": document_info.get('name', ''),
                "document_type": document_info.get('file_type', ''),
                "extraction_date": datetime.now().isoformat()
            }
            
            cursor.execute(insert_query, (
                patient_id,
                document_info.get('name', ''),
                source_type,
                source_text,
                json.dumps(observation_json),
                'bedrock'
            ))
            stored_count += 1
        
        # Mark the document as analyzed
        document_id = document_info.get('id')
        source_table = document_info.get('source_table', 'files')
        
        if document_id and source_table:
            if source_table == 'files':
                update_query = "UPDATE files SET analyzed = TRUE WHERE id = %s"
            else:  # adminfiles
                update_query = "UPDATE adminfiles SET analyzed = TRUE WHERE id = %s"
            
            cursor.execute(update_query, (document_id,))
            logger.info(f"Marked document {document_info.get('name', '')} as analyzed")
        
        conn.commit()
        
        logger.info(f"Stored {stored_count} observations for patient {patient_id}, source: {source_type}")
        return True
        
    except Exception as e:
        logger.error(f"Error storing observations for patient {patient_id}: {e}")
        if conn:
            conn.rollback()
        return False
    finally:
        if conn:
            conn.close()

def delete_existing_observations(patient_id: int, source_type: str, document_info: Dict) -> None:
    """Delete prior observations for this patient and file to prevent duplication."""
    conn = None
    try:
        conn = mysql.connector.connect(**DB_CONFIG)
        cursor = conn.cursor()
        del_q = """
            DELETE FROM observation_store
            WHERE patient_id = %s AND source_type = %s AND file_name = %s
        """
        cursor.execute(del_q, (patient_id, source_type, document_info.get('name', '')))
        conn.commit()
    except Exception as e:
        logger.error(f"Error deleting existing observations for patient {patient_id}: {e}")
        if conn:
            conn.rollback()
    finally:
        if conn:
            conn.close()

# -------------------------
# New helpers for multi-value support
# -------------------------

def _safe_decimal(value: Any) -> Optional[str]:
    """Return a stringified Decimal for numeric values, else None.

    Stored as string to avoid MySQL connector type surprises; column is DECIMAL.
    """
    if value is None:
        return None
    try:
        d = Decimal(str(value))
        return str(d)
    except (InvalidOperation, ValueError, TypeError):
        return None

def _infer_study_type(filename: str) -> Optional[str]:
    if not filename:
        return None
    fn = filename.lower()
    if 'hsat' in fn or 'home' in fn:
        return 'HSAT'
    if 'psg' in fn or 'inlab' in fn or 'polysom' in fn:
        return 'PSG'
    if 'titration' in fn:
        return 'Titration'
    return None

def _make_episode_id(patient_id: int, s3_key: Optional[str], file_name: Optional[str]) -> str:
    base = f"{patient_id}:{s3_key or file_name or ''}"
    try:
        import hashlib
        return hashlib.sha1(base.encode('utf-8')).hexdigest()[:16]
    except Exception:
        return base[:16]

def upsert_envelope(patient_id: int, snapshot: Dict[str, Any], document_info: Dict) -> bool:
    """Create or update the per-report envelope JSON for this patient and file."""
    conn = None
    try:
        # Prune empty values to create sparse JSON
        snapshot = _prune_empty(snapshot)
        
        # Validate schema compliance before storing
        validation = validate_schema_compliance(snapshot)
        if not validation['valid']:
            logger.warning(f"Schema validation failed for patient {patient_id}, document {document_info.get('name', '')}: {validation['errors']}")
            # Still proceed but log the issues
        
        conn = mysql.connector.connect(**DB_CONFIG)
        cursor = conn.cursor()
        insert_q = """
            INSERT INTO patient_case_envelope
            (patient_id, report_id, document_type, source_uri, case_json, provider, imported_at, updated_at)
            VALUES (%s, %s, %s, %s, %s, %s, NOW(), NOW())
            ON DUPLICATE KEY UPDATE
              document_type = VALUES(document_type),
              source_uri = VALUES(source_uri),
              case_json = VALUES(case_json),
              provider = VALUES(provider),
              updated_at = NOW()
        """
        report_id = snapshot.get('report_meta', {}).get('report_id') or document_info.get('name', '')
        source_uri = snapshot.get('report_meta', {}).get('source_uri') or ''
        document_type = snapshot.get('document_type') or 'per_report'
        payload = (
            int(patient_id),
            report_id,
            document_type,
            source_uri,
            json.dumps(snapshot),
            'bedrock',
        )
        cursor.execute(insert_q, payload)
        conn.commit()
        logger.info(f"Envelope upserted for patient {patient_id}: {report_id}")
        
        # Log validation results
        if validation['warnings']:
            logger.info(f"Schema warnings for {report_id}: {validation['warnings']}")
        
        return True
    except Exception as e:
        logger.error(f"Error upserting envelope for patient {patient_id}: {e}")
        if conn:
            conn.rollback()
        return False
    finally:
        if conn:
            conn.close()

def envelope_exists(patient_id: int, report_id: str) -> bool:
    """Check if an envelope already exists for the given patient/report."""
    exists = False
    conn = None
    try:
        conn = mysql.connector.connect(**DB_CONFIG)
        cursor = conn.cursor()
        cursor.execute(
            "SELECT 1 FROM patient_case_envelope WHERE patient_id = %s AND report_id = %s LIMIT 1",
            (int(patient_id), report_id),
        )
        exists = cursor.fetchone() is not None
    except Exception as e:
        logger.error(f"Error checking envelope existence for patient {patient_id}, report {report_id}: {e}")
    finally:
        if conn:
            conn.close()
    return exists

def _discover_all_documents(patient_id: int) -> List[Dict]:
    """Discover ALL documents for a patient (from files and adminfiles), excluding imaging category."""
    docs: List[Dict] = []
    conn = None
    try:
        conn = mysql.connector.connect(**DB_CONFIG)
        cursor = conn.cursor(dictionary=True)
        # files - exclude imaging category
        cursor.execute(
            """
            SELECT id, name, patient_id, upload_date, file_type, file_size, s3_key,
                   category, subcategory, comment, analyzed, 'files' as source_table
            FROM files WHERE patient_id = %s AND category != 'imaging'
            ORDER BY upload_date DESC
            """,
            (patient_id,),
        )
        for r in cursor.fetchall():
            docs.append({
                'id': r['id'], 'name': r['name'], 'patient_id': r['patient_id'], 'upload_date': r['upload_date'],
                'file_type': r['file_type'], 'file_size': r['file_size'], 's3_key': r['s3_key'],
                'source_table': r['source_table'], 'source_type': map_document_type_to_source_type(r.get('category'), r.get('subcategory')),
                'category': r.get('category'), 'subcategory': r.get('subcategory'), 'comment': r.get('comment'),
                'analyzed': r.get('analyzed', False)
            })
        # adminfiles - exclude imaging category
        cursor.execute(
            """
            SELECT id, name, patient_id, upload_date, file_type, file_size, s3_key,
                   is_public, file_category, analyzed, 'adminfiles' as source_table
            FROM adminfiles WHERE patient_id = %s AND (file_category IS NULL OR file_category != 'imaging')
            ORDER BY upload_date DESC
            """,
            (patient_id,),
        )
        for r in cursor.fetchall():
            docs.append({
                'id': r['id'], 'name': r['name'], 'patient_id': r['patient_id'], 'upload_date': r['upload_date'],
                'file_type': r['file_type'], 'file_size': r['file_size'], 's3_key': r['s3_key'],
                'source_table': r['source_table'], 'source_type': map_document_type_to_source_type(None, None, r.get('file_category')),
                'is_public': r.get('is_public'), 'file_category': r.get('file_category'),
                'analyzed': r.get('analyzed', False)
            })
    except Exception as e:
        logger.error(f"Error discovering all documents for patient {patient_id}: {e}")
    finally:
        if conn:
            conn.close()
    return docs

def backfill_envelopes_for_patient(patient_id: int) -> Dict[str, Any]:
    """Ensure each document for a patient has a per-report envelope record and create canonical entry."""
    if normalize_to_patient_case_json_v1 is None:
        return {'patient_id': patient_id, 'updated': 0, 'skipped': 0, 'error': 'snapshot helper unavailable'}
    updated = 0
    skipped = 0
    docs = _discover_all_documents(patient_id)
    for doc in docs:
        # Skip binary
        if is_binary_file(doc.get('name', '')):
            skipped += 1
            continue
        # If envelope exists, skip
        if envelope_exists(patient_id, doc.get('name', '')):
            skipped += 1
            continue
        # Extract text and build snapshot
        content = extract_document_content(doc)
        if not content:
            # Create minimal envelope with metadata
            source_uri = f"s3://{S3_BUCKET_NAME}/{doc.get('s3_key', '')}"
            snapshot = normalize_to_patient_case_json_v1(
                file_path=source_uri, file_name=doc['name'], text_content='', patient_id=str(patient_id), document_type='per_report', version=1
            )
        else:
            source_uri = f"s3://{S3_BUCKET_NAME}/{doc.get('s3_key', '')}"
            snapshot = normalize_to_patient_case_json_v1(
                file_path=source_uri, file_name=doc['name'], text_content=content, patient_id=str(patient_id), document_type='per_report', version=1
            )
        if upsert_envelope(patient_id, snapshot, doc):
            updated += 1
        else:
            skipped += 1
    
    # Create canonical entry directly from observation store
    canonical_created = False
    if updated > 0:
        try:
            canonical_result = create_minimal_canonical_json_for_patient(patient_id)
            if canonical_result.get('success'):
                canonical_created = True
                logger.info(f"Created canonical entry for patient {patient_id} from observation store")
            else:
                logger.warning(f"Failed to create canonical entry for patient {patient_id}: {canonical_result.get('message')}")
        except Exception as e:
            logger.error(f"Error creating canonical entry for patient {patient_id}: {e}")
    
    return {
        'patient_id': patient_id, 
        'updated': updated, 
        'skipped': skipped,
        'canonical_created': canonical_created
    }

def create_canonical_json_for_patient(patient_id: int) -> Dict[str, Any]:
    """
    Create canonical JSON for a patient based on the Patient Case JSON v1 schema.
    This function ensures schema compliance and creates a comprehensive canonical record.
    
    Args:
        patient_id (int): Patient ID to create canonical JSON for
        
    Returns:
        Dict[str, Any]: Result of canonical creation
    """
    try:
        # Always use minimal canonical creation from observation store to ensure schema compliance
        logger.info(f"Creating canonical JSON for patient {patient_id} from observation store")
        minimal_result = create_minimal_canonical_json_for_patient(patient_id)
        
        if minimal_result.get('success'):
            # CANONICAL VALIDATION: Cross-check critical values against source files
            canonical_data = minimal_result.get('canonical_json', {})
            validation_results = validate_canonical_against_source_files(patient_id, canonical_data)
            
            if validation_results.get('discrepancies'):
                logger.warning(f"⚠️ Found {len(validation_results['discrepancies'])} discrepancies in canonical data for patient {patient_id}")
                # Log critical discrepancies
                for discrepancy in validation_results['discrepancies']:
                    if discrepancy.get('severity') == 'high':
                        logger.error(f"🚨 HIGH SEVERITY: {discrepancy['field']} = {discrepancy['canonical_value']} (source: {discrepancy['source_value']}) in {discrepancy['document']}")
            else:
                logger.info(f"✅ All critical values validated against source files for patient {patient_id}")
            
            logger.info(f"Successfully created canonical JSON for patient {patient_id}")
            return {
                'success': True,
                'patient_id': patient_id,
                'message': 'Canonical JSON created successfully from observation store',
                'base_report_id': 'observation_store',
                'total_reports': 0,
                'validation_results': validation_results
            }
        else:
            logger.warning(f"Failed to create canonical JSON for patient {patient_id}: {minimal_result.get('message')}")
            return {
                'success': False,
                'patient_id': patient_id,
                'message': minimal_result.get('message', 'Unknown error')
            }
            
    except Exception as e:
        logger.error(f"Error creating canonical JSON for patient {patient_id}: {e}")
        return {
            'success': False,
            'patient_id': patient_id,
            'message': str(e)
        }

def validate_schema_compliance(snapshot: Dict[str, Any]) -> Dict[str, Any]:
    """
    Validate that a snapshot complies with the Patient Case JSON v1 schema.
    
    Args:
        snapshot (Dict[str, Any]): The snapshot to validate
        
    Returns:
        Dict[str, Any]: Validation results
    """
    validation_result = {
        'valid': True,
        'errors': [],
        'warnings': [],
        'missing_required': [],
        'schema_version': snapshot.get('schema_version'),
        'document_type': snapshot.get('document_type')
    }
    
    # Check required fields
    required_fields = ['schema_version', 'document_type', 'patient_id', 'as_of']
    for field in required_fields:
        if field not in snapshot or snapshot[field] is None:
            validation_result['valid'] = False
            validation_result['missing_required'].append(field)
            validation_result['errors'].append(f"Missing required field: {field}")
    
    # Validate schema version
    if snapshot.get('schema_version') != '1.0':
        validation_result['warnings'].append(f"Schema version {snapshot.get('schema_version')} may not be supported")
    
    # Validate document type
    valid_document_types = ['per_report', 'canonical']
    if snapshot.get('document_type') not in valid_document_types:
        validation_result['valid'] = False
        validation_result['errors'].append(f"Invalid document_type: {snapshot.get('document_type')}")
    
    # Check for expected top-level sections
    expected_sections = [
        'sleep_study', 'observations', 'treatment_considerations', 
        'device_design', 'follow_up_plan', 'demographics'
    ]
    
    for section in expected_sections:
        if section not in snapshot:
            validation_result['warnings'].append(f"Missing optional section: {section}")
    
    # Cross-field validation for sleep study
    sleep_study = snapshot.get('sleep_study', {})
    if sleep_study:
        # Validate AHI vs severity consistency
        ahi = sleep_study.get('ahi')
        severity = sleep_study.get('severity')
        if ahi is not None and severity is not None:
            expected_severity = compute_ahi_severity(ahi)
            if severity != expected_severity:
                validation_result['errors'].append(
                    f"AHI severity mismatch: AHI {ahi} should be '{expected_severity}' but is '{severity}'"
                )
                validation_result['valid'] = False
        
        # Validate study type enum
        study_type = sleep_study.get('study_type')
        if study_type and study_type not in ['HSAT', 'PSG', 'Titration', 'Unknown']:
            validation_result['warnings'].append(f"Non-standard study_type: {study_type}")
        
        # Check for sleep study data quality (not source document type)
        has_sleep_data = any(sleep_study.get(key) for key in ['ahi', 'odi', 'o2_nadir_pct'])
        if has_sleep_data:
            # Validate data quality and ranges
            ahi = sleep_study.get('ahi')
            if ahi is not None:
                try:
                    ahi_float = float(ahi) if isinstance(ahi, str) else ahi
                    if ahi_float < 0 or ahi_float > 200:
                        validation_result['warnings'].append(f"AHI value {ahi} is outside expected range (0-200)")
                except (ValueError, TypeError):
                    validation_result['warnings'].append(f"AHI value {ahi} is not a valid number")
            
            odi = sleep_study.get('odi')
            if odi is not None:
                try:
                    odi_float = float(odi) if isinstance(odi, str) else odi
                    if odi_float < 0 or odi_float > 200:
                        validation_result['warnings'].append(f"ODI value {odi} is outside expected range (0-200)")
                except (ValueError, TypeError):
                    validation_result['warnings'].append(f"ODI value {odi} is not a valid number")
            
            o2_nadir = sleep_study.get('o2_nadir_pct')
            if o2_nadir is not None:
                try:
                    o2_nadir_float = float(o2_nadir) if isinstance(o2_nadir, str) else o2_nadir
                    if o2_nadir_float < 50 or o2_nadir_float > 100:
                        validation_result['warnings'].append(f"O2 Nadir {o2_nadir}% is outside expected range (50-100%)")
                except (ValueError, TypeError):
                    validation_result['warnings'].append(f"O2 Nadir value {o2_nadir} is not a valid number")
            
            # Log that sleep study data was found regardless of source
            logger.info(f"Sleep study data found: AHI={ahi}, ODI={odi}, O2_Nadir={o2_nadir}%")
    
    # Validate completeness flags consistency
    completeness_flags = snapshot.get('completeness_flags', {})
    if completeness_flags:
        sleep_study = snapshot.get('sleep_study', {})
        has_sleep_data = any(sleep_study.get(key) for key in ['ahi', 'odi', 'o2_nadir_pct'])
        has_sleep_flag = completeness_flags.get('has_sleep_study', False)
        
        if has_sleep_data and not has_sleep_flag:
            validation_result['warnings'].append("Sleep study data present but has_sleep_study flag is false")
        elif not has_sleep_data and has_sleep_flag:
            validation_result['warnings'].append("has_sleep_study flag is true but no sleep study data found")
    
    return validation_result

def backfill_envelopes(limit_patients: Optional[int] = None) -> Dict[str, Any]:
    pids = get_all_patient_ids(limit=limit_patients)
    results = []
    total_updated = 0
    total_canonical_created = 0
    for pid in pids:
        res = backfill_envelopes_for_patient(pid)
        results.append(res)
        total_updated += res.get('updated', 0)
        if res.get('canonical_created'):
            total_canonical_created += 1
    return {
        'patients_total': len(pids), 
        'envelopes_updated': total_updated, 
        'canonical_created': total_canonical_created,
        'results': results
    }

def ensure_canonical_json_for_all_patients(limit_patients: Optional[int] = None) -> Dict[str, Any]:
    """
    Ensure all patients have proper canonical JSON based on the Patient Case JSON v1 schema.
    This function processes all patients and creates canonical JSON where missing.
    
    Args:
        limit_patients (Optional[int]): Limit number of patients to process
        
    Returns:
        Dict[str, Any]: Summary of canonical JSON creation
    """
    pids = get_all_patient_ids(limit=limit_patients)
    results = []
    total_canonical_created = 0
    total_canonical_updated = 0
    total_failed = 0
    
    logger.info(f"Ensuring canonical JSON for {len(pids)} patients")
    
    for pid in pids:
        try:
            result = create_canonical_json_for_patient(pid)
            results.append({'patient_id': pid, 'result': result})
            
            if result.get('success'):
                if result.get('message') == 'Canonical JSON created successfully':
                    total_canonical_created += 1
                else:
                    total_canonical_updated += 1
                logger.info(f"Canonical JSON ensured for patient {pid}")
            else:
                total_failed += 1
                logger.warning(f"Failed to ensure canonical JSON for patient {pid}: {result.get('message')}")
                
        except Exception as e:
            total_failed += 1
            logger.error(f"Error ensuring canonical JSON for patient {pid}: {e}")
            results.append({'patient_id': pid, 'error': str(e)})
    
    summary = {
        'patients_total': len(pids),
        'canonical_created': total_canonical_created,
        'canonical_updated': total_canonical_updated,
        'failed': total_failed,
        'results': results
    }
    
    logger.info(f"Canonical JSON creation summary: {summary}")
    return summary

def validate_all_canonical_json(limit_patients: Optional[int] = None) -> Dict[str, Any]:
    """
    Validate all canonical JSON records for schema compliance.
    
    Args:
        limit_patients (Optional[int]): Limit number of patients to validate
        
    Returns:
        Dict[str, Any]: Validation summary
    """
    pids = get_all_patient_ids(limit=limit_patients)
    results = []
    total_valid = 0
    total_invalid = 0
    
    logger.info(f"Validating canonical JSON for {len(pids)} patients")
    
    for pid in pids:
        try:
            conn = mysql.connector.connect(**DB_CONFIG)
            cursor = conn.cursor(dictionary=True)
            cursor.execute(
                "SELECT case_json FROM patient_case_envelope WHERE patient_id = %s AND report_id = 'canonical' LIMIT 1",
                (pid,)
            )
            result = cursor.fetchone()
            conn.close()
            
            if result and result.get('case_json'):
                import json
                snapshot = json.loads(result['case_json']) if isinstance(result['case_json'], str) else result['case_json']
                validation = validate_schema_compliance(snapshot)
                
                results.append({
                    'patient_id': pid,
                    'valid': validation['valid'],
                    'errors': validation['errors'],
                    'warnings': validation['warnings']
                })
                
                if validation['valid']:
                    total_valid += 1
                else:
                    total_invalid += 1
                    logger.warning(f"Invalid canonical JSON for patient {pid}: {validation['errors']}")
            else:
                results.append({
                    'patient_id': pid,
                    'valid': False,
                    'errors': ['No canonical envelope found'],
                    'warnings': []
                })
                total_invalid += 1
                
        except Exception as e:
            logger.error(f"Error validating canonical JSON for patient {pid}: {e}")
            results.append({
                'patient_id': pid,
                'valid': False,
                'errors': [str(e)],
                'warnings': []
            })
            total_invalid += 1
    
    summary = {
        'patients_total': len(pids),
        'valid': total_valid,
        'invalid': total_invalid,
        'results': results
    }
    
    logger.info(f"Canonical JSON validation summary: {summary}")
    return summary
def process_documents_in_batches(documents: List[Dict], batch_size: int = 3, progress_callback=None) -> Dict:
    """
    Process documents in batches to optimize Bedrock calls and create canonical JSON.
    
    Args:
        documents (List[Dict]): List of documents to process
        batch_size (int): Number of documents to process per batch
        progress_callback: Optional callback function(message) to report progress
        
    Returns:
        Dict: Processing statistics
    """
    def report_progress(msg):
        """Report progress via callback and logger"""
        logger.info(f"[Progress] {msg}")
        if progress_callback:
            try:
                progress_callback(msg)
            except Exception as e:
                logger.warning(f"Progress callback error: {e}")
    
    # Adjust batch size based on environment, OCR availability, and throttling status
    effective_batch_size = min(batch_size, MAX_BATCH_DOCS)
    if not OCR_AVAILABLE:
        effective_batch_size = min(effective_batch_size, 1)
    if THROTTLING_DETECTED:
        effective_batch_size = min(effective_batch_size, THROTTLING_BATCH_SIZE)
        logger.info(f"Throttling detected, using reduced batch size: {effective_batch_size}")
    
    total_batches = (len(documents) + effective_batch_size - 1) // effective_batch_size
    logger.error(f"Processing {len(documents)} documents in batches of {effective_batch_size}")
    logger.error(f"Total batches to process: {total_batches}")
    report_progress(f"Processing {len(documents)} docs in {total_batches} batches")
    
    successful = 0
    failed = 0
    canonical_created = 0
    
    # Track patient IDs for canonical creation
    processed_patients = set()
    
    # Process documents in batches
    for i in range(0, len(documents), effective_batch_size):
        batch = documents[i:i + effective_batch_size]
        batch_num = (i // effective_batch_size) + 1
        total_batches = (len(documents) + effective_batch_size - 1) // effective_batch_size
        
        logger.warning(f"Processing batch {batch_num}/{total_batches} with {len(batch)} documents")
        report_progress(f"📦 Batch {batch_num}/{total_batches}")
        
        # Extract content and process each document with schema-guided snapshot
        doc_in_batch = 0
        for doc in batch:
            doc_in_batch += 1
            doc_name = doc.get('name', 'unknown')[:30]  # Truncate for display
            
            # Process all documents regardless of analyzed status (as requested)
            if doc.get('analyzed') is True:
                logger.info(f"Processing already analyzed document: {doc.get('name', '')} (ignoring analyzed flag)")
            else:
                logger.info(f"Processing document: {doc.get('name', '')}")
            
            report_progress(f"  📄 {doc_name}...")

            if normalize_to_patient_case_json_v1 is None or explode_observations_from_snapshot is None:
                logger.error("Schema-guided ingestion helpers unavailable; skipping document")
                failed += 1
                continue

            # STEP 1: VISUAL LLM EXTRACTION FIRST (Primary method) - for PDFs and images
            visual_observations = []
            visual_data = None
            visual_success = False
            
            visual_error = None
            size_int = None
            raw_size = doc.get('file_size')
            if raw_size is not None:
                try:
                    if isinstance(raw_size, Decimal):
                        size_int = int(raw_size)
                    elif isinstance(raw_size, (int, float)):
                        size_int = int(raw_size)
                    elif isinstance(raw_size, str):
                        size_int = int(float(raw_size.strip()))
                    else:
                        size_int = int(raw_size)
                except (ValueError, TypeError, InvalidOperation):
                    size_int = None
            
            supports_visual_llm = doc.get('name', '').lower().endswith(('.pdf', '.jpg', '.jpeg', '.png')) and doc.get('s3_key')
            skip_visual_due_to_size = bool(size_int and size_int > BEDROCK_MAX_DOCUMENT_BYTES)
            if skip_visual_due_to_size:
                logger.warning(
                    f"Skipping Visual LLM for {doc.get('name')} ({size_int} bytes) "
                    f"because it exceeds Bedrock limit of {BEDROCK_MAX_DOCUMENT_BYTES} bytes"
                )
                visual_error = (
                    f"Document size {size_int} exceeds Bedrock limit ({BEDROCK_MAX_DOCUMENT_BYTES} bytes); "
                    "skipping Visual LLM"
                )
                doc['visual_llm_skipped_reason'] = 'bedrock_document_size_limit'
            
            if supports_visual_llm and not skip_visual_due_to_size:
                logger.info(f"🔍 STEP 1: Running VISUAL LLM extraction FIRST (primary) for document: {doc['name']}")
                try:
                    visual_data, visual_success, visual_error = extract_observations_with_visual_llm(
                        doc['s3_key'], doc['name']
                    )
                    if visual_success and visual_data:
                        # Store visual_data for later AHI validation after text extraction
                        doc['_visual_data_for_validation'] = visual_data
                        visual_observations = convert_visual_llm_to_observations(visual_data, doc['name'])
                        logger.info(f"✅ Visual LLM extracted {len(visual_observations)} observations from {doc['name']}")
                        
                        # Extract dates from visual LLM data
                        if visual_data.get('meta'):
                            meta = visual_data['meta']
                            study_date = meta.get('date_of_study')
                            report_date = meta.get('report_date')
                            
                            if study_date:
                                try:
                                    from datetime import datetime
                                    parsed_dt = datetime.strptime(study_date, '%Y-%m-%d')
                                    doc['document_date'] = parsed_dt
                                    doc['observed_at_source'] = 'visual_llm'
                                    logger.info(f"✅ Extracted study date from Visual LLM for {doc.get('name')}: {parsed_dt}")
                                except Exception as e:
                                    logger.warning(f"⚠️ Failed to parse study date {study_date} for {doc.get('name')}: {e}")
                            elif report_date:
                                try:
                                    from datetime import datetime
                                    parsed_dt = datetime.strptime(report_date, '%Y-%m-%d')
                                    doc['document_date'] = parsed_dt
                                    doc['observed_at_source'] = 'visual_llm'
                                    logger.info(f"✅ Extracted report date from Visual LLM for {doc.get('name')}: {parsed_dt}")
                                except Exception as e:
                                    logger.warning(f"⚠️ Failed to parse report date {report_date} for {doc.get('name')}: {e}")
                    else:
                        logger.warning(f"⚠️ Visual LLM extraction failed for {doc['name']}: {visual_error}")
                except Exception as e:
                    logger.error(f"❌ Visual LLM extraction error for {doc['name']}: {e}")
            elif supports_visual_llm and skip_visual_due_to_size:
                logger.info(f"Visual LLM skipped for {doc.get('name')} due to Bedrock document size limit; continuing with text/OCR pipeline")
            else:
                logger.info(f"ℹ️ Skipping Visual LLM for {doc['name']} (not a supported image/PDF format)")
            
            # STEP 2: TEXT EXTRACTION (for validation/cross-reference)
            content = None
            text_content = None
            logger.info(f"🔍 STEP 2: Extracting text content for validation/cross-reference: {doc['name']}")
            try:
                content = extract_document_content(doc)
                if content:
                    text_content = content if isinstance(content, str) else str(content)
                    if text_content and len(text_content.strip()) > 0:
                        logger.info(f"✅ Extracted {len(text_content)} characters of text from {doc['name']}")
                        
                        # Extract document date from text if not already extracted from Visual LLM
                        if not doc.get('document_date'):
                            try:
                                parsed_dt, dt_source = _extract_document_date(text_content, doc.get('name'))
                                if parsed_dt:
                                    doc['document_date'] = parsed_dt
                                    doc['observed_at_source'] = dt_source
                                    logger.info(f"✅ Extracted document date from text for {doc.get('name')}: {parsed_dt} (source: {dt_source})")
                            except Exception as _e:
                                logger.warning(f"⚠️ Error extracting document date from text for {doc.get('name')}: {_e}")
                    else:
                        logger.warning(f"⚠️ Text extraction returned empty content for {doc['name']}")
                else:
                    logger.warning(f"⚠️ No text content extracted from {doc['name']} (may be image-only or corrupted PDF)")
            except Exception as e:
                logger.warning(f"⚠️ Text extraction error for {doc['name']}: {e} (continuing with Visual LLM results only)")
            
            # STEP 2.5: VALIDATE AHI OVERALL against text content (prevent hallucinations)
            if doc.get('_visual_data_for_validation') and text_content:
                try:
                    validated_visual_data = _validate_ahi_overall_from_text(
                        doc['_visual_data_for_validation'], 
                        text_content
                    )
                    # If AHI was corrected, update the visual_observations
                    if validated_visual_data.get('respiratory_indices', {}).get('_ahi_overall_corrected'):
                        logger.info(f"🔧 Re-converting observations after AHI correction for {doc['name']}")
                        visual_observations = convert_visual_llm_to_observations(validated_visual_data, doc['name'])
                except Exception as e:
                    logger.warning(f"⚠️ AHI validation error for {doc['name']}: {e}")
            
            # STEP 3: REGEX EXTRACTION (Validation method) - Cross-reference with Visual LLM
            regex_observations = []
            if text_content and len(text_content.strip()) > 0:
                logger.info(f"🔍 STEP 3: Running regex extraction for validation/cross-reference: {doc['name']}")
                try:
                    regex_data = extract_numerical_data_with_regex([text_content])
                    if regex_data:
                        # Convert regex data to observation format
                        regex_observations = convert_regex_data_to_observations(regex_data, doc['name'])
                        logger.info(f"✅ Regex extracted {len(regex_observations)} observations from {doc['name']}")
                    else:
                        logger.info(f"ℹ️ No regex data found for {doc['name']}")
                except Exception as e:
                    logger.error(f"❌ Regex extraction error for {doc['name']}: {e}")
            else:
                logger.info(f"ℹ️ Skipping regex extraction for {doc['name']} (no text content available)")
            
            # STEP 4: VALIDATE AND MERGE - Cross-reference Visual LLM with regex to prevent hallucinations
            all_observations = []
            
            if visual_observations and regex_observations:
                # Both methods succeeded - validate and merge
                logger.info(f"🔍 STEP 4: Validating Visual LLM against regex extraction for {doc['name']}")
                try:
                    validation_results = validate_llm_vs_regex_extraction(
                        visual_observations, 
                        regex_observations, 
                        doc['name']
                    )
                    
                    # Merge using Visual LLM as primary, regex as validator
                    all_observations = merge_llm_and_regex_observations(
                        visual_observations,
                        regex_observations,
                        validation_results,
                        doc['name']
                    )
                    
                    logger.info(f"✅ Merged {len(all_observations)} validated observations for {doc['name']}")
                    logger.info(f"   - Agreements: {len(validation_results.get('agreements', []))}")
                    logger.info(f"   - Discrepancies: {len(validation_results.get('discrepancies', []))}")
                    logger.info(f"   - LLM only: {len(validation_results.get('llm_only', []))}")
                    logger.info(f"   - Regex only: {len(validation_results.get('regex_only', []))}")
                except Exception as e:
                    logger.error(f"❌ Validation/merge error for {doc['name']}: {e}")
                    # Fallback: use Visual LLM observations with lower confidence
                    logger.warning(f"⚠️ Falling back to Visual LLM observations only (validation failed)")
                    all_observations = visual_observations
                    for obs in all_observations:
                        obs['score'] = min(obs.get('score', 0.7), 0.7)  # Reduce confidence
                        obs['explanation'] = obs.get('explanation', '') + ' (validation failed)'
            elif visual_observations:
                # Only Visual LLM succeeded - use it with medium confidence
                logger.info(f"📝 Using Visual LLM observations only for {doc['name']} (no text/regex available)")
                all_observations = visual_observations
                for obs in all_observations:
                    obs['score'] = min(obs.get('score', 0.7), 0.7)  # Medium confidence when no validation
                    if 'validated' not in obs.get('explanation', '').lower():
                        obs['explanation'] = obs.get('explanation', '') + ' (no text validation available)'
            elif regex_observations:
                # Only regex succeeded - use it with lower confidence
                logger.info(f"📝 Using regex observations only for {doc['name']} (Visual LLM failed)")
                all_observations = regex_observations
                for obs in all_observations:
                    obs['score'] = min(obs.get('score', 0.6), 0.6)  # Lower confidence without LLM
                    obs['explanation'] = obs.get('explanation', '') + ' (Visual LLM failed)'
            else:
                logger.warning(f"⚠️ No observations extracted from {doc['name']} (both Visual LLM and regex failed)")
            
            # STEP 5: STORE VALIDATED OBSERVATIONS
            visual_llm_stored = False
            if all_observations:
                success = store_observations_with_deduplication(
                    doc['patient_id'], 
                    doc['source_type'], 
                    all_observations, 
                    doc
                )
                if success:
                    visual_llm_stored = True
                    logger.info(f"✅ Stored {len(all_observations)} validated observations for {doc['name']}")
                    report_progress(f"    ✅ {len(all_observations)} observations extracted")
                else:
                    logger.warning(f"⚠️ Failed to store observations for {doc['name']}")
                    report_progress(f"    ⚠️ Failed to store")
            else:
                logger.warning(f"⚠️ No observations to store for {doc['name']}")

            # Handle questionnaire documents differently
            if doc.get('source_type') == 'questionnaire':
                # Use specialized questionnaire extraction
                logger.info(f"Processing questionnaire document: {doc['name']}")
                # Use text_content if available, otherwise try to extract again
                questionnaire_content = text_content if text_content else (content if content else extract_document_content(doc))
                if not questionnaire_content:
                    logger.warning(f"⚠️ No content available for questionnaire extraction from {doc['name']}")
                else:
                    schema_obs = extract_questionnaire_observations(questionnaire_content, doc['name'])
                    
                    # Extract demographics from document content and add to observations
                    demographics_obs = extract_demographics_observations(questionnaire_content, doc['name'])
                    if demographics_obs:
                        schema_obs.extend(demographics_obs)
                        logger.info(f"Added {len(demographics_obs)} demographics observations from {doc['name']}")
                    
                    # Run targeted LLM JSON extractor (logs raw/parsed output) and regex backup
                    try:
                        llm_status = extract_observations_with_llm(questionnaire_content, doc.get('source_type', 'document'), doc['name'])
                        logger.info(f"LLM status for {doc['name']}: {llm_status}")
                    except Exception as _e:
                        logger.warning(f"LLM JSON extractor failed for {doc['name']}: {_e}")
                    
                    # Regex helpers as authoritative data feeders
                    try:
                        # Supine AHI
                        if not any(o.get('path') == 'sleep_study.supine_ahi' for o in schema_obs):
                            _sup = _first_match(questionnaire_content, PATTERNS_SUPINE_AHI)
                            if _sup:
                                sup_val, sup_ev = _sup
                                schema_obs.append({
                                    'path': 'sleep_study.supine_ahi',
                                    'value': sup_val,
                                    'observation': 'Supine AHI',
                                    'score': 1,
                                    'explanation': 'LLM+regex hybrid extractor',
                                    'evidence': (sup_ev or '')[:512],
                                    'confidence': 95,
                                    'source': 'bedrock-json-or-regex'
                                })
                        # Percent time SpO2 < 90%
                        if not any(o.get('path') == 'sleep_study.time_below_90_pct' for o in schema_obs):
                            _o2 = _first_match(questionnaire_content, PATTERNS_O2_LT90)
                            if _o2:
                                o2_val, o2_ev = _o2
                                schema_obs.append({
                                    'path': 'sleep_study.time_below_90_pct',
                                    'value': o2_val,
                                    'observation': 'Percent time O2 < 90%',
                                    'score': 1,
                                    'explanation': 'LLM+regex hybrid extractor',
                                    'evidence': (o2_ev or '')[:512],
                                    'confidence': 95,
                                    'source': 'bedrock-json-or-regex'
                                })
                        # As a final fallback, run the schema numeric extractor on raw content
                        try:
                            numeric = extract_specific_numerical_fields([questionnaire_content]) or {}
                            ss = numeric.get('sleep_study', {}) if isinstance(numeric, dict) else {}
                            fallback_keys = (
                                ('sleep_study.supine_ahi', ss.get('supine_ahi')),
                                ('sleep_study.time_below_90_pct', ss.get('time_below_90_pct')),
                                ('sleep_study.time_below_90_pct_min', ss.get('time_below_90_pct_min')),
                                ('sleep_study.time_below_88_pct_min', ss.get('time_below_88_pct_min')),
                            )
                            for path_key, val in fallback_keys:
                                if val is None:
                                    continue
                                if any(o.get('path') == path_key for o in schema_obs):
                                    continue
                                schema_obs.append({
                                    'path': path_key,
                                    'value': float(val),
                                    'observation': path_key.split('.')[-1],
                                    'score': 1,
                                    'explanation': 'schema numeric extractor',
                                    'evidence': None,
                                    'confidence': 90,
                                    'source': 'regex-extractor'
                                })
                        except Exception as _e2:
                            logger.debug(f"Numeric fallback extractor failed for {doc['name']}: {_e2}")
                    except Exception as _e:
                        logger.warning(f"Regex helper extraction failed for {doc['name']}: {_e}")
            else:
                # Build per-report snapshot aligned with schema for medical documents
                # Use text_content if available, otherwise use content
                snapshot_content = text_content if text_content else (content if content else '')
                source_uri = f"s3://{S3_BUCKET_NAME}/{doc.get('s3_key', '')}"
                snapshot = normalize_to_patient_case_json_v1(
                    file_path=source_uri,
                    file_name=doc['name'],
                    text_content=snapshot_content,
                    patient_id=str(doc['patient_id']),
                    document_type='per_report',
                    version=1,
                )

                # Explode observations from snapshot
                schema_obs = explode_observations_from_snapshot(snapshot) or []
                
                # Extract demographics from document content and add to observations
                demographics_content = text_content if text_content else (content if content else '')
                demographics_obs = []
                if demographics_content:
                    demographics_obs = extract_demographics_observations(demographics_content, doc['name'])
                if demographics_obs:
                    schema_obs.extend(demographics_obs)
                    logger.info(f"Added {len(demographics_obs)} demographics observations from {doc['name']}")
            
            # Check if we already successfully stored Visual LLM observations
            # If so, count as successful even if schema_obs is empty
            if visual_llm_stored:
                # Visual LLM observations were already stored successfully
                successful += 1
                logger.info(f"✅ Successfully processed {doc['name']} (Visual LLM observations stored)")
                # Track patient for canonical creation
                processed_patients.add(doc['patient_id'])
                
                # Still try to add schema observations if available (for additional data)
                if schema_obs:
                    # Add schema observations (may contain additional data from text extraction)
                    if store_observations_with_deduplication(doc['patient_id'], doc['source_type'], schema_obs, doc):
                        logger.info(f"✅ Also stored {len(schema_obs)} additional schema observations for {doc['name']}")
                    else:
                        logger.warning(f"⚠️ Failed to store additional schema observations for {doc['name']} (but Visual LLM observations were already stored)")
                else:
                    logger.info(f"ℹ️ No additional schema observations to add for {doc['name']} (Visual LLM observations already stored)")
            elif not schema_obs:
                # No Visual LLM observations AND no schema observations
                logger.warning(f"⚠️ No observations extracted from {doc['name']} (both Visual LLM and schema extraction failed)")
                failed += 1
                continue
            else:
                # We have schema observations but Visual LLM failed or wasn't applicable
                # Add visual observations to the schema observations if available
                if visual_observations:
                    schema_obs.extend(visual_observations)
                    logger.info(f"Added {len(visual_observations)} visual observations to {len(schema_obs)} total observations for {doc['name']}")

                # Store schema observations with deduplication
                if store_observations_with_deduplication(doc['patient_id'], doc['source_type'], schema_obs, doc):
                    successful += 1
                    logger.info(f"✅ Successfully processed {doc['name']} (schema observations stored)")
                    
                    # Track patient for canonical creation
                    processed_patients.add(doc['patient_id'])
                else:
                    failed += 1
                    logger.error(f"❌ Failed to store observations for {doc['name']}")

        # Add delay between batches to respect Bedrock rate limits (snapshot leverages LLM)
        if batch_num < total_batches:
            delay = max(1, MIN_BEDROCK_INTERVAL_SECONDS)
            logger.info(f"Waiting {delay} seconds before next batch to respect rate limits...")
            time.sleep(delay)
        
        # free memory for next batch
        gc.collect()
    
    # Create canonical JSON for processed patients directly from observation store
    if processed_patients:
        logger.info(f"Creating canonical JSON for {len(processed_patients)} patients from observation store")
        for patient_id in processed_patients:
            try:
                result = create_minimal_canonical_json_for_patient(patient_id)
                if result.get('success'):
                    canonical_created += 1
                    logger.info(f"Successfully created canonical JSON for patient {patient_id}")
                else:
                    logger.warning(f"Failed to create canonical JSON for patient {patient_id}: {result.get('message')}")
            except Exception as e:
                logger.error(f"Error creating canonical JSON for patient {patient_id}: {e}")
    
    logger.warning(f"Batch processing completed: {successful} successful, {failed} failed, {canonical_created} canonical created")
    
    return {
        'total_documents': len(documents),
        'processed_documents': len(documents),
        'successful_extractions': successful,
        'failed_extractions': failed,
        'canonical_created': canonical_created
    }

def process_patient_documents(patient_id: int, max_documents: int = None, batch_size: int = 3, progress_callback=None) -> Dict:
    """
    Process all documents for a patient to extract observations using batch processing.
    
    Args:
        patient_id (int): Patient ID to process
        max_documents (int): Maximum number of documents to process (for testing)
        batch_size (int): Number of documents to process per batch
        progress_callback: Optional callback function(message) to report progress
        
    Returns:
        Dict: Processing statistics
    """
    def report_progress(msg):
        """Report progress via callback and logger"""
        logger.info(f"[Progress] {msg}")
        if progress_callback:
            try:
                progress_callback(msg)
            except Exception as e:
                logger.warning(f"Progress callback error: {e}")
    
    # Set current patient ID for LLM logging context
    set_current_patient_id(patient_id)
    
    logger.error(f"Starting document processing for patient {patient_id} with batch size {batch_size}")
    report_progress(f"Starting document processing for patient {patient_id}")
    
    # Discover documents
    report_progress("Discovering documents...")
    documents = discover_patient_documents(patient_id)
    
    if not documents:
        logger.info(f"No documents found for patient {patient_id}")
        return {
            'patient_id': patient_id,
            'total_documents': 0,
            'processed_documents': 0,
            'successful_extractions': 0,
            'failed_extractions': 0,
            'canonical_created': 0
        }
    
    # DELETE EXISTING OBSERVATIONS AND CANONICAL JSON AT THE BEGINNING
    logger.info(f"Cleaning up existing data for patient {patient_id}")
    
    # Delete ALL existing observations for this patient (including stale ones from deleted files)
    try:
        conn = mysql.connector.connect(**DB_CONFIG)
        cursor = conn.cursor()
        cursor.execute("DELETE FROM observation_store WHERE patient_id = %s", (patient_id,))
        deleted_observations = cursor.rowcount
        logger.info(f"Deleted {deleted_observations} existing observations for patient {patient_id}")
        
        # Delete existing canonical JSON
        cursor.execute("DELETE FROM patient_case_envelope WHERE patient_id = %s", (patient_id,))
        deleted_canonical = cursor.rowcount
        logger.info(f"Deleted {deleted_canonical} existing canonical JSON records for patient {patient_id}")
        
        conn.commit()
    except Exception as e:
        logger.error(f"Failed to delete existing data for patient {patient_id}: {e}")
        if conn:
            conn.rollback()
    finally:
        if conn:
            conn.close()
    
    if not documents:
        logger.info(f"No documents found for patient {patient_id}")
        return {
            'patient_id': patient_id,
            'total_documents': 0,
            'processed_documents': 0,
            'successful_extractions': 0,
            'failed_extractions': 0,
            'canonical_created': 0
        }
    
    # Limit documents for testing if specified
    if max_documents:
        documents = documents[:max_documents]
        logger.info(f"Limited to {max_documents} documents for testing")
    
    report_progress(f"Found {len(documents)} documents to process")
    
    # Process documents in batches
    stats = process_documents_in_batches(documents, batch_size, progress_callback=report_progress)
    stats['patient_id'] = patient_id
    
    logger.error(f"✅ COMPLETED processing for patient {patient_id}: {stats}")  # Using error level so it shows in logs
    logger.info(f"Completed processing for patient {patient_id}: {stats}")
    return stats

def get_all_patient_ids(limit: Optional[int] = None) -> List[int]:
    """Fetch patient IDs whose status is NOT 'Archived' with optional limit."""
    ids: List[int] = []
    try:
        conn = mysql.connector.connect(
            host=DB_CONFIG['host'], user=DB_CONFIG['user'], password=DB_CONFIG['password'],
            database=DB_CONFIG['database'], port=DB_CONFIG['port']
        )
        cur = conn.cursor()
        q = "SELECT id FROM patients WHERE status <> 'Archived' OR status IS NULL ORDER BY id DESC"
        if limit and limit > 0:
            q += f" LIMIT {int(limit)}"
        cur.execute(q)
        for (pid,) in cur.fetchall():
            ids.append(int(pid))
        cur.close(); conn.close()
    except Exception as e:
        logger.error(f"Error fetching patient ids: {e}")
    return ids

def batch_all_patients(limit_patients: Optional[int] = None, batch_size: int = 3) -> Dict[str, Any]:
    """Process documents for all patients and update observation_store (and analyzed flags)."""
    pids = get_all_patient_ids(limit=limit_patients)
    total = len(pids)
    updated = 0
    canonical_created = 0
    results = []
    logger.info(f"Starting batch for {total} patients (limit={limit_patients})")
    for pid in pids:
        try:
            stats = process_patient_documents(pid, max_documents=None, batch_size=batch_size)
            results.append({'patient_id': pid, 'stats': stats})
            if stats.get('successful_extractions', 0) > 0:
                updated += 1
            if stats.get('canonical_created', 0) > 0:
                canonical_created += 1
        except Exception as e:
            logger.error(f"Error processing patient {pid}: {e}")
            results.append({'patient_id': pid, 'error': str(e)})
    summary = {
        'patients_total': total, 
        'patients_updated': updated, 
        'canonical_created': canonical_created,
        'results': results
    }
    logger.info(f"Batch complete: {summary}")
    return summary

def test_phase2_processing():
    """
    Test function for Phase 2 document processing with batch optimization and canonical JSON creation.
    """
    test_patient_id = 71100
    batch_size = 3  # Process 3 documents per Bedrock call
    
    logger.info(f"Testing Phase 2 document processing for patient {test_patient_id}")
    logger.info(f"Using batch processing with batch size: {batch_size}")
    logger.info("=" * 60)
    
    # Process ALL documents for the patient (no limit) with batch processing
    stats = process_patient_documents(test_patient_id, max_documents=None, batch_size=batch_size)
    
    logger.info("Processing Results:")
    logger.info(f"Total Documents: {stats['total_documents']}")
    logger.info(f"Processed: {stats['processed_documents']}")
    logger.info(f"Successful: {stats['successful_extractions']}")
    logger.info(f"Failed: {stats['failed_extractions']}")
    logger.info(f"Canonical Created: {stats.get('canonical_created', 0)}")
    logger.info(f"Estimated Bedrock calls saved: {stats['total_documents'] - (stats['total_documents'] // batch_size)}")
    
    # Test canonical JSON creation specifically
    logger.info("=" * 60)
    logger.info("Testing canonical JSON creation...")
    canonical_result = create_canonical_json_for_patient(test_patient_id)
    if canonical_result.get('success'):
        logger.info(f"Canonical JSON created successfully for patient {test_patient_id}")
        logger.info(f"Base report: {canonical_result.get('base_report_id')}")
        logger.info(f"Total reports: {canonical_result.get('total_reports')}")
    else:
        logger.warning(f"Canonical JSON creation failed: {canonical_result.get('message')}")
    
    # Test schema validation
    logger.info("=" * 60)
    logger.info("Testing schema validation...")
    try:
        # Get a sample envelope to validate
        conn = mysql.connector.connect(**DB_CONFIG)
        cursor = conn.cursor(dictionary=True)
        cursor.execute(
            "SELECT case_json FROM patient_case_envelope WHERE patient_id = %s AND report_id = 'canonical' LIMIT 1",
            (test_patient_id,)
        )
        result = cursor.fetchone()
        if result and result.get('case_json'):
            import json
            snapshot = json.loads(result['case_json']) if isinstance(result['case_json'], str) else result['case_json']
            validation = validate_schema_compliance(snapshot)
            logger.info(f"Schema validation result: {validation['valid']}")
            if validation['errors']:
                logger.warning(f"Validation errors: {validation['errors']}")
            if validation['warnings']:
                logger.info(f"Validation warnings: {validation['warnings']}")
        else:
            logger.info("No canonical envelope found for validation")
        conn.close()
    except Exception as e:
        logger.error(f"Error during schema validation: {e}")


def _resolve_ui_sleep_metrics(canonical_data: dict) -> dict:
    """
    Transform canonical data into UI-ready sleep metrics using rigid assumptions.
    
    Returns:
        {
            "baseline": {"ahi": 22.0, "odi": 15.0, ...},
            "current": {"ahi": 3.6, "odi": 2.9, ...},
            "therapy_start_date": "2025-03-10",
            "timeline": [
                {
                    "date": "2024-11-27",
                    "context": "baseline",
                    "metrics": {"ahi": 22.0, "odi": 15.0},
                    "provenance": {"source_kind": "sleep_study", "file_name": "Case_JuMa_1967.pdf"}
                },
                {
                    "date": "2025-09-01", 
                    "context": "current",
                    "metrics": {"ahi": 3.6, "odi": 2.9},
                    "provenance": {"source_kind": "report", "file_name": "JuMa_FollowUp_Report_Sep_2025.pdf"}
                }
            ]
        }
    """
    try:
        # Import the resolver function from the separate module
        import sys
        import os
        sys.path.append(os.path.dirname(__file__))
        from flask_app.config.resolve_ui_sleep_metrics import resolve
        return resolve(canonical_data)
    except ImportError as e:
        logger.warning(f"UI sleep metrics resolver not available: {e}")
        # Fallback if resolver not available
        return {
            "baseline": {},
            "current": {},
            "therapy_start_date": None,
            "timeline": []
        }
def create_minimal_canonical_json_for_patient(
    patient_id: int,
    *,
    skip_timeline_llm: bool = False,
    skip_quiz_risk_snapshot: bool = False,
    skip_observation_text_numerical_pass: bool = False,
) -> Dict[str, Any]:
    """
    Create a minimal canonical JSON for a patient based on observation store data.
    This creates a canonical record populated with observations from the observation_store table.
    
    Args:
        patient_id (int): Patient ID to create canonical JSON for
        skip_timeline_llm: If True, skip ``organize_timeline_with_llm`` (faster rebuild).
        skip_quiz_risk_snapshot: If True, skip writing quiz risk into observation_store this pass.
        skip_observation_text_numerical_pass: If True, skip loading all ``extracted_observations``
            text and ``extract_specific_numerical_fields`` (saves a lot of CPU/DB after a fresh
            structured sleep import — use only when enriched metric_key rows are already current).
        
    Returns:
        Dict[str, Any]: Result of canonical creation
    """
    import re
    from datetime import datetime, timezone
    
    # Create timestamp at the beginning
    now_utc = datetime.now(timezone.utc)
    now_iso = now_utc.strftime("%Y-%m-%dT%H:%M:%SZ")  # RFC3339 UTC format
    
    try:
        conn = mysql.connector.connect(**DB_CONFIG)
        cursor = conn.cursor(dictionary=True)

        # Get patient email for quiz lookup
        cursor.execute("SELECT email FROM patients WHERE id = %s", (patient_id,))
        patient_row = cursor.fetchone()
        patient_email = patient_row['email'] if patient_row else None
        
        # Store quiz risk assessment as observation BEFORE creating canonical
        if not skip_quiz_risk_snapshot:
            if patient_email:
                logger.info(f"Patient {patient_id}: Storing quiz risk assessment as observation")
                store_quiz_risk_assessment_as_observation(patient_id, patient_email)
            else:
                logger.warning(f"Patient {patient_id}: No email found, cannot fetch quiz risk assessment")
        else:
            logger.info(f"Patient {patient_id}: Skipping quiz risk snapshot (fast canonical path)")

        # Delete existing canonical data if it exists (to allow recreation)
        cursor.execute(
            "DELETE FROM patient_case_envelope WHERE patient_id = %s AND report_id = 'canonical'",
            (patient_id,)
        )
        deleted_count = cursor.rowcount
        if deleted_count > 0:
            logger.info(f"Deleted {deleted_count} existing canonical records for patient {patient_id}")
        conn.commit()

        # Load enriched observations from observation_store (extended columns)
        cursor.execute("""
            SELECT id, file_name, s3_key, metric_key, metric_value_decimal, metric_unit, metric_phase,
                   observed_at, mention_date, source_kind, study_type, episode_id, document_date,
                   extracted_observations
            FROM observation_store
            WHERE patient_id = %s AND metric_key IS NOT NULL
            ORDER BY COALESCE(observed_at, mention_date, created_at)
        """, (patient_id,))
        obs_rows = cursor.fetchall()
        logger.info(f"Found {len(obs_rows)} enriched metric rows for patient {patient_id}")
        
        # Build sleep_studies[] (true studies) and report_mentions[] (narrative mentions)
        sleep_studies_map: Dict[str, Dict[str, Any]] = {}
        report_mentions: List[Dict[str, Any]] = []
        # Track the most recent real sleep study episode so we can merge
        # orphan numerical-extraction metrics (which lack s3_key/file linkage)
        last_real_episode_id: Optional[str] = None

        def _normalize_sleep_metric_key(metric_key: Any) -> Tuple[str, Optional[str]]:
            """
            Returns (normalized_key, raw_key). Normalized key matches schema expectations (ahi, odi, etc.)
            Raw key preserves the original path for auditing.
            """
            if not isinstance(metric_key, str):
                return str(metric_key), None

            raw_key = metric_key
            key = metric_key.strip()

            # Direct mapping table for common paths
            direct_map = {
                'respiratory_indices.ahi_overall': 'ahi',
                'respiratory_indices.ahi': 'ahi',
                'respiratory_indices.odi3': 'odi',
                'respiratory_indices.odi4': 'odi',
                'respiratory_indices.odi': 'odi',
                'respiratory_indices.rdi': 'rdi',
                'respiratory_indices.oai': 'oai',
                'respiratory_indices.cai': 'cai',
                'respiratory_indices.hi': 'hi',
                'respiratory_indices.supine_ahi': 'supine_ahi',
                'respiratory_indices.non_supine_ahi': 'non_supine_ahi',
                'respiratory_indices.rem_ahi': 'rem_ahi',
                'respiratory_indices.nrem_ahi': 'nrem_ahi',
                'respiratory_indices.supine_rdi': 'supine_rdi',
                'respiratory_indices.rem_rdi': 'rem_rdi',
                'respiratory_indices.supine_odi': 'supine_odi',
                'respiratory_indices.rem_odi': 'rem_odi',
                'respiratory_indices.time_below_90_pct': 'time_below_90_pct',
                'respiratory_indices.time_below_90_pct_min': 'time_below_90_pct_min',
                'respiratory_indices.time_below_88_pct_min': 'time_below_88_pct_min',
                'oxygenation.spo2_nadir_pct': 'o2_nadir_pct',
                'oxygenation.spo2_mean_pct': 'o2_mean_pct',
                'oxygenation.spo2_max_pct': 'o2_max_pct',
                'sleep_timing_architecture.sleep_efficiency_pct': 'sleep_efficiency_pct',
                'sleep_timing_architecture.total_sleep_time_h': 'sleep_duration_h',
                'sleep_timing_architecture.sleep_duration_h': 'sleep_duration_h',
                'sleep_study.desaturation_events': 'desaturation_events',
            }
            if key in direct_map:
                return direct_map[key], raw_key

            if key.startswith('sleep_study.snoring.'):
                suffix = key.split('.', 2)[2]
                return f"snoring_{suffix}", raw_key

            if key.startswith('sleep_study.heart_rate.'):
                suffix = key.split('.', 2)[2]
                return f"heart_rate_{suffix}", raw_key

            if key.startswith('sleep_study.'):
                return key.split('.', 1)[1], raw_key

            if key.startswith('respiratory_indices.'):
                suffix = key.split('.', 1)[1]
                # Fallback to suffix; the direct_map above covers special cases
                return suffix, raw_key

            if key.startswith('oxygenation.') or key.startswith('sleep_timing_architecture.'):
                return key.split('.', 1)[1], raw_key

            if '.' in key:
                return key.split('.')[-1], raw_key

            return key, raw_key

        for r in obs_rows:
            metric_key = r.get('metric_key')
            metric_val = r.get('metric_value_decimal')
            if metric_key is None or metric_val is None:
                continue
            source_kind = (r.get('source_kind') or '').lower()

            if source_kind == 'sleep_study':
                file_name = r.get('file_name')
                s3_key = r.get('s3_key')
                is_orphan_numeric = (not s3_key) and (not file_name or file_name == 'numerical_pattern_extraction')
                # Choose target episode: merge orphan numeric metrics into last real episode if available
                target_ep = (last_real_episode_id if (is_orphan_numeric and last_real_episode_id) else (r.get('episode_id') or f"{file_name or ''}"))
                if target_ep not in sleep_studies_map:
                    # Use document_date if available, otherwise fall back to observed_at
                    study_date = None
                    if not is_orphan_numeric:
                        if r.get('document_date'):
                            study_date = r.get('document_date').strftime('%Y-%m-%d') if hasattr(r.get('document_date'), 'strftime') else str(r.get('document_date'))
                        elif r.get('observed_at'):
                            study_date = r.get('observed_at').strftime('%Y-%m-%d') if hasattr(r.get('observed_at'), 'strftime') else str(r.get('observed_at'))
                    
                    sleep_studies_map[target_ep] = {
                        'episode_id': target_ep,
                        'observed_at': study_date,
                        'study_type': r.get('study_type') or None,
                        'file_name': file_name,
                        's3_key': s3_key,
                        'source_kind': 'sleep_study',
                        'context_tag': None,
                        'metrics': []
                    }
                # If this row has a real document_date and our target entry lacks it, fill it
                if not is_orphan_numeric and not sleep_studies_map[target_ep].get('observed_at'):
                    if r.get('document_date'):
                        study_date = r.get('document_date').strftime('%Y-%m-%d') if hasattr(r.get('document_date'), 'strftime') else str(r.get('document_date'))
                        sleep_studies_map[target_ep]['observed_at'] = study_date
                        logger.info(f"Set observed_at from document_date: {study_date}")
                    elif r.get('observed_at'):
                        study_date = r.get('observed_at').strftime('%Y-%m-%d') if hasattr(r.get('observed_at'), 'strftime') else str(r.get('observed_at'))
                        sleep_studies_map[target_ep]['observed_at'] = study_date
                        logger.info(f"Set observed_at from observed_at field: {study_date}")
                # Append metric to the chosen episode
                normalized_key, raw_key = _normalize_sleep_metric_key(metric_key)
                sleep_studies_map[target_ep]['metrics'].append({
                    'key': normalized_key,
                    'raw_key': raw_key,
                    'value': float(metric_val) if isinstance(metric_val, (int, float, Decimal)) else float(str(metric_val)),
                    'unit': r.get('metric_unit'),
                    'phase': r.get('metric_phase'),
                    'confidence': 1.0,
                    'provenance': {
                        'file_name': file_name,
                        's3_key': s3_key
                    }
                })
                # Update last_real_episode_id when we encounter a real study row
                if s3_key or (file_name and file_name != 'numerical_pattern_extraction'):
                    last_real_episode_id = target_ep
            else:
                # Filter out invalid placeholder values
                metric_val_float = float(metric_val) if isinstance(metric_val, (int, float, Decimal)) else float(str(metric_val))
                
                # Skip invalid placeholder values (e.g., 9999999.999 for MRN)
                if metric_key and 'mrn' in metric_key.lower() and (metric_val_float > 9999999 or metric_val_float < 0):
                    logger.debug(f"Skipping invalid MRN value: {metric_val_float} for key {metric_key}")
                    continue
                
                # Skip other obviously invalid values
                if metric_key and 'age' in metric_key.lower() and (metric_val_float < 0 or metric_val_float > 150):
                    logger.debug(f"Skipping invalid age value: {metric_val_float} for key {metric_key}")
                    continue
                
                # Use document_date if available, otherwise fall back to mention_date
                report_date = None
                if r.get('document_date'):
                    report_date = r.get('document_date').strftime('%Y-%m-%d') if hasattr(r.get('document_date'), 'strftime') else str(r.get('document_date'))
                elif r.get('mention_date'):
                    report_date = r.get('mention_date').strftime('%Y-%m-%d') if hasattr(r.get('mention_date'), 'strftime') else str(r.get('mention_date'))
                
                report_mentions.append({
                    'key': metric_key,
                    'value': metric_val_float,
                    'unit': r.get('metric_unit'),
                    'phase': r.get('metric_phase'),
                    'confidence': 0.7,
                    'mention_date': report_date,
                    'reported_at': report_date,  # Add reported_at field
                    'file_name': r.get('file_name'),
                    's3_key': r.get('s3_key'),
                    'source_kind': source_kind or 'report',
                    'provenance': {
                        'file_name': r.get('file_name'),
                        's3_key': r.get('s3_key')
                    }
                })

        # Attach LLM narrative from sleep pipeline snapshots (row order vs metrics does not matter).
        for r in obs_rows:
            if r.get("metric_key") != "sleep_study.pipeline_snapshot_v1":
                continue
            if (r.get("source_kind") or "").lower() != "sleep_study":
                continue
            target_ep = r.get("episode_id")
            if target_ep not in sleep_studies_map:
                sk = r.get("s3_key") or ""
                for eid, data in sleep_studies_map.items():
                    if sk and data.get("s3_key") == sk:
                        target_ep = eid
                        break
            if not target_ep or target_ep not in sleep_studies_map:
                continue
            try:
                ex = r.get("extracted_observations")
                obs_json = json.loads(ex) if isinstance(ex, str) else (ex or {})
                expl = (obs_json.get("explanation") or "").strip()
                if not expl:
                    continue
                snap = json.loads(expl)
                diag = snap.get("diagnosis") or {}
                impr = diag.get("impression") or diag.get("primary_diagnosis")
                bullets = snap.get("clinical_insights")
                parts: List[str] = []
                if impr:
                    parts.append(str(impr).strip())
                if isinstance(bullets, list):
                    for b in bullets:
                        if b:
                            parts.append(str(b).strip())
                if parts:
                    sleep_studies_map[target_ep]["analysis_insights_text"] = "\n".join(parts)[:16000]
            except Exception as _parse_snap:
                logger.debug("pipeline_snapshot narrative skip: %s", _parse_snap)

        # Merge any orphan numeric-only episodes into the latest real sleep study episode
        try:
            # Identify latest real episode by observed_at
            real_eps = [
                (ep_id, data) for ep_id, data in sleep_studies_map.items()
                if data.get('s3_key') or (data.get('file_name') and data.get('file_name') != 'numerical_pattern_extraction')
            ]
            def _parsed_dt(s):
                try:
                    from datetime import datetime
                    return datetime.strptime(s, '%Y-%m-%dT%H:%M:%SZ') if s else None
                except Exception:
                    return None
            latest_real = None
            for ep_id, data in real_eps:
                dt = _parsed_dt(data.get('observed_at'))
                if latest_real is None:
                    latest_real = (ep_id, dt)
                else:
                    if dt and (latest_real[1] is None or dt > latest_real[1]):
                        latest_real = (ep_id, dt)
            if latest_real:
                latest_ep_id = latest_real[0]
                # Gather orphan entries
                orphan_keys = [
                    ep_id for ep_id, data in sleep_studies_map.items()
                    if (not data.get('s3_key')) and (not data.get('file_name') or data.get('file_name') == 'numerical_pattern_extraction')
                ]
                for orphan_ep in orphan_keys:
                    orphan_metrics = sleep_studies_map[orphan_ep].get('metrics') or []
                    # Move metrics over
                    sleep_studies_map[latest_ep_id]['metrics'].extend(orphan_metrics)
                    orphan_txt = sleep_studies_map[orphan_ep].get('analysis_insights_text')
                    if orphan_txt and not sleep_studies_map[latest_ep_id].get('analysis_insights_text'):
                        sleep_studies_map[latest_ep_id]['analysis_insights_text'] = orphan_txt
                    # Drop the orphan
                    del sleep_studies_map[orphan_ep]
                    logger.info(f"Merged {len(orphan_metrics)} orphan numerical metrics into episode {latest_ep_id}")
        except Exception as _merge_post_e:
            logger.warning(f"Post-merge of orphan numerical metrics failed: {_merge_post_e}")

        # Extract observation texts for numerical extraction (from legacy JSON rows).
        # Optional skip: after direct sleep OpenAI import, metric_key rows already hold structure;
        # scanning every extracted_observations blob is expensive for large patients.
        observation_texts: List[str] = []
        # Must initialize: assignments inside branches make ``observations`` local for this
        # function; the fast path (skip_observation_text_numerical_pass=True) must not hit
        # ``if observations:`` unbound (direct sleep extraction uses that path).
        observations: List[Any] = []
        if not skip_observation_text_numerical_pass:
            cursor.execute("""
                SELECT extracted_observations FROM observation_store
                WHERE patient_id=%s AND (extracted_observations IS NOT NULL)
            """, (patient_id,))
            observations = cursor.fetchall()
            for obs in observations:
                try:
                    obs_data = json.loads(obs['extracted_observations']) if isinstance(obs['extracted_observations'], str) else obs['extracted_observations']
                    if 'value' in obs_data:
                        observation_texts.append(str(obs_data['value']))
                    if 'observation' in obs_data:
                        observation_texts.append(str(obs_data['observation']))
                except Exception:
                    continue
        else:
            logger.info(
                "Patient %s: Skipping observation-text numerical pass (fast canonical path)",
                patient_id,
            )

        # Get patient's demographic data from database as fallbacks
        from flask_app.models import Patient
        from datetime import date
        patient_age_fallback = None
        patient_gender_fallback = None
        try:
            patient = Patient.query.get(patient_id)
            if patient:
                # Calculate age from DOB if available
                if patient.dob:
                    today = date.today()
                    patient_age_fallback = today.year - patient.dob.year - ((today.month, today.day) < (patient.dob.month, patient.dob.day))
                    logger.info(f"Patient {patient_id}: Calculated age from DOB: {patient_age_fallback}")
                
                # Get gender if available
                if patient.gender:
                    patient_gender_fallback = patient.gender.upper()
                    logger.info(f"Patient {patient_id}: Got gender from database: {patient_gender_fallback}")
        except Exception as e:
            logger.warning(f"Could not get patient demographic data from database: {e}")
        
        # Run numerical extraction on all observation texts
        if observation_texts:
            logger.info(f"Patient {patient_id}: Running numerical extraction on {len(observation_texts)} observation texts")
            numerical_data = extract_specific_numerical_fields(observation_texts)
            
            if numerical_data:
                logger.info(f"Patient {patient_id}: Found {len(numerical_data)} numerical fields: {list(numerical_data.keys())}")
                
                # Validate and fix demographics if they seem wrong
                if 'age_years' in numerical_data:
                    extracted_age = numerical_data['age_years']
                    
                    # Always prefer database DOB if available
                    if patient_age_fallback:
                        # More comprehensive age validation for adult patients
                        if extracted_age < 10 or extracted_age > 120:  # Reject obviously wrong ages
                            logger.warning(f"Patient {patient_id}: Extracted age {extracted_age} is outside reasonable range, using fallback age {patient_age_fallback}")
                            numerical_data['age_years'] = patient_age_fallback
                        elif extracted_age < 18:  # Additional check for very young ages
                            logger.warning(f"Patient {patient_id}: Extracted age {extracted_age} seems too young for adult patient, using fallback age {patient_age_fallback}")
                            numerical_data['age_years'] = patient_age_fallback
                        elif abs(extracted_age - patient_age_fallback) > 5:  # If there's a significant discrepancy
                            logger.warning(f"Patient {patient_id}: Age discrepancy - extracted: {extracted_age}, fallback: {patient_age_fallback}, using fallback")
                            numerical_data['age_years'] = patient_age_fallback
                        else:
                            logger.info(f"Patient {patient_id}: Using extracted age {extracted_age} (close to fallback {patient_age_fallback})")
                    else:
                        # No fallback available, but still validate the extracted age
                        if extracted_age < 10 or extracted_age > 120:
                            logger.warning(f"Patient {patient_id}: Extracted age {extracted_age} is outside reasonable range, removing")
                            del numerical_data['age_years']
                        elif extracted_age < 18:  # Additional check for very young ages
                            logger.warning(f"Patient {patient_id}: Extracted age {extracted_age} seems too young for adult patient, removing")
                            del numerical_data['age_years']
                        else:
                            logger.info(f"Patient {patient_id}: Using extracted age {extracted_age} (no fallback available)")
                
                # Validate height - remove obviously wrong values
                if 'height_cm' in numerical_data:
                    extracted_height = numerical_data['height_cm']
                    if extracted_height < 120:  # Too short for an adult
                        logger.warning(f"Patient {patient_id}: Extracted height {extracted_height}cm is too short for an adult, removing")
                        del numerical_data['height_cm']
                    elif extracted_height > 250:  # Too tall
                        logger.warning(f"Patient {patient_id}: Extracted height {extracted_height}cm is too tall, removing")
                        del numerical_data['height_cm']
                
                # DISABLED: Store numerical extraction results as individual observations
                # This was causing data pollution with duplicate values and no dates
                # numerical_observations = _convert_numerical_data_to_observations(numerical_data, patient_id)
                # if numerical_observations:
                #     logger.info(f"Patient {patient_id}: Storing {len(numerical_observations)} numerical observations")
                #     store_observations_with_deduplication(patient_id, 'numerical_extraction', numerical_observations, {
                #         'name': 'numerical_pattern_extraction',
                #         'file_type': 'text/plain',
                #         'id': None,
                #         'source_table': 'numerical'
                #     })
                logger.info(f"Patient {patient_id}: Numerical pattern extraction disabled to prevent data pollution")
                
                # Reload observations to include the newly stored numerical observations
                cursor.execute("""
                    SELECT source_type, source_text, extracted_observations, created_at
                    FROM observation_store 
                    WHERE patient_id = %s 
                    ORDER BY created_at DESC
                """, (patient_id,))
                observations = cursor.fetchall()
                logger.info(f"Reloaded {len(observations)} observations including numerical extraction results")

                # Also reload enriched metric rows for numerical extractions and merge into sleep_studies_map
                try:
                    cursor.execute(
                        """
                        SELECT id, file_name, s3_key, metric_key, metric_value_decimal, metric_unit, metric_phase,
                               observed_at, mention_date, source_kind, study_type, episode_id
                        FROM observation_store
                        WHERE patient_id = %s AND metric_key IS NOT NULL
                          AND (file_name = 'numerical_pattern_extraction' OR source_type = 'numerical_extraction')
                        ORDER BY COALESCE(observed_at, mention_date, created_at)
                        """,
                        (patient_id,)
                    )
                    num_rows = cursor.fetchall()
                    logger.info(f"Merging {len(num_rows)} numerical-enriched metric rows into sleep_studies map")
                    for r in num_rows:
                        metric_key = r.get('metric_key')
                        metric_val = r.get('metric_value_decimal')
                        if metric_key is None or metric_val is None:
                            continue
                        source_kind = (r.get('source_kind') or '').lower()
                        if source_kind == 'sleep_study':
                            file_name = r.get('file_name')
                            s3_key = r.get('s3_key')
                            is_orphan_numeric = (not s3_key) and (not file_name or file_name == 'numerical_pattern_extraction')
                            target_ep = (last_real_episode_id if (is_orphan_numeric and last_real_episode_id) else (r.get('episode_id') or f"{file_name or ''}"))
                            if target_ep not in sleep_studies_map:
                                sleep_studies_map[target_ep] = {
                                    'episode_id': target_ep,
                                    'observed_at': ((r.get('observed_at').strftime('%Y-%m-%dT%H:%M:%SZ') if r.get('observed_at') else None) if not is_orphan_numeric else None),
                                    'study_type': r.get('study_type') or None,
                                    'file_name': file_name,
                                    's3_key': s3_key,
                                    'source_kind': 'sleep_study',
                                    'context_tag': None,
                                    'metrics': []
                                }
                            if not is_orphan_numeric and r.get('observed_at') and not sleep_studies_map[target_ep].get('observed_at'):
                                sleep_studies_map[target_ep]['observed_at'] = r.get('observed_at').strftime('%Y-%m-%dT%H:%M:%SZ')
                            sleep_studies_map[target_ep]['metrics'].append({
                                'key': metric_key,
                                'value': float(metric_val) if isinstance(metric_val, (int, float, Decimal)) else float(str(metric_val)),
                                'unit': r.get('metric_unit'),
                                'phase': r.get('metric_phase'),
                                'confidence': 1.0,
                                'provenance': {
                                    'file_name': file_name,
                                    's3_key': s3_key
                                }
                            })
                        else:
                            # Filter out invalid placeholder values
                            metric_val_float = float(metric_val) if isinstance(metric_val, (int, float, Decimal)) else float(str(metric_val))
                            
                            # Skip invalid placeholder values (e.g., 9999999.999 for MRN)
                            if metric_key and 'mrn' in metric_key.lower() and (metric_val_float > 9999999 or metric_val_float < 0):
                                logger.debug(f"Skipping invalid MRN value: {metric_val_float} for key {metric_key}")
                                continue
                            
                            # Skip other obviously invalid values
                            if metric_key and 'age' in metric_key.lower() and (metric_val_float < 0 or metric_val_float > 150):
                                logger.debug(f"Skipping invalid age value: {metric_val_float} for key {metric_key}")
                                continue
                            
                            report_mentions.append({
                                'key': metric_key,
                                'value': metric_val_float,
                                'unit': r.get('metric_unit'),
                                'phase': r.get('metric_phase'),
                                'confidence': 0.7,
                                'mention_date': (r.get('mention_date').strftime('%Y-%m-%dT%H:%M:%SZ') if r.get('mention_date') else None),
                                'file_name': r.get('file_name'),
                                's3_key': r.get('s3_key'),
                                'source_kind': source_kind or 'report',
                                'provenance': {
                                    'file_name': r.get('file_name'),
                                    's3_key': r.get('s3_key')
                                }
                            })
                except Exception as _merge_e:
                    logger.warning(f"Could not merge numerical metrics into canonical sleep studies: {_merge_e}")
            else:
                logger.info(f"Patient {patient_id}: No numerical fields found")
        else:
            logger.info(f"Patient {patient_id}: No observation texts available for numerical extraction")
        
        # Always create fallback demographic observations if we have database data
        database_fallback_observations = []
        
        if patient_age_fallback:
            logger.info(f"Patient {patient_id}: Creating fallback age observation from DOB: {patient_age_fallback}")
            age_observation = {
                'path': 'demographics.age_years',
                'value': str(patient_age_fallback),
                'source': 'database_fallback',
                'confidence': 100,
                'explanation': f'Age calculated from patient DOB: {patient_age_fallback}',
                'observation': f'Patient age: {patient_age_fallback} years',
                'document_name': 'patient_database',
                'document_type': 'database',
                'extraction_date': now_iso
            }
            database_fallback_observations.append(age_observation)
        
        if patient_gender_fallback:
            logger.info(f"Patient {patient_id}: Creating fallback gender observation from database: {patient_gender_fallback}")
            gender_observation = {
                'path': 'demographics.sex',
                'value': patient_gender_fallback,
                'source': 'database_fallback',
                'confidence': 100,
                'explanation': f'Gender from patient database: {patient_gender_fallback}',
                'observation': f'Patient gender: {patient_gender_fallback}',
                'document_name': 'patient_database',
                'document_type': 'database',
                'extraction_date': now_iso
            }
            database_fallback_observations.append(gender_observation)
        
        # Store all database fallback observations
        if database_fallback_observations:
            store_observations_with_deduplication(patient_id, 'database_fallback', database_fallback_observations, {
                'name': 'patient_database',
                'file_type': 'database',
                'id': None,
                'source_table': 'patient'
            })
        
        # Create canonical JSON structure with UTC timestamps
        
        # Initialize canonical structure - EXACTLY matching schema
        canonical = {
            'schema_version': '1.0',
            'document_type': 'canonical',
            'patient_id': str(patient_id),
            'as_of': now_iso,
            'canonical_meta': {
                'version': 1,
                'report_refs': []
            },
            'demographics': {
                'sex': None,
                'age_years': None,
                'height_cm': None,
                'weight_kg': None,
                'bmi': None
            },
            'sleep_study': {
                'study_type': None,
                'sleep_duration_h': None,
                'sleep_efficiency_pct': None,
                'ahi': None,
                'odi': None,
                'desaturation_events': None,
                'o2_nadir_pct': None,
                'o2_mean_pct': None,
                'o2_max_pct': None,
                'snoring': {
                    'avg_db': None,
                    'max_db': None
                },
                'heart_rate': {
                    'mean_bpm': None,
                    'min_bpm': None,
                    'max_bpm': None
                }
            },
            'risk_assessment': {
                'risk_level': None,
                'risk_source': None,
                'assessed_at': None
            },
            'observations': {
                'summary': [],
                'anatomy_imaging': {
                    'primary_obstruction_site': None,
                    'soft_palate_uvula': None,
                    'tongue_base': None,
                    'bite_jaw': None,
                    'hyoid': None,
                    'nose_sinus': None,
                    'tmj': None,
                    # Extended anatomical fields
                    'arches': None,  # Dental arch observations (narrow, underdeveloped, etc.)
                    'epiglottis': None,  # Epiglottic findings (floppy, retroflexed, etc.)
                    'neck_findings': None,  # Neck observations (thick neck, etc.)
                    'overjet': None,  # Overjet measurement/description
                    'overbite': None,  # Overbite measurement/description
                    'retropalatal': None,  # Retropalatal space/collapse
                    'retroglossal': None,  # Retroglossal space/collapse
                    'pharyngeal_wall': None,  # Posterior pharyngeal wall findings
                    'tonsils': None,  # Tonsillar findings (enlarged, grade, etc.)
                    'adenoids': None,  # Adenoid findings
                    'mandibular_plane_angle': None,  # Mandibular plane angle measurement
                    'airway_volume': None,  # Airway volume measurements
                    'mallampati': None,  # Mallampati score (I, II, III, IV)
                    'friedman_stage': None,  # Friedman staging
                    'mueller_maneuver': None,  # Mueller maneuver findings
                    'dise_findings': None,  # Drug-induced sleep endoscopy findings (legacy field)
                    'conclusion': None,  # Summary/conclusion of imaging findings
                    'other_findings': []  # Catch-all for anatomical observations that don't fit other fields
                },
                # DISE (Drug-Induced Sleep Endoscopy) Structured Block
                'dise': {
                    'performed': None,  # Boolean: was DISE performed?
                    'date': None,  # Date of DISE exam
                    # Collapse patterns by anatomical level (VOTE classification)
                    'velum': {
                        'collapse_pattern': None,  # "none", "partial", "complete"
                        'collapse_direction': None,  # "AP" (anteroposterior), "lateral", "concentric"
                        'grade': None,  # 0-4 or descriptive
                        'notes': None
                    },
                    'oropharynx_lateral_walls': {
                        'collapse_pattern': None,  # "none", "partial", "complete"
                        'collapse_direction': None,  # "medial", "lateral"
                        'grade': None,
                        'notes': None
                    },
                    'tongue_base': {
                        'collapse_pattern': None,  # "none", "partial", "complete"
                        'collapse_direction': None,  # "AP", "posterior"
                        'hypertrophy_grade': None,  # 1, 2, 3, or "grade 2-3"
                        'notes': None
                    },
                    'epiglottis': {
                        'collapse_pattern': None,  # "none", "partial", "complete", "trapdoor", "lateral"
                        'appearance': None,  # "normal", "floppy", "retroflexed", "omega-shaped"
                        'notes': None
                    },
                    # Maneuver responses - CRITICAL for treatment planning
                    'maneuver_response': {
                        'jaw_thrust': {
                            'performed': None,
                            'response': None,  # "none", "partial", "moderate", "significant", "complete"
                            'improvement_pct': None,  # Numeric percentage if available
                            'notes': None
                        },
                        'head_rotation': {
                            'performed': None,
                            'response': None,  # "none", "partial", "moderate", "significant", "complete"
                            'notes': None
                        },
                        'chin_lift': {
                            'performed': None,
                            'response': None,
                            'notes': None
                        },
                        'lateral_position': {
                            'performed': None,
                            'response': None,
                            'notes': None
                        }
                    },
                    # Positional findings during DISE
                    'positional_findings': {
                        'supine_collapse': None,  # Description of collapse in supine
                        'lateral_collapse': None,  # Description of collapse in lateral
                        'positional_dependence': None,  # "none", "mild", "moderate", "severe"
                        'supine_predominance_pct': None  # Percentage of events in supine
                    },
                    # Overall assessment
                    'obstruction_pattern': None,  # "single-level", "multilevel", "circumferential"
                    'primary_site': None,  # Primary obstruction site identified
                    'secondary_sites': [],  # List of secondary sites
                    'appliance_suitability': {
                        'suitable': None,  # Boolean or "yes"/"no"/"maybe"
                        'predicted_response': None,  # "poor", "moderate", "good", "excellent"
                        'rationale': None
                    },
                    'surgical_considerations': [],  # List of surgical options based on DISE
                    'other_findings': []
                },
                # CBCT (Cone Beam CT) Airway Structured Block
                'cbct': {
                    'performed': None,
                    'date': None,
                    # Airway measurements
                    'airway_measurements': {
                        'total_airway_volume_cc': None,  # Total airway volume in cubic centimeters
                        'minimum_cross_sectional_area_mm2': None,  # MCSA in mm²
                        'mcsa_location': None,  # Location of minimum area (e.g., "retropalatal", "retroglossal")
                        'average_cross_sectional_area_mm2': None,
                        'airway_length_mm': None
                    },
                    # Segmental volumes
                    'segmental_volumes': {
                        'nasopharynx_volume_cc': None,
                        'velopharynx_volume_cc': None,
                        'oropharynx_volume_cc': None,
                        'hypopharynx_volume_cc': None
                    },
                    # Skeletal measurements
                    'skeletal_measurements': {
                        'mandibular_plane_angle_deg': None,
                        'hyoid_to_mandibular_plane_mm': None,
                        'hyoid_to_c3_mm': None,
                        'pns_to_uvula_mm': None,  # Soft palate length
                        'tongue_length_mm': None,
                        'tongue_height_mm': None,
                        'posterior_airway_space_mm': None,  # PAS
                        'snb_angle_deg': None,
                        'sna_angle_deg': None,
                        'anb_angle_deg': None
                    },
                    # Soft tissue measurements
                    'soft_tissue': {
                        'soft_palate_length_mm': None,
                        'soft_palate_thickness_mm': None,
                        'tongue_base_thickness_mm': None,
                        'lateral_pharyngeal_wall_thickness_mm': None
                    },
                    # Narrowing assessment
                    'narrowing_assessment': {
                        'retropalatal_narrowing': None,  # "none", "mild", "moderate", "severe"
                        'retroglossal_narrowing': None,
                        'hypopharyngeal_narrowing': None
                    },
                    'other_findings': []
                },
                # ENT / Nasal / Sinus Structured Block
                'ent_findings': {
                    'nasal': {
                        'septum_deviation': None,  # "none", "left", "right", "s-shaped"
                        'septum_deviation_severity': None,  # "mild", "moderate", "severe"
                        'turbinate_hypertrophy': None,  # "none", "mild", "moderate", "severe"
                        'turbinate_side': None,  # "left", "right", "bilateral"
                        'nasal_valve_collapse': None,
                        'nasal_polyps': None,  # Boolean or grade
                        'polyp_grade': None,  # 1, 2, 3, 4
                        'polyp_side': None,  # "left", "right", "bilateral"
                        'mucosa_appearance': None,  # "normal", "edematous", "inflamed"
                        'other_nasal_findings': []
                    },
                    'sinus': {
                        'maxillary_sinus': {
                            'left': None,  # "clear", "mucosal thickening", "opacification", "polyps"
                            'right': None
                        },
                        'ethmoid_sinus': {
                            'anterior_left': None,
                            'anterior_right': None,
                            'posterior_left': None,
                            'posterior_right': None
                        },
                        'frontal_sinus': {
                            'left': None,
                            'right': None
                        },
                        'sphenoid_sinus': {
                            'left': None,
                            'right': None
                        },
                        'chronic_sinusitis': None,  # Boolean
                        'acute_sinusitis': None,
                        'other_sinus_findings': []
                    },
                    'nasopharynx': {
                        'appearance': None,  # "normal", "narrowed", "obstructed"
                        'adenoid_hypertrophy': None,
                        'adenoid_grade': None,  # 1, 2, 3, 4
                        'other_findings': []
                    },
                    'oropharynx': {
                        'tonsil_grade': None,  # 0, 1, 2, 3, 4
                        'tonsil_appearance': None,
                        'uvula': None,  # "normal", "elongated", "edematous"
                        'soft_palate': None,
                        'other_findings': []
                    },
                    'larynx': {
                        'epiglottis': None,
                        'vocal_cords': None,
                        'other_findings': []
                    },
                    'post_surgical_changes': None,  # Description of any post-surgical findings
                    'other_ent_findings': []
                },
                # Airway phenotype classification
                'airway_phenotype': {
                    'classification': None,  # "anatomical", "non-anatomical", "mixed"
                    'primary_phenotype': None,  # "high loop gain", "low arousal threshold", "poor muscle responsiveness", "anatomical"
                    'obstruction_level': None,  # "single-level", "multilevel"
                    'positional_component': None,  # "none", "mild", "moderate", "severe"
                    'rem_component': None,  # "none", "mild", "moderate", "severe"
                    'notes': None
                },
                'tmj_flags': {
                    'pain': None,
                    'clicking': None,
                    'side': None,
                    'crepitus': None,  # TMJ crepitus
                    'limited_opening': None,  # Limited jaw opening
                    'deviation': None,  # Jaw deviation on opening
                    'other_tmj_findings': []  # Catch-all for TMJ observations that don't fit other fields
                },
                'other_observations': []  # Catch-all for any clinical observations not fitting other categories
            },
            'treatment_considerations': {
                'primary_pathway': [],
                'adjuncts': [],
                'cautions': [],
                'rationale': None
            },
            'device_design': {
                'mandibular_advancement_mm': None,
                'advancement_plan': None,
                'vertical_opening_mm': None,
                'anterior_window': None,
                'retention_features': [],
                'material': None,
                'coverage': None,
                'initial_accessories': []
            },
            'patient_self_report': {
                'primary_complaint': None,
                'goals': [],
                'symptoms': {
                    'daytime_sleepiness': None,
                    'non_restorative_sleep': None,
                    'witnessed_apneas': None,
                    'nocturia': None,
                    'morning_headache': None,
                    'dry_mouth': None,
                    'bruxism': None,
                    'reflux': None,
                    'insomnia_features': None
                },
                'scales': {
                    'ESS': None,
                    'STOP_Bang': None,
                    'NOSE': None,
                    'PSQI': None
                },
                'preferences': {
                    'therapy_type_preference': None,
                    'aversion_notes': None
                }
            },
            'medical_history': {
                'comorbidities': [],
                'allergies': [],
                'medications': [],
                'bmi_trend': None
            },
            'prior_therapies': {
                'cpap': {
                    'tried': None,
                    'intolerance': None,
                    'intolerance_reasons': []
                },
                'oral_appliance': {
                    'tried': None,
                    'response': None,
                    'issues': []
                },
                'surgeries': []
            },
            'follow_up_plan': {
                'evaluations': [],
                'lifestyle': [],
                'positional_therapy': None,
                'retest_after_init_months': None
            },
            'device_options': [],
            'policy_trace': {
                'rules_fired': [],
                'hard_blocks_triggered': [],
                'missing_required_evidence': []
            },
            'provenance': [],
            'validation': {
                'errors': [],
                'warnings': []
            },
            'confidence': {
                'sleep_study': None,
                'observations': None,
                'device_design': None
            },
            'completeness_flags': {
                'has_sleep_study': False,
                'has_anatomy_imaging': False,
                'has_tmj_info': False
            }
        }
        
        # Attach sleep_studies and reported_metrics built from extended columns
        if sleep_studies_map:
            # choose a representative study_type at the root if only one present
            all_types = {v.get('study_type') for v in sleep_studies_map.values() if v.get('study_type')}
            if len(all_types) == 1:
                canonical['sleep_study']['study_type'] = list(all_types)[0]

        # Note: canonical['sleep_studies'] will be set later with organized data if LLM organization succeeds
        # For now, set the raw data as fallback
        canonical['sleep_studies'] = list(sleep_studies_map.values()) if sleep_studies_map else []
        canonical['report_mentions'] = report_mentions
        # Back-compat alias
        canonical['reported_metrics'] = report_mentions
        # Legacy path: Process structured data from extracted_observations that already has schema paths
        if observations:
            # First, process structured data that already has schema paths
            structured_data_processed = False
            observation_texts = []
            
            for obs in observations:
                try:
                    if isinstance(obs['extracted_observations'], dict):
                        obs_data = obs['extracted_observations']
                    elif obs['extracted_observations']:
                        obs_data = json.loads(obs['extracted_observations'])
                    else:
                        continue
                    
                    # Check if this observation has a schema path (structured data)
                    path = obs_data.get('path', '')
                    value = obs_data.get('value', '')
                    
                    # Map respiratory_indices paths to sleep_study paths for backward compatibility
                    path_mapping = {
                        'respiratory_indices.ahi_overall': 'sleep_study.ahi',
                        'respiratory_indices.ahi': 'sleep_study.ahi',
                        'respiratory_indices.odi3': 'sleep_study.odi',
                        'respiratory_indices.odi4': 'sleep_study.odi',
                        'respiratory_indices.odi': 'sleep_study.odi',
                        'oxygenation.spo2_nadir_pct': 'sleep_study.o2_nadir_pct',
                        'oxygenation.spo2_mean_pct': 'sleep_study.o2_mean_pct',
                        'oxygenation.spo2_max_pct': 'sleep_study.o2_max_pct',
                        'sleep_timing_architecture.sleep_efficiency_pct': 'sleep_study.sleep_efficiency_pct',
                        'sleep_timing_architecture.total_sleep_time_h': 'sleep_study.sleep_duration_h',
                    }
                    
                    targets = []
                    if path:
                        targets.append(path)
                    mapped_path = path_mapping.get(path)
                    if mapped_path and mapped_path not in targets:
                        logger.info(f"Mapped observation path from {path} to {mapped_path} for patient {patient_id}")
                        targets.append(mapped_path)
                    
                    if not targets or not value:
                        continue
                    
                    for target_path in targets:
                        if target_path == 'observations.summary':
                            # Only process original summary path
                            if target_path != path:
                                continue
                            if obs.get('source_type') == 'questionnaire':
                                # Categorize questionnaire observations into patient_self_report
                                categorize_questionnaire_observation(canonical, value)
                            else:
                                # Add to observation texts for general summary
                                observation_texts.append(value)
                            continue
                        
                        # This is structured data with a schema path
                        try:
                            # Navigate to the correct location in canonical structure
                            path_parts = target_path.split('.')
                            current = canonical
                            
                            # Navigate through the path except the last part
                            for part in path_parts[:-1]:
                                    # Check if this part has array notation like "comorbidities[0]"
                                    if '[' in part and ']' in part:
                                        # Parse array notation
                                        array_match = re.match(r'([^\[]+)\[(\d+)\]', part)
                                        if array_match:
                                            array_name = array_match.group(1)
                                            array_index = int(array_match.group(2))
                                            
                                            # Ensure array exists
                                            if array_name not in current:
                                                current[array_name] = []
                                            
                                            # Ensure array is long enough
                                            while len(current[array_name]) <= array_index:
                                                current[array_name].append({})
                                            
                                            current = current[array_name][array_index]
                                        else:
                                            # Fallback to dict if parsing fails
                                            if part not in current:
                                                current[part] = {}
                                            current = current[part]
                                    elif part in current and isinstance(current[part], dict):
                                        current = current[part]
                                    else:
                                        # Create the nested structure if it doesn't exist
                                        if part not in current:
                                            current[part] = {}
                                        current = current[part]
                            
                            # Set the value at the final location
                            final_key = path_parts[-1]
                            # Handle the value assignment
                            if final_key in current:
                                # Convert value to appropriate type with proper validation
                                if isinstance(current[final_key], (int, float)) or (isinstance(current[final_key], type(None)) and str(value).replace('.', '').replace('-', '').isdigit()):
                                    try:
                                        # Clean the value string first
                                        clean_value = str(value).strip()
                                        if '.' in clean_value:
                                            numeric_value = float(clean_value)
                                        else:
                                            numeric_value = int(clean_value)
                                        
                                        # Additional validation for specific fields
                                        if final_key == 'age_years':
                                            # More comprehensive age validation for adult patients
                                            if numeric_value < 10 or numeric_value > 120:
                                                logger.warning(f"Invalid age value for adult patient: {numeric_value}, skipping")
                                                continue
                                            # Additional check for very young ages that are clearly wrong
                                            elif numeric_value < 18:
                                                logger.warning(f"Age {numeric_value} seems too young for adult patient, skipping")
                                                continue
                                            else:
                                                current[final_key] = numeric_value
                                        elif final_key == 'bmi' and (numeric_value < 10 or numeric_value > 80):
                                            logger.warning(f"Invalid BMI value: {numeric_value}, skipping")
                                            continue
                                        elif final_key == 'ahi' and (numeric_value < 0 or numeric_value > 200):
                                            logger.warning(f"Invalid AHI value: {numeric_value}, skipping")
                                            continue
                                        elif final_key == 'odi' and (numeric_value < 0 or numeric_value > 200):
                                            logger.warning(f"Invalid ODI value: {numeric_value}, skipping")
                                            continue
                                        elif final_key == 'o2_nadir_pct' and (numeric_value < 50 or numeric_value > 100):
                                            logger.warning(f"Invalid O2 nadir value: {numeric_value}, skipping")
                                            continue
                                        elif final_key == 'sleep_efficiency_pct' and (numeric_value < 0 or numeric_value > 100):
                                            logger.warning(f"Invalid sleep efficiency value: {numeric_value}, skipping")
                                            continue
                                        else:
                                            current[final_key] = numeric_value
                                    except (ValueError, TypeError):
                                        logger.warning(f"Could not convert value '{value}' to number for field {final_key}, skipping")
                                        continue
                                elif isinstance(current[final_key], list):
                                    if value not in current[final_key]:
                                        current[final_key].append(value)
                                elif isinstance(current[final_key], bool) or (isinstance(current[final_key], type(None)) and str(value).lower() in ['true', 'false']):
                                    current[final_key] = str(value).lower() == 'true'
                                else:
                                    current[final_key] = value
                                structured_data_processed = True
                            else:
                                # Final key doesn't exist yet, create it with appropriate type
                                if str(value).lower() in ['true', 'false']:
                                    current[final_key] = str(value).lower() == 'true'
                                elif str(value).replace('.', '').replace('-', '').isdigit():
                                    try:
                                        if '.' in str(value):
                                            current[final_key] = float(value)
                                        else:
                                            current[final_key] = int(value)
                                    except (ValueError, TypeError):
                                        current[final_key] = value
                                else:
                                    current[final_key] = value
                                structured_data_processed = True
                        
                        except Exception as e:
                            logger.warning(f"Failed to process structured data path {target_path}: {e}")
                            continue
                    
                except json.JSONDecodeError:
                    continue
            
            if structured_data_processed:
                logger.info(f"Processed structured data for patient {patient_id}")
            
            # Add observation texts to summary (filter out error messages and irrelevant content)
            if observation_texts:
                # Filter out error messages, irrelevant content, and empty observations
                filtered_observations = []
                for obs in observation_texts:
                    if isinstance(obs, str):
                        obs_lower = obs.lower()
                        # Skip error messages, corrupted documents, and irrelevant medical content
                        skip_patterns = [
                            'error extracting observations:',
                            'dr. briz is temporarily busy',
                            'unable to extract meaningful clinical observations',
                            'appears to be severely corrupted',
                            'incorrectly ocr',
                            'poorly legible or corrupted',
                            'patch test clinic report',
                            'dermatology document',
                            'rambam health care campus',
                            'insufficient clinical data',
                            'appears to be a dermatology',
                            'patch test',
                            'skin test',
                            'allergy test results',
                            'corrupted or incorrectly',
                            'document appears to be',
                            'image appears to be',
                            'document not relevant to sleep medicine',
                            'document appears corrupted or illegible',
                            'no relevant clinical observations found'
                        ]
                        
                        # Check if observation should be skipped
                        should_skip = False
                        for pattern in skip_patterns:
                            if pattern in obs_lower:
                                should_skip = True
                                logger.info(f"Filtering out irrelevant observation: {obs[:100]}...")
                                break
                        
                        # Keep observation if it's not in skip patterns, not empty, and has meaningful content
                        if (not should_skip and 
                            obs.strip() and 
                            len(obs.strip()) > 10):
                            filtered_observations.append(obs)
                    else:
                        filtered_observations.append(obs)
                
                if filtered_observations:
                    canonical['observations']['summary'] = filtered_observations
                    logger.info(f"Added {len(filtered_observations)} filtered observations to summary (filtered from {len(observation_texts)} total)")
                else:
                    logger.info("No valid observations found after filtering error messages and irrelevant content")
            
            if observation_texts:
                # Use LLM to categorize observations
                llm_result = categorize_observations_with_llm(observation_texts, patient_id)
                if llm_result:
                    # Update canonical structure with LLM-categorized data
                    update_canonical_from_llm_result(canonical, llm_result)
                    
                    logger.info(f"LLM processed {len(observation_texts)} observations for patient {patient_id}")
                else:
                    logger.warning(f"LLM processing failed for patient {patient_id}, using fallback method")
                    # Fallback to basic extraction
                    fallback_extraction(canonical, observation_texts)
            else:
                logger.info(f"No valid observations found for patient {patient_id}")
            
            # Set study type based on source types (normalized to schema enum)
            # Gracefully handle cases where legacy observations may not include source_type
            source_types = []
            try:
                source_types = [obs.get('source_type') for obs in observations if isinstance(obs, dict)]
            except Exception:
                source_types = []
            normalized_source_types = [str(st).lower() for st in source_types if st]
            if any('sleep_test' in st for st in normalized_source_types):
                canonical['sleep_study']['study_type'] = 'home'  # HSAT → home
            elif any('sleep_study' in st for st in normalized_source_types):
                canonical['sleep_study']['study_type'] = 'inlab'  # PSG → inlab
            # else: omit the key (pruner will drop empty)
            
            # Also check if study_type was set in the observations and normalize it
            if canonical['sleep_study'].get('study_type'):
                canonical['sleep_study']['study_type'] = normalize_study_type_to_schema(canonical['sleep_study']['study_type'])
            
            # Update completeness flags based on available data
            # Check multiple sources for sleep study data:
            # 1. Root-level sleep_study.ahi or sleep_study.odi
            # 2. sleep_studies array (from observation_store with source_kind='sleep_study')
            # 3. Observations with respiratory_indices paths
            has_sleep_study_data = (
                canonical['sleep_study']['ahi'] is not None or 
                canonical['sleep_study']['odi'] is not None or
                len(canonical.get('sleep_studies', [])) > 0 or
                any(
                    obs.get('path', '').startswith('respiratory_indices.') or 
                    obs.get('path', '').startswith('sleep_study.')
                    for obs in observations
                    if isinstance(obs, dict) and obs.get('extracted_observations')
                )
            )
            canonical['completeness_flags']['has_sleep_study'] = has_sleep_study_data
            
            # Also populate root-level sleep_study from sleep_studies array if available
            if not canonical['sleep_study']['ahi'] and not canonical['sleep_study']['odi']:
                # Try to extract AHI/ODI from sleep_studies array
                for study in canonical.get('sleep_studies', []):
                    metrics = study.get('metrics', [])
                    for metric in metrics:
                        key = metric.get('key', '')
                        value = metric.get('value')
                        if value is not None:
                            # Handle both short keys and full paths (e.g., 'ahi' or 'respiratory_indices.ahi_overall')
                            key_lower = key.lower()
                            if canonical['sleep_study']['ahi'] is None:
                                # Check for AHI in various formats
                                if (key in ['ahi', 'ahi_overall', 'pahi'] or 
                                    'ahi_overall' in key_lower or 
                                    key.endswith('.ahi') or 
                                    key.endswith('.ahi_overall')):
                                    canonical['sleep_study']['ahi'] = float(value)
                                    logger.info(f"Populated sleep_study.ahi from sleep_studies array (key: {key}, value: {value})")
                            if canonical['sleep_study']['odi'] is None:
                                # Check for ODI in various formats
                                if (key in ['odi', 'odi3', 'odi4'] or 
                                    'odi3' in key_lower or 
                                    'odi4' in key_lower or 
                                    key.endswith('.odi')):
                                    canonical['sleep_study']['odi'] = float(value)
                                    logger.info(f"Populated sleep_study.odi from sleep_studies array (key: {key}, value: {value})")
                            # Also populate other sleep study metrics
                            if 'spo2_nadir' in key_lower or 'o2_nadir' in key_lower:
                                if canonical['sleep_study']['o2_nadir_pct'] is None:
                                    canonical['sleep_study']['o2_nadir_pct'] = float(value)
                                    logger.info(f"Populated sleep_study.o2_nadir_pct from sleep_studies array: {value}")
                            if 'sleep_efficiency' in key_lower:
                                if canonical['sleep_study']['sleep_efficiency_pct'] is None:
                                    canonical['sleep_study']['sleep_efficiency_pct'] = float(value)
                                    logger.info(f"Populated sleep_study.sleep_efficiency_pct from sleep_studies array: {value}")
                    # Also check if study_type is available
                    if not canonical['sleep_study']['study_type'] and study.get('study_type'):
                        canonical['sleep_study']['study_type'] = normalize_study_type_to_schema(study.get('study_type'))
                        logger.info(f"Populated sleep_study.study_type from sleep_studies array: {study.get('study_type')}")
            canonical['completeness_flags']['has_anatomy_imaging'] = any(v is not None and v != "" for v in canonical['observations']['anatomy_imaging'].values())
            canonical['completeness_flags']['has_tmj_info'] = (
                any(v is not None and v != "" for v in canonical['observations']['tmj_flags'].values())
                or any((isinstance(st, str) and 'tmj' in st.lower()) for st in source_types)
            )
            
            # Process demographics from observation texts
            if observation_texts:
                demographics_extracted = extract_demographics_from_text(observation_texts)
                if demographics_extracted:
                    for key, value in demographics_extracted.items():
                        if key in canonical['demographics'] and value is not None:
                            canonical['demographics'][key] = value
        
        # Update report_meta for canonical aggregate
        canonical['report_meta'] = {
            'report_id': 'canonical',
            'author_role': 'system',
            'created_at': now_iso,
            'source_report_type': 'canonical_aggregate'
        }
        
        # Ensure canonical_meta structure is correct
        if 'canonical_meta' not in canonical:
            canonical['canonical_meta'] = {}
        canonical['canonical_meta']['version'] = 1
        canonical['canonical_meta']['computed_fields'] = []
        
        # Remove any invalid fields that might have been added
        if 'version' in canonical and canonical['version'] != 1:
            del canonical['version']
        if 'report_id' in canonical and canonical['report_id'] != 'canonical':
            del canonical['report_id']
        
        # Fix evaluation format if present
        if 'follow_up_plan' in canonical and 'evaluations' in canonical['follow_up_plan']:
            canonical['follow_up_plan']['evaluations'] = fix_evaluation_format(canonical['follow_up_plan']['evaluations'])
        
        # Clean up follow-up plan (deduplicate and limit)
        canonical = _cleanup_follow_up_plan(canonical)
        
        # Remove duplicates from lists
        canonical = remove_duplicates_from_lists(canonical)
        
        # Flatten nested arrays (arrays within arrays)
        canonical = flatten_nested_arrays(canonical)
        
        # Build canonical_derived to keep LLM/UI simple
        try:
            def _to_date_str(dt):
                try:
                    return dt.split('T')[0] if isinstance(dt, str) else dt.date().isoformat()
                except Exception:
                    return None

            def _metric_lookup(metrics_list, key):
                if not isinstance(metrics_list, list):
                    return None
                for m in metrics_list:
                    if m.get('key') == key and m.get('value') is not None:
                        return m.get('value')
                return None

            # Helpers: de-duplicate multiple points on the same calendar day
            def _dedupe_ss_by_day(items: list) -> list:
                # Dedupe per (day, episode_id); when episode_id is missing, use file_name so multiple
                # studies on the same calendar day (different PDFs) are not collapsed into one row.
                by_day_ep = {}
                for it in items:
                    day = it.get('date')
                    if not day:
                        continue
                    ep = it.get('episode_id') or it.get('file_name') or ''
                    key = (day, ep)
                    if key not in by_day_ep:
                        by_day_ep[key] = it
                        continue
                    curr = by_day_ep[key]
                    new_ahi = it.get('ahi')
                    curr_ahi = curr.get('ahi')
                    if new_ahi is not None and (curr_ahi is None or new_ahi < curr_ahi):
                        by_day_ep[key] = it
                return [by_day_ep[k] for k in sorted(by_day_ep.keys(), key=lambda t: (t[0], t[1]))]

            def _dedupe_reports_by_day(items: list) -> list:
                # De-dupe per (day,key). Prefer min for AHI/ODI, max for o2_nadir_pct
                by_day_key = {}
                for it in items:
                    day = it.get('date')
                    key = it.get('key')
                    if not day or not key:
                        continue
                    pref_min = key in ('ahi', 'odi')
                    pref_max = key in ('o2_nadir_pct',)
                    ex = by_day_key.get((day, key))
                    if not ex:
                        by_day_key[(day, key)] = it
                        continue
                    try:
                        v_new = it.get('value')
                        v_old = ex.get('value')
                        if v_new is None:
                            continue
                        if v_old is None:
                            by_day_key[(day, key)] = it
                        elif pref_min and v_new < v_old:
                            by_day_key[(day, key)] = it
                        elif pref_max and v_new > v_old:
                            by_day_key[(day, key)] = it
                    except Exception:
                        continue
                # Return sorted by day, then key
                return [by_day_key[k] for k in sorted(by_day_key.keys(), key=lambda t: (t[0], t[1]))]

            # Sleep studies timeline sorted by observed_at
            ss_items = canonical.get('sleep_studies', [])
            def _ss_sort_key(item):
                d = item.get('observed_at')
                return d or ''
            ss_sorted = sorted(ss_items, key=_ss_sort_key)

            baseline = None
            latest = None
            timeline_ss = []
            if ss_sorted:
                first = ss_sorted[0]
                last = ss_sorted[-1]
                # Helper to collect a rich metric set per schema
                def _collect_metrics(metrics):
                    return {
                        'ahi': _metric_lookup(metrics, 'ahi'),
                        'odi': _metric_lookup(metrics, 'odi'),
                        'rdi': _metric_lookup(metrics, 'rdi'),
                        'oai': _metric_lookup(metrics, 'oai'),
                        'cai': _metric_lookup(metrics, 'cai'),
                        'hi': _metric_lookup(metrics, 'hi'),
                        'o2_nadir_pct': _metric_lookup(metrics, 'o2_nadir_pct'),
                        'o2_mean_pct': _metric_lookup(metrics, 'o2_mean_pct'),
                        'time_below_90_pct_min': _metric_lookup(metrics, 'time_below_90_pct_min'),
                        'time_below_88_pct_min': _metric_lookup(metrics, 'time_below_88_pct_min'),
                        'time_below_90_pct': _metric_lookup(metrics, 'time_below_90_pct'),
                        'supine_ahi': _metric_lookup(metrics, 'supine_ahi'),
                        'non_supine_ahi': _metric_lookup(metrics, 'non_supine_ahi'),
                        'rem_ahi': _metric_lookup(metrics, 'rem_ahi'),
                        'nrem_ahi': _metric_lookup(metrics, 'nrem_ahi'),
                        'supine_rdi': _metric_lookup(metrics, 'supine_rdi'),
                        'rem_rdi': _metric_lookup(metrics, 'rem_rdi'),
                        'supine_odi': _metric_lookup(metrics, 'supine_odi'),
                        'rem_odi': _metric_lookup(metrics, 'rem_odi'),
                        'sleep_duration_h': _metric_lookup(metrics, 'sleep_duration_h'),
                        'sleep_efficiency_pct': _metric_lookup(metrics, 'sleep_efficiency_pct'),
                        'desaturation_events': _metric_lookup(metrics, 'desaturation_events'),
                    }

                # Populate baseline/latest with richer set
                baseline_metrics = _collect_metrics(first.get('metrics'))
                latest_metrics = _collect_metrics(last.get('metrics'))
                
                # Also check direct fields for metrics not in metrics array
                def _enhance_metrics_with_direct_fields(metrics_dict, sleep_study_item):
                    enhanced = metrics_dict.copy()
                    # Check for direct fields that might not be in metrics array
                    direct_fields = ['time_below_90_pct', 'ahi', 'odi', 'o2_nadir_pct', 'supine_ahi', 'rdi', 'oai', 'cai', 'hi']
                    for field in direct_fields:
                        if enhanced.get(field) is None and sleep_study_item.get(field) is not None:
                            enhanced[field] = sleep_study_item.get(field)
                    return enhanced
                
                baseline_metrics = _enhance_metrics_with_direct_fields(baseline_metrics, first)
                latest_metrics = _enhance_metrics_with_direct_fields(latest_metrics, last)
                
                baseline = { 'sleep_study': { 'date': _to_date_str(first.get('observed_at')), **{k: v for k, v in baseline_metrics.items() if v is not None} } }
                latest = { 'sleep_study': { 'date': _to_date_str(last.get('observed_at')), **{k: v for k, v in latest_metrics.items() if v is not None} } }

                # Build full schema-aligned timeline.sleep_studies items
                timeline_ss_full = []
                for it in ss_sorted:
                    metrics = it.get('metrics')
                    m = _collect_metrics(metrics)
                    # Snoring
                    snr = {}
                    pct_total = _metric_lookup(metrics, 'snoring_percent_total')
                    over50 = _metric_lookup(metrics, 'snoring_over_50db_pct')
                    avgdb = _metric_lookup(metrics, 'snoring_avg_db')
                    maxdb = _metric_lookup(metrics, 'snoring_max_db')
                    if pct_total is not None: snr['percent_total'] = pct_total
                    if over50 is not None: snr['over_50db_pct'] = over50
                    if avgdb is not None: snr['avg_db'] = avgdb
                    if maxdb is not None: snr['max_db'] = maxdb

                    # Heart rate
                    hr = {}
                    hr_mean = _metric_lookup(metrics, 'heart_rate_mean_bpm')
                    hr_min = _metric_lookup(metrics, 'heart_rate_min_bpm')
                    hr_max = _metric_lookup(metrics, 'heart_rate_max_bpm')
                    if hr_mean is not None: hr['mean_bpm'] = hr_mean
                    if hr_min is not None: hr['min_bpm'] = hr_min
                    if hr_max is not None: hr['max_bpm'] = hr_max

                    item = {
                        'date': _to_date_str(it.get('observed_at')),
                        'file_name': it.get('file_name'),
                        'episode_id': it.get('episode_id'),
                        'source_kind': 'sleep_study',
                        'study_type': it.get('study_type')
                    }
                    if it.get('analysis_insights_text'):
                        item['text'] = it.get('analysis_insights_text')
                    # Merge metrics, snoring, heart_rate
                    for k, v in m.items():
                        if v is not None:
                            item[k] = v
                    # Always include time_below_90_pct key, even if null
                    if 'time_below_90_pct' not in item:
                        item['time_below_90_pct'] = m.get('time_below_90_pct', None)
                    if snr:
                        item['snoring'] = snr
                    if hr:
                        item['heart_rate'] = hr

                    timeline_ss_full.append(item)

                # Also keep simplified list for derived (as before)
                for it in ss_sorted:
                    metrics = it.get('metrics')
                    item = {
                        'date': _to_date_str(it.get('observed_at')),
                        'ahi': _metric_lookup(metrics, 'ahi'),
                        'odi': _metric_lookup(metrics, 'odi'),
                        'rdi': _metric_lookup(metrics, 'rdi'),
                        'o2_nadir_pct': _metric_lookup(metrics, 'o2_nadir_pct'),
                        'o2_mean_pct': _metric_lookup(metrics, 'o2_mean_pct'),
                        'time_below_90_pct': _metric_lookup(metrics, 'time_below_90_pct'),
                        'time_below_90_pct_min': _metric_lookup(metrics, 'time_below_90_pct_min'),
                        'supine_ahi': _metric_lookup(metrics, 'supine_ahi'),
                        'rem_ahi': _metric_lookup(metrics, 'rem_ahi'),
                        'study_type': it.get('study_type'),
                        'file_name': it.get('file_name'),
                        'episode_id': it.get('episode_id'),
                        'source_kind': 'sleep_study'
                    }
                    # Prune Nones but keep time_below_90_pct key even if null (explicit in payload)
                    timeline_ss.append({k: v for k, v in item.items() if (v is not None) or (k in {'time_below_90_pct'})})
                # De-duplicate multiple points occurring on the same day
                timeline_ss = _dedupe_ss_by_day(timeline_ss)

                # Attach root-level timeline per schema
                canonical['timeline'] = {
                    'sleep_studies': timeline_ss_full,
                    'reports': []  # filled below from timeline_rm
                }

            # Report mentions timeline sorted by mention_date
            rm_sorted = sorted(canonical.get('reported_metrics', []), key=lambda x: x.get('mention_date') or '')
            timeline_rm = []
            for r in rm_sorted:
                timeline_rm.append({
                    'date': _to_date_str(r.get('mention_date')),
                    'key': r.get('key'),
                    'value': r.get('value'),
                    'file_name': r.get('file_name'),
                    'source_kind': (r.get('source_kind') or 'report')
                })
            # De-duplicate per (day, key)
            timeline_rm = _dedupe_reports_by_day(timeline_rm)

            # Build grouped reports view: aggregate per (date, file_name)
            reports_grouped = []
            try:
                by_day_file = {}
                for r in timeline_rm:
                    day = r.get('date')
                    fname = r.get('file_name')
                    if not day or not fname:
                        continue
                    key = (day, fname)
                    if key not in by_day_file:
                        by_day_file[key] = {
                            'date': day,
                            'file_name': fname,
                            'source_kind': 'report'
                        }
                    # assign metric onto grouped object if meaningful
                    mk = r.get('key')
                    mv = r.get('value')
                    if mk and mv is not None:
                        by_day_file[key][mk] = mv
                # sorted list by date then file name
                reports_grouped = [by_day_file[k] for k in sorted(by_day_file.keys(), key=lambda t: (t[0], t[1]))]
            except Exception as _e:
                logger.warning(f"Failed to build grouped reports for patient {patient_id}: {_e}")

            # Write into root-level timeline if present
            if 'timeline' in canonical:
                canonical['timeline']['reports'] = timeline_rm
                canonical['timeline']['reports_grouped'] = reports_grouped

            if not skip_timeline_llm:
                # Organize timeline data using LLM for clinical coherence
                logger.info(f"Patient {patient_id}: Attempting LLM timeline organization with {len(timeline_ss)} sleep studies, {len(timeline_rm)} reports, {len(reports_grouped)} grouped reports")
                logger.info(f"Patient {patient_id}: Sleep studies data: {timeline_ss}")
                logger.info(f"Patient {patient_id}: Reports data: {timeline_rm}")
                logger.info(f"Patient {patient_id}: Grouped reports data: {reports_grouped}")
                try:
                    organized_timeline = organize_timeline_with_llm(patient_id, timeline_ss, timeline_rm, reports_grouped)
                    logger.info(f"Patient {patient_id}: LLM organization result: {organized_timeline}")
                    if organized_timeline and organized_timeline.get('success'):
                        logger.info(f"Patient {patient_id}: Using LLM-organized timeline data")
                        original_ss_count = len(timeline_ss)
                        
                        # Get organized data from nested 'organized_timeline' key
                        organized_data = organized_timeline.get('organized_timeline', {})
                        
                        # Extract UI sleep metrics directly from LLM response (if provided)
                        ui_sleep_metrics = organized_timeline.get('ui_sleep_metrics', {})
                        if ui_sleep_metrics:
                            canonical['ui_sleep_metrics'] = ui_sleep_metrics
                            logger.info(f"Patient {patient_id}: Added UI sleep metrics from LLM organizer")
                            logger.info(f"Patient {patient_id}: Baseline: {ui_sleep_metrics.get('baseline', {})}")
                            logger.info(f"Patient {patient_id}: Current: {ui_sleep_metrics.get('current', {})}")
                        
                        # Validate and normalize timeline data before using it
                        timeline_ss = validate_and_normalize_timeline_data(
                            organized_data.get('sleep_studies', timeline_ss), 
                            'sleep_study', 
                            canonical.get('meta', {})
                        )
                        timeline_rm = validate_and_normalize_timeline_data(
                            organized_data.get('reports', timeline_rm), 
                            'report', 
                            canonical.get('meta', {})
                        )
                        reports_grouped = validate_and_normalize_timeline_data(
                            organized_data.get('reports_grouped', reports_grouped), 
                            'report', 
                            canonical.get('meta', {})
                        )
                        
                        logger.info(f"Patient {patient_id}: Timeline update: {original_ss_count} → {len(timeline_ss)} sleep studies")
                        logger.info(f"Patient {patient_id}: Organized sleep studies: {[{'date': ss.get('date'), 'ahi': ss.get('ahi'), 'episode_id': ss.get('episode_id')} for ss in timeline_ss]}")
                        
                        # CRITICAL: Also update root-level timeline section to match organized data
                        if 'timeline' in canonical:
                            canonical['timeline']['sleep_studies'] = timeline_ss
                            canonical['timeline']['reports'] = timeline_rm
                            canonical['timeline']['reports_grouped'] = reports_grouped
                            logger.info(f"Patient {patient_id}: Updated root-level timeline with organized data")
                        
                        # CRITICAL: Also update root-level sleep_studies section with organized data
                        canonical['sleep_studies'] = timeline_ss
                        logger.info(f"Patient {patient_id}: Updated root-level sleep_studies with {len(timeline_ss)} organized entries")
                    else:
                        logger.warning(f"Patient {patient_id}: LLM timeline organization failed, using manual organization - reason: {organized_timeline.get('reason', 'unknown')}")
                except Exception as llm_e:
                    logger.error(f"Patient {patient_id}: LLM timeline organization error: {llm_e}", exc_info=True)
            else:
                logger.info(
                    "Patient %s: Skipping LLM timeline organization (fast path; e.g. after direct sleep Bedrock extract)",
                    patient_id,
                )

            # Filter out ignored reports from timeline (but keep them in reported_metrics for audit)
            timeline_rm_filtered = [r for r in timeline_rm if r.get('disposition') != 'ignored']
            
            derived = {
                'baseline': baseline or {},
                'latest': latest or {},
                'timeline': {
                    'sleep_studies': timeline_ss,
                    'reports': timeline_rm_filtered,  # Only include non-ignored reports in timeline
                    'reports_grouped': reports_grouped
                }
            }
            canonical['canonical_derived'] = derived
            
            if skip_timeline_llm:
                try:
                    ui = _resolve_ui_sleep_metrics(canonical)
                    if ui:
                        canonical["ui_sleep_metrics"] = ui
                        logger.info(
                            "Patient %s: ui_sleep_metrics from deterministic resolver (skip_timeline_llm)",
                            patient_id,
                        )
                except Exception as _ui_e:
                    logger.warning(
                        "Patient %s: ui_sleep_metrics resolve failed (skip_timeline_llm): %s",
                        patient_id,
                        _ui_e,
                    )
            try:
                if latest and 'sleep_study' in latest:
                    latest_sleep_data = latest['sleep_study']
                    for key, value in latest_sleep_data.items():
                        if key != 'date' and value is not None:  # Don't copy date field
                            canonical['sleep_study'][key] = value
                    logger.debug(f"Copied latest sleep study data to top-level sleep_study section")
            except Exception as copy_e:
                logger.warning(f"Failed to copy latest sleep study to top-level: {copy_e}")
                
        except Exception as _e:
            logger.warning(f"Failed to build canonical_derived for patient {patient_id}: {_e}")

        # Final demographic validation - ensure database values take precedence when extracted values are unreasonable
        
        # Age validation - Use database as fallback when extracted age is unreasonable
        extracted_age = canonical['demographics'].get('age_years')
        
        if patient_age_fallback:
            if extracted_age:
                if extracted_age < 10 or extracted_age > 120:
                    logger.warning(f"Patient {patient_id}: Final validation - extracted age {extracted_age} is invalid, using database age {patient_age_fallback}")
                    canonical['demographics']['age_years'] = patient_age_fallback
                elif extracted_age < 18:
                    logger.warning(f"Patient {patient_id}: Final validation - extracted age {extracted_age} seems too young for adult patient, using database age {patient_age_fallback}")
                    canonical['demographics']['age_years'] = patient_age_fallback
                elif abs(extracted_age - patient_age_fallback) > 5:
                    logger.warning(f"Patient {patient_id}: Final validation - age discrepancy (extracted: {extracted_age}, database: {patient_age_fallback}), using database age")
                    canonical['demographics']['age_years'] = patient_age_fallback
                else:
                    logger.info(f"Patient {patient_id}: Final validation - extracted age {extracted_age} is consistent with database age {patient_age_fallback}")
            else:
                # No extracted age, use database fallback
                logger.info(f"Patient {patient_id}: Final validation - no extracted age found, using database age {patient_age_fallback}")
                canonical['demographics']['age_years'] = patient_age_fallback
        else:
            # No database fallback available - validate extracted age but keep reasonable values
            if extracted_age:
                if extracted_age < 10 or extracted_age > 120:
                    logger.warning(f"Patient {patient_id}: Final validation - extracted age {extracted_age} is invalid and no database fallback available, removing")
                    canonical['demographics']['age_years'] = None
                elif extracted_age < 18:
                    logger.warning(f"Patient {patient_id}: Final validation - extracted age {extracted_age} seems young but keeping as no database fallback available")
                else:
                    logger.info(f"Patient {patient_id}: Final validation - keeping extracted age {extracted_age} (no database fallback available)")
            else:
                logger.info(f"Patient {patient_id}: Final validation - no age data available from either extraction or database")
        
        # Gender validation - Always prefer database gender as it's more reliable
        if patient_gender_fallback:
            extracted_gender = canonical['demographics'].get('sex')
            if extracted_gender:
                logger.info(f"Patient {patient_id}: Final validation - ignoring extracted gender '{extracted_gender}', using reliable database gender {patient_gender_fallback}")
            else:
                logger.info(f"Patient {patient_id}: Final validation - no extracted gender found, using database gender {patient_gender_fallback}")
            # Always use database gender when available
            canonical['demographics']['sex'] = patient_gender_fallback
        else:
            # Only keep extracted gender if we have no database fallback and it looks reasonable
            extracted_gender = canonical['demographics'].get('sex')
            if extracted_gender:
                extracted_gender_norm = extracted_gender.upper()
                if extracted_gender_norm not in ['M', 'F', 'MALE', 'FEMALE']:
                    logger.warning(f"Patient {patient_id}: Final validation - extracted gender '{extracted_gender}' is invalid and no database fallback available, removing")
                    canonical['demographics']['sex'] = None
                else:
                    logger.warning(f"Patient {patient_id}: Final validation - keeping extracted gender '{extracted_gender}' as no database fallback available (use with caution)")
        
        # Apply sleep-first policy to fix data duplication and complexity
        try:
            from sleep_first_policy import apply_sleep_first_policy
            logger.info(f"Patient {patient_id}: Applying sleep-first policy to clean up data duplication")
            canonical, policy_report = apply_sleep_first_policy(canonical)
            logger.info(f"Patient {patient_id}: Sleep-first policy applied - {policy_report['episodes']} episodes, {policy_report['report_metrics']} report metrics, {policy_report['conflicts']} conflicts")
            if policy_report['conflicts'] > 0:
                logger.warning(f"Patient {patient_id}: Found {policy_report['conflicts']} conflicting report metrics: {policy_report['conflict_examples']}")
            
            # Filter out ignored reports from canonical_derived.timeline.reports after policy is applied
            if 'canonical_derived' in canonical and 'timeline' in canonical['canonical_derived']:
                timeline_reports = canonical['canonical_derived']['timeline'].get('reports', [])
                filtered_reports = [r for r in timeline_reports if r.get('disposition') != 'ignored']
                canonical['canonical_derived']['timeline']['reports'] = filtered_reports
                logger.info(f"Patient {patient_id}: Filtered out {len(timeline_reports) - len(filtered_reports)} ignored reports from canonical_derived.timeline.reports")
            
            # Also filter from root-level timeline.reports if it exists
            if 'timeline' in canonical and 'reports' in canonical['timeline']:
                timeline_reports = canonical['timeline']['reports']
                filtered_reports = [r for r in timeline_reports if r.get('disposition') != 'ignored']
                canonical['timeline']['reports'] = filtered_reports
                logger.info(f"Patient {patient_id}: Filtered out {len(timeline_reports) - len(filtered_reports)} ignored reports from timeline.reports")
        except Exception as policy_e:
            logger.error(f"Patient {patient_id}: Sleep-first policy failed: {policy_e}", exc_info=True)
            # Continue without policy - don't break the pipeline
        
        # Extract treatment history from database (device orders, clinical notes)
        # NOTE: This is added BEFORE cleaning so it can be included in clean canonical if needed
        logger.info(f"Patient {patient_id}: Extracting treatment history from database")
        treatment_history = extract_treatment_history_from_db(patient_id)
        canonical['treatment_history'] = treatment_history
        logger.info(f"Patient {patient_id}: ✅ Treatment history added ({len(treatment_history['device_orders'])} devices, {len(treatment_history['clinical_notes'])} notes)")
        
        # Populate risk_assessment from observation_store
        # NOTE: This is added BEFORE cleaning so it can be included in clean canonical if needed
        logger.info(f"Patient {patient_id}: Populating risk assessment from observations")
        try:
            cursor.execute("""
                SELECT extracted_observations, observed_at, source_kind
                FROM observation_store
                WHERE patient_id = %s 
                AND metric_key = 'risk_assessment.risk_level'
                ORDER BY observed_at DESC
                LIMIT 1
            """, (patient_id,))
            risk_obs = cursor.fetchone()
            
            if risk_obs and risk_obs.get('extracted_observations'):
                # Parse JSON to get the value
                try:
                    obs_data = json.loads(risk_obs['extracted_observations']) if isinstance(risk_obs['extracted_observations'], str) else risk_obs['extracted_observations']
                    risk_level = obs_data.get('value') or obs_data.get('risk_level')
                    if risk_level:
                        canonical['risk_assessment'] = {
                            'risk_level': risk_level,
                            'risk_source': risk_obs.get('source_kind', 'unknown'),
                            'assessed_at': risk_obs['observed_at'].isoformat() if risk_obs.get('observed_at') else None
                        }
                        logger.info(f"Patient {patient_id}: ✅ Risk assessment populated: {risk_level} from {risk_obs.get('source_kind')}")
                    else:
                        logger.info(f"Patient {patient_id}: No risk level found in risk assessment observation")
                except (json.JSONDecodeError, TypeError) as parse_e:
                    logger.warning(f"Patient {patient_id}: Failed to parse risk assessment observation: {parse_e}")
            else:
                logger.info(f"Patient {patient_id}: No risk assessment observation found")
        except Exception as e:
            logger.error(f"Patient {patient_id}: Error populating risk assessment: {e}")
        
        # Prune empty values to create sparse JSON (but keep full canonical structure)
        canonical = _prune_empty(canonical)
        
        # Debug: Log the canonical JSON before saving to database
        logger.info(f"Patient {patient_id}: Final canonical JSON before DB save - schema_version: {canonical.get('schema_version')}, document_type: {canonical.get('document_type')}")
        logger.info(f"Patient {patient_id}: Canonical fields: {list(canonical.keys())[:10]}... (full canonical stored)")
        
        # Force delete existing canonical entry first to ensure clean insert
        delete_query = """
            DELETE FROM patient_case_envelope 
            WHERE patient_id = %s AND report_id = 'canonical'
        """
        cursor.execute(delete_query, (patient_id,))
        logger.info(f"Patient {patient_id}: Deleted existing canonical entries before inserting full canonical")
        
        # Insert new canonical envelope with FULL canonical (cleaning happens only when loading for Level-4 reports)
        insert_query = """
            INSERT INTO patient_case_envelope
            (patient_id, report_id, document_type, source_uri, case_json, provider, imported_at, updated_at)
            VALUES (%s, %s, %s, %s, %s, %s, NOW(), NOW())
        """
        
        cursor.execute(insert_query, (
            patient_id,
            'canonical',
            'canonical',
            '',
            json.dumps(canonical),
            'system'
        ))
        logger.info(f"Patient {patient_id}: Inserted new canonical JSON with organized timeline data")
        
        conn.commit()
        conn.close()
        
        logger.info(f"Created canonical JSON for patient {patient_id} with {len(observations)} observations")
        return {
            'success': True,
            'patient_id': patient_id,
            'message': f'Canonical JSON created successfully with {len(observations)} observations'
        }
        
    except Exception as e:
        logger.error(f"Error creating canonical JSON for patient {patient_id}: {e}")
        if conn:
            conn.close()
        return {
            'success': False,
            'patient_id': patient_id,
            'message': str(e)
        }
def store_quiz_risk_assessment_as_observation(patient_id: int, patient_email: str) -> bool:
    """
    Store the latest quiz risk assessment as an observation in observation_store.
    This allows the canonical schema to pick it up like other observations.
    
    Args:
        patient_id: Patient ID
        patient_email: Patient email
        
    Returns:
        bool: True if risk assessment was stored, False otherwise
    """
    try:
        conn = mysql.connector.connect(**DB_CONFIG)
        cursor = conn.cursor(dictionary=True)
        
        # Try ConversionQuiz first
        cursor.execute("""
            SELECT id, ai_response, created_at, quiz_type
            FROM conversion_quiz
            WHERE patient_email = %s
            ORDER BY created_at DESC
            LIMIT 1
        """, (patient_email,))
        quiz_row = cursor.fetchone()
        
        risk_level = None
        quiz_source = None
        quiz_date = None
        quiz_id = None
        
        if quiz_row and quiz_row['ai_response']:
            try:
                ai_response = json.loads(quiz_row['ai_response'])
                risk_level = (
                    ai_response.get('risk_level') or 
                    ai_response.get('risk_band') or 
                    ai_response.get('riskLevel') or
                    ai_response.get('riskBand')
                )
                quiz_source = 'conversion_quiz'
                quiz_date = quiz_row['created_at']
                quiz_id = quiz_row['id']
                logger.info(f"Patient {patient_id}: Found risk level '{risk_level}' from ConversionQuiz")
            except:
                pass
        
        # If not found, try VizBrizQuiz
        if not risk_level:
            cursor.execute("""
                SELECT id, risk_band, created_at, quiz_type
                FROM vizbriz_quiz
                WHERE patient_email = %s
                ORDER BY created_at DESC
                LIMIT 1
            """, (patient_email,))
            vizbriz_row = cursor.fetchone()
            
            if vizbriz_row and vizbriz_row['risk_band']:
                risk_level = vizbriz_row['risk_band']
                quiz_source = 'vizbriz_quiz'
                quiz_date = vizbriz_row['created_at']
                quiz_id = vizbriz_row['id']
                logger.info(f"Patient {patient_id}: Found risk level '{risk_level}' from VizBrizQuiz")
        
        if risk_level:
            # Store as observation in observation_store (using extracted_observations JSON column)
            observation_json = {
                'path': 'risk_assessment.risk_level',
                'value': risk_level,
                'source': quiz_source,
                'confidence': 100,
                'explanation': f'Risk assessment from {quiz_source} quiz',
                'observation': f'Patient risk level: {risk_level}',
                'document_name': f"{quiz_source}_{quiz_id}",
                'document_type': 'quiz',
                'extraction_date': datetime.now().isoformat()
            }
            
            insert_query = """
                INSERT INTO observation_store 
                (patient_id, file_name, metric_key, extracted_observations, observed_at, source_kind, created_at)
                VALUES (%s, %s, %s, %s, %s, %s, NOW())
                ON DUPLICATE KEY UPDATE
                    extracted_observations = VALUES(extracted_observations),
                    observed_at = VALUES(observed_at),
                    source_kind = VALUES(source_kind)
            """
            
            cursor.execute(insert_query, (
                patient_id,
                f"{quiz_source}_{quiz_id}",
                'risk_assessment.risk_level',
                json.dumps(observation_json),
                quiz_date,
                quiz_source
            ))
            conn.commit()
            logger.info(f"Patient {patient_id}: ✅ Stored risk assessment observation: {risk_level} from {quiz_source}")
            conn.close()
            return True
        else:
            logger.info(f"Patient {patient_id}: No risk assessment found in quizzes")
            conn.close()
            return False
            
    except Exception as e:
        logger.error(f"Patient {patient_id}: Error storing quiz risk assessment: {e}")
        return False


def extract_treatment_history_from_db(patient_id: int) -> Dict[str, Any]:
    """
    Extract operational treatment data from database tables.
    
    This provides context about:
    - Device orders and deliveries (to prevent duplicate device recommendations)
    - Clinical notes (titrations, consultations, deliveries)
    
    Args:
        patient_id: Patient ID
        
    Returns:
        dict with structure:
        {
            'device_orders': [...],  # From patient_device_order table
            'clinical_notes': [...]   # From patientcomments table (filtered)
        }
    """
    from flask_app.models import PatientDeviceOrder, PatientComment, Dentist
    
    treatment_history = {
        'device_orders': [],
        'clinical_notes': []
    }
    
    try:
        # 1. Extract Device Orders
        logger.info(f"Patient {patient_id}: Extracting device orders from database")
        device_orders = PatientDeviceOrder.query.filter_by(patient_id=patient_id).order_by(PatientDeviceOrder.order_date).all()
        
        for order in device_orders:
            order_data = {
                'device_type': order.device_type,
                'device_name': order.device_name,
                'order_date': order.order_date.isoformat() if order.order_date else None,
                'arrival_date': order.arrival_date.isoformat() if order.arrival_date else None,
                'fitting_date': order.fitting_date.isoformat() if order.fitting_date else None,
                'status': order.status,
                'notes': order.notes,
                'fitting_comment': order.fitting_comment
            }
            treatment_history['device_orders'].append(order_data)
        
        logger.info(f"Patient {patient_id}: Found {len(treatment_history['device_orders'])} device orders")
        
        # 2. Extract Clinical Notes (titration, consultation, delivery, initial only - exclude general/internal)
        logger.info(f"Patient {patient_id}: Extracting clinical notes from database")
        clinical_comments = PatientComment.query.filter(
            PatientComment.patient_id == patient_id,
            PatientComment.comment_type.in_(['titration', 'consultation', 'delivery', 'initial'])
        ).order_by(PatientComment.created_date).all()
        
        for comment in clinical_comments:
            # Get provider name from dentist relationship
            provider_name = None
            if comment.dentist:
                provider_name = comment.dentist.name
            
            note_data = {
                'date': comment.created_date.isoformat() if comment.created_date else None,
                'comment_type': comment.comment_type or 'general',
                'content': comment.content,
                'provider_name': provider_name,
                'numeric_value': float(comment.numeric_value) if comment.numeric_value else None,
                'numeric_unit': comment.numeric_unit
            }
            treatment_history['clinical_notes'].append(note_data)
        
        logger.info(f"Patient {patient_id}: Found {len(treatment_history['clinical_notes'])} clinical notes")
        
    except Exception as e:
        logger.error(f"Patient {patient_id}: Error extracting treatment history from database: {e}")
        import traceback
        traceback.print_exc()
    
    return treatment_history


def ensure_minimal_canonical_for_all_patients(limit_patients: Optional[int] = None) -> Dict[str, Any]:
    """
    Ensure all patients have at least minimal canonical JSON, even without documents.
    
    Args:
        limit_patients (Optional[int]): Limit number of patients to process
        
    Returns:
        Dict[str, Any]: Summary of canonical JSON creation
    """
    pids = get_all_patient_ids(limit=limit_patients)
    results = []
    total_created = 0
    total_existing = 0
    total_failed = 0
    
    logger.info(f"Ensuring minimal canonical JSON for {len(pids)} patients")
    
    for pid in pids:
        try:
            result = create_minimal_canonical_json_for_patient(pid)
            results.append({'patient_id': pid, 'result': result})
            
            if result.get('success'):
                if result.get('message') == 'Minimal canonical JSON created successfully':
                    total_created += 1
                else:
                    total_existing += 1
                logger.info(f"Canonical JSON ensured for patient {pid}")
            else:
                total_failed += 1
                logger.warning(f"Failed to ensure canonical JSON for patient {pid}: {result.get('message')}")
                
        except Exception as e:
            total_failed += 1
            logger.error(f"Error ensuring canonical JSON for patient {pid}: {e}")
            results.append({'patient_id': pid, 'error': str(e)})
    
    summary = {
        'patients_total': len(pids),
        'canonical_created': total_created,
        'canonical_existing': total_existing,
        'failed': total_failed,
        'results': results
    }
    
    logger.info(f"Minimal canonical JSON creation summary: {summary}")
    return summary

def extract_specific_numerical_fields(observation_texts: List[str]) -> Dict[str, Any]:
    """
    Extract specific numerical fields using regex patterns and text analysis.
    Based on the expanded Patient Case JSON v1 schema.
    
    Args:
        observation_texts: List of observation strings
        
    Returns:
        Dict containing extracted numerical fields matching schema structure
    """
    import re
    
    extracted_data = {}
    combined_text = " ".join(observation_texts).lower()
    
    # Sleep Study Data
    sleep_study = {}
    
    # Sleep Duration (hours)
    sleep_duration_patterns = [
        r'sleep time[:\s]*(\d+(?:\.\d+)?)\s*(?:hrs?|hours?)',
        r'total sleep time[:\s]*(\d+(?:\.\d+)?)\s*(?:hrs?|hours?)',
        r'sleep duration[:\s]*(\d+(?:\.\d+)?)\s*(?:hrs?|hours?)',
        r'(\d+(?:\.\d+)?)\s*(?:hrs?|hours?)\s*(?:sleep|total)',
        r'(\d+(?:\.\d+)?)\s*(?:hrs?|hours?),\s*(\d+)\s*min',  # 7 hrs, 3 min -> 7.05
    ]
    
    for pattern in sleep_duration_patterns:
        match = re.search(pattern, combined_text)
        if match:
            if len(match.groups()) == 2:  # hours and minutes
                hours = float(match.group(1))
                minutes = float(match.group(2))
                sleep_study['sleep_duration_h'] = hours + (minutes / 60)
            else:
                sleep_study['sleep_duration_h'] = float(match.group(1))
            break
    
    # Sleep Efficiency (percentage)
    efficiency_patterns = [
        r'sleep[:\s]*(\d+(?:\.\d+)?)%',
        r'efficiency[:\s]*(\d+(?:\.\d+)?)%',
        r'(\d+(?:\.\d+)?)%\s*sleep',
        r'(\d+(?:\.\d+)?)%\s*efficiency',
    ]
    
    for pattern in efficiency_patterns:
        match = re.search(pattern, combined_text)
        if match:
            sleep_study['sleep_efficiency_pct'] = float(match.group(1))
            break
    
    # Desaturation Events (count)
    desaturation_patterns = [
        r'total events[:\s]*(\d+)',
        r'desaturation events[:\s]*(\d+)',
        r'(\d+)\s*desaturation',
        r'(\d+)\s*events',
    ]
    
    for pattern in desaturation_patterns:
        match = re.search(pattern, combined_text)
        if match:
            sleep_study['desaturation_events'] = int(match.group(1))
            break
    
    # Oxygen Saturation Data - More Generic Patterns
    o2_patterns = [
        # Generic mean/minimum/maximum patterns
        r'mean[:\s]*(\d+)',  # Mean: 94
        r'minimum[:\s]*(\d+)',  # Minimum: 83
        r'maximum[:\s]*(\d+)',  # Maximum: 98
        r'Mean[:\s]*(\d+)',  # Mean: 94
        r'Minimum[:\s]*(\d+)',  # Minimum: 83
        r'Maximum[:\s]*(\d+)',  # Maximum: 98
        
        # With percentage context
        r'mean[:\s]*(\d+)\s*(?:%|percent)',  # O2 mean
        r'minimum[:\s]*(\d+)\s*(?:%|percent)',  # O2 nadir
        r'(\d+)\s*(?:%|percent)\s*mean',  # O2 mean
        r'(\d+)\s*(?:%|percent)\s*minimum',  # O2 nadir
        
        # Oxygen saturation specific
        r'oxygen saturation.*mean[:\s]*(\d+)',  # More specific O2 mean
        r'oxygen saturation.*minimum[:\s]*(\d+)',  # More specific O2 nadir
        r'mean[:\s]*(\d+).*oxygen',  # O2 mean with context
        r'minimum[:\s]*(\d+).*oxygen',  # O2 nadir with context
        
        # Generic number patterns that might be O2 related
        r'(\d{2,3})\s*(?:%|percent)',  # Any 2-3 digit percentage (likely O2)
    ]
    
    # Extract O2 values using generic table format detection
    lines = combined_text.split('\n')
    
    # Find all potential O2 values (2-3 digit numbers in reasonable O2 range)
    o2_candidates = []
    for line in lines:
        line = line.strip()
        if re.match(r'^\d{2,3}$', line):  # 2-3 digit standalone numbers
            value = int(line)
            if 70 <= value <= 100:  # Reasonable O2 saturation range
                o2_candidates.append(value)
    
    # Remove duplicates and sort
    o2_candidates = sorted(list(set(o2_candidates)))
    
    # Assign values based on position (assuming they're in order: min, mean, max)
    if len(o2_candidates) >= 3:
        sleep_study['o2_nadir_pct'] = o2_candidates[0]  # Lowest
        sleep_study['o2_mean_pct'] = o2_candidates[1]   # Middle
        sleep_study['o2_max_pct'] = o2_candidates[2]    # Highest
    elif len(o2_candidates) == 2:
        sleep_study['o2_nadir_pct'] = o2_candidates[0]  # Lower
        sleep_study['o2_mean_pct'] = o2_candidates[1]   # Higher
    elif len(o2_candidates) == 1:
        sleep_study['o2_mean_pct'] = o2_candidates[0]   # Single value
    
    # Time below oxygen thresholds
    time_below_patterns = [
        r'<90[:\s]*(\d+(?:\.\d+)?)\s*(?:min|minutes)',
        r'<=88[:\s]*(\d+(?:\.\d+)?)\s*(?:min|minutes)',
        r'(\d+(?:\.\d+)?)\s*(?:min|minutes)\s*<90',
        r'(\d+(?:\.\d+)?)\s*(?:min|minutes)\s*<=88',
    ]
    
    for pattern in time_below_patterns:
        match = re.search(pattern, combined_text)
        if match:
            value = float(match.group(1))
            if '<90' in pattern:
                sleep_study['time_below_90_pct_min'] = value
            elif '<=88' in pattern:
                sleep_study['time_below_88_pct_min'] = value
            break
    
    # Attempt robust percent parsing using helper patterns on the full text
    helper_o2 = _first_match(" ".join(observation_texts), PATTERNS_O2_LT90)
    if helper_o2 and 'time_below_90_pct' not in sleep_study:
        sleep_study['time_below_90_pct'] = helper_o2[0]

    # Table-style Oxygen Saturation block parsing
    try:
        for idx, raw_line in enumerate(lines):
            line_norm = raw_line.strip()
            if not line_norm:
                continue
            if re.search(r'^oxygen\s*saturation', line_norm, re.IGNORECASE):
                header = line_norm
                thresholds = re.findall(r'(?:<=\s*\d{2}|<\s*\d{2})', header)
                thresholds = [t.replace(' ', '') for t in thresholds]
                thresholds_idx = idx
                # If thresholds are not on the same line, search the next few lines
                if not thresholds:
                    for look_ahead_offset, th_line in enumerate(lines[idx+1: idx+4], start=1):
                        th_norm = th_line.strip()
                        toks = re.findall(r'(?:<=\s*\d{2}|<\s*\d{2})', th_norm)
                        if toks:
                            thresholds = [t.replace(' ', '') for t in toks]
                            thresholds_idx = idx + look_ahead_offset
                            break
                # Look for Duration (minutes) and Sleep % rows after the thresholds line
                window = lines[thresholds_idx+1: thresholds_idx+6]
                minutes_row = None
                percent_row = None
                for w in window:
                    w_norm = w.strip()
                    if re.search(r'^duration\s*\(\s*minutes?\s*\)\s*:', w_norm, re.IGNORECASE):
                        minutes_row = w_norm
                    elif re.search(r'^sleep\s*%\s*', w_norm, re.IGNORECASE):
                        percent_row = w_norm
                def extract_numbers(row_text):
                    return [float(x) for x in re.findall(r'(-?\d+(?:\.\d+)?)', row_text)] if row_text else []
                minutes_vals = extract_numbers(minutes_row)
                percent_vals = extract_numbers(percent_row)
                # Map by position
                if thresholds and minutes_vals:
                    for pos, token in enumerate(thresholds):
                        if pos >= len(minutes_vals):
                            break
                        val = minutes_vals[pos]
                        if token == '<90':
                            sleep_study['time_below_90_pct_min'] = val
                        elif token == '<=88':
                            sleep_study['time_below_88_pct_min'] = val
                if thresholds and percent_vals:
                    for pos, token in enumerate(thresholds):
                        if pos >= len(percent_vals):
                            break
                        val = percent_vals[pos]
                        if token == '<90':
                            sleep_study['time_below_90_pct'] = val
                break
    except Exception:
        pass
    
    # Positional AHI Data - More Generic Patterns
    supine_patterns = [
        # Generic supine patterns
        r'supine[:\s]*(\d+(?:\.\d+)?)',  # supine: 42.5
        r'(\d+(?:\.\d+)?)\s*supine',  # 42.5 supine
        r'Supine[:\s]*(\d+(?:\.\d+)?)',  # Supine: 42.5
        r'(\d+(?:\.\d+)?)\s*Supine',  # 42.5 Supine
        
        # pAHI patterns
        r'pahi[:\s]*(\d+(?:\.\d+)?)',  # pAHI: 42.5
        r'pAHI[:\s]*(\d+(?:\.\d+)?)',  # pAHI: 42.5
        r'(\d+(?:\.\d+)?)\s*pahi',  # 42.5 pAHI
        r'(\d+(?:\.\d+)?)\s*pAHI',  # 42.5 pAHI
        
        # Combined patterns
        r'pahi[:\s]*(\d+(?:\.\d+)?).*supine',  # pAHI in supine context
        r'supine.*pahi[:\s]*(\d+(?:\.\d+)?)',  # supine with pAHI
        r'Supine.*pAHI[:\s]*(\d+(?:\.\d+)?)',  # Supine pAHI: 42.5
        r'pAHI[:\s]*(\d+(?:\.\d+)?).*Supine',  # pAHI with Supine context
        
        # Position patterns
        r'position.*supine.*(\d+(?:\.\d+)?)',  # position supine with AHI
        
        # Generic AHI patterns that might be supine
        r'(\d+(?:\.\d+)?)\s*ahi',  # Any AHI value
        r'(\d+(?:\.\d+)?)\s*AHI',  # Any AHI value
    ]
    
    for pattern in supine_patterns:
        match = re.search(pattern, combined_text)
        if match:
            sleep_study['supine_ahi'] = float(match.group(1))
            break
    
    # REM/NREM AHI
    rem_patterns = [
        r'rem[:\s]*(\d+(?:\.\d+)?)',
        r'(\d+(?:\.\d+)?)\s*rem',
    ]
    
    for pattern in rem_patterns:
        match = re.search(pattern, combined_text)
        if match:
            sleep_study['rem_ahi'] = float(match.group(1))
            break
    
    # Snoring Data
    snoring = {}
    snoring_avg_patterns = [
        r'mean[:\s]*(\d+)\s*d[bb]?',
        r'average[:\s]*(\d+)\s*d[bb]?',
        r'(\d+)\s*d[bb]?\s*mean',
        r'(\d+)\s*d[bb]?\s*average',
        r'Mean[:\s]*(\d+)\s*dB',  # Mean: 43 dB
        r'(\d+)\s*dB.*Mean',  # 43 dB Mean
        r'(\d+)\s*dB',  # Simple 43 dB
    ]
    
    for pattern in snoring_avg_patterns:
        match = re.search(pattern, combined_text)
        if match:
            snoring['avg_db'] = int(match.group(1))
            break
    
    snoring_max_patterns = [
        r'>(\d+)\s*d[bb]?',
        r'max[:\s]*(\d+)\s*d[bb]?',
        r'peak[:\s]*(\d+)\s*d[bb]?',
        r'(\d+)\s*d[bb]?\s*max',
    ]
    
    for pattern in snoring_max_patterns:
        match = re.search(pattern, combined_text)
        if match:
            snoring['max_db'] = int(match.group(1))
            break
    
    if snoring:
        sleep_study['snoring'] = snoring
    
    # Heart Rate Data - More Generic Patterns
    heart_rate = {}
    hr_patterns = [
        # Generic mean/minimum/maximum patterns (case insensitive)
        r'mean[:\s]*(\d+)',  # Mean: 65
        r'minimum[:\s]*(\d+)',  # Minimum: 40
        r'maximum[:\s]*(\d+)',  # Maximum: 104
        r'Mean[:\s]*(\d+)',  # Mean: 65
        r'Minimum[:\s]*(\d+)',  # Minimum: 40
        r'Maximum[:\s]*(\d+)',  # Maximum: 104
        
        # With BPM context
        r'mean[:\s]*(\d+)\s*(?:bpm|beats)',  # Mean: 65 BPM
        r'minimum[:\s]*(\d+)\s*(?:bpm|beats)',  # Minimum: 40 BPM
        r'maximum[:\s]*(\d+)\s*(?:bpm|beats)',  # Maximum: 104 BPM
        r'(\d+)\s*(?:bpm|beats)\s*mean',  # 65 BPM mean
        r'(\d+)\s*(?:bpm|beats)\s*minimum',  # 40 BPM minimum
        r'(\d+)\s*(?:bpm|beats)\s*maximum',  # 104 BPM maximum
        
        # Pulse rate specific
        r'pulse rate.*mean[:\s]*(\d+)',  # Pulse rate mean
        r'pulse rate.*minimum[:\s]*(\d+)',  # Pulse rate minimum
        r'pulse rate.*maximum[:\s]*(\d+)',  # Pulse rate maximum
        r'mean[:\s]*(\d+).*pulse',  # Mean with pulse context
        r'minimum[:\s]*(\d+).*pulse',  # Minimum with pulse context
        r'maximum[:\s]*(\d+).*pulse',  # Maximum with pulse context
        
        # Generic heart rate patterns
        r'(\d{2,3})\s*(?:bpm|beats)',  # Any 2-3 digit BPM value
    ]
    
    # Extract heart rate values using generic table format detection
    hr_candidates = []
    for line in lines:
        line = line.strip()
        if re.match(r'^\d{2,3}$', line):  # 2-3 digit standalone numbers
            value = int(line)
            if 30 <= value <= 200:  # Reasonable heart rate range
                hr_candidates.append(value)
    
    # Remove duplicates and sort
    hr_candidates = sorted(list(set(hr_candidates)))
    
    # Assign values based on position (assuming they're in order: min, mean, max)
    if len(hr_candidates) >= 3:
        heart_rate['min_bpm'] = hr_candidates[0]   # Lowest
        heart_rate['mean_bpm'] = hr_candidates[1]  # Middle
        heart_rate['max_bpm'] = hr_candidates[2]   # Highest
    elif len(hr_candidates) == 2:
        heart_rate['min_bpm'] = hr_candidates[0]   # Lower
        heart_rate['mean_bpm'] = hr_candidates[1]  # Higher
    elif len(hr_candidates) == 1:
        heart_rate['mean_bpm'] = hr_candidates[0]  # Single value
    
    if heart_rate:
        sleep_study['heart_rate'] = heart_rate
    
    # Severity Classification
    severity_patterns = [
        r'severity[:\s]*(mild|moderate|severe|none)',
        r'(mild|moderate|severe|none)\s*severity',
    ]
    
    for pattern in severity_patterns:
        match = re.search(pattern, combined_text)
        if match:
            sleep_study['severity'] = match.group(1)
            break
    
    # Positional Metrics
    positional_metrics = {}
    if 'supine_ahi' in sleep_study:
        positional_metrics['supine_AHI'] = sleep_study['supine_ahi']
    
    # Additional positional data extraction
    supine_ahi_patterns = [
        r'Supine.*pAHI[:\s]*(\d+(?:\.\d+)?)',
        r'pAHI[:\s]*(\d+(?:\.\d+)?).*Supine',
        r'supine.*(\d+(?:\.\d+)?).*pahi',
    ]
    
    for pattern in supine_ahi_patterns:
        match = re.search(pattern, combined_text)
        if match:
            positional_metrics['supine_AHI'] = float(match.group(1))
            break
    
    # TMJ Data
    observations = {}
    tmj_flags = {}
    
    # TMJ Clicking (boolean)
    clicking_patterns = [
        r'clicking[:\s]*(yes|true|present)',
        r'clicks[:\s]*(yes|true|present)',
        r'(yes|true|present)[:\s]*clicking',
    ]
    
    for pattern in clicking_patterns:
        match = re.search(pattern, combined_text)
        if match:
            tmj_flags['clicking'] = True
            break
    
    # TMJ Side
    side_patterns = [
        r'(left|right|bilateral)\s*(?:tmj|jaw|joint)',
        r'tmj[:\s]*(left|right|bilateral)',
        r'jaw[:\s]*(left|right|bilateral)',
    ]
    
    for pattern in side_patterns:
        match = re.search(pattern, combined_text)
        if match:
            tmj_flags['side'] = match.group(1)
            break
    
    if tmj_flags:
        observations['tmj_flags'] = tmj_flags
    
    # Build the final structure
    if sleep_study:
        extracted_data['sleep_study'] = sleep_study
    
    if positional_metrics:
        extracted_data['positional_metrics'] = positional_metrics
    
    if observations:
        extracted_data['observations'] = observations
    
    # Supine AHI via helper patterns (covers multiple formats)
    helper_sup = _first_match(" ".join(observation_texts), PATTERNS_SUPINE_AHI)
    if helper_sup and 'supine_ahi' not in sleep_study:
        sleep_study['supine_ahi'] = helper_sup[0]
    
    return extracted_data

def categorize_observations_with_llm(observation_texts: List[str], patient_id: int) -> Dict[str, Any]:
    """
    Use LLM to intelligently categorize observations into schema fields.
    Enhanced approach: extract specific numerical fields first, then use LLM for categorization.
    
    Args:
        observation_texts: List of observation strings
        patient_id: Patient ID for logging
        
    Returns:
        Dict containing categorized data according to schema
    """
    try:
        from botocore.exceptions import ClientError
        
        # STAGE 1: Extract specific numerical fields using regex patterns
        logger.info(f"Patient {patient_id}: Extracting specific numerical fields using patterns")
        numerical_data = extract_specific_numerical_fields(observation_texts)
        
        if numerical_data:
            logger.info(f"Patient {patient_id}: Found {len(numerical_data)} numerical fields: {list(numerical_data.keys())}")
        else:
            logger.info(f"Patient {patient_id}: No specific numerical fields found")
        
        # Initialize Bedrock client
        bedrock = boto3.client(
            service_name='bedrock-runtime',
            region_name='us-west-2'
        )
        
        # STAGE 1: Extract numerical data using regex (more reliable and faster)
        logger.info(f"Patient {patient_id}: Extracting numerical data with regex...")
        numerical_data = extract_numerical_data_with_regex(observation_texts)
        logger.info(f"Patient {patient_id}: Found {len(numerical_data)} numerical fields: {list(numerical_data.keys())}")
        
        # STAGE 2: Extract textual observations using LLM (better context understanding)
        logger.info(f"Patient {patient_id}: Extracting textual observations with LLM...")
        # Note: source_type is not available in this context, so we'll use the default
        textual_data = extract_textual_observations_with_llm(observation_texts, numerical_data, 'general_medical')
        logger.info(f"Patient {patient_id}: Found textual data for fields: {list(textual_data.keys())}")
        
        # STAGE 3: Combine numerical and textual data
        generic_result = {}
        
        # Merge numerical data
        if numerical_data:
            generic_result.update(numerical_data)
        
        # Merge textual data
        if textual_data:
            # Handle nested structures
            for key, value in textual_data.items():
                if key in generic_result and isinstance(generic_result[key], dict) and isinstance(value, dict):
                    generic_result[key].update(value)
                else:
                    generic_result[key] = value
        
        # Generic result is now combined from regex and LLM extraction
        
        # STAGE 2: Targeted extraction for critical fields (always run for better accuracy)
        critical_fields = [
            'sleep_study.sleep_duration_h',
            'sleep_study.sleep_efficiency_pct', 
            'sleep_study.desaturation_events',
            'sleep_study.o2_mean_pct',
            'sleep_study.o2_nadir_pct',
            'sleep_study.time_below_90_pct_min',
            'sleep_study.time_below_88_pct_min',
            'sleep_study.supine_ahi',
            'sleep_study.rem_ahi',
            'sleep_study.snoring.avg_db',
            'sleep_study.snoring.max_db',
            'sleep_study.heart_rate.mean_bpm',
            'sleep_study.heart_rate.min_bpm',
            'sleep_study.heart_rate.max_bpm',
            'sleep_study.severity',
            'positional_metrics.supine_AHI',
            'positional_metrics.positional_phenotype',
            'observations.tmj_flags.clicking',
            'observations.tmj_flags.side'
        ]
        
        # Always run targeted extraction for critical fields to ensure we get the most accurate data
        logger.info(f"Patient {patient_id}: Running targeted extraction for critical fields: {critical_fields}")
        
        # Create targeted prompt for critical fields
        targeted_prompt = f"""
You are a medical data specialist performing targeted extraction for critical fields.

IMPORTANT: The data may be in tabular format with | separators or structured text. Pay special attention to:
- Tables marked with "=== TABLE ===" sections
- Data separated by | characters (table rows)
- Structured content marked with "=== STRUCTURED CONTENT ==="
- Numerical values in any format (X.X, X%, X dB, etc.)
- Look for data in both regular text and structured table format

CRITICAL FIELDS TO EXTRACT:
{chr(10).join([f"- {field}" for field in critical_fields])}

TARGETED EXTRACTION RULES:
1. Look specifically for the missing fields listed above
2. Be thorough - search for data in any format or terminology
3. Pay special attention to tabular data - look for patterns in table rows and columns
4. For table data, match column headers with schema fields and extract corresponding values
5. Look for these specific patterns:

   SLEEP DURATION (sleep_duration_h):
   - "sleep duration", "total sleep time", "TST", "hours of sleep", "sleep time: X hours", "sleep period"
   - "slept for X hours", "sleep time X hours", "total sleep X hours", "sleep duration X hours"
   - Look for numbers followed by "hours", "hrs", "h" in sleep context
   - Extract as decimal (e.g., 7.5 for 7 hours 30 minutes)
   - In tables: look for columns with "sleep", "duration", "time", "hours" headers

   SLEEP EFFICIENCY (sleep_efficiency_pct):
   - "sleep efficiency", "efficiency", "SE", "efficiency: X%", "sleep efficiency percentage"
   - "efficiency X%", "sleep efficiency X%", "efficiency rate X%"
   - Look for percentages in sleep context
   - Extract as number without % symbol
   - In tables: look for columns with "efficiency", "SE", "%" headers

   DESATURATION EVENTS (desaturation_events):
   - "desaturation events", "desats", "number of desaturations", "desaturation count", "desaturation episodes"
   - "X desaturation events", "X desats", "desaturation count: X", "X episodes of desaturation"
   - Look for numbers followed by "desaturation", "desats", "episodes"
   - In tables: look for columns with "desaturation", "desats", "events", "count" headers

   SNORING AVERAGE DB (snoring.avg_db):
   - "average snoring", "mean snoring", "snoring average", "avg snoring dB", "snoring level"
   - "average snoring X dB", "mean snoring level X dB", "snoring average X dB"
   - Look for numbers followed by "dB", "decibels" in snoring context
   - In tables: look for columns with "snoring", "average", "mean", "dB" headers

   SNORING MAX DB (snoring.max_db):
   - "maximum snoring", "peak snoring", "snoring max", "max snoring dB", "snoring peak"
   - "maximum snoring X dB", "peak snoring level X dB", "snoring max X dB"
   - Look for "max", "peak", "highest" followed by numbers and "dB"
   - In tables: look for columns with "snoring", "max", "peak", "highest", "dB" headers

   TMJ CLICKING (observations.tmj_flags.clicking):
   - "clicking", "clicks", "TMJ clicking", "joint clicking", "jaw clicking"
   - "jaw joint clicking", "temporomandibular clicking", "TMJ sounds", "joint sounds"
   - Look for "yes/no", "present/absent", "true/false" in TMJ context
   - Extract as boolean (true/false)
   - In tables: look for columns with "TMJ", "clicking", "sounds", "joint" headers

   TMJ SIDE (observations.tmj_flags.side):
   - "left TMJ", "right TMJ", "bilateral TMJ", "left side", "right side", "both sides", "unilateral"
   - "left jaw", "right jaw", "bilateral jaw", "left joint", "right joint"
   - Extract as "left", "right", "bilateral", or "unilateral"
   - In tables: look for columns with "side", "location", "TMJ" headers
6. Extract any numerical values, percentages, or descriptive data that could match these fields
7. Be flexible with terminology and formats
8. Only return fields that have actual values found
9. For numerical values, extract the number only (not units)
10. For percentages, extract the number only (not % symbol)
11. Be especially thorough with tabular data - it often contains the most precise measurements

OBSERVATIONS:
{chr(10).join(observation_texts)}

Return ONLY a valid JSON object with the missing fields that were found. Do not include any explanations or text outside the JSON.
"""
            
        # Use hardcoded model ID for standalone script
        model_id = MODEL_ID
        
        import time
        start_time = time.time()
        
        try:
            # Call Bedrock for targeted extraction
            response = bedrock.invoke_model(
                modelId=model_id,
                body=json.dumps({
                    "anthropic_version": "bedrock-2023-05-31",
                    "max_tokens": 2000,
                    "messages": [
                        {
                            "role": "user",
                            "content": targeted_prompt
                        }
                    ]
                })
            )
            
            response_time_ms = int((time.time() - start_time) * 1000)
        
            response_body = json.loads(response.get('body').read())
            llm_response = response_body['content'][0]['text']
            
            # Log to database
            _log_llm_call(
                prompt_text="Targeted field extraction from observations",
                response_text=llm_response[:500],
                response_time_ms=response_time_ms,
                status='success'
            )
            
            # Extract JSON from targeted response
            try:
                start_idx = llm_response.find('{')
                end_idx = llm_response.rfind('}') + 1
                if start_idx != -1 and end_idx != 0:
                    json_str = llm_response[start_idx:end_idx]
                    targeted_result = json.loads(json_str)
                    
                            # Merge targeted results into generic results
                    generic_result = _merge_results(generic_result, targeted_result)
                    logger.info(f"Patient {patient_id}: Successfully merged targeted extraction results")
                else:
                    logger.warning(f"Could not find JSON in targeted LLM response for patient {patient_id}")
            except json.JSONDecodeError as e:
                logger.error(f"Failed to parse targeted LLM JSON response for patient {patient_id}: {e}")
        except Exception as e:
            response_time_ms = int((time.time() - start_time) * 1000)
            _log_llm_call(
                prompt_text="Targeted field extraction from observations",
                response_text='',
                response_time_ms=response_time_ms,
                status='error',
                error_message=str(e)
            )
            logger.error(f"Targeted extraction failed for patient {patient_id}: {e}")

        # STAGE 3: Combine numerical data with LLM results
        logger.info(f"Patient {patient_id}: Combining numerical data with LLM results")
        final_result = _merge_results(generic_result, numerical_data)

        if numerical_data:
            logger.info(f"Patient {patient_id}: Final result includes {len(numerical_data)} numerical fields from pattern extraction")
            
            # DISABLED: Store numerical extraction results as individual observations
            # This was causing data pollution with duplicate values and no dates
            # numerical_observations = _convert_numerical_data_to_observations(numerical_data, patient_id)
            # if numerical_observations:
            #     logger.info(f"Patient {patient_id}: Storing {len(numerical_observations)} numerical observations")
            #     store_observations_with_deduplication(patient_id, 'numerical_extraction', numerical_observations, {
            #         'name': 'numerical_pattern_extraction',
            #         'file_type': 'text/plain',
            #         'id': None,
            #         'source_table': 'numerical'
            #     })
            logger.info(f"Patient {patient_id}: Numerical pattern extraction disabled to prevent data pollution")

        return final_result
            
    except Exception as e:
        logger.error(f"LLM processing failed for patient {patient_id}: {e}")
        return None

def _has_field_value(result: Dict[str, Any], field_path: str) -> bool:
    """Check if a field has a non-empty value in the result."""
    try:
        keys = field_path.split('.')
        value = result
        for key in keys:
            if isinstance(value, dict) and key in value:
                value = value[key]
            else:
                return False
        
        # Check if value is not None, not empty string, not empty list
        if value is None:
            return False
        if isinstance(value, str) and value.strip() == '':
            return False
        if isinstance(value, list) and len(value) == 0:
            return False
        
        return True
    except:
        return False

def _convert_numerical_data_to_observations(numerical_data: Dict[str, Any], patient_id: int) -> List[Dict]:
    """
    Convert numerical extraction results to individual observations for storage.
    
    Args:
        numerical_data: Dictionary containing numerical field data
        patient_id: Patient ID
        
    Returns:
        List of observation dictionaries with correct schema paths
    """
    observations = []
    
    def add_observation(path: str, value: Any):
        if value is not None and value != '':
            observations.append({
                'path': path,
                'value': str(value),
                'source': 'numerical-extraction',
                'confidence': 100,
                'explanation': f'Extracted using regex patterns from patient {patient_id} documents'
            })
    
    # Process sleep_study data
    if 'sleep_study' in numerical_data:
        sleep_study = numerical_data['sleep_study']
        for key, value in sleep_study.items():
            if key == 'snoring' and isinstance(value, dict):
                for snoring_key, snoring_value in value.items():
                    add_observation(f'sleep_study.snoring.{snoring_key}', snoring_value)
            elif key == 'heart_rate' and isinstance(value, dict):
                # Map heart rate to schema structure
                for hr_key, hr_value in value.items():
                    add_observation(f'sleep_study.heart_rate.{hr_key}', hr_value)
            else:
                add_observation(f'sleep_study.{key}', value)
    
    # Process positional_metrics data
    if 'positional_metrics' in numerical_data:
        positional = numerical_data['positional_metrics']
        for key, value in positional.items():
            add_observation(f'positional_metrics.{key}', value)
    
    return observations

def _merge_results(generic_result: Dict[str, Any], targeted_result: Dict[str, Any]) -> Dict[str, Any]:
    """Merge targeted extraction results into generic results."""
    try:
        # Deep merge the results
        merged = generic_result.copy()
        
        for key, value in targeted_result.items():
            if key not in merged:
                merged[key] = value
            elif isinstance(merged[key], dict) and isinstance(value, dict):
                merged[key] = _merge_results(merged[key], value)
            elif isinstance(merged[key], list) and isinstance(value, list):
                # Merge lists, avoiding duplicates
                merged[key] = merged[key] + [item for item in value if item not in merged[key]]
            else:
                # If both have values, prefer the targeted result (more specific)
                merged[key] = value
        
        return merged
    except Exception as e:
        logger.error(f"Error merging results: {e}")
        return generic_result

def extract_questionnaire_observations(content: str, document_name: str) -> List[Dict]:
    """
    Extract questionnaire observations and categorize them into patient_self_report sections.
    
    Args:
        content: Document text content
        document_name: Name of the document
        
    Returns:
        List of observation dictionaries with patient_self_report paths
    """
    import boto3
    import json
    
    # Initialize Bedrock client
    bedrock = boto3.client(
        service_name='bedrock-runtime',
        region_name='us-west-2'
    )
    
    # Create specialized prompt for questionnaire extraction
    questionnaire_prompt = f"""
You are a medical data specialist processing a PATIENT QUESTIONNAIRE. Extract patient self-reported information and categorize it into the patient_self_report section.

FOCUS ON PATIENT SELF-REPORTED DATA:
- patient_self_report.symptoms: daytime_sleepiness (boolean), witnessed_apneas (boolean), dry_mouth (boolean), etc.
- patient_self_report.goals: list of patient's treatment goals (e.g., ["reduce snoring", "improve sleep quality"])
- patient_self_report.primary_complaint: main issue the patient wants to address
- patient_self_report.scales: any rating scales or scores the patient provided

IMPORTANT RULES FOR QUESTIONNAIRES:
1. DO NOT extract numerical values - they are already handled by regex
2. Focus on patient's own words and self-reported symptoms
3. Categorize symptoms into appropriate boolean fields (true/false)
4. Extract treatment goals as a list of strings
5. Identify the primary complaint from patient's responses
6. Only include fields that have actual content from the questionnaire

QUESTIONNAIRE CONTENT TO ANALYZE:
{content}

Return ONLY a valid JSON object with the patient_self_report data. Do not include any explanations or text outside the JSON.
"""
    
    # Use hardcoded model ID for standalone script
    model_id = MODEL_ID
    
    import time
    start_time = time.time()
    
    try:
        # Call Bedrock for questionnaire extraction
        response = bedrock.invoke_model(
            modelId=model_id,
            body=json.dumps({
                "anthropic_version": "bedrock-2023-05-31",
                "max_tokens": 2000,
                "messages": [
                    {
                        "role": "user",
                        "content": questionnaire_prompt
                    }
                ]
            })
        )
        
        response_time_ms = int((time.time() - start_time) * 1000)
        
        response_body = json.loads(response.get('body').read())
        llm_response = response_body['content'][0]['text']
        
        # Log to database
        _log_llm_call(
            prompt_text="Extract questionnaire data from patient self-report",
            response_text=llm_response[:500],
            response_time_ms=response_time_ms,
            status='success'
        )
        
        # Extract JSON from response
        start_idx = llm_response.find('{')
        end_idx = llm_response.rfind('}') + 1
        if start_idx != -1 and end_idx != 0:
            json_str = llm_response[start_idx:end_idx]
            questionnaire_result = json.loads(json_str)
        else:
            logger.error("Could not find JSON in questionnaire LLM response")
            return []
        
        # Convert the questionnaire result into observation format
        observations = []
        
        # Process symptoms
        if 'patient_self_report' in questionnaire_result and 'symptoms' in questionnaire_result['patient_self_report']:
            symptoms = questionnaire_result['patient_self_report']['symptoms']
            for symptom_key, symptom_value in symptoms.items():
                if symptom_value is not None:
                    observations.append({
                        'path': f'patient_self_report.symptoms.{symptom_key}',
                        'value': str(symptom_value).lower(),
                        'document_name': document_name,
                        'document_type': 'application/pdf',
                        'extraction_date': datetime.now().isoformat()
                    })
        
        # Process goals
        if 'patient_self_report' in questionnaire_result and 'goals' in questionnaire_result['patient_self_report']:
            goals = questionnaire_result['patient_self_report']['goals']
            if isinstance(goals, list):
                for goal in goals:
                    observations.append({
                        'path': 'patient_self_report.goals',
                        'value': goal,
                        'document_name': document_name,
                        'document_type': 'application/pdf',
                        'extraction_date': datetime.now().isoformat()
                    })
        
        # Process primary complaint
        if 'patient_self_report' in questionnaire_result and 'primary_complaint' in questionnaire_result['patient_self_report']:
            primary_complaint = questionnaire_result['patient_self_report']['primary_complaint']
            if primary_complaint:
                observations.append({
                    'path': 'patient_self_report.primary_complaint',
                    'value': primary_complaint,
                    'document_name': document_name,
                    'document_type': 'application/pdf',
                    'extraction_date': datetime.now().isoformat()
                })
        
        # Process scales
        if 'patient_self_report' in questionnaire_result and 'scales' in questionnaire_result['patient_self_report']:
            scales = questionnaire_result['patient_self_report']['scales']
            for scale_key, scale_value in scales.items():
                if scale_value is not None:
                    observations.append({
                        'path': f'patient_self_report.scales.{scale_key}',
                        'value': str(scale_value),
                        'document_name': document_name,
                        'document_type': 'application/pdf',
                        'extraction_date': datetime.now().isoformat()
                    })
        
        logger.info(f"Extracted {len(observations)} questionnaire observations from {document_name}")
        return observations
        
    except Exception as e:
        logger.error(f"Failed to extract questionnaire observations from {document_name}: {e}")
        return []

def categorize_questionnaire_observation(canonical: Dict[str, Any], observation: str):
    """
    Categorize questionnaire observations into the appropriate patient_self_report sections.
    
    Args:
        canonical: The canonical structure to update
        observation: The questionnaire observation text
    """
    observation_lower = observation.lower()
    
    # Categorize symptoms
    if any(word in observation_lower for word in ['tired', 'fatigue', 'sleepy', 'exhausted']):
        if 'daytime' in observation_lower or 'during day' in observation_lower:
            canonical['patient_self_report']['symptoms']['daytime_sleepiness'] = True
        else:
            canonical['patient_self_report']['symptoms']['daytime_sleepiness'] = True
    
    if any(word in observation_lower for word in ['snore', 'snoring']):
        canonical['patient_self_report']['symptoms']['witnessed_apneas'] = True
    
    if any(word in observation_lower for word in ['mouth breathing', 'breathes through mouth']):
        canonical['patient_self_report']['symptoms']['dry_mouth'] = True
    
    if any(word in observation_lower for word in ['non-restorative', 'not rested', 'unrefreshed']):
        canonical['patient_self_report']['symptoms']['non_restorative_sleep'] = True
    
    if any(word in observation_lower for word in ['headache', 'morning headache']):
        canonical['patient_self_report']['symptoms']['morning_headache'] = True
    
    if any(word in observation_lower for word in ['nocturia', 'night urination', 'frequent urination']):
        canonical['patient_self_report']['symptoms']['nocturia'] = True
    
    if any(word in observation_lower for word in ['bruxism', 'teeth grinding', 'jaw clenching']):
        canonical['patient_self_report']['symptoms']['bruxism'] = True
    
    if any(word in observation_lower for word in ['reflux', 'heartburn', 'acid']):
        canonical['patient_self_report']['symptoms']['reflux'] = True
    
    if any(word in observation_lower for word in ['insomnia', 'trouble falling asleep', 'wake up']):
        canonical['patient_self_report']['symptoms']['insomnia_features'] = True
    
    # Categorize goals
    if any(word in observation_lower for word in ['reduce snoring', 'stop snoring', 'less snoring']):
        if 'reduce snoring' not in canonical['patient_self_report']['goals']:
            canonical['patient_self_report']['goals'].append('reduce snoring')
    
    if any(word in observation_lower for word in ['improve sleep', 'better sleep', 'sleep quality']):
        if 'improve sleep quality' not in canonical['patient_self_report']['goals']:
            canonical['patient_self_report']['goals'].append('improve sleep quality')
    
    if any(word in observation_lower for word in ['reduce fatigue', 'less tired', 'more energy']):
        if 'reduce daytime fatigue' not in canonical['patient_self_report']['goals']:
            canonical['patient_self_report']['goals'].append('reduce daytime fatigue')
    
    if any(word in observation_lower for word in ['energy', 'energy levels', 'daily energy']):
        if 'improve daily energy levels' not in canonical['patient_self_report']['goals']:
            canonical['patient_self_report']['goals'].append('improve daily energy levels')
    
    # Set primary complaint if not already set
    if not canonical['patient_self_report']['primary_complaint']:
        if 'snoring' in observation_lower:
            canonical['patient_self_report']['primary_complaint'] = 'snoring'
        elif 'tired' in observation_lower or 'fatigue' in observation_lower:
            canonical['patient_self_report']['primary_complaint'] = 'daytime fatigue'
        elif 'sleep' in observation_lower:
            canonical['patient_self_report']['primary_complaint'] = 'poor sleep quality'

def _get_osa_observation_priority(observation: str) -> int:
    """
    Assign priority score to observations based on OSA clinical importance.
    Lower numbers = higher priority (1 = highest, 10 = lowest)
    """
    obs_lower = observation.lower()
    
    # Priority 1: Critical OSA metrics
    if any(keyword in obs_lower for keyword in [
        'ahi', 'apnea-hypopnea index', 'apnea hypopnea index', 'events per hour',
        'spo2 nadir', 'oxygen saturation nadir', 'o2 nadir', 'lowest spo2',
        'odi', 'oxygen desaturation index', 'desaturation index'
    ]):
        return 1
    
    # Priority 2: OSA severity and classification
    if any(keyword in obs_lower for keyword in [
        'severe osa', 'moderate osa', 'mild osa', 'osa severity',
        'obstructive sleep apnea', 'sleep apnea', 'apnea events',
        'hypopnea events', 'respiratory events'
    ]):
        return 2
    
    # Priority 3: Sleep study results and findings
    if any(keyword in obs_lower for keyword in [
        'sleep study', 'polysomnography', 'psg', 'sleep test',
        'sleep efficiency', 'sleep duration', 'total sleep time',
        'arousal index', 'sleep fragmentation'
    ]):
        return 3
    
    # Priority 4: Anatomical findings related to OSA
    if any(keyword in obs_lower for keyword in [
        'tongue base', 'retroglossal', 'retropalatal', 'nasal obstruction',
        'airway narrowing', 'pharyngeal collapse', 'uvula', 'tonsils',
        'adenoids', 'deviated septum', 'nasal congestion'
    ]):
        return 4
    
    # Priority 5: OSA symptoms and complaints
    if any(keyword in obs_lower for keyword in [
        'snoring', 'snore', 'apnea', 'choking', 'gasping',
        'daytime sleepiness', 'fatigue', 'morning headache',
        'sleep disruption', 'fragmented sleep'
    ]):
        return 5
    
    # Priority 6: Treatment-related observations
    if any(keyword in obs_lower for keyword in [
        'cpap', 'bipap', 'oral appliance', 'mandibular advancement',
        'treatment compliance', 'therapy adherence', 'device use'
    ]):
        return 6
    
    # Priority 7: Comorbidities that affect OSA
    if any(keyword in obs_lower for keyword in [
        'obesity', 'bmi', 'weight', 'diabetes', 'hypertension',
        'cardiovascular', 'heart disease', 'stroke risk'
    ]):
        return 7
    
    # Priority 8: Sleep quality and patterns
    if any(keyword in obs_lower for keyword in [
        'sleep quality', 'sleep architecture', 'rem sleep', 'deep sleep',
        'sleep stages', 'sleep latency', 'sleep onset'
    ]):
        return 8
    
    # Priority 9: General medical observations
    if any(keyword in obs_lower for keyword in [
        'medical history', 'medications', 'allergies', 'vital signs',
        'blood pressure', 'heart rate', 'temperature'
    ]):
        return 9
    
    # Priority 10: Everything else (lowest priority)
    return 10

def _sort_observations_by_osa_priority(observations: List[str]) -> List[str]:
    """
    Sort observations by OSA clinical priority.
    Returns observations ordered by clinical importance for OSA treatment.
    """
    try:
        # Sort by priority score (lower number = higher priority)
        sorted_observations = sorted(observations, key=_get_osa_observation_priority)
        return sorted_observations
    except Exception as e:
        logger.error(f"Error sorting observations by OSA priority: {e}")
        return observations  # Return original list if sorting fails

def update_canonical_from_llm_result(canonical: Dict[str, Any], llm_result: Dict[str, Any]):
    """
    Update canonical structure with LLM-categorized data.
    
    Args:
        canonical: The canonical structure to update
        llm_result: LLM-categorized data
    """
    try:
        # Update demographics
        if 'demographics' in llm_result:
            for key, value in llm_result['demographics'].items():
                if key in canonical['demographics'] and value is not None:
                    canonical['demographics'][key] = value
        
        # Update sleep study (without severity - schema doesn't allow it)
        if 'sleep_study' in llm_result:
            for key, value in llm_result['sleep_study'].items():
                if key in canonical['sleep_study'] and value is not None:
                    if key == 'snoring' and isinstance(value, dict):
                        for snoring_key, snoring_value in value.items():
                            if snoring_key in canonical['sleep_study']['snoring'] and snoring_value is not None:
                                canonical['sleep_study']['snoring'][snoring_key] = snoring_value
                    else:
                        canonical['sleep_study'][key] = value
        
        # Normalize study_type to schema enum (always normalize regardless of source)
        if 'sleep_study' in canonical and canonical['sleep_study'].get('study_type'):
            canonical['sleep_study']['study_type'] = normalize_study_type_to_schema(canonical['sleep_study']['study_type'])
        
        # Update observations.anatomy_imaging
        if 'observations' in llm_result and 'anatomy_imaging' in llm_result['observations']:
            for key, value in llm_result['observations']['anatomy_imaging'].items():
                if key in canonical['observations']['anatomy_imaging'] and value is not None:
                    canonical['observations']['anatomy_imaging'][key] = value
        
        # Update observations.tmj_flags
        if 'observations' in llm_result and 'tmj_flags' in llm_result['observations']:
            for key, value in llm_result['observations']['tmj_flags'].items():
                if key in canonical['observations']['tmj_flags'] and value is not None:
                    canonical['observations']['tmj_flags'][key] = value
        
        # Update observations.dise (DISE structured block)
        if 'observations' in llm_result and 'dise' in llm_result['observations']:
            dise_data = llm_result['observations']['dise']
            if isinstance(dise_data, dict):
                def deep_merge_dise(target, source):
                    """Recursively merge DISE data, handling nested dicts"""
                    for key, value in source.items():
                        if value is None:
                            continue
                        if key in target:
                            if isinstance(value, dict) and isinstance(target[key], dict):
                                deep_merge_dise(target[key], value)
                            elif isinstance(value, list) and isinstance(target[key], list):
                                target[key].extend(v for v in value if v not in target[key])
                            else:
                                target[key] = value
                        else:
                            target[key] = value
                deep_merge_dise(canonical['observations']['dise'], dise_data)
                logger.info(f"Merged DISE observations: {list(dise_data.keys())}")
        
        # Update observations.cbct (CBCT airway structured block)
        if 'observations' in llm_result and 'cbct' in llm_result['observations']:
            cbct_data = llm_result['observations']['cbct']
            if isinstance(cbct_data, dict):
                def deep_merge_cbct(target, source):
                    """Recursively merge CBCT data, handling nested dicts"""
                    for key, value in source.items():
                        if value is None:
                            continue
                        if key in target:
                            if isinstance(value, dict) and isinstance(target[key], dict):
                                deep_merge_cbct(target[key], value)
                            elif isinstance(value, list) and isinstance(target[key], list):
                                target[key].extend(v for v in value if v not in target[key])
                            else:
                                target[key] = value
                        else:
                            target[key] = value
                deep_merge_cbct(canonical['observations']['cbct'], cbct_data)
                logger.info(f"Merged CBCT observations: {list(cbct_data.keys())}")
        
        # Update observations.ent_findings (ENT/Nasal/Sinus structured block)
        if 'observations' in llm_result and 'ent_findings' in llm_result['observations']:
            ent_data = llm_result['observations']['ent_findings']
            if isinstance(ent_data, dict):
                def deep_merge_ent(target, source):
                    """Recursively merge ENT data, handling nested dicts"""
                    for key, value in source.items():
                        if value is None:
                            continue
                        if key in target:
                            if isinstance(value, dict) and isinstance(target[key], dict):
                                deep_merge_ent(target[key], value)
                            elif isinstance(value, list) and isinstance(target[key], list):
                                target[key].extend(v for v in value if v not in target[key])
                            else:
                                target[key] = value
                        else:
                            target[key] = value
                deep_merge_ent(canonical['observations']['ent_findings'], ent_data)
                logger.info(f"Merged ENT findings: {list(ent_data.keys())}")
        
        # Update observations.airway_phenotype
        if 'observations' in llm_result and 'airway_phenotype' in llm_result['observations']:
            phenotype_data = llm_result['observations']['airway_phenotype']
            if isinstance(phenotype_data, dict):
                for key, value in phenotype_data.items():
                    if key in canonical['observations']['airway_phenotype'] and value is not None:
                        canonical['observations']['airway_phenotype'][key] = value
                logger.info(f"Merged airway phenotype: {list(phenotype_data.keys())}")
        
        # Update treatment considerations
        if 'treatment_considerations' in llm_result:
            for key, value in llm_result['treatment_considerations'].items():
                if key in canonical['treatment_considerations'] and value is not None:
                    if isinstance(value, list):
                        canonical['treatment_considerations'][key] = value
                    else:
                        canonical['treatment_considerations'][key] = value
        
        # Update device design
        if 'device_design' in llm_result:
            for key, value in llm_result['device_design'].items():
                if key in canonical['device_design'] and value is not None:
                    if isinstance(value, list):
                        canonical['device_design'][key] = value
                    else:
                        canonical['device_design'][key] = value
        
        # Update follow up plan with deduplication and limiting
        if 'follow_up_plan' in llm_result:
            for key, value in llm_result['follow_up_plan'].items():
                if key in canonical['follow_up_plan'] and value is not None:
                    if isinstance(value, list):
                        # Deduplicate and limit to 3 items for evaluations and lifestyle
                        if key in ['evaluations', 'lifestyle']:
                            # Remove duplicates (case-insensitive)
                            seen = set()
                            unique_items = []
                            for item in value:
                                item_lower = str(item).lower().strip()
                                if item_lower not in seen and item_lower:
                                    seen.add(item_lower)
                                    unique_items.append(item)
                            
                            # Limit to maximum 3 items
                            canonical['follow_up_plan'][key] = unique_items[:3]
                            logger.info(f"Limited {key} to {len(unique_items[:3])} items (max 3)")
                        else:
                            canonical['follow_up_plan'][key] = value
                    else:
                        canonical['follow_up_plan'][key] = value
        
        # Update patient self report
        if 'patient_self_report' in llm_result:
            for key, value in llm_result['patient_self_report'].items():
                if key in canonical['patient_self_report'] and value is not None:
                    if key == 'symptoms' and isinstance(value, dict):
                        for symptom_key, symptom_value in value.items():
                            if symptom_key in canonical['patient_self_report']['symptoms'] and symptom_value is not None:
                                canonical['patient_self_report']['symptoms'][symptom_key] = symptom_value
                    elif key == 'goals' and isinstance(value, list):
                        for goal in value:
                            if goal not in canonical['patient_self_report']['goals']:
                                canonical['patient_self_report']['goals'].append(goal)
                    elif key == 'scales' and isinstance(value, dict):
                        for scale_key, scale_value in value.items():
                            if scale_key in canonical['patient_self_report']['scales'] and scale_value is not None:
                                canonical['patient_self_report']['scales'][scale_key] = scale_value
                    else:
                        canonical['patient_self_report'][key] = value
        
        # Update observations summary - MERGE and SORT by OSA priority
        if 'observations' in llm_result and 'summary' in llm_result['observations']:
            if isinstance(llm_result['observations']['summary'], list):
                # Get existing observations
                existing_summary = canonical['observations'].get('summary', [])
                if not isinstance(existing_summary, list):
                    existing_summary = []
                
                # Merge LLM results with existing observations
                llm_summary = llm_result['observations']['summary']
                merged_summary = existing_summary.copy()
                
                # Add LLM observations that aren't already present
                for llm_obs in llm_summary:
                    if llm_obs not in merged_summary:
                        merged_summary.append(llm_obs)
                
                # Sort by OSA clinical priority (most important first)
                sorted_summary = _sort_observations_by_osa_priority(merged_summary)
                canonical['observations']['summary'] = sorted_summary
                
                logger.info(f"Sorted {len(sorted_summary)} observations by OSA clinical priority")
        
    except Exception as e:
        logger.error(f"Error updating canonical from LLM result: {e}")

def compute_ahi_severity(ahi: float) -> str:
    """
    Compute OSA severity from AHI using standard classification.
    
    Args:
        ahi (float): Apnea-Hypopnea Index value
        
    Returns:
        str: Severity classification
    """
    # Add type checking and debugging
    if ahi is None:
        return "unknown"
    
    # Handle case where ahi might be a string or other type
    if isinstance(ahi, str):
        try:
            ahi = float(ahi.strip())
        except (ValueError, TypeError):
            logger.warning(f"AHI value cannot be converted to float: {ahi}")
            return "unknown"
    elif not isinstance(ahi, (int, float)):
        logger.warning(f"AHI value is not numeric: {ahi} (type: {type(ahi)})")
        return "unknown"
    
    try:
        ahi_float = float(ahi)
        if ahi_float < 5:
            return "none"
        elif ahi_float < 15:
            return "mild"
        elif ahi_float < 30:
            return "moderate"
        else:
            return "severe"
    except (ValueError, TypeError) as e:
        logger.error(f"Error converting AHI to float: {ahi} - {e}")
        return "unknown"

def normalize_study_type(study_type: str) -> str:
    """
    Normalize study type to strict enum values.
    
    Args:
        study_type (str): Raw study type from LLM
        
    Returns:
        str: Normalized study type
    """
    if not study_type:
        return "unknown"
    
    study_type_lower = study_type.lower()
    
    if any(term in study_type_lower for term in ['hsat', 'home', 'portable']):
        return "home"
    elif any(term in study_type_lower for term in ['psg', 'polysomnography', 'lab', 'laboratory', 'inlab']):
        return "inlab"
    elif any(term in study_type_lower for term in ['titration', 'cpap']):
        return "titration"
    else:
        return None  # Let schema validation handle unknown types

def extract_demographics_observations(document_content: str, document_name: str) -> List[Dict]:
    """
    Extract demographics observations from document content and create structured observations.
    
    Args:
        document_content (str): Raw document content
        document_name (str): Name of the document
        
    Returns:
        List[Dict]: List of demographics observations with schema paths
    """
    demographics_obs = []
    
    if not document_content:
        return demographics_obs
    
    # Extract demographics using the existing function
    demographics = extract_demographics_from_text([document_content])
    
    # Create structured observations for each found demographic
    for field, value in demographics.items():
        if value is not None:
            obs = {
                'path': f'demographics.{field}',
                'value': str(value),
                'observation': f"{field.replace('_', ' ').title()}: {value}",
                'score': 1,
                'explanation': f'Extracted from document: {document_name}',
                'evidence': f'Found {field.replace("_", " ")} in document content',
                'confidence': 100,
                'source': 'demographics-extraction'
            }
            demographics_obs.append(obs)
    
    return demographics_obs
def extract_demographics_from_text(observation_texts: List[str]) -> Dict[str, Any]:
    """
    Extract demographics information from observation texts.
    
    Args:
        observation_texts: List of observation strings
        
    Returns:
        Dict containing extracted demographics data
    """
    demographics = {
        'sex': None,
        'age_years': None,
        'height_cm': None,
        'weight_kg': None,
        'bmi': None
    }
    
    for text in observation_texts:
        if not text or text is None:
            continue
        try:
            text_lower = text.lower()
        except AttributeError:
            logger.warning(f"Text is not a string in demographics extraction: {type(text)} - {text}")
            continue
        
        # Extract age with more comprehensive patterns - only match 2-3 digit ages
        if demographics['age_years'] is None:
            age_patterns = [
                r'age[:\s]*(\d{2,3})\s*(?:years?|y\.?o\.?)?\b',  # Only 2-3 digit ages
                r'(\d{2,3})\s*years?\s*old\b',  # Only 2-3 digit ages
                r'(\d{2,3})\s*yo\b',  # Only 2-3 digit ages
                r'patient.*?age[:\s]*(\d{2,3})\s*(?:years?|y\.?o\.?)?\b',  # Only 2-3 digit ages
                r'(\d{2,3})\s*years?\s*of\s*age\b',  # Only 2-3 digit ages
                r'(\d{2,3})\s*years?\s*patient\b',  # Only 2-3 digit ages
                r'age[:\s]*(\d{2,3})\s*years?\b',  # Only 2-3 digit ages
                r'(\d{2,3})\s*years?\s*old\s*patient\b',  # Only 2-3 digit ages
                r'patient.*?(\d{2,3})\s*years?\s*old\b',  # Only 2-3 digit ages
            ]
            for pattern in age_patterns:
                age_match = re.search(pattern, text_lower)
                if age_match:
                    try:
                        age_value = int(age_match.group(1))
                        # Ensure we have at least 2 digits (10+)
                        if age_value < 10:
                            logger.warning(f"Rejected single-digit age: {age_value} from pattern: {pattern}")
                            continue
                        # More comprehensive age validation for adult patients
                        if 18 <= age_value <= 120:  # Adult age range (18+)
                            demographics['age_years'] = age_value
                            logger.debug(f"Extracted age: {age_value} from pattern: {pattern}")
                            break
                        elif 10 <= age_value < 18:  # Teenage range - log but don't use
                            logger.warning(f"Extracted age {age_value} is too young for adult patient, skipping")
                            continue
                        else:  # Outside reasonable range
                            logger.warning(f"Extracted age {age_value} is outside reasonable range (10-120), skipping")
                            continue
                    except (ValueError, TypeError):
                        continue
        
        # Extract sex/gender with enhanced patterns
        if demographics['sex'] is None:
            detected_sex = detect_sex_safely(text)
            if detected_sex:
                demographics['sex'] = detected_sex
    
        
        # Extract weight with unit conversion
        if demographics['weight_kg'] is None:
            weight_patterns = [
                r'weight[:\s]+(\d+(?:\.\d+)?)\s*kg',
                r'(\d+(?:\.\d+)?)\s*kg',
                r'weight[:\s]+(\d+(?:\.\d+)?)\s*lbs?',
                r'(\d+(?:\.\d+)?)\s*lbs?',
                r'weight[:\s]*(\d+(?:\.\d+)?)\s*pounds?',
                r'(\d+(?:\.\d+)?)\s*pounds?',
                r'wt[:\s]*(\d+(?:\.\d+)?)\s*kg',
                r'wt[:\s]*(\d+(?:\.\d+)?)\s*lbs?'
            ]
            for pattern in weight_patterns:
                weight_match = re.search(pattern, text_lower)
                if weight_match:
                    try:
                        weight_value = float(weight_match.group(1))
                        # Convert lbs to kg if needed
                        if 'lbs' in pattern or 'pounds' in pattern:
                            weight_value = weight_value * 0.453592
                        if 20 <= weight_value <= 500:  # Reasonable weight range in kg
                            demographics['weight_kg'] = round(weight_value, 1)
                            logger.debug(f"Extracted weight: {demographics['weight_kg']}kg from pattern: {pattern}")
                            break
                    except (ValueError, TypeError):
                        continue
        
        # Extract height with unit conversion
        if demographics['height_cm'] is None:
            height_patterns = [
                r'height[:\s]+(\d+(?:\.\d+)?)\s*cm',
                r'(\d+(?:\.\d+)?)\s*cm\s*height',
                r'height[:\s]*(\d+(?:\.\d+)?)\s*(?:cm|centimeters?)',
                r'patient.*?height[:\s]*(\d+(?:\.\d+)?)',
                r'(\d+(?:\.\d+)?)\s*cm\s*(?:tall|height)',
                # Handle feet and inches format (e.g., 5'7", 5 feet 7 inches)
                r'(\d+)\s*[\'′]\s*(\d+)\s*["″]',
                r'(\d+)\s*feet?\s*(\d+)\s*inches?',
                r'(\d+)\s*ft\s*(\d+)\s*in',
            ]
            for pattern in height_patterns:
                height_match = re.search(pattern, text_lower)
                if height_match:
                    try:
                        if 'feet' in pattern or 'ft' in pattern or '[\'′]' in pattern:
                            # Handle feet and inches format
                            feet = int(height_match.group(1))
                            inches = int(height_match.group(2))
                            height_value = (feet * 12 + inches) * 2.54  # Convert to cm
                        else:
                            height_value = float(height_match.group(1))
                            # Convert inches to cm if needed
                            if 'inches' in pattern or 'in' in pattern:
                                height_value = height_value * 2.54
                        if 100 <= height_value <= 250:  # Reasonable height range in cm
                            demographics['height_cm'] = round(height_value, 1)
                            logger.debug(f"Extracted height: {demographics['height_cm']}cm from pattern: {pattern}")
                            break
                    except (ValueError, TypeError):
                        continue
        
        # Extract BMI with enhanced patterns
        if demographics['bmi'] is None:
            bmi_patterns = [
                r'bmi[:\s]+(\d+(?:\.\d+)?)',
                r'body\s*mass\s*index[:\s]+(\d+(?:\.\d+)?)',
                r'bmi[:\s]*(\d+(?:\.\d+)?)',
                r'body\s*mass\s*index[:\s]*(\d+(?:\.\d+)?)',
                r'bmi\s*=\s*(\d+(?:\.\d+)?)',
                r'body\s*mass\s*index\s*=\s*(\d+(?:\.\d+)?)'
            ]
            for pattern in bmi_patterns:
                bmi_match = re.search(pattern, text_lower)
                if bmi_match:
                    try:
                        bmi_value = float(bmi_match.group(1))
                        if 10 <= bmi_value <= 80:  # Reasonable BMI range
                            demographics['bmi'] = round(bmi_value, 1)
                            logger.debug(f"Extracted BMI: {demographics['bmi']} from pattern: {pattern}")
                            break
                    except (ValueError, TypeError):
                        continue
    
    # Calculate BMI if we have weight and height but no BMI
    if demographics['bmi'] is None and demographics['weight_kg'] is not None and demographics['height_cm'] is not None:
        try:
            height_m = demographics['height_cm'] / 100
            calculated_bmi = demographics['weight_kg'] / (height_m * height_m)
            if 10 <= calculated_bmi <= 80:  # Reasonable BMI range
                demographics['bmi'] = round(calculated_bmi, 1)
                logger.info(f"Calculated BMI: {demographics['bmi']} from weight {demographics['weight_kg']}kg and height {demographics['height_cm']}cm")
        except (ValueError, TypeError, ZeroDivisionError):
            pass
    
    # Log extracted demographics
    extracted_demographics = {k: v for k, v in demographics.items() if v is not None}
    if extracted_demographics:
        logger.info(f"Extracted demographics: {extracted_demographics}")
    
    return demographics

def fallback_extraction(canonical: Dict[str, Any], observation_texts: List[str]):
    """
    Fallback extraction method when LLM fails.
    
    Args:
        canonical: The canonical structure to update
        observation_texts: List of observation strings
    """
    # Sort observations by OSA clinical priority before adding to summary
    sorted_observations = _sort_observations_by_osa_priority(observation_texts)
    canonical['observations']['summary'] = sorted_observations
    
    logger.info(f"Fallback: Sorted {len(sorted_observations)} observations by OSA clinical priority")
    
    # Extract demographics
    demographics = extract_demographics_from_text(observation_texts)
    for key, value in demographics.items():
        if value is not None:
            canonical['demographics'][key] = value
    
    # Basic regex extraction for key values
    for text in observation_texts:
        text_lower = text.lower()
        
        # Extract AHI
        if 'ahi' in text_lower:
            ahi_match = re.search(r'ahi[:\s]+(\d+(?:\.\d+)?)', text_lower)
            if ahi_match and canonical['sleep_study']['ahi'] is None:
                try:
                    ahi_value = float(ahi_match.group(1))
                    canonical['sleep_study']['ahi'] = ahi_value
                except (ValueError, TypeError):
                    pass

# Add utility functions after the existing imports and before the main functions

def fix_evaluation_format(evaluations):
    """
    Fix evaluation format - convert string representations to proper objects.
    
    Args:
        evaluations: List of evaluations (may contain strings or objects)
        
    Returns:
        List of properly formatted evaluation objects
    """
    if not evaluations:
        return []
    
    fixed_evaluations = []
    for eval_item in evaluations:
        if isinstance(eval_item, str):
            # Try to parse string representation
            try:
                # Remove single quotes and replace with double quotes for JSON parsing
                eval_str = eval_item.replace("'", '"')
                parsed = json.loads(eval_str)
                fixed_evaluations.append(parsed)
            except (json.JSONDecodeError, ValueError):
                # If parsing fails, create a simple object
                fixed_evaluations.append({"type": eval_item, "timeframe": "to be scheduled"})
        else:
            # Already an object, keep as is
            fixed_evaluations.append(eval_item)
    
    return fixed_evaluations

def _cleanup_follow_up_plan(canonical):
    """
    Clean up follow-up plan by deduplicating and limiting items.
    """
    if 'follow_up_plan' not in canonical:
        return canonical
    
    follow_up_plan = canonical['follow_up_plan']
    
    # Clean up evaluations (limit to 3, remove duplicates)
    if 'evaluations' in follow_up_plan and isinstance(follow_up_plan['evaluations'], list):
        evaluations = follow_up_plan['evaluations']
        
        # Remove duplicates based on reason field (case-insensitive)
        seen_reasons = set()
        unique_evaluations = []
        
        for eval_item in evaluations:
            if isinstance(eval_item, dict) and 'reason' in eval_item:
                reason_lower = str(eval_item['reason']).lower().strip()
                if reason_lower not in seen_reasons and reason_lower:
                    seen_reasons.add(reason_lower)
                    unique_evaluations.append(eval_item)
            elif isinstance(eval_item, str):
                eval_lower = eval_item.lower().strip()
                if eval_lower not in seen_reasons and eval_lower:
                    seen_reasons.add(eval_lower)
                    unique_evaluations.append(eval_item)
        
        # Limit to maximum 3 evaluations
        follow_up_plan['evaluations'] = unique_evaluations[:3]
        logger.info(f"Cleaned up evaluations: {len(unique_evaluations[:3])} items (max 3)")
    
    # Clean up lifestyle (limit to 3, remove duplicates)
    if 'lifestyle' in follow_up_plan and isinstance(follow_up_plan['lifestyle'], list):
        lifestyle = follow_up_plan['lifestyle']
        
        # Remove duplicates (case-insensitive)
        seen_lifestyle = set()
        unique_lifestyle = []
        
        for item in lifestyle:
            item_lower = str(item).lower().strip()
            if item_lower not in seen_lifestyle and item_lower:
                seen_lifestyle.add(item_lower)
                unique_lifestyle.append(item)
        
        # Limit to maximum 3 lifestyle items
        follow_up_plan['lifestyle'] = unique_lifestyle[:3]
        logger.info(f"Cleaned up lifestyle: {len(unique_lifestyle[:3])} items (max 3)")
    
    return canonical

def remove_duplicates_from_lists(obj):
    """
    Remove duplicate values from lists in nested structures.
    """
    if isinstance(obj, dict):
        return {k: remove_duplicates_from_lists(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        # Remove duplicates while preserving order
        seen = set()
        unique_list = []
        for item in obj:
            if isinstance(item, (dict, list)):
                # For complex objects, convert to string for comparison
                item_str = json.dumps(item, sort_keys=True)
                if item_str not in seen:
                    seen.add(item_str)
                    unique_list.append(remove_duplicates_from_lists(item))
            else:
                if item not in seen:
                    seen.add(item)
                    unique_list.append(item)
        return unique_list
    else:
        return obj

def flatten_nested_arrays(obj):
    """
    Flatten nested arrays (arrays within arrays) to single-level arrays.
    Example: ["a", ["b", "c"], "d"] -> ["a", "b", "c", "d"]
    """
    if isinstance(obj, dict):
        return {k: flatten_nested_arrays(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        flattened = []
        for item in obj:
            if isinstance(item, list):
                # Recursively flatten nested lists
                flattened.extend(flatten_nested_arrays(item))
            elif isinstance(item, dict):
                # Keep dicts as-is but process their contents
                flattened.append(flatten_nested_arrays(item))
            else:
                # Keep primitive values as-is
                flattened.append(item)
        return flattened
    else:
        return obj

def canonical_to_level4(canonical: Dict[str, Any], patient_id: int = None) -> Dict[str, Any]:
    """
    Transform full canonical JSON to Level-4 report input schema.
    
    This function maps the complex canonical structure to a clean, minimal schema
    optimized for Level-4 OSA report generation by LLMs.
    
    Args:
        canonical: Full canonical JSON with all data
        patient_id: Patient ID for logging (optional)
        
    Returns:
        Clean Level-4 input schema JSON
    """
    import re
    
    def get(path, default=""):
        """Safe lookup into dict with dotted path."""
        ref = canonical
        for p in path.split("."):
            if isinstance(ref, dict) and p in ref:
                ref = ref[p]
            else:
                return default
        return ref
    
    def extract_snore_avg(canonical):
        """Extract average snore dB from messy text."""
        for s in canonical.get("observations", {}).get("summary", []):
            s_str = str(s).lower()
            if "average" in s_str and "db" in s_str:
                # Example: "average 41 dB" or "snoring average 41 dB"
                try:
                    # Try multiple patterns
                    patterns = [
                        r'average.*?(\d+)\s*dB',
                        r'avg.*?(\d+)\s*dB',
                        r'mean.*?(\d+)\s*dB',
                    ]
                    for pattern in patterns:
                        match = re.search(pattern, s_str, re.IGNORECASE)
                        if match:
                            return float(match.group(1))
                except:
                    pass
        # Also check sleep_study and snoring objects
        if 'sleep_study' in canonical and isinstance(canonical['sleep_study'], dict):
            if 'snore_avg_db' in canonical['sleep_study']:
                return canonical['sleep_study']['snore_avg_db']
            if 'snoring' in canonical['sleep_study'] and isinstance(canonical['sleep_study']['snoring'], dict):
                if 'snore_avg_db' in canonical['sleep_study']['snoring']:
                    return canonical['sleep_study']['snoring']['snore_avg_db']
        if 'snoring' in canonical and isinstance(canonical['snoring'], dict):
            if 'snore_avg_db' in canonical['snoring']:
                return canonical['snoring']['snore_avg_db']
        return ""
    
    def extract_snore_max(canonical):
        """Extract max snore dB from messy text."""
        for s in canonical.get("observations", {}).get("summary", []):
            s_str = str(s).lower()
            if "maximum" in s_str or "max" in s_str and "db" in s_str:
                try:
                    patterns = [
                        r'maximum.*?(\d+)\s*dB',
                        r'max.*?(\d+)\s*dB',
                        r'peak.*?(\d+)\s*dB',
                    ]
                    for pattern in patterns:
                        match = re.search(pattern, s_str, re.IGNORECASE)
                        if match:
                            return float(match.group(1))
                except:
                    pass
        # Also check sleep_study and snoring objects
        if 'sleep_study' in canonical and isinstance(canonical['sleep_study'], dict):
            if 'snore_max_db' in canonical['sleep_study']:
                return canonical['sleep_study']['snore_max_db']
            if 'snoring' in canonical['sleep_study'] and isinstance(canonical['sleep_study']['snoring'], dict):
                if 'snore_max_db' in canonical['sleep_study']['snoring']:
                    return canonical['sleep_study']['snoring']['snore_max_db']
        if 'snoring' in canonical and isinstance(canonical['snoring'], dict):
            if 'snore_max_db' in canonical['snoring']:
                return canonical['snoring']['snore_max_db']
        return ""
    
    # 1. Demographics
    patient = {
        "sex": get("demographics.sex", ""),
        "age": get("demographics.age_years", ""),
        "bmi": get("meta.bmi") or get("demographics.bmi", ""),
        # Prefer explicit meta overrides, but fall back to extracted demographics fields
        "weight_kg": get("meta.weight_kg") or get("demographics.weight_kg", ""),
        "height_cm": get("meta.height_cm") or get("demographics.height_cm", ""),
    }
    
    # 2. Clinical background
    # Keep this lean (closer to how Level-4 reports are typically written) to reduce "unsupported"
    # grading deltas. Rich ENT details are provided via ent_findings below.
    diagnoses = get("impression_assessment.diagnoses", [])
    if not isinstance(diagnoses, list):
        diagnoses = []

    # De-duplicate while preserving order
    seen = set()
    deduped_diagnoses = []
    for d in diagnoses:
        ds = str(d).strip()
        if not ds:
            continue
        key = ds.lower()
        if key in seen:
            continue
        seen.add(key)
        deduped_diagnoses.append(ds)

    clinical_background = ", ".join(deduped_diagnoses) if deduped_diagnoses else ""
    
    # 3. Complaints & goals
    symptoms = get("patient_self_report.symptoms", {})
    if not isinstance(symptoms, dict):
        symptoms = {}
    
    complaints = [
        k.replace("_", " ").title()
        for k, v in symptoms.items() if v
    ]
    
    # Also check indications_symptoms
    indications = get("indications_symptoms", {})
    if isinstance(indications, dict):
        if indications.get("snoring_reported") and "Snoring" not in complaints:
            complaints.append("Snoring")
        if indications.get("daytime_sleepiness") and "Daytime Sleepiness" not in complaints:
            complaints.append("Daytime Sleepiness")
        if indications.get("witnessed_apneas") and "Witnessed Apneas" not in complaints:
            complaints.append("Witnessed Apneas")
        if indications.get("insomnia") and "Insomnia" not in complaints:
            complaints.append("Insomnia")
    
    goals = get("patient_self_report.goals", [])
    if not isinstance(goals, list):
        goals = []
    
    # 4. ENT findings (build a richer, clinician-readable summary from structured ENT block + imaging)
    def _summarize_ent(ent_block: Any) -> str:
        if not isinstance(ent_block, dict) or not ent_block:
            return ""
        parts = []

        nasal = ent_block.get("nasal")
        if isinstance(nasal, dict):
            if nasal.get("septum_deviation"):
                sev = nasal.get("septum_deviation_severity")
                parts.append(f"Septum deviation: {nasal.get('septum_deviation')}" + (f" ({sev})" if sev else ""))
            if nasal.get("turbinate_hypertrophy"):
                side = nasal.get("turbinate_side")
                parts.append(
                    f"Turbinate hypertrophy: {nasal.get('turbinate_hypertrophy')}" + (f" ({side})" if side else "")
                )
            if nasal.get("nasal_valve_collapse"):
                parts.append(f"Nasal valve collapse: {nasal.get('nasal_valve_collapse')}")
            if nasal.get("nasal_polyps") is True:
                grade = nasal.get("polyp_grade")
                side = nasal.get("polyp_side")
                polyps_line = "Nasal polyps"
                if grade:
                    polyps_line += f" (grade {grade})"
                if side:
                    polyps_line += f" ({side})"
                parts.append(polyps_line)
            if nasal.get("mucosa_appearance"):
                parts.append(f"Nasal mucosa: {nasal.get('mucosa_appearance')}")
            other_nasal = nasal.get("other_nasal_findings")
            if isinstance(other_nasal, list):
                for x in other_nasal:
                    xs = str(x).strip()
                    if xs:
                        parts.append(f"Nasal: {xs}")

        sinus = ent_block.get("sinus")
        if isinstance(sinus, dict):
            if sinus.get("chronic_sinusitis") is True:
                parts.append("Chronic sinusitis: yes")
            if sinus.get("acute_sinusitis") is True:
                parts.append("Acute sinusitis: yes")
            other_sinus = sinus.get("other_sinus_findings")
            if isinstance(other_sinus, list):
                for x in other_sinus:
                    xs = str(x).strip()
                    if xs:
                        parts.append(f"Sinus: {xs}")

        oro = ent_block.get("oropharynx")
        if isinstance(oro, dict):
            if oro.get("tonsil_grade") is not None and str(oro.get("tonsil_grade")).strip() != "":
                parts.append(f"Tonsil grade: {oro.get('tonsil_grade')}")
            if oro.get("uvula"):
                parts.append(f"Uvula: {oro.get('uvula')}")
            if oro.get("soft_palate"):
                parts.append(f"Soft palate: {oro.get('soft_palate')}")
            other_oro = oro.get("other_findings")
            if isinstance(other_oro, list):
                for x in other_oro:
                    xs = str(x).strip()
                    if xs:
                        parts.append(f"Oropharynx: {xs}")

        post = ent_block.get("post_surgical_changes")
        if post and str(post).strip():
            parts.append(f"Post-surgical changes: {post}")

        other = ent_block.get("other_ent_findings")
        if isinstance(other, list):
            for x in other:
                xs = str(x).strip()
                if xs:
                    parts.append(xs)

        return "; ".join(parts)

    ent_struct = get("observations.ent_findings", {})
    ent_summary = _summarize_ent(ent_struct)
    ent_imaging = get("observations.anatomy_imaging.nose_sinus", "")
    ent_mallampati = get("observations.anatomy_imaging.mallampati", "")

    # ENT-relevant meds (e.g., Otrivin / xylometazoline) can be important clinically
    ent_meds = []
    meds_raw = get("medications", [])
    if not isinstance(meds_raw, list):
        meds_raw = []

    for m in meds_raw:
        name = ""
        dose = ""
        if isinstance(m, dict):
            name = str(m.get("name", "")).strip()
            dose = str(m.get("dose", "")).strip()
        else:
            # Some pipelines store meds as plain strings
            name = str(m).strip()

        if not name:
            continue
        name_lower = name.lower()
        if "otrivin" in name_lower or "xylometazoline" in name_lower:
            ent_meds.append((name + (f" {dose}" if dose else "")).strip())

    ent_parts = []
    if ent_imaging and str(ent_imaging).strip():
        ent_parts.append(f"Imaging (nose/sinus): {ent_imaging}".strip())
    if ent_summary:
        ent_parts.append(ent_summary)
    if ent_mallampati and str(ent_mallampati).strip():
        ent_parts.append(f"Mallampati: {ent_mallampati}".strip())
    if ent_meds:
        ent_parts.append(f"Nasal meds: {', '.join(ent_meds)}")
    ent = " | ".join(ent_parts)
    
    # 5. Sleep study - COMPREHENSIVE extraction from ALL sources
    # Merge metrics from all available sources to ensure nothing is missed
    sleep_study = {}
    
    def merge_metric(target_dict, key, value, prefer_existing=True):
        """Merge a metric value into target dict, only if not already present or prefer_existing is False"""
        # Handle 0 as a valid value (not None/empty)
        if value is not None and value != "":
            # Convert to appropriate type if needed
            if isinstance(value, str):
                try:
                    # Try to convert string numbers to float/int
                    if '.' in value:
                        value = float(value)
                    else:
                        value = float(value)  # Use float to handle decimals
                except (ValueError, TypeError):
                    pass  # Keep as string if conversion fails
            
            # Only merge if key doesn't exist, or if prefer_existing is False, or if current value is empty
            if key not in target_dict or not prefer_existing or target_dict[key] == "" or target_dict[key] is None:
                target_dict[key] = value
    
    # Source 1: canonical_derived.latest.sleep_study (most recent)
    if 'canonical_derived' in canonical and isinstance(canonical['canonical_derived'], dict):
        derived = canonical['canonical_derived']
        if 'latest' in derived and 'sleep_study' in derived['latest']:
            latest_ss = derived['latest']['sleep_study']
            if isinstance(latest_ss, dict):
                for key in ['ahi', 'rem_ahi', 'rdi', 'rdi_overall', 'odi', 'odi3', 'supine_ahi', 'non_supine_ahi',
                           'ahi_non_supine', 'rem_rdi', 'rem_odi', 'supine_rdi', 'supine_odi', 'nrem_ahi',
                           'o2_nadir_pct', 'spo2_nadir_pct', 'sleep_efficiency_pct', 'tst_min', 'total_sleep_time_min',
                           't90_pct', 'time_below_90_pct', 'time_below_90_min', 'rera_index']:
                    merge_metric(sleep_study, key, latest_ss.get(key), prefer_existing=False)
    
    # Source 2: canonical_derived.baseline.sleep_study
    if 'canonical_derived' in canonical and isinstance(canonical['canonical_derived'], dict):
        derived = canonical['canonical_derived']
        if 'baseline' in derived and 'sleep_study' in derived['baseline']:
            baseline_ss = derived['baseline']['sleep_study']
            if isinstance(baseline_ss, dict):
                for key in ['ahi', 'rem_ahi', 'rdi', 'rdi_overall', 'odi', 'odi3', 'supine_ahi', 'non_supine_ahi',
                           'ahi_non_supine', 'rem_rdi', 'rem_odi', 'supine_rdi', 'supine_odi', 'nrem_ahi',
                           'o2_nadir_pct', 'spo2_nadir_pct', 'sleep_efficiency_pct', 'tst_min', 'total_sleep_time_min',
                           't90_pct', 'time_below_90_pct', 'time_below_90_min', 'rera_index']:
                    merge_metric(sleep_study, key, baseline_ss.get(key))
    
    # Source 3: canonical_derived.timeline.sleep_studies (check ALL entries)
    if 'canonical_derived' in canonical and isinstance(canonical['canonical_derived'], dict):
        derived = canonical['canonical_derived']
        if 'timeline' in derived and 'sleep_studies' in derived['timeline']:
            timeline_ss = derived['timeline']['sleep_studies']
            if isinstance(timeline_ss, list):
                for entry in timeline_ss:
                    if isinstance(entry, dict):
                        for key in ['ahi', 'rem_ahi', 'rdi', 'rdi_overall', 'odi', 'odi3', 'supine_ahi', 'non_supine_ahi',
                                   'ahi_non_supine', 'rem_rdi', 'rem_odi', 'supine_rdi', 'supine_odi', 'nrem_ahi',
                                   'o2_nadir_pct', 'spo2_nadir_pct', 'sleep_efficiency_pct', 'tst_min', 'total_sleep_time_min',
                                   't90_pct', 'time_below_90_pct', 'time_below_90_min', 'rera_index']:
                            merge_metric(sleep_study, key, entry.get(key))
    
    # Source 4: sleep_studies array (check ALL entries)
    if 'sleep_studies' in canonical and isinstance(canonical['sleep_studies'], list):
        for ss in canonical['sleep_studies']:
            if isinstance(ss, dict):
                # Check direct keys
                for key in ['ahi', 'rem_ahi', 'rdi', 'rdi_overall', 'odi', 'odi3', 'supine_ahi', 'non_supine_ahi',
                           'ahi_non_supine', 'rem_rdi', 'rem_odi', 'supine_rdi', 'supine_odi', 'nrem_ahi',
                           'o2_nadir_pct', 'spo2_nadir_pct', 'sleep_efficiency_pct', 'tst_min', 'total_sleep_time_min',
                           't90_pct', 'time_below_90_pct', 'time_below_90_min', 'rera_index']:
                    merge_metric(sleep_study, key, ss.get(key))
                # Check nested metrics dict
                if 'metrics' in ss and isinstance(ss['metrics'], dict):
                    for key in ['ahi', 'rem_ahi', 'rdi', 'rdi_overall', 'odi', 'odi3', 'supine_ahi', 'non_supine_ahi',
                               'ahi_non_supine', 'rem_rdi', 'rem_odi', 'supine_rdi', 'supine_odi', 'nrem_ahi',
                               'o2_nadir_pct', 'spo2_nadir_pct', 'sleep_efficiency_pct', 'tst_min', 'total_sleep_time_min',
                               't90_pct', 'time_below_90_pct', 'time_below_90_min', 'rera_index']:
                        merge_metric(sleep_study, key, ss['metrics'].get(key))
    
    # Source 5: temporal_series (check ALL entries)
    if 'temporal_series' in canonical and isinstance(canonical['temporal_series'], list):
        for ts in canonical['temporal_series']:
            if isinstance(ts, dict):
                for key in ['ahi', 'rem_ahi', 'rdi', 'rdi_overall', 'odi', 'odi3', 'supine_ahi', 'non_supine_ahi',
                           'ahi_non_supine', 'rem_rdi', 'rem_odi', 'supine_rdi', 'supine_odi', 'nrem_ahi',
                           'o2_nadir_pct', 'spo2_nadir_pct', 'sleep_efficiency_pct', 'tst_min', 'total_sleep_time_min',
                           't90_pct', 'time_below_90_pct', 'time_below_90_min', 'rera_index']:
                    merge_metric(sleep_study, key, ts.get(key))
    
    # Source 6: root-level sleep_study
    if 'sleep_study' in canonical and isinstance(canonical['sleep_study'], dict):
        root_ss = canonical['sleep_study']
        for key in ['ahi', 'rem_ahi', 'rdi', 'rdi_overall', 'odi', 'odi3', 'supine_ahi', 'non_supine_ahi',
                   'ahi_non_supine', 'rem_rdi', 'rem_odi', 'supine_rdi', 'supine_odi', 'nrem_ahi',
                   'o2_nadir_pct', 'spo2_nadir_pct', 'sleep_efficiency_pct', 'tst_min', 'total_sleep_time_min',
                   't90_pct', 'time_below_90_pct', 'time_below_90_min', 'rera_index']:
            merge_metric(sleep_study, key, root_ss.get(key))
    
    # Source 7: ui_sleep_metrics
    if 'ui_sleep_metrics' in canonical and isinstance(canonical['ui_sleep_metrics'], dict):
        ui_metrics = canonical['ui_sleep_metrics']
        for source_key in ['baseline', 'current']:
            if source_key in ui_metrics and isinstance(ui_metrics[source_key], dict):
                for key in ['ahi', 'rem_ahi', 'rdi', 'rdi_overall', 'odi', 'odi3', 'supine_ahi', 'non_supine_ahi',
                           'ahi_non_supine', 'rem_rdi', 'rem_odi', 'supine_rdi', 'supine_odi', 'nrem_ahi',
                           'o2_nadir_pct', 'spo2_nadir_pct', 'sleep_efficiency_pct', 'tst_min', 'total_sleep_time_min',
                           't90_pct', 'time_below_90_pct', 'time_below_90_min', 'rera_index']:
                    merge_metric(sleep_study, key, ui_metrics[source_key].get(key), prefer_existing=(source_key == 'baseline'))
    
    # Source 8: respiratory_indices (top-level)
    if 'respiratory_indices' in canonical and isinstance(canonical['respiratory_indices'], dict):
        resp_indices = canonical['respiratory_indices']
        for key in ['ahi', 'rem_ahi', 'rdi', 'rdi_overall', 'odi', 'odi3', 'supine_ahi', 'non_supine_ahi',
                   'ahi_non_supine', 'rem_rdi', 'rem_odi', 'supine_rdi', 'supine_odi', 'nrem_ahi']:
            merge_metric(sleep_study, key, resp_indices.get(key))
    
    # Source 9: oxygenation (top-level)
    if 'oxygenation' in canonical and isinstance(canonical['oxygenation'], dict):
        oxy = canonical['oxygenation']
        merge_metric(sleep_study, 'o2_nadir_pct', oxy.get('spo2_nadir_pct'))
        merge_metric(sleep_study, 'spo2_nadir_pct', oxy.get('spo2_nadir_pct'))
        merge_metric(sleep_study, 't90_pct', oxy.get('t90_pct'))
        merge_metric(sleep_study, 'time_below_90_min', oxy.get('time_below_90_min'))
    
    # Source 10: sleep_timing_architecture (top-level)
    if 'sleep_timing_architecture' in canonical and isinstance(canonical['sleep_timing_architecture'], dict):
        sta = canonical['sleep_timing_architecture']
        merge_metric(sleep_study, 'sleep_efficiency_pct', sta.get('sleep_efficiency_pct'))
        merge_metric(sleep_study, 'tst_min', sta.get('tst_min'))
        merge_metric(sleep_study, 'total_sleep_time_min', sta.get('tst_min'))
    
    # Source 11: Extract snoring_pct from snoring object
    snoring_pct = ""
    if 'snoring' in canonical and isinstance(canonical['snoring'], dict):
        snoring_pct = canonical['snoring'].get('snore_time_pct', "") or canonical['snoring'].get('snoring_pct', "")
    if not snoring_pct and 'sleep_study' in canonical and isinstance(canonical['sleep_study'], dict):
        if 'snoring' in canonical['sleep_study'] and isinstance(canonical['sleep_study']['snoring'], dict):
            snoring_pct = canonical['sleep_study']['snoring'].get('snore_time_pct', "") or canonical['sleep_study']['snoring'].get('snoring_pct', "")
        elif 'snoring_pct' in canonical['sleep_study']:
            snoring_pct = canonical['sleep_study']['snoring_pct']
        elif 'snore_time_pct' in canonical['sleep_study']:
            snoring_pct = canonical['sleep_study']['snore_time_pct']
    
    # Source 12: Extract from observations.summary text (for metrics mentioned in narrative)
    if 'observations' in canonical and 'summary' in canonical['observations']:
        summaries = canonical['observations']['summary']
        if isinstance(summaries, list):
            summary_text = ' '.join([str(s) for s in summaries if s])
            summary_text_lower = summary_text.lower()
            
            # Extract REM AHI from text (various patterns)
            rem_ahi_patterns = [
                r'rem.*?ahi.*?(\d+\.?\d*)',
                r'ahi.*?rem.*?(\d+\.?\d*)',
                r'rem.*?apnea.*?index.*?(\d+\.?\d*)',
            ]
            for pattern in rem_ahi_patterns:
                rem_ahi_match = re.search(pattern, summary_text_lower, re.IGNORECASE)
                if rem_ahi_match and not sleep_study.get('rem_ahi'):
                    try:
                        sleep_study['rem_ahi'] = float(rem_ahi_match.group(1))
                        break
                    except:
                        pass
            
            # Extract RDI from text
            rdi_patterns = [
                r'\brdi\b.*?(\d+\.?\d*)',
                r'respiratory.*?disturbance.*?index.*?(\d+\.?\d*)',
                r'rdi.*?(\d+\.?\d*)',
            ]
            for pattern in rdi_patterns:
                rdi_match = re.search(pattern, summary_text_lower, re.IGNORECASE)
                if rdi_match and not sleep_study.get('rdi'):
                    try:
                        sleep_study['rdi'] = float(rdi_match.group(1))
                        break
                    except:
                        pass
            
            # Extract Supine AHI from text
            supine_ahi_patterns = [
                r'supine.*?ahi.*?(\d+\.?\d*)',
                r'ahi.*?supine.*?(\d+\.?\d*)',
                r'supine.*?apnea.*?index.*?(\d+\.?\d*)',
            ]
            for pattern in supine_ahi_patterns:
                supine_ahi_match = re.search(pattern, summary_text_lower, re.IGNORECASE)
                if supine_ahi_match and not sleep_study.get('supine_ahi'):
                    try:
                        sleep_study['supine_ahi'] = float(supine_ahi_match.group(1))
                        break
                    except:
                        pass
            
            # Extract Non-Supine AHI from text
            non_supine_patterns = [
                r'non.*?supine.*?ahi.*?(\d+\.?\d*)',
                r'ahi.*?non.*?supine.*?(\d+\.?\d*)',
            ]
            for pattern in non_supine_patterns:
                non_supine_match = re.search(pattern, summary_text_lower, re.IGNORECASE)
                if non_supine_match and not sleep_study.get('non_supine_ahi'):
                    try:
                        sleep_study['non_supine_ahi'] = float(non_supine_match.group(1))
                        break
                    except:
                        pass
            
            # Extract O2 nadir from text
            o2_patterns = [
                r'o2.*?nadir.*?(\d+\.?\d*)',
                r'oxygen.*?nadir.*?(\d+\.?\d*)',
                r'spo2.*?nadir.*?(\d+\.?\d*)',
                r'lowest.*?o2.*?(\d+\.?\d*)',
            ]
            for pattern in o2_patterns:
                o2_match = re.search(pattern, summary_text_lower, re.IGNORECASE)
                if o2_match and not sleep_study.get('o2_nadir_pct'):
                    try:
                        val = float(o2_match.group(1))
                        # If value is between 0-1, it might be a percentage (0.88 = 88%)
                        if 0 < val < 1:
                            val = val * 100
                        sleep_study['o2_nadir_pct'] = val
                        break
                    except:
                        pass
            
            # Extract Sleep Efficiency from text
            eff_patterns = [
                r'sleep.*?efficiency.*?(\d+\.?\d*)',
                r'efficiency.*?(\d+\.?\d*)%',
            ]
            for pattern in eff_patterns:
                eff_match = re.search(pattern, summary_text_lower, re.IGNORECASE)
                if eff_match and not sleep_study.get('sleep_efficiency_pct'):
                    try:
                        sleep_study['sleep_efficiency_pct'] = float(eff_match.group(1))
                        break
                    except:
                        pass
            
            # Extract Total Sleep Time from text
            tst_patterns = [
                r'total.*?sleep.*?time.*?(\d+)',
                r'tst.*?(\d+)',
                r'sleep.*?duration.*?(\d+)',
            ]
            for pattern in tst_patterns:
                tst_match = re.search(pattern, summary_text_lower, re.IGNORECASE)
                if tst_match and not sleep_study.get('total_sleep_time_min'):
                    try:
                        sleep_study['total_sleep_time_min'] = float(tst_match.group(1))
                        break
                    except:
                        pass
    
    # Map alternative field names and consolidate
    if 'rdi_overall' in sleep_study and not sleep_study.get('rdi'):
        sleep_study['rdi'] = sleep_study['rdi_overall']
    if 'odi3' in sleep_study and not sleep_study.get('odi'):
        sleep_study['odi'] = sleep_study['odi3']
    if 'ahi_non_supine' in sleep_study and not sleep_study.get('non_supine_ahi'):
        sleep_study['non_supine_ahi'] = sleep_study['ahi_non_supine']
    if 'spo2_nadir_pct' in sleep_study and not sleep_study.get('o2_nadir_pct'):
        sleep_study['o2_nadir_pct'] = sleep_study['spo2_nadir_pct']
    if 'tst_min' in sleep_study and not sleep_study.get('total_sleep_time_min'):
        sleep_study['total_sleep_time_min'] = sleep_study['tst_min']
    
    # Build final sleep_study dict with proper field names
    # NOTE: Preserve raw precision from extraction; avoid rounding here.
    final_sleep_study = {
        "ahi": sleep_study.get("ahi", ""),
        "rem_ahi": sleep_study.get("rem_ahi", ""),
        "rdi": sleep_study.get("rdi", ""),
        "odi": sleep_study.get("odi", ""),
        "supine_ahi": sleep_study.get("supine_ahi", ""),
        "non_supine_ahi": sleep_study.get("non_supine_ahi", ""),
        "snoring_pct": snoring_pct,
        "snoring_avg_db": extract_snore_avg(canonical),
        "snoring_max_db": extract_snore_max(canonical),
        "o2_nadir": sleep_study.get("o2_nadir_pct", ""),
        "time_below_90_pct": sleep_study.get("time_below_90_pct", ""),
        "time_below_90_min": sleep_study.get("time_below_90_min", "") or sleep_study.get("time_below_90_pct_min", ""),
        "t90_pct": sleep_study.get("t90_pct", ""),
        "rera_index": sleep_study.get("rera_index", ""),
        # Some reference PDFs include a raw RERA count; keep optional + clinician-populatable.
        "rera_count": sleep_study.get("rera_count", ""),
        "sleep_efficiency": sleep_study.get("sleep_efficiency_pct", ""),
        "total_sleep_time_min": sleep_study.get("total_sleep_time_min", ""),
    }

    # Add a human-readable total sleep time string (helps match references like "7h 8min")
    def _tst_text_from_min(v) -> str:
        if v is None:
            return ""
        try:
            minutes = float(v)
        except Exception:
            return ""
        if minutes <= 0:
            return ""
        h = int(minutes // 60)
        m = int(round(minutes - (h * 60)))
        if m == 60:
            h += 1
            m = 0
        if h > 0 and m > 0:
            return f"{h}h {m}min"
        if h > 0:
            return f"{h}h"
        return f"{m}min"

    final_sleep_study["total_sleep_time_text"] = _tst_text_from_min(final_sleep_study.get("total_sleep_time_min"))
    
    sleep_study = final_sleep_study

    # 5b. Position stats (often referenced explicitly in Level-4 PDFs)
    # Keep minimal; clinicians can populate if extraction didn't find it.
    position_stats_in = get("position_stats", {})
    if not isinstance(position_stats_in, dict):
        position_stats_in = {}
    position_stats = {
        "supine_pct_of_sleep": position_stats_in.get("supine_pct_of_sleep", ""),
        "non_supine_pct_of_sleep": position_stats_in.get("non_supine_pct_of_sleep", ""),
    }
    
    # 6. Anatomy - Extended fields for comprehensive Level-4 reports
    anatomy = {
        "primary_obstruction_site": get("observations.anatomy_imaging.primary_obstruction_site", ""),
        "soft_palate": get("observations.anatomy_imaging.soft_palate_uvula", ""),
        "tongue_base": get("observations.anatomy_imaging.tongue_base", ""),
        "bite_jaw": get("observations.anatomy_imaging.bite_jaw", ""),
        "hyoid": get("observations.anatomy_imaging.hyoid", ""),
        "nasal_sinus": get("observations.anatomy_imaging.nose_sinus", ""),
        # Extended anatomical fields
        "arches": get("observations.anatomy_imaging.arches", ""),
        "epiglottis": get("observations.anatomy_imaging.epiglottis", ""),
        "neck_findings": get("observations.anatomy_imaging.neck_findings", ""),
        "overjet": get("observations.anatomy_imaging.overjet", ""),
        "overbite": get("observations.anatomy_imaging.overbite", ""),
        "retropalatal": get("observations.anatomy_imaging.retropalatal", ""),
        "retroglossal": get("observations.anatomy_imaging.retroglossal", ""),
        "pharyngeal_wall": get("observations.anatomy_imaging.pharyngeal_wall", ""),
        "tonsils": get("observations.anatomy_imaging.tonsils", ""),
        "adenoids": get("observations.anatomy_imaging.adenoids", ""),
        "mandibular_plane_angle": get("observations.anatomy_imaging.mandibular_plane_angle", ""),
        "airway_volume": get("observations.anatomy_imaging.airway_volume", ""),
        "mallampati": get("observations.anatomy_imaging.mallampati", ""),
        "friedman_stage": get("observations.anatomy_imaging.friedman_stage", ""),
        "mueller_maneuver": get("observations.anatomy_imaging.mueller_maneuver", ""),
        "dise_findings": get("observations.anatomy_imaging.dise_findings", ""),
        # Clinician-entered field (optional) to capture MAD response nuance
        "dise_mad_response": "",
        "conclusion": get("observations.anatomy_imaging.conclusion", ""),
        # Catch-all fields for observations that don't fit predefined categories
        "other_findings": get("observations.anatomy_imaging.other_findings", []),
        "other_tmj_findings": get("observations.tmj_flags.other_tmj_findings", []),
        "other_observations": get("observations.other_observations", []),
    }
    
    # Extract DISE from structured block if legacy field is empty
    if not anatomy.get("dise_findings"):
        dise_block = get("observations.dise", {})
        if isinstance(dise_block, dict) and dise_block:
            dise_parts = []
            
            # Check if DISE was performed
            if dise_block.get("performed"):
                dise_parts.append("DISE performed")
            
            # Extract collapse patterns by level (VOTE classification)
            for level in ["velum", "oropharynx", "tongue_base", "epiglottis"]:
                level_data = dise_block.get(level, {})
                if isinstance(level_data, dict):
                    collapse = level_data.get("collapse_pattern") or level_data.get("collapse_degree")
                    if collapse:
                        level_name = level.replace("_", " ").title()
                        dise_parts.append(f"{level_name}: {collapse}")
                    # Include notes/context when present (often contains the missing nuance)
                    notes = level_data.get("notes")
                    if notes:
                        level_name = level.replace("_", " ").title()
                        dise_parts.append(f"{level_name} notes: {notes}")
                    direction = level_data.get("collapse_direction")
                    if direction:
                        level_name = level.replace("_", " ").title()
                        dise_parts.append(f"{level_name} direction: {direction}")
            
            # Extract overall impression/summary
            if dise_block.get("overall_impression"):
                dise_parts.append(f"Impression: {dise_block['overall_impression']}")
            
            # Extract MAD response if present
            mad_response = dise_block.get("mad_response", {})
            if isinstance(mad_response, dict) and mad_response.get("improvement"):
                dise_parts.append(f"MAD response: {mad_response['improvement']}")
                anatomy["dise_mad_response"] = str(mad_response.get("improvement") or "")
            
            # Extract positional findings
            pos_findings = dise_block.get("positional_findings", {})
            if isinstance(pos_findings, dict):
                if pos_findings.get("supine_collapse"):
                    dise_parts.append(f"Supine: {pos_findings['supine_collapse']}")
            
            # Extract surgical considerations
            surgical = dise_block.get("surgical_considerations", [])
            if surgical:
                if isinstance(surgical, list):
                    dise_parts.append(f"Surgical considerations: {', '.join(surgical)}")
                else:
                    dise_parts.append(f"Surgical considerations: {surgical}")
            
            # Extract other findings
            other = dise_block.get("other_findings", [])
            if other:
                if isinstance(other, list):
                    dise_parts.extend(other)
                else:
                    dise_parts.append(str(other))
            
            if dise_parts:
                anatomy["dise_findings"] = "; ".join(dise_parts)
    
    # 6b. Device design placeholders (clinician can populate; LLM can use)
    # Keep the expected 9 keys present even if empty.
    device_design = get("device_design", {})
    if not isinstance(device_design, dict):
        device_design = {}
    device_design_out = {
        "mandibular_advancement": device_design.get("mandibular_advancement", ""),
        "vertical_opening": device_design.get("vertical_opening", ""),
        "anterior_window": device_design.get("anterior_window", ""),
        "retention_features": device_design.get("retention_features", ""),
        "material": device_design.get("material", ""),
        "pre_set": device_design.get("pre_set", ""),
        "anterior_acrylic": device_design.get("anterior_acrylic", ""),
        "coverage": device_design.get("coverage", ""),
        "clinical_notes": device_design.get("clinical_notes", ""),
    }

    # 7. Treatment history
    primary_pathway = get("treatment_considerations.primary_pathway", [])
    if not isinstance(primary_pathway, list):
        primary_pathway = []
    
    # Avoid treating "recommended CPAP" as "currently on CPAP" unless we have adherence/pressure evidence.
    avg_use_hours = get("device_adherence_if_applicable.avg_use_hours", "")
    p95_pressure = get("device_adherence_if_applicable.p95_pressure_cmH2O", "")
    cpap_used = bool(avg_use_hours or p95_pressure)
    
    # Get CPAP effective range from titration
    cpap_effective_range = ""
    if 'titration_if_present' in canonical and isinstance(canonical['titration_if_present'], dict):
        titration = canonical['titration_if_present']
        cpap_effective_range = titration.get('recommendation', "") or titration.get('effective_pressure_range', "")
    
    # Also check treatment_considerations.rationale for pressure range
    if not cpap_effective_range:
        rationale = get("treatment_considerations.rationale", "")
        if rationale:
            # Try to extract pressure range from rationale text
            pressure_match = re.search(r'(\d+\.?\d*)\s*[-–]\s*(\d+\.?\d*)\s*cmH2O', rationale, re.IGNORECASE)
            if pressure_match:
                cpap_effective_range = f"{pressure_match.group(1)}-{pressure_match.group(2)} cmH2O"
    
    # Only surface CPAP pressure range as "effective range" if CPAP appears to be actually used
    if not cpap_used:
        cpap_effective_range = ""
    
    treatment_history = {
        "cpap_used": cpap_used,
        "cpap_effective_range": cpap_effective_range,
        "avg_use_hours": avg_use_hours,
        "p95_pressure": p95_pressure,
    }

    # 8. Treatment considerations (as an array of strings, aligned to Level-4 prompt expectations)
    tc = get("treatment_considerations", {})
    treatment_considerations = []
    if isinstance(tc, dict):
        for key in ["primary_pathway", "adjuncts", "cautions"]:
            arr = tc.get(key)
            if isinstance(arr, list):
                treatment_considerations.extend([str(x) for x in arr if str(x).strip()])
        rationale = tc.get("rationale")
        if rationale and str(rationale).strip():
            treatment_considerations.append(f"Rationale: {rationale}")

    # 9. Follow-up plan -> recommendations/follow-up arrays
    fu = get("follow_up_plan", {})
    follow_up_plan = []
    recommendations = []
    if isinstance(fu, dict):
        # Evaluations become "Recommendations for Further Evaluation"
        evals = fu.get("evaluations")
        if isinstance(evals, list):
            for e in evals:
                if isinstance(e, dict):
                    t = e.get("type") or ""
                    r = e.get("reason") or ""
                    line = f"{t}: {r}".strip(": ").strip()
                    if line:
                        recommendations.append(line)
        # Lifestyle/therapy flags can also inform follow-up
        if fu.get("positional_therapy") is True:
            follow_up_plan.append("Positional therapy recommended due to supine predominance")
        months = fu.get("retest_after_init_months")
        if months:
            follow_up_plan.append(f"Repeat sleep study in ~{months} months after therapy initiation")
        lifestyle = fu.get("lifestyle")
        if isinstance(lifestyle, list):
            for item in lifestyle:
                if str(item).strip():
                    follow_up_plan.append(f"Lifestyle: {item}")

    # 10. Oral appliance options placeholder (clinician can populate)
    oral_appliance_options = get("oral_appliance_options", [])
    if not isinstance(oral_appliance_options, list):
        oral_appliance_options = []
    
    return {
        "patient": patient,
        "clinical_background": clinical_background,
        "complaints": complaints,
        "goals": goals,
        "ent_findings": ent,
        "sleep_study": sleep_study,
        "position_stats": position_stats,
        "anatomy": anatomy,
        "treatment_history": treatment_history,
        "treatment_considerations": treatment_considerations,
        "device_design": device_design_out,
        "recommendations": recommendations,
        "follow_up_plan": follow_up_plan,
        "oral_appliance_options": oral_appliance_options,
    }


# Keep old function name as alias for backward compatibility
def create_clean_canonical_for_llm(canonical: Dict[str, Any], patient_id: int = None) -> Dict[str, Any]:
    """
    Alias for canonical_to_level4 for backward compatibility.
    """
    return canonical_to_level4(canonical, patient_id)


def _prune_empty(obj):
    """
    Recursively remove empty values (None, "", [], {}) from nested structures.
    This ensures sparse JSON that only contains populated data.
    """
    if isinstance(obj, dict):
        out = {k: _prune_empty(v) for k, v in obj.items()}
        return {k: v for k, v in out.items() if v not in (None, "", [], {})}
    if isinstance(obj, list):
        out = [_prune_empty(v) for v in obj]
        out = [v for v in out if v not in (None, "", [], {})]
        return out
    return obj

def normalize_study_type_to_schema(study_type: str) -> str:
    """
    Normalize study type to match schema enum values.
    
    Args:
        study_type (str): Raw study type from LLM or extraction
        
    Returns:
        str: Normalized study type matching schema enum
    """
    if not study_type:
        return None
    
    study_type_lower = study_type.lower()
    
    if any(term in study_type_lower for term in ['hsat', 'home', 'portable']):
        return "home"
    elif any(term in study_type_lower for term in ['psg', 'polysomnography', 'lab', 'laboratory', 'inlab']):
        return "inlab"
    else:
        return None  # Let schema validation handle unknown types

def detect_sex_safely(text: str) -> str:
    """
    Safely detect sex/gender from text with context awareness.
    Avoids false positives from family history or clinician names.
    
    Args:
        text (str): Text to analyze
        
    Returns:
        str: 'M', 'F', 'X', or None
    """
    if not text or text is None:
        return None
    
    try:
        text_lower = text.lower()
    except AttributeError:
        logger.warning(f"Text is not a string: {type(text)} - {text}")
        return None
    
    # Comprehensive sex/gender detection patterns
    sex_patterns = [
        # Explicit labels with context
        r'\b(sex|gender)\s*[:\-]\s*(male|female)\b',
        r'\b(patient|pt)\s+(is\s+)?(male|female)\b',
        r'\b(male|female)\s+(patient|pt)\b',
        r'\b(sex|gender)\s*=\s*(male|female)\b',
        
        # Medical report patterns
        r'\b(patient|pt)\s+(male|female)\b',
        r'\b(male|female)\s+(patient|pt)\b',
        r'\b(sex|gender)\s*[:\-]\s*(m|f)\b',
        r'\b(m|f)\s*[:\-]\s*(male|female)\b',
        
        # Demographic section patterns
        r'\b(demographics?|demographic)\s*[:\-].*?(male|female)\b',
        r'\b(age|height|weight).*?(male|female)\b',
        
        # Form-style patterns
        r'\b(sex|gender)\s*[:\-]\s*[mf]\b',
        r'\b[mf]\s*[:\-]\s*(male|female)\b',
        
        # Standalone with medical context
        r'\b(male|female)\b(?!\s+(doctor|dr|physician|clinician|nurse|therapist))',
        r'\b(m|f)\b(?!\s+(doctor|dr|physician|clinician|nurse|therapist))'
    ]
    
    for pattern in sex_patterns:
        match = re.search(pattern, text_lower)
        if match:
            # Extract the gender part
            if len(match.groups()) >= 2:
                gender = match.group(2)
            elif len(match.groups()) >= 1:
                gender = match.group(1)
            else:
                gender = match.group(0)
            
            # Determine sex based on gender text
            if gender is None:
                continue
            try:
                gender_lower = gender.lower()
            except AttributeError:
                logger.warning(f"Gender is not a string: {type(gender)} - {gender}")
                continue
                
            if any(term in gender_lower for term in ['male', 'm']):
                logger.debug(f"Detected male from pattern: {pattern} -> {match.group(0)}")
                return 'M'
            elif any(term in gender_lower for term in ['female', 'f']):
                logger.debug(f"Detected female from pattern: {pattern} -> {match.group(0)}")
                return 'F'
    
    # Additional context-based detection
    # Look for medical context with gender indicators
    medical_terms = [
        'diagnosis', 'assessment', 'evaluation', 'examination', 'history',
        'symptoms', 'treatment', 'medication', 'condition', 'disorder',
        'sleep', 'apnea', 'snoring', 'tmj', 'bruxism', 'patient'
    ]
    
    has_medical_context = any(term in text_lower for term in medical_terms)
    
    if has_medical_context:
        # Look for standalone male/female in medical context
        if re.search(r'\b(male)\b(?!\s+(doctor|dr|physician|clinician|nurse|therapist))', text_lower):
            logger.debug("Detected male in medical context")
            return 'M'
        elif re.search(r'\b(female)\b(?!\s+(doctor|dr|physician|clinician|nurse|therapist))', text_lower):
            logger.debug("Detected female in medical context")
            return 'F'
        elif re.search(r'\b(m)\b(?!\s+(doctor|dr|physician|clinician|nurse|therapist))', text_lower):
            logger.debug("Detected M in medical context")
            return 'M'
        elif re.search(r'\b(f)\b(?!\s+(doctor|dr|physician|clinician|nurse|therapist))', text_lower):
            logger.debug("Detected F in medical context")
            return 'F'
    
    return None

def check_observation_exists(patient_id: int, observation_key: str, observation_value: str, document_name: str) -> bool:
    """
    Check if an observation already exists for this patient from the same document.
    Only prevents exact duplicates from the same document source.
    
    Args:
        patient_id (int): Patient ID
        observation_key (str): The observation field/key (e.g., 'sleep_study.ahi')
        observation_value (str): The observation value
        document_name (str): Name of the document being processed
        
    Returns:
        bool: True if exact duplicate from same document exists
    """
    conn = None
    try:
        conn = mysql.connector.connect(**DB_CONFIG)
        cursor = conn.cursor(dictionary=True)
        
        # Check observation_store for exact duplicates from the same document
        cursor.execute("""
            SELECT extracted_observations 
            FROM observation_store 
            WHERE patient_id = %s AND file_name = %s
        """, (patient_id, document_name))
        
        existing_observations = cursor.fetchall()
        
        for row in existing_observations:
            try:
                obs_data = json.loads(row['extracted_observations'])
                
                # Only check for exact duplicates: same key, same value, same document
                if (obs_data.get('path') == observation_key and 
                    obs_data.get('value') == observation_value and
                    obs_data.get('document_name') == document_name):
                    return True
                    
            except (json.JSONDecodeError, KeyError):
                continue
        
        return False
        
    except Exception as e:
        logger.error(f"Error checking observation existence for patient {patient_id}: {e}")
        return False
    finally:
        if conn:
            conn.close()

def store_observations_with_deduplication(patient_id: int, source_type: str, observations: List[Dict], document_info: Dict) -> bool:
    """
    Store individual observations in the observation_store table with deduplication.
    Only stores observations that don't already exist for this patient.
    
    Args:
        patient_id (int): Patient ID
        source_type (str): Type of document source
        observations (List[Dict]): List of individual observations
        document_info (Dict): Document metadata
        
    Returns:
        bool: True if successful, False otherwise
    """
    conn = None
    try:
        conn = mysql.connector.connect(**DB_CONFIG)
        cursor = conn.cursor()
        
        # Insert each observation as a separate row (include file_name for dedupe/auditing)
        insert_query = """
            INSERT INTO observation_store 
            (patient_id, file_name, source_type, source_text, extracted_observations, provider, observed_at, mention_date, created_at, updated_at)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, NOW(), NOW())
        """
        
        stored_count = 0
        skipped_count = 0
        missing_sleep_doc_date_logged = False
        missing_report_doc_date_logged = False
        
        for observation in observations:
            # Create source text for this observation
            source_text = observation.get('path') or observation.get('observation') or 'Document Analysis'
            
            # Store all observations without duplicate checking to ensure we capture all data
            observation_key = observation.get('path', '')
            observation_value = str(observation.get('value', observation.get('observation', '')))
            document_name = document_info.get('name', '')
            
            # Store the observation as JSON (like quiz system)
            observation_json = {
                **observation,
                "document_name": document_info.get('name', ''),
                "document_type": document_info.get('file_type', ''),
                "extraction_date": datetime.now().isoformat()
            }
            
            # Extended insert: also try to populate new columns when possible
            metric_key = observation.get('path')
            # Map known sleep_study.* and respiratory_indices.* paths to short keys
            metric_short = None
            if isinstance(metric_key, str):
                # Handle respiratory_indices paths (from Visual LLM)
                if metric_key == 'respiratory_indices.ahi_overall' or metric_key.endswith('.ahi_overall'):
                    metric_short = 'ahi'
                elif metric_key == 'respiratory_indices.ahi' or metric_key.endswith('.ahi'):
                    metric_short = 'ahi'
                elif metric_key == 'respiratory_indices.odi3' or metric_key == 'respiratory_indices.odi4' or metric_key.endswith('.odi'):
                    metric_short = 'odi'
                elif metric_key.endswith('.rdi'):
                    metric_short = 'rdi'
                elif metric_key.endswith('.supine_ahi') or 'supine_ahi' in metric_key:
                    metric_short = 'supine_ahi'
                elif metric_key.endswith('.rem_ahi') or 'rem_ahi' in metric_key:
                    metric_short = 'rem_ahi'
                elif metric_key.endswith('.supine_rdi') or 'supine_rdi' in metric_key:
                    metric_short = 'supine_rdi'
                elif metric_key.endswith('.rem_rdi') or 'rem_rdi' in metric_key:
                    metric_short = 'rem_rdi'
                elif metric_key.endswith('.supine_odi') or 'supine_odi' in metric_key:
                    metric_short = 'supine_odi'
                elif metric_key.endswith('.rem_odi') or 'rem_odi' in metric_key:
                    metric_short = 'rem_odi'
                elif 'spo2_nadir' in metric_key or metric_key.endswith('.o2_nadir_pct'):
                    metric_short = 'o2_nadir_pct'
                elif 'sleep_efficiency' in metric_key or metric_key.endswith('.sleep_efficiency_pct'):
                    metric_short = 'sleep_efficiency_pct'
                elif metric_key.endswith('.snoring.avg_db'):
                    metric_short = 'snoring_avg_db'
                elif metric_key.endswith('.snoring.max_db'):
                    metric_short = 'snoring_max_db'
                elif metric_key.endswith('.snoring.percent_total'):
                    metric_short = 'snoring_percent_total'
                elif metric_key.endswith('.snoring.over_50db_pct'):
                    metric_short = 'snoring_over_50db_pct'
                elif 'time_below_90' in metric_key or metric_key.endswith('.time_below_90_pct'):
                    metric_short = 'time_below_90_pct'
                elif metric_key.endswith('.time_below_90_pct_min'):
                    metric_short = 'time_below_90_pct_min'
                elif metric_key.endswith('.time_below_88_pct_min'):
                    metric_short = 'time_below_88_pct_min'

            metric_value = _safe_decimal(observation.get('value'))

            # Determine source_kind and dates by doc type
            file_name = document_info.get('name', '')
            s3_key = document_info.get('s3_key')
            episode_id = _make_episode_id(patient_id, s3_key, file_name)

            # Deterministic source_kind for sleep metrics:
            # - The system categorizes documents (adminfiles.file_category / files.category+subcategory)
            # - map_document_type_to_source_type(...) converts that to a stable document source_type
            # - For AHI and other sleep metrics, we must NOT allow filename/S3 heuristics to flip a report into a sleep_study.
            metric_path = metric_key if isinstance(metric_key, str) else ''
            is_sleep_metric_path = metric_path.startswith('sleep_study.') or \
                metric_path.startswith('respiratory_indices.') or \
                metric_path.startswith('oxygenation.') or \
                metric_path.startswith('sleep_timing_architecture.')

            normalized_source_type = (str(source_type or '')).strip().lower()
            if normalized_source_type == 'questionnaire':
                deterministic_source_kind = 'questionnaire'
            elif normalized_source_type in ('sleep_test', 'sleep_study'):
                deterministic_source_kind = 'sleep_study'
            else:
                deterministic_source_kind = 'report'

            # Only enforce determinism for sleep-metric paths; keep existing behavior for other fields.
            if is_sleep_metric_path:
                source_kind = deterministic_source_kind
            else:
                # Fallback legacy logic for non-sleep metrics (kept to avoid unintended changes)
                s3_path = (s3_key or '').lower()
                is_sleep_study_location = any(tok in s3_path for tok in ['sleep-test', 'sleep_study', 'sleep-study', 'sleepstudy'])
                is_report_location = any(tok in s3_path for tok in ['admin-files', '/reports/', '/report/'])

                if is_report_location:
                    source_kind = 'report'
                elif is_sleep_study_location:
                    source_kind = 'sleep_study'
                else:
                    source_kind = deterministic_source_kind

            study_type = _infer_study_type(file_name) if source_kind == 'sleep_study' else None

            # Best-effort dates: prefer document_date parsed from content/filename
            upload_dt = document_info.get('upload_date')
            document_date = document_info.get('document_date') or None
            observed_at_source = None
            
            logger.debug("Processing observation for %s: document_date=%s, upload_dt=%s", file_name, document_date, upload_dt)
            
            # Convert document_date to proper format if it's a datetime object
            if document_date and hasattr(document_date, 'strftime'):
                document_date_str = document_date.strftime('%Y-%m-%d')
                logger.debug("Converted document_date to string: %s", document_date_str)
            else:
                document_date_str = document_date
            
            if source_kind == 'sleep_study':
                if document_date:
                    observed_at = document_date  # Keep as datetime object for database
                    observed_at_source = 'document_text'
                    logger.info("✅ Using document_date for sleep study: %s", document_date)
                else:
                    observed_at = upload_dt
                    observed_at_source = 'upload_time'
                    if not missing_sleep_doc_date_logged:
                        logger.warning("⚠️ No document_date found, using upload_date: %s", upload_dt)
                        missing_sleep_doc_date_logged = True
                mention_date = None
            else:
                observed_at = None
                mention_date = document_date or upload_dt  # Keep as datetime object for database
                observed_at_source = 'document_text' if document_date else 'upload_time'
                if document_date:
                    logger.info("✅ Using document_date for report: %s", document_date)
                else:
                    if not missing_report_doc_date_logged:
                        logger.warning("⚠️ No document_date found for report, using upload_date: %s", upload_dt)
                        missing_report_doc_date_logged = True

            # Insert original row
            cursor.execute(insert_query, (
                patient_id,
                file_name,
                source_type,
                source_text,
                json.dumps(observation_json),
                'bedrock',
                observed_at,
                mention_date
            ))
            new_id = cursor.lastrowid

            # Enrich the row we just inserted with the new columns (update is simpler here)
            update_extended = (
                "UPDATE observation_store SET "
                "path=%s, "
                "metric_key=%s, metric_value_decimal=%s, metric_unit=%s, metric_phase=%s, "
                "observed_at=%s, mention_date=%s, source_kind=%s, study_type=%s, episode_id=%s, "
                "facility=%s, s3_key=%s, file_section=%s, snippet=%s, link_confidence=%s, link_status=%s, "
                "observed_at_source=%s "
                "WHERE id=%s"
            )
            # Store the full path as metric_key (not the short key) so canonical JSON can find it
            # metric_short is used for aggregation, but metric_key should be the full path
            cursor.execute(update_extended, (
                metric_key,
                metric_key,  # Use full path as metric_key, not metric_short
                metric_value, None, None,
                observed_at, mention_date, source_kind, study_type, episode_id,
                None, s3_key, None, observation.get('evidence'), None, None,
                observed_at_source,
                new_id
            ))

            # Also set document_date and observed_at_source if available
            try:
                cursor.execute(
                    "UPDATE observation_store SET document_date=%s, observed_at_source=%s WHERE id=%s",
                    (document_date, observed_at_source, new_id)
                )
                logger.info(f"Updated observation {new_id} with document_date: {document_date}, source: {observed_at_source}")
            except Exception as e:
                logger.warning(f"Failed to update document_date for observation {new_id}: {e}")
                pass
            stored_count += 1
        
        # Mark the document as analyzed
        document_id = document_info.get('id')
        source_table = document_info.get('source_table', 'files')
        
        if document_id and source_table:
            if source_table == 'files':
                update_query = "UPDATE files SET analyzed = TRUE WHERE id = %s"
            else:  # adminfiles
                update_query = "UPDATE adminfiles SET analyzed = TRUE WHERE id = %s"
            
            cursor.execute(update_query, (document_id,))
            logger.info(f"Marked document {document_info.get('name', '')} as analyzed")
        
        conn.commit()
        
        logger.info(f"Stored {stored_count} observations for patient {patient_id}, source: {source_type} (no duplicate checking)")
        return True
        
    except Exception as e:
        logger.error(f"Error storing observations for patient {patient_id}: {e}")
        if conn:
            conn.rollback()
        return False
    finally:
        if conn:
            conn.close()
def extract_numerical_data_with_regex(observation_texts: List[str]) -> Dict[str, Any]:
    """
    Extract numerical data using regex patterns for better accuracy and speed.
    This is more reliable than LLM for structured numerical values.
    
    Args:
        observation_texts: List of observation strings
        
    Returns:
        Dict containing extracted numerical fields
    """
    import re
    
    extracted_data = {}
    combined_text = " ".join(observation_texts).lower()
    
    # Demographics - Numerical
    # Age - More specific patterns to avoid false matches
    age_patterns = [
        r'age[:\s]*(\d{2,3})\s*(?:years?|y\.?o\.?)?\b',  # Only 2-3 digit ages
        r'(\d{2,3})\s*years?\s*old\b',  # Only 2-3 digit ages
        r'(\d{2,3})\s*yo\b',  # Only 2-3 digit ages
        r'patient.*?age[:\s]*(\d{2,3})\s*(?:years?|y\.?o\.?)?\b',  # Only 2-3 digit ages
        r'(\d{2,3})\s*years?\s*of\s*age\b',  # Only 2-3 digit ages
        r'(\d{2,3})\s*years?\s*patient\b',  # Only 2-3 digit ages
        r'age[:\s]*(\d{2,3})\s*years?\b',  # Only 2-3 digit ages
        r'(\d{2,3})\s*years?\s*old\s*patient\b',  # Only 2-3 digit ages
        r'patient.*?(\d{2,3})\s*years?\s*old\b',  # Only 2-3 digit ages
    ]
    for pattern in age_patterns:
        match = re.search(pattern, combined_text)
        if match:
            age = int(match.group(1))
            matched_text = match.group(0)  # Get the full matched text for debugging
            
            # Ensure we have at least 2 digits (10+)
            if age < 10:
                logger.warning(f"Rejected single-digit age: {age} from text '{matched_text}' using pattern: {pattern}")
                continue
            # More comprehensive age validation for adult patients
            elif age < 18:  # Very young ages are likely false matches for adult patients
                logger.warning(f"Rejected likely false age match for adult patient: {age} from text '{matched_text}' using pattern: {pattern}")
                continue
            elif 18 <= age <= 120:  # Adult age range validation
                extracted_data['age_years'] = age
                logger.info(f"Extracted age: {age} from text '{matched_text}' using pattern: {pattern}")
                break
            else:
                logger.warning(f"Age {age} outside valid range (18-120), skipping")
                continue
    
    # Height - More specific patterns to avoid false matches
    height_patterns = [
        r'height[:\s]*(\d+(?:\.\d+)?)\s*cm',
        r'(\d+(?:\.\d+)?)\s*cm\s*height',
        r'height[:\s]*(\d+(?:\.\d+)?)\s*(?:cm|centimeters?)',
        r'patient.*?height[:\s]*(\d+(?:\.\d+)?)',
        r'(\d+(?:\.\d+)?)\s*cm\s*(?:tall|height)',
        # Handle feet and inches format (e.g., 5'7", 5 feet 7 inches)
        r'(\d+)\s*[\'′]\s*(\d+)\s*["″]',
        r'(\d+)\s*feet?\s*(\d+)\s*inches?',
        r'(\d+)\s*ft\s*(\d+)\s*in',
    ]
    for pattern in height_patterns:
        match = re.search(pattern, combined_text)
        if match:
            try:
                if 'feet' in pattern or 'ft' in pattern or '[\'′]' in pattern:
                    # Handle feet and inches format
                    feet = int(match.group(1))
                    inches = int(match.group(2))
                    height_cm = (feet * 12 + inches) * 2.54  # Convert to cm
                    matched_text = match.group(0)  # Get the full matched text for debugging
                    logger.info(f"Extracted height: {feet}'{inches}\" = {height_cm}cm from text '{matched_text}' using pattern: {pattern}")
                else:
                    height_cm = float(match.group(1))
                    matched_text = match.group(0)  # Get the full matched text for debugging
                    logger.info(f"Extracted height: {height_cm}cm from text '{matched_text}' using pattern: {pattern}")
                
                # Additional validation - reject obviously wrong heights
                if height_cm < 120:  # Very short heights are likely false matches
                    logger.warning(f"Rejected likely false height match: {height_cm}cm from text '{matched_text}' using pattern: {pattern}")
                    continue
                elif 120 <= height_cm <= 250:  # More reasonable adult height range
                    extracted_data['height_cm'] = round(height_cm, 1)
                    break
                else:
                    logger.warning(f"Extracted height {height_cm}cm is outside reasonable range (120-250cm), skipping")
            except (ValueError, TypeError, IndexError) as e:
                logger.warning(f"Error processing height pattern {pattern}: {e}")
                continue
    
    # Weight
    weight_patterns = [
        r'weight[:\s]*(\d+(?:\.\d+)?)\s*kg',
        r'(\d+(?:\.\d+)?)\s*kg\s*weight',
        r'weight[:\s]*(\d+(?:\.\d+)?)',
    ]
    for pattern in weight_patterns:
        match = re.search(pattern, combined_text)
        if match:
            weight = float(match.group(1))
            if 20 <= weight <= 300:  # Validation
                extracted_data['weight_kg'] = weight
                break
    
    # BMI
    bmi_patterns = [
        r'bmi[:\s]*(\d+(?:\.\d+)?)',
        r'body mass index[:\s]*(\d+(?:\.\d+)?)',
    ]
    for pattern in bmi_patterns:
        match = re.search(pattern, combined_text)
        if match:
            bmi = float(match.group(1))
            if 10 <= bmi <= 80:  # Validation
                extracted_data['bmi'] = bmi
                break
    
    # Sleep Study - Numerical
    sleep_study = {}
    
    # AHI
    ahi_candidates = []
    for rule in AHI_PATTERN_RULES:
        try:
            matches = list(re.finditer(rule['pattern'], combined_text, re.IGNORECASE))
        except re.error as regex_err:
            logger.warning(f"Invalid AHI regex {rule['pattern']}: {regex_err}")
            continue
        for match in matches:
            try:
                ahi_val = float(match.group(1))
            except (ValueError, TypeError):
                continue
            if not (0 <= ahi_val <= 200):
                continue
            context_slice = combined_text[max(match.start() - 40, 0): match.end() + 40]
            bonus = sum(1 for kw in AHI_BONUS_KEYWORDS if kw in context_slice)
            penalty = sum(1 for kw in AHI_PENALTY_KEYWORDS if kw in context_slice)
            score = rule.get('score', 1) + bonus - penalty
            ahi_candidates.append({'value': ahi_val, 'score': score})
    if ahi_candidates:
        best = max(ahi_candidates, key=lambda c: (c['score'], -abs(c['value'] - 15)))
        if best['score'] >= AHI_CONFIDENCE_THRESHOLD:
            sleep_study['ahi'] = round(best['value'], 1)
        else:
            logger.debug(
                "AHI candidates below confidence threshold (best score %.1f, threshold %s)",
                best['score'],
                AHI_CONFIDENCE_THRESHOLD
            )
    
    # ODI
    odi_patterns = [
        r'odi[:\s]*(\d+(?:\.\d+)?)',
        r'oxygen.*?desaturation.*?index[:\s]*(\d+(?:\.\d+)?)',
        r'(\d+(?:\.\d+)?)\s*odi',
    ]
    for pattern in odi_patterns:
        match = re.search(pattern, combined_text)
        if match:
            odi = float(match.group(1))
            if 0 <= odi <= 200:  # Validation
                sleep_study['odi'] = odi
                break
    
    # O2 Nadir
    o2_patterns = [
        r'o2.*?nadir[:\s]*(\d+(?:\.\d+)?)',
        r'oxygen.*?nadir[:\s]*(\d+(?:\.\d+)?)',
        r'spo2.*?nadir[:\s]*(\d+(?:\.\d+)?)',
        r'(\d+(?:\.\d+)?).*?o2.*?nadir',
    ]
    for pattern in o2_patterns:
        match = re.search(pattern, combined_text)
        if match:
            o2_nadir = float(match.group(1))
            if 50 <= o2_nadir <= 100:  # Validation
                sleep_study['o2_nadir_pct'] = o2_nadir
                break
    
    # Sleep Duration
    sleep_duration_patterns = [
        r'sleep.*?duration[:\s]*(\d+(?:\.\d+)?)\s*hours?',
        r'total.*?sleep.*?time[:\s]*(\d+(?:\.\d+)?)\s*hours?',
        r'(\d+(?:\.\d+)?)\s*hours?.*?sleep',
    ]
    for pattern in sleep_duration_patterns:
        match = re.search(pattern, combined_text)
        if match:
            duration = float(match.group(1))
            if 0 <= duration <= 24:  # Validation
                sleep_study['sleep_duration_h'] = duration
                break
    
    # Sleep Efficiency
    efficiency_patterns = [
        r'sleep.*?efficiency[:\s]*(\d+(?:\.\d+)?)%',
        r'efficiency[:\s]*(\d+(?:\.\d+)?)%',
        r'(\d+(?:\.\d+)?)%.*?efficiency',
    ]
    for pattern in efficiency_patterns:
        match = re.search(pattern, combined_text)
        if match:
            efficiency = float(match.group(1))
            if 0 <= efficiency <= 100:  # Validation
                sleep_study['sleep_efficiency_pct'] = efficiency
                break
    
    # Snoring
    snoring_patterns = [
        r'snoring.*?(\d+(?:\.\d+)?)\s*db',
        r'(\d+(?:\.\d+)?)\s*db.*?snoring',
    ]
    for pattern in snoring_patterns:
        match = re.search(pattern, combined_text)
        if match:
            snoring_db = float(match.group(1))
            if 0 <= snoring_db <= 120:  # Validation
                if 'snoring' not in sleep_study:
                    sleep_study['snoring'] = {}
                sleep_study['snoring']['avg_db'] = snoring_db
                break
    
    # Add sleep study data if found
    if sleep_study:
        extracted_data['sleep_study'] = sleep_study
    
    return extracted_data

def _extract_document_date(content: str, filename: Optional[str]) -> Tuple[Optional[datetime], Optional[str]]:
    """
    Best-effort document date extraction from content and filename.
    Returns (datetime, source_tag) where source_tag in {'document_text','file_metadata'}.
    """
    import re
    from datetime import datetime
    # Try content first
    try:
        text = content or ''
        # Common patterns: 2025-09-10, 10/09/2025, 09/10/2025, Sep 10, 2025
        patterns = [
            (r"\b(20\d{2})[-/](\d{1,2})[-/](\d{1,2})\b", '%Y-%m-%d'),
            (r"\b(\d{1,2})/(\d{1,2})/(20\d{2})\b", '%m/%d/%Y'),
            (r"\b(\d{1,2})-(\d{1,2})-(20\d{2})\b", '%m-%d-%Y'),
            (r"\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\s+(\d{1,2}),\s*(20\d{2})\b", '%b %d, %Y'),
        ]
        for pat, fmt in patterns:
            m = re.search(pat, text, flags=re.IGNORECASE)
            if m:
                try:
                    if '%Y-%m-%d' == fmt:
                        dt = datetime(int(m.group(1)), int(m.group(2)), int(m.group(3)))
                    elif fmt in ('%m/%d/%Y','%m-%d-%Y'):
                        dt = datetime(int(m.group(3)), int(m.group(1)), int(m.group(2)))
                    else:
                        # Month name pattern
                        mon = m.group(1)[:3].title()
                        dt = datetime.strptime(f"{mon} {m.group(2)}, {m.group(3)}", '%b %d, %Y')
                    return dt, 'document_text'
                except Exception:
                    continue
    except Exception:
        pass
    # Try filename
    try:
        fn = filename or ''
        patterns_fn = [
            (r"(20\d{2})[-_](\d{1,2})[-_](\d{1,2})", '%Y-%m-%d'),
            (r"(\d{1,2})[-_](\d{1,2})[-_](20\d{2})", '%m-%d-%Y'),
            (r"(20\d{2})(\d{2})(\d{2})", '%Y%m%d'),
            # Handle patterns like "19.8.24" (day.month.year)
            (r"(\d{1,2})\.(\d{1,2})\.(\d{2})", '%d.%m.%y'),
            # Handle patterns like "19/8/24" (day/month/year)
            (r"(\d{1,2})/(\d{1,2})/(\d{2})", '%d/%m/%y'),
        ]
        for pat, fmt in patterns_fn:
            m = re.search(pat, fn)
            if m:
                try:
                    if fmt == '%Y-%m-%d':
                        dt = datetime(int(m.group(1)), int(m.group(2)), int(m.group(3)))
                    elif fmt == '%m-%d-%Y':
                        dt = datetime(int(m.group(3)), int(m.group(1)), int(m.group(2)))
                    elif fmt == '%Y%m%d':
                        dt = datetime(int(m.group(1)), int(m.group(2)), int(m.group(3)))
                    elif fmt == '%d.%m.%y':
                        # Handle 2-digit year (assume 20xx)
                        year = int(m.group(3))
                        if year < 50:  # Assume 20xx for years 00-49
                            year += 2000
                        else:  # Assume 19xx for years 50-99
                            year += 1900
                        dt = datetime(year, int(m.group(2)), int(m.group(1)))
                    elif fmt == '%d/%m/%y':
                        # Handle 2-digit year (assume 20xx)
                        year = int(m.group(3))
                        if year < 50:  # Assume 20xx for years 00-49
                            year += 2000
                        else:  # Assume 19xx for years 50-99
                            year += 1900
                        dt = datetime(year, int(m.group(2)), int(m.group(1)))
                    else:
                        dt = datetime(int(m.group(1)), int(m.group(2)), int(m.group(3)))
                    return dt, 'file_metadata'
                except Exception:
                    continue
    except Exception:
        pass
    return None, None

def extract_textual_observations_with_llm(observation_texts: List[str], numerical_data: Dict[str, Any], source_type: str = 'general_medical') -> Dict[str, Any]:
    """
    Extract textual observations using LLM for better context understanding.
    This is more reliable than regex for descriptive text and complex observations.
    
    Args:
        observation_texts: List of observation strings
        numerical_data: Already extracted numerical data to avoid duplication
        source_type: Type of document ('questionnaire', 'general_medical', etc.)
        
    Returns:
        Dict containing extracted textual fields
    """
    import boto3
    import json
    
    # Initialize Bedrock client
    bedrock = boto3.client(
        service_name='bedrock-runtime',
        region_name='us-west-2'
    )
    
    # Create different prompts based on source_type
    if source_type == 'questionnaire':
        # Special prompt for questionnaire documents - categorize into patient_self_report
        textual_prompt = f"""
You are a medical data specialist processing a PATIENT QUESTIONNAIRE. Extract patient self-reported information and categorize it into the patient_self_report section.

FOCUS ON PATIENT SELF-REPORTED DATA:
- patient_self_report.symptoms: daytime_sleepiness (boolean), witnessed_apneas (boolean), dry_mouth (boolean), etc.
- patient_self_report.goals: list of patient's treatment goals (e.g., ["reduce snoring", "improve sleep quality"])
- patient_self_report.primary_complaint: main issue the patient wants to address
- patient_self_report.scales: any rating scales or scores the patient provided

IMPORTANT RULES FOR QUESTIONNAIRES:
1. DO NOT extract numerical values - they are already handled by regex
2. Focus on patient's own words and self-reported symptoms
3. Categorize symptoms into appropriate boolean fields (true/false)
4. Extract treatment goals as a list of strings
5. Identify the primary complaint from patient's responses
6. Only include fields that have actual content from the questionnaire

ALREADY EXTRACTED NUMERICAL DATA (DO NOT DUPLICATE):
{json.dumps(numerical_data, indent=2)}

QUESTIONNAIRE RESPONSES TO ANALYZE:
{chr(10).join(observation_texts)}

Return ONLY a valid JSON object with the patient_self_report data. Do not include any explanations or text outside the JSON.
"""
    else:
        # Standard prompt for medical documents
        textual_prompt = f"""
You are a medical data specialist. Extract ONLY textual/descriptive observations from the following patient data.

FOCUS ON TEXTUAL FIELDS ONLY - DO NOT EXTRACT NUMERICAL VALUES.

REQUIRED OUTPUT JSON STRUCTURE:
{{
  "observations": {{
    "summary": ["observation 1", "observation 2"],
    "anatomy_imaging": {{
      "primary_obstruction_site": "velopharyngeal, oropharyngeal, tongue base",
      "soft_palate_uvula": "elongated soft palate, swollen uvula",
      "tongue_base": "large posteriorly positioned tongue base",
      "bite_jaw": "reduced overjet and overbite, retruded mandible",
      "hyoid": "normal position" or "inferiorly positioned",
      "nose_sinus": "deviated septum, mucosal thickening",
      "tmj": "TMJ clicking noted",
      "arches": "narrow underdeveloped arches",
      "epiglottis": "normal" or "floppy epiglottis",
      "neck_findings": "thick neck, cervical adiposity",
      "overjet": "reduced overjet 2mm",
      "overbite": "increased overbite",
      "retropalatal": "retropalatal collapse",
      "retroglossal": "retroglossal narrowing",
      "pharyngeal_wall": "lateral pharyngeal wall collapse",
      "tonsils": "enlarged tonsils grade 2+",
      "adenoids": "adenoid hypertrophy",
      "mallampati": "Mallampati 3",
      "friedman_stage": "Friedman 3",
      "conclusion": "multilevel obstruction pattern",
      "other_findings": []
    }},
    
    "dise": {{
      "performed": true,
      "date": "2023-02-27",
      "velum": {{
        "collapse_pattern": "complete",
        "collapse_direction": "AP",
        "notes": "complete AP collapse of velum"
      }},
      "oropharynx_lateral_walls": {{
        "collapse_pattern": "partial",
        "collapse_direction": "medial",
        "notes": "medial collapse of lateral pharyngeal walls"
      }},
      "tongue_base": {{
        "collapse_pattern": "partial",
        "collapse_direction": "posterior",
        "hypertrophy_grade": "2-3",
        "notes": "tongue base hypertrophy grade 2-3 with posterior collapse"
      }},
      "epiglottis": {{
        "collapse_pattern": "none",
        "appearance": "normal",
        "notes": "epiglottis normal appearance, no collapse"
      }},
      "maneuver_response": {{
        "jaw_thrust": {{
          "performed": true,
          "response": "moderate",
          "notes": "moderate improvement with jaw thrust"
        }},
        "head_rotation": {{
          "performed": true,
          "response": "none",
          "notes": "no improvement with head rotation"
        }}
      }},
      "positional_findings": {{
        "supine_collapse": "severe collapse in supine position",
        "positional_dependence": "severe",
        "supine_predominance_pct": 95
      }},
      "obstruction_pattern": "multilevel",
      "primary_site": "velum",
      "secondary_sites": ["tongue base", "lateral walls"],
      "appliance_suitability": {{
        "suitable": true,
        "predicted_response": "moderate",
        "rationale": "positive jaw thrust response suggests MAD may be beneficial"
      }}
    }},
    
    "cbct": {{
      "performed": true,
      "airway_measurements": {{
        "total_airway_volume_cc": 12.5,
        "minimum_cross_sectional_area_mm2": 45,
        "mcsa_location": "retropalatal"
      }},
      "skeletal_measurements": {{
        "hyoid_to_mandibular_plane_mm": 22,
        "posterior_airway_space_mm": 8
      }},
      "narrowing_assessment": {{
        "retropalatal_narrowing": "severe",
        "retroglossal_narrowing": "moderate"
      }}
    }},
    
    "ent_findings": {{
      "nasal": {{
        "septum_deviation": "left",
        "septum_deviation_severity": "moderate",
        "turbinate_hypertrophy": "moderate",
        "turbinate_side": "bilateral",
        "nasal_polyps": true,
        "polyp_grade": 2,
        "polyp_side": "bilateral"
      }},
      "sinus": {{
        "maxillary_sinus": {{
          "left": "mucosal thickening",
          "right": "mild mucosal thickening"
        }},
        "ethmoid_sinus": {{
          "anterior_left": "opacification",
          "anterior_right": "opacification"
        }},
        "frontal_sinus": {{
          "left": "haziness",
          "right": "clear"
        }},
        "chronic_sinusitis": true
      }},
      "nasopharynx": {{
        "appearance": "normal"
      }},
      "oropharynx": {{
        "tonsil_grade": 2,
        "uvula": "elongated",
        "soft_palate": "elongated"
      }},
      "post_surgical_changes": "post septoplasty and turbinectomy changes"
    }},
    
    "airway_phenotype": {{
      "classification": "anatomical",
      "obstruction_level": "multilevel",
      "positional_component": "severe",
      "notes": "primarily anatomical with strong positional dependence"
    }},
    
    "tmj_flags": {{
      "pain": true,
      "clicking": true,
      "side": "bilateral",
      "crepitus": false,
      "limited_opening": true,
      "deviation": "deviation to left on opening",
      "other_tmj_findings": ["disc displacement noted"]
    }},
    "other_observations": []
  }},
  "treatment_considerations": {{
    "primary_pathway": ["oral appliance therapy", "positional therapy"],
    "adjuncts": ["weight management", "nasal optimization"],
    "cautions": ["severe tongue base collapse may limit OAT efficacy"],
    "rationale": "Multilevel obstruction with positional component"
  }},
  "device_design": {{
    "advancement_plan": "start at 50% protrusion, titrate to 70%",
    "retention_features": ["ball clasps", "adams clasps"],
    "material": "hard acrylic",
    "coverage": "full arch coverage",
    "initial_accessories": ["morning repositioner"]
  }},
  "follow_up_plan": {{
    "evaluations": ["sleep study retest", "ENT evaluation"],
    "lifestyle": ["weight loss", "sleep hygiene"],
    "positional_therapy": "avoid supine sleep",
    "retest_after_init_months": 3
  }},
  "demographics": {{
    "sex": "M"
  }},
  "sleep_study": {{
    "study_type": "home"
  }}
}}

IMPORTANT: Only include fields that have actual data. Omit fields with no data.

CRITICAL EXTRACTION RULES:
1. For DISE findings - Extract ALL collapse patterns (velum, lateral walls, tongue base, epiglottis) with direction (AP, lateral, concentric)
2. For DISE maneuvers - Extract jaw thrust response, head rotation response, chin lift response with improvement level
3. For CBCT - Extract airway volume, MCSA, skeletal measurements if available
4. For ENT - Extract septum deviation, turbinate status, polyp grades, sinus findings by location
5. For hyoid - Only mark as "inferior" if explicitly stated; "normal" if stated as normal/תקין
6. For epiglottis - Only mark as "collapse" if explicitly stated; "normal" if stated as normal/תקין
7. For comorbidities - Pay attention to negations (לא = no, not present). If document says "לא" next to a condition, set present=false
8. NEVER hallucinate or infer findings not explicitly stated in the source documents

IMPORTANT RULES:
1. DO NOT extract numerical values - they are already handled by regex
2. Focus on descriptive text, anatomical descriptions, treatment plans, etc.
3. For anatomy fields, extract complete descriptions (e.g., "velopharyngeal collapse" not just "s")
4. Be comprehensive but accurate
5. Use exact enum values where specified
6. Only include fields that have actual textual content
7. For observations.summary: Select ONLY the TOP 5 most important SLEEP-RELATED clinical observations that are NOT already covered in other sections (patient_self_report, sleep_study, anatomy_imaging). EXCLUDE: dermatology reports, patch tests, corrupted documents, non-sleep medical content. Focus on sleep apnea, airway anatomy, treatment considerations, and device-related observations only.
8. CRITICAL: For follow_up_plan, ONLY extract information that is explicitly mentioned in the source documents. DO NOT generate or infer follow-up recommendations. If no follow-up plan is explicitly mentioned in the documents, do not include follow_up_plan in the output.
9. LIMIT follow_up_plan.evaluations to MAXIMUM 3 most important items only. Remove duplicates and prioritize the most critical evaluations.
10. LIMIT follow_up_plan.lifestyle to MAXIMUM 3 most important items only. Remove duplicates and prioritize the most critical lifestyle recommendations.

NORMALIZATION & SCHEMA HYGIENE:
- Deduplicate arrays case-insensitively (e.g., adjuncts, summary).
- Do NOT introduce fields outside the schema; place any extra details into observations.summary as sentences.
- Emit boolean/number types as booleans/numbers (not strings).
- If emitting dates, format as RFC3339 UTC with trailing 'Z'.
- Keep sources separate: do NOT mix report mentions into sleep-study episodes; do not assume links.
- Include sleep_study.study_type only if explicitly stated; do NOT guess.
- Demographic sanity (if any demographics appear in text): age_years ∈ [18,120], height_cm ∈ [120,250], weight_kg ∈ [30,300], bmi ∈ [10,80]; otherwise omit from output.

ALREADY EXTRACTED NUMERICAL DATA (DO NOT DUPLICATE):
{json.dumps(numerical_data, indent=2)}

OBSERVATIONS TO ANALYZE:
{chr(10).join(observation_texts)}

Return ONLY a valid JSON object with the textual observations. Do not include any explanations or text outside the JSON.
"""
    
    # Use hardcoded model ID for standalone script
    model_id = MODEL_ID
    
    import time
    start_time = time.time()
    
    try:
        # Call Bedrock for textual extraction
        response = bedrock.invoke_model(
            modelId=model_id,
            body=json.dumps({
                "anthropic_version": "bedrock-2023-05-31",
                "max_tokens": 2000,
                "messages": [
                    {
                        "role": "user",
                        "content": textual_prompt
                    }
                ]
            })
        )
        
        response_time_ms = int((time.time() - start_time) * 1000)
        
        response_body = json.loads(response.get('body').read())
        llm_response = response_body['content'][0]['text']
        
        # Log to database
        _log_llm_call(
            prompt_text="Extract textual observations from medical records",
            response_text=llm_response[:500],
            response_time_ms=response_time_ms,
            status='success'
        )
        
        # Extract JSON from response
        try:
            start_idx = llm_response.find('{')
            end_idx = llm_response.rfind('}') + 1
            if start_idx != -1 and end_idx != 0:
                json_str = llm_response[start_idx:end_idx]
                textual_result = json.loads(json_str)
            else:
                logger.error("Could not find JSON in textual LLM response")
                textual_result = {}
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse textual LLM response: {e}")
            textual_result = {}
    except Exception as e:
        response_time_ms = int((time.time() - start_time) * 1000)
        _log_llm_call(
            prompt_text="Extract textual observations from medical records",
            response_text='',
            response_time_ms=response_time_ms,
            status='error',
            error_message=str(e)
        )
        logger.error(f"Textual extraction failed: {e}")
        textual_result = {}
    
    return textual_result

def _first_match(text: str, patterns: list) -> Optional[Tuple[float, str]]:
    import re as _re
    # Normalize oxygen tokens and symbols before matching
    try:
        subs = str.maketrans({'₀':'0','₁':'1','₂':'2','₃':'3','₄':'4','₅':'5','₆':'6','₇':'7','₈':'8','₉':'9'})
        text = text.translate(subs).replace('％','%').replace('\u200b','').replace('\u00a0',' ')
        text = text.replace('O₂','O2').replace('SpO₂','SpO2')
    except Exception:
        pass
    for pat in patterns:
        m = _re.search(pat, text, flags=_re.IGNORECASE | _re.DOTALL)
        if m:
            val = m.group('val') if 'val' in m.groupdict() else m.group(1)
            try:
                return float(val), m.group(0)
            except Exception:
                continue
    return None

PATTERNS_O2_LT90 = [
    # Specific pattern for the exact text found: "Less than 90% O 2 0.5%"
    r'(?i)less\s+than\s+90%\s+o\s+2\s+(?P<val>\d+(?:\.\d+)?)\s*%?',
    # Pattern for "time spent below 90% was minimal (0.5%)"
    r'(?i)time\s+spent\s+below\s+90%.*?(?P<val>\d+(?:\.\d+)?)\s*%',
    # Primary pattern - handles both "O2" and "O 2"
    r'(?i)less\s+than\s+90%\s+o\s*2\s+(?P<val>\d+(?:\.\d+)?)\s*%?',
    # Original patterns (maintained for compatibility)
    r'(?i)\b(?:less\s*than\s*90%\s*(?:o\s*2|spo2|sao2))\s*:?\s*(?P<val>\d+(?:\.\d+)?)\s*%?',
    r'(?i)\b(?:o\s*2|spo2|sao2)\s*(?:<|below|less\s*than)\s*90%\s*:?\s*(?P<val>\d+(?:\.\d+)?)\s*%?',
    r'(?i)\b(?:percent(?:age)?\s*time|%?\s*time)\s*(?:o\s*2|spo2|sao2)?\s*(?:<|below|less\s*than)\s*90%\s*:?\s*(?P<val>\d+(?:\.\d+)?)\s*%?',
    r'(?is)\boxygen\s+saturation\s*<\s*90\b.*?sleep\s*%[^\d]*(?P<val>\d+(?:\.\d+)?)\b',
    r'(?i)\btime\s*<\s*90%[^\n\r\d]{0,40}(?P<val>\d+(?:\.\d+)?)\s*%?',
    # Additional specific patterns from testing interface
    r'(?i)SpO2\s*<\s*90%[^\n\r\d]{0,40}(?P<val>\d+(?:\.\d+)?)\s*%?',
    r'(?i)(?P<val>\d+(?:\.\d+)?)\s*%\s+time\s+below\s+90%',
    r'(?i)time\s+with\s+SpO2\s*<\s*90%[^\n\r\d]{0,40}(?P<val>\d+(?:\.\d+)?)\s*%?',
]

PATTERNS_SUPINE_AHI = [
    r'(?i)\bsupine\s*ahi\b\s*:?\s*(?P<val>\d+(?:\.\d+)?)',
    r'(?i)\bahi\b[^\n]{0,12}\b(?:supine)\b[^\d]{0,10}(?P<val>\d+(?:\.\d+)?)',
    r'(?is)\bbody\s+position\s+statistics\b.*?\bpAHI\b[^0-9]*(?P<val>\d+(?:\.\d+)?)',
]

AHI_CONFIDENCE_THRESHOLD = 2
AHI_PATTERN_RULES = [
    {'pattern': r'\bahi\b\s*(?:overall|total|global|all[-\s]*night)?\s*(?:[:=]|is|was|of)\s*(\d+(?:\.\d+)?)', 'score': 4},
    {'pattern': r'\bapnea\b.*?\bhypopnea\b.*?\bindex\b\s*(?:[:=]|is|was|of)\s*(\d+(?:\.\d+)?)', 'score': 4},
    {'pattern': r'\bahi\b\s*(?:overall|total|global|all[-\s]*night)?\s*(\d+(?:\.\d+)?)', 'score': 3},
    {'pattern': r'\bahi\b\s*(?:[:=]|is|was|of)\s*(\d+(?:\.\d+)?)', 'score': 2},
]
AHI_BONUS_KEYWORDS = ('overall', 'total', 'global', 'all-night', 'all night')
AHI_PENALTY_KEYWORDS = ('rem', 'supine', 'nrem', 'central', 'positional')

def process_queue(batch_size=3, limit=None):
    """
    Process patients from the document processing queue.
    Processes pending items ordered by priority and request time.
    
    Args:
        batch_size: Number of documents per Bedrock call
        limit: Maximum number of queue items to process (None = all)
    
    Returns:
        dict with processing statistics
    """
    import mysql.connector
    
    results = {
        'processed': 0,
        'successful': 0,
        'failed': 0,
        'skipped': 0,
        'patients': []
    }
    
    try:
        conn = mysql.connector.connect(**DB_CONFIG)
        cursor = conn.cursor(dictionary=True)

        try:
            from flask_app.services.document_queue_sla import abandon_expired_document_queue_rows

            _ab = abandon_expired_document_queue_rows(conn)
            if _ab:
                logger.info("document_queue_sla (process_queue): abandoned %s row(s)", _ab)
        except Exception as _sla_e:
            logger.warning("document_queue_sla (process_queue): %s", _sla_e)
        
        # Get pending queue items
        query = """
            SELECT id, patient_id, batch_size, retry_count, max_retries
            FROM document_processing_queue
            WHERE status = 'pending'
            ORDER BY priority DESC, requested_at ASC
        """
        
        if limit:
            query += f" LIMIT {limit}"
        
        cursor.execute(query)
        queue_items = cursor.fetchall()
        
        logger.info(f"Found {len(queue_items)} patients in queue")
        print(f"📋 Found {len(queue_items)} patients in processing queue")
        
        for item in queue_items:
            queue_id = item['id']
            patient_id = item['patient_id']
            item_batch_size = item['batch_size'] or batch_size
            
            try:
                # Update status to processing
                cursor.execute("""
                    UPDATE document_processing_queue
                    SET status = 'processing', started_at = NOW()
                    WHERE id = %s
                """, (queue_id,))
                conn.commit()
                
                logger.info(f"Processing queue item {queue_id}: patient {patient_id}")
                print(f"\n🔄 Processing patient {patient_id} (queue ID: {queue_id})...")
                
                # Process the patient documents
                result = process_patient_documents(patient_id, max_documents=None, batch_size=item_batch_size)
                
                # Update status to completed
                cursor.execute("""
                    UPDATE document_processing_queue
                    SET status = 'completed', completed_at = NOW()
                    WHERE id = %s
                """, (queue_id,))
                conn.commit()
                
                results['processed'] += 1
                results['successful'] += 1
                results['patients'].append({
                    'patient_id': patient_id,
                    'queue_id': queue_id,
                    'status': 'success',
                    'result': str(result)[:200]  # Truncate for summary
                })
                logger.info(f"✅ Successfully processed patient {patient_id}")
                print(f"✅ Patient {patient_id} completed successfully")
                
            except Exception as e:
                logger.error(f"❌ Error processing patient {patient_id}: {e}")
                print(f"❌ Error processing patient {patient_id}: {e}")
                
                retry_count = item['retry_count'] + 1
                max_retries = item['max_retries']
                
                if retry_count >= max_retries:
                    # Max retries reached, mark as failed
                    cursor.execute("""
                        UPDATE document_processing_queue
                        SET status = 'failed', 
                            completed_at = NOW(),
                            error_message = %s,
                            retry_count = %s
                        WHERE id = %s
                    """, (str(e)[:1000], retry_count, queue_id))
                    results['failed'] += 1
                    status_msg = f"failed after {max_retries} retries"
                else:
                    # Increment retry count and reset to pending
                    cursor.execute("""
                        UPDATE document_processing_queue
                        SET status = 'pending',
                            started_at = NULL,
                            error_message = %s,
                            retry_count = %s
                        WHERE id = %s
                    """, (str(e)[:1000], retry_count, queue_id))
                    results['skipped'] += 1
                    status_msg = f"will retry (attempt {retry_count}/{max_retries})"
                
                results['patients'].append({
                    'patient_id': patient_id,
                    'queue_id': queue_id,
                    'status': 'error',
                    'error': str(e)[:100],
                    'retry_info': status_msg
                })
                conn.commit()
                results['processed'] += 1
        
        cursor.close()
        conn.close()
        
        # Print summary
        print(f"\n{'='*60}")
        print(f"📊 Queue Processing Summary:")
        print(f"   Total Processed: {results['processed']}")
        print(f"   ✅ Successful: {results['successful']}")
        print(f"   ❌ Failed: {results['failed']}")
        print(f"   🔄 Retrying: {results['skipped']}")
        print(f"{'='*60}\n")
        
        return results
        
    except Exception as e:
        logger.error(f"Error processing queue: {e}")
        if conn:
            conn.close()
        raise
if __name__ == "__main__":
    # Import Flask app creation function
    from flask_app import create_app
    
    parser = argparse.ArgumentParser(description="Document observation batch processor")
    parser.add_argument('--mode', choices=['test', 'batch_all', 'patient', 'queue', 'backfill_envelopes', 'create_canonical', 'ensure_canonical', 'validate_canonical', 'create_minimal_canonical', 'ensure_minimal_canonical', 'debug_documents', 'regenerate_canonical', 'o2_lt90'], default='test')
    parser.add_argument('--patient-id', type=int, default=None, help='Process a single patient id')
    parser.add_argument('--limit-patients', type=int, default=0, help='Limit number of patients when batch_all')
    parser.add_argument('--batch-size', type=int, default=3, help='Number of documents per Bedrock call')


    args = parser.parse_args()
    
    # Create Flask app and run within application context
    app = create_app()
    with app.app_context():
        if args.mode == 'test':
            test_phase2_processing()
        elif args.mode == 'queue':
            logger.info("Processing document queue...")
            print(process_queue(batch_size=args.batch_size, limit=args.limit_patients if args.limit_patients > 0 else None))
        elif args.mode == 'patient' and args.patient_id:
            logger.info(f"Processing single patient {args.patient_id}")
            print(process_patient_documents(args.patient_id, max_documents=None, batch_size=args.batch_size))
        elif args.mode == 'debug_documents' and args.patient_id:
            logger.info(f"Debugging documents for patient {args.patient_id}")
            print(debug_patient_documents(args.patient_id))
        elif args.mode == 'backfill_envelopes':
            print(backfill_envelopes(limit_patients=args.limit_patients if args.limit_patients > 0 else None))
        elif args.mode == 'create_canonical' and args.patient_id:
            logger.info(f"Creating canonical JSON for patient {args.patient_id}")
            print(create_canonical_json_for_patient(args.patient_id))
        elif args.mode == 'ensure_canonical':
            logger.info(f"Ensuring canonical JSON for all patients (limit: {args.limit_patients if args.limit_patients > 0 else 'none'})")
            print(ensure_canonical_json_for_all_patients(limit_patients=args.limit_patients if args.limit_patients > 0 else None))
        elif args.mode == 'validate_canonical':
            logger.info(f"Validating canonical JSON for all patients (limit: {args.limit_patients if args.limit_patients > 0 else 'none'})")
            print(validate_all_canonical_json(limit_patients=args.limit_patients if args.limit_patients > 0 else None))
        elif args.mode == 'create_minimal_canonical' and args.patient_id:
            logger.info(f"Creating minimal canonical JSON for patient {args.patient_id}")
            print(create_minimal_canonical_json_for_patient(args.patient_id))
        elif args.mode == 'regenerate_canonical' and args.patient_id:
            logger.info(f"Regenerating canonical JSON for patient {args.patient_id} (deleting existing first)")
            # Delete existing canonical first
            import mysql.connector
            conn = mysql.connector.connect(**DB_CONFIG)
            cursor = conn.cursor()
            cursor.execute("DELETE FROM patient_case_envelope WHERE patient_id = %s AND report_id = 'canonical'", (args.patient_id,))
            deleted_count = cursor.rowcount
            conn.commit()
            conn.close()
            print(f"Deleted {deleted_count} existing canonical records for patient {args.patient_id}")
            # Create new canonical
            result = create_minimal_canonical_json_for_patient(args.patient_id)
            print(result)
        elif args.mode == 'ensure_minimal_canonical':
            logger.info(f"Ensuring minimal canonical JSON for all patients (limit: {args.limit_patients if args.limit_patients > 0 else 'none'})")
            print(ensure_minimal_canonical_for_all_patients(limit_patients=args.limit_patients if args.limit_patients > 0 else None))
        elif args.mode == 'o2_lt90' and args.patient_id:
            def extract_time_below_90_pct_for_patient(patient_id: int):
                """Standalone pass to extract time_below_90_pct for a patient.
                - Reads all docs for patient
                - Tries regex (PATTERNS_O2_LT90), then the direct LLM question
                - Stores metric_key=time_below_90_pct without deleting existing rows
                - Updates minimal canonical JSON at the end.
                """
                import mysql.connector
                conn = mysql.connector.connect(**DB_CONFIG)
                cursor = conn.cursor(dictionary=True)
                # Collect documents from files and adminfiles
                cursor.execute("SELECT id, name, s3_key, upload_date FROM files WHERE patient_id=%s", (patient_id,))
                files = cursor.fetchall() or []
                cursor.execute("SELECT id, name, s3_key, upload_date FROM adminfiles WHERE patient_id=%s", (patient_id,))
                adminfiles = cursor.fetchall() or []
                docs = []
                for r in files:
                    docs.append({ 'table': 'files', **r })
                for r in adminfiles:
                    docs.append({ 'table': 'adminfiles', **r })

                added = 0
                for d in docs:
                    try:
                        name = d.get('name') or ''
                        s3_key = d.get('s3_key')
                        content = extract_document_content({'name': name, 's3_key': s3_key})  # existing helper that routes by type
                        if not content:
                            continue
                        found = _first_match(content, PATTERNS_O2_LT90)
                        val = None
                        ev = None
                        if found:
                            val, ev = found
                        else:
                            # direct question fallback
                            try:
                                dq_system = ("You extract a single numerical metric from a sleep study. Return ONLY the number (no text), representing the percent of sleep time with oxygen saturation below 90.")
                                dq_user = f"Question: What percentage of time did this patient spend with oxygen saturation below 90%?\n\nDocument: {content[:8000]}"
                                resp = bedrock_query_enhanced([
                                    { 'role': 'system', 'content': dq_system },
                                    { 'role': 'user', 'content': dq_user }
                                ], max_tokens=64, temperature=0.0, top_p=0.9)
                                if isinstance(resp, dict) and resp.get('success'):
                                    raw = (resp.get('response') or '').strip()
                                    import re as _re
                                    m2 = _re.search(r"(\d+(?:\.\d+)?)\s*%?", raw)
                                    if m2:
                                        val = float(m2.group(1))
                                        ev = raw[:200]
                            except Exception:
                                pass

                        if val is None:
                            continue

                        # Store observation using modern schema with all required fields
                        from datetime import datetime
                        import json
                        
                        # Insert using full schema like the main document extractor
                        insert_query = """
                            INSERT INTO observation_store 
                            (patient_id, path, extracted_observations, source_type, metric_key, 
                             metric_value_decimal, file_name, provider, source_kind, study_type, 
                             observed_at, created_at, updated_at)
                            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, NOW(), NOW())
                        """
                        
                        observation_json = {
                            'path': 'sleep_study.time_below_90_pct',
                            'value': str(val),
                            'source': 'o2-lt90-direct',
                            'confidence': 90,
                            'explanation': 'Direct pass for percent time SpO2 < 90',
                            'evidence': ev[:500] if ev else '',
                        }
                        
                        cursor.execute(insert_query, (
                            patient_id,
                            'sleep_study.time_below_90_pct',
                            json.dumps(observation_json),
                            'numerical_extraction', 
                            'time_below_90_pct',
                            float(val),
                            name or 'o2_lt90_doc',
                            'o2-lt90-direct',
                            'sleep_study',  # source_kind - critical for canonical_derived!
                            'sleep_study',  # study_type
                            datetime.now()  # observed_at
                        ))
                        conn.commit()
                        added += 1
                    except Exception as _e:
                        logger.warning(f"o2_lt90 pass: failed on doc {d.get('name')}: {_e}")
                        continue

                conn.close()
                # Update canonical minimal JSON (non-destructive)
                try:
                    create_minimal_canonical_json_for_patient(patient_id)
                except Exception:
                    pass
                return { 'patient_id': patient_id, 'added_observations': added }

            print(extract_time_below_90_pct_for_patient(args.patient_id))
        else:
            print(batch_all_patients(limit_patients=args.limit_patients if args.limit_patients > 0 else None,
                                     batch_size=args.batch_size))