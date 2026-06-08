from flask import Blueprint, request, render_template, jsonify, current_app
import json
import re
import os
import tempfile
from typing import Optional
from werkzeug.utils import secure_filename

# Create blueprint
o2_extraction = Blueprint('o2_extraction', __name__, url_prefix='/o2-extraction')

@o2_extraction.route('/')
def index():
    """Main O2 extraction testing page"""
    return render_template('o2_extraction_test.html')

@o2_extraction.route('/load-sample', methods=['POST'])
def load_sample():
    """Load a sample sleep study report for testing"""
    sample_report = """Mallampati IV, Modified Mallampati IV, hypertrophic muscular tongue base, 100% 
concentric soft palate collapse on Müller maneuver, nasal obstruction, mouth breathing.

Sleep Study Data
AHI 28.2 (Moderate OSA)
RDI 32.8
Supine AHI 42.5 (62% of sleep time) Supine RDI 45.7
REM AHI 30.1 (27.5% of sleep time) REM RDI 36.8
O2 Nadir 83% ODI 16.1
Less than 90% O2 0.5% Supine ODI 25.6
Snoring 56%, >50dB=13.2% REM ODI 19.7"""
    
    return jsonify({"success": True, "sample_text": sample_report})

@o2_extraction.route('/test-regex', methods=['POST'])
def test_regex_patterns():
    """Test regex patterns against the provided text"""
    try:
        from flask_app.config.extract_o2_lt90_only import normalize_text
        
        data = request.get_json()
        text = data.get('text', '')
        
        if not text:
            return jsonify({"success": False, "error": "No text provided"})
        
        # Normalize the text first
        normalized_text = normalize_text(text)
        
        # Test various regex patterns for O2 < 90%
        patterns = [
            (r"Less than 90%\s+O2\s+(\d+(?:\.\d+)?)\s*%?", "Pattern 1: 'Less than 90% O2 X%'"),
            (r"time\s*<\s*90%[^\n\r\d]{0,40}(\d+(?:\.\d+)?)\s*%?", "Pattern 2: 'time < 90% ... X%'"),
            (r"SpO2\s*<\s*90%[^\n\r\d]{0,40}(\d+(?:\.\d+)?)\s*%?", "Pattern 3: 'SpO2 < 90% ... X%'"),
            (r"oxygen\s+saturation\s+below\s+90%[^\n\r\d]{0,40}(\d+(?:\.\d+)?)\s*%?", "Pattern 4: 'oxygen saturation below 90% ... X%'"),
            (r"(\d+(?:\.\d+)?)\s*%\s+time\s+below\s+90%", "Pattern 5: 'X% time below 90%'"),
            (r"time\s+with\s+SpO2\s*<\s*90%[^\n\r\d]{0,40}(\d+(?:\.\d+)?)\s*%?", "Pattern 6: 'time with SpO2 < 90% ... X%'"),
        ]
        
        results = []
        for pattern, description in patterns:
            matches = re.findall(pattern, normalized_text, re.IGNORECASE)
            results.append({
                "pattern": pattern,
                "description": description,
                "matches": matches,
                "found": len(matches) > 0
            })
        
        return jsonify({
            "success": True,
            "normalized_text": normalized_text,
            "results": results
        })
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})

@o2_extraction.route('/simulate-llm', methods=['POST'])
def simulate_llm_extraction():
    """LLM comprehensive document analysis for all sleep study parameters"""
    try:
        from flask_app.config.bedrock_config import query_bedrock_claude_enhanced as bedrock_query_enhanced
        from flask_app.config.extract_o2_lt90_only import normalize_text
        
        data = request.get_json()
        text = data.get('text', '')
        filename = data.get('filename', 'document')
        
        if not text:
            return jsonify({"success": False, "error": "No text provided"})
        
        # Normalize the text
        normalized_text = normalize_text(text)
        
        # Comprehensive LLM analysis
        analysis_result = comprehensive_sleep_study_analysis(normalized_text, filename)
        
        return jsonify({
            "success": True,
            "normalized_text": normalized_text[:1000] + "..." if len(normalized_text) > 1000 else normalized_text,
            "analysis_result": analysis_result
        })
        
    except Exception as e:
        current_app.logger.error(f"Error in simulate_llm_extraction: {str(e)}")
        return jsonify({"success": False, "error": str(e)})

def comprehensive_sleep_study_analysis(text: str, filename: str) -> dict:
    """Perform comprehensive LLM analysis of sleep study document"""
    try:
        # Import with better error handling
        try:
            from flask_app.config.bedrock_config import query_bedrock_claude_enhanced as bedrock_query_enhanced
        except ImportError as e:
            current_app.logger.error(f"Failed to import bedrock function: {e}")
            return {
                "success": False,
                "error": f"LLM service not available: {str(e)}",
                "method": "Comprehensive LLM Document Analysis"
            }
        
        system_prompt = """You are a medical data extraction specialist analyzing sleep study reports. 
        Extract ALL relevant sleep study metrics and parameters from the provided document.
        
        Return a JSON object with the following structure:
        {
            "patient_demographics": {
                "age": "number or null",
                "sex": "M/F or null",
                "bmi": "number or null",
                "height": "string or null",
                "weight": "string or null"
            },
            "sleep_metrics": {
                "ahi": "number or null",
                "rdi": "number or null", 
                "odi": "number or null",
                "supine_ahi": "number or null",
                "rem_ahi": "number or null",
                "nrem_ahi": "number or null"
            },
            "oxygen_metrics": {
                "o2_nadir": "number or null",
                "time_below_90_percent": "number or null",
                "time_below_85_percent": "number or null",
                "time_below_80_percent": "number or null",
                "average_spo2": "number or null",
                "baseline_spo2": "number or null"
            },
            "sleep_architecture": {
                "total_sleep_time": "number or null",
                "sleep_efficiency": "number or null",
                "rem_percent": "number or null",
                "deep_sleep_percent": "number or null",
                "light_sleep_percent": "number or null",
                "sleep_onset_latency": "number or null",
                "rem_latency": "number or null"
            },
            "respiratory_events": {
                "total_apneas": "number or null",
                "total_hypopneas": "number or null",
                "central_apneas": "number or null",
                "obstructive_apneas": "number or null",
                "mixed_apneas": "number or null"
            },
            "other_metrics": {
                "arousal_index": "number or null",
                "limb_movement_index": "number or null",
                "snoring_events": "number or null",
                "heart_rate_min": "number or null",
                "heart_rate_max": "number or null",
                "heart_rate_avg": "number or null"
            },
            "diagnosis": {
                "severity": "Normal/Mild/Moderate/Severe or null",
                "primary_diagnosis": "string or null",
                "recommendations": "string or null"
            },
            "study_info": {
                "study_date": "string or null",
                "study_type": "string or null",
                "recording_time": "string or null",
                "analysis_software": "string or null"
            }
        }
        
        IMPORTANT: 
        - Extract exact numerical values when found
        - Use null for any metrics not found in the document
        - Pay special attention to percentages and units
        - For "time_below_90_percent", look for phrases like "Less than 90% O2", "time < 90%", "SpO2 < 90%"
        """
        
        user_prompt = f"""Analyze this sleep study document and extract all metrics:

Document: {text[:8000]}

Return ONLY the JSON object with extracted values."""
        
        # Combine system and user prompts for Bedrock compatibility
        combined_prompt = f"{system_prompt}\n\n{user_prompt}"
        messages = [
            {"role": "user", "content": combined_prompt}
        ]
        
        response = bedrock_query_enhanced(messages, max_tokens=1500, temperature=0.1, top_p=0.9)
        
        if isinstance(response, dict) and response.get("success"):
            raw_response = response.get("response", "")
            try:
                import json
                # Try to parse the JSON response
                extracted_data = json.loads(raw_response)
                
                return {
                    "success": True,
                    "extracted_metrics": extracted_data,
                    "raw_llm_response": raw_response,
                    "method": "Comprehensive LLM Document Analysis"
                }
            except json.JSONDecodeError as e:
                return {
                    "success": False,
                    "error": f"Failed to parse LLM JSON response: {str(e)}",
                    "raw_llm_response": raw_response,
                    "method": "Comprehensive LLM Document Analysis"
                }
        else:
            return {
                "success": False,
                "error": "LLM query failed",
                "method": "Comprehensive LLM Document Analysis"
            }
            
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "method": "Comprehensive LLM Document Analysis"
        }

@o2_extraction.route('/run-full-test', methods=['POST'])
def run_full_test():
    """Run the full O2 extraction test on a patient"""
    try:
        from flask_app.config.extract_o2_lt90_only import run as run_o2_extraction
        
        data = request.get_json()
        patient_id = data.get('patient_id')
        
        if not patient_id:
            return jsonify({"success": False, "error": "No patient ID provided"})
        
        try:
            patient_id = int(patient_id)
        except ValueError:
            return jsonify({"success": False, "error": "Invalid patient ID format"})
        
        # Run the extraction
        result = run_o2_extraction(patient_id)
        
        return jsonify({
            "success": True,
            "result": result
        })
        
    except Exception as e:
        current_app.logger.error(f"Error in run_full_test: {str(e)}")
        return jsonify({"success": False, "error": str(e)})

@o2_extraction.route('/test-document-extraction', methods=['POST'])
def test_document_extraction():
    """Test document content extraction from S3"""
    try:
        from flask_app.config.document_observation_extractor_phase2 import extract_document_content
        from flask_app.config.extract_o2_lt90_only import normalize_text
        
        data = request.get_json()
        s3_key = data.get('s3_key', '')
        file_type = data.get('file_type', 'application/pdf')
        file_name = data.get('file_name', 'test_document')
        
        if not s3_key:
            return jsonify({"success": False, "error": "No S3 key provided"})
        
        # Create document dict for extraction
        document = {
            's3_key': s3_key,
            'file_type': file_type,
            'name': file_name
        }
        
        # Extract document content
        content = extract_document_content(document)
        
        if content:
            normalized_content = normalize_text(content)
            return jsonify({
                "success": True,
                "raw_content": content[:2000] + "..." if len(content) > 2000 else content,
                "normalized_content": normalized_content[:2000] + "..." if len(normalized_content) > 2000 else normalized_content,
                "content_length": len(content)
            })
        else:
            return jsonify({"success": False, "error": "Failed to extract document content"})
        
    except Exception as e:
        current_app.logger.error(f"Error in test_document_extraction: {str(e)}")
        return jsonify({"success": False, "error": str(e)})

@o2_extraction.route('/upload-file', methods=['POST'])
def upload_file():
    """Handle local file upload for testing"""
    try:
        if 'file' not in request.files:
            return jsonify({"success": False, "error": "No file uploaded"})
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({"success": False, "error": "No file selected"})
        
        # Save file temporarily
        filename = secure_filename(file.filename)
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(filename)[1]) as tmp_file:
            file.save(tmp_file.name)
            
            # Extract content based on file type
            content = extract_content_from_local_file(tmp_file.name, filename)
            
            # Clean up temporary file
            os.unlink(tmp_file.name)
            
            if content:
                from flask_app.config.extract_o2_lt90_only import normalize_text
                normalized_content = normalize_text(content)
                
                return jsonify({
                    "success": True,
                    "filename": filename,
                    "raw_content": content[:2000] + "..." if len(content) > 2000 else content,
                    "normalized_content": normalized_content[:2000] + "..." if len(normalized_content) > 2000 else normalized_content,
                    "full_normalized_content": normalized_content,  # For extraction testing
                    "content_length": len(content)
                })
            else:
                return jsonify({"success": False, "error": "Failed to extract content from file"})
        
    except Exception as e:
        current_app.logger.error(f"Error in upload_file: {str(e)}")
        return jsonify({"success": False, "error": str(e)})

def extract_content_from_local_file(file_path: str, filename: str) -> str:
    """Extract content from a local file"""
    try:
        file_ext = os.path.splitext(filename)[1].lower()
        
        if file_ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        elif file_ext == '.pdf':
            import PyPDF2
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text
        elif file_ext in ['.doc', '.docx']:
            from docx import Document
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        elif file_ext in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
            # For images, we'd need OCR - for now, return empty
            return "Image file uploaded - OCR not implemented in local testing"
        else:
            # Try to read as text
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
                
    except Exception as e:
        current_app.logger.error(f"Error extracting content from {filename}: {str(e)}")
        return ""

@o2_extraction.route('/direct-llm-file', methods=['POST'])
def direct_llm_file():
    """Direct file to LLM analysis - no preprocessing, exactly like Claude"""
    try:
        from flask_app.config.bedrock_config import query_bedrock_claude_enhanced as bedrock_query_enhanced
        
        if 'file' not in request.files:
            return jsonify({"success": False, "error": "No file uploaded"})
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({"success": False, "error": "No file selected"})
        
        filename = file.filename
        
        # Extract content from the raw file
        content = ""
        try:
            if filename.lower().endswith('.txt'):
                content = file.read().decode('utf-8')
            elif filename.lower().endswith('.pdf'):
                import PyPDF2
                from io import BytesIO
                pdf_reader = PyPDF2.PdfReader(BytesIO(file.read()))
                for page in pdf_reader.pages:
                    content += page.extract_text() + "\n"
            elif filename.lower().endswith(('.doc', '.docx')):
                from docx import Document
                from io import BytesIO
                doc = Document(BytesIO(file.read()))
                for paragraph in doc.paragraphs:
                    content += paragraph.text + "\n"
            elif filename.lower().endswith(('.jpg', '.jpeg', '.png', '.tiff', '.tif', '.bmp')):
                # Image OCR extraction
                try:
                    import pytesseract
                    from PIL import Image
                    from io import BytesIO
                    
                    # Open image and perform OCR
                    image = Image.open(BytesIO(file.read()))
                    content = pytesseract.image_to_string(image)
                    
                    if not content.strip():
                        return jsonify({"success": False, "error": "No text could be extracted from image via OCR"})
                        
                except ImportError:
                    return jsonify({"success": False, "error": "OCR support not available. Install pytesseract and Pillow for image processing."})
                except Exception as ocr_error:
                    return jsonify({"success": False, "error": f"OCR extraction failed: {str(ocr_error)}"})
            else:
                # Try to read as text
                content = file.read().decode('utf-8', errors='ignore')
        except Exception as e:
            return jsonify({"success": False, "error": f"Could not read file: {str(e)}"})
        
        if not content.strip():
            return jsonify({"success": False, "error": "No content could be extracted from file"})
        
        # Send directly to LLM - exactly like Claude interface
        system_instructions = """You are a sleep medicine specialist. Analyze this sleep study document and extract ALL sleep metrics WITH DATES.

        CRITICAL: Look for comparison tables, baseline vs follow-up data, and temporal progression.
        Extract EACH time point as a separate entry with its date.

        CRITICAL: Return ONLY a valid JSON object, no explanatory text before or after.

        JSON format - array of time points:
        [
            {
                "date": "YYYY-MM-DD or descriptive (e.g., 'Baseline 2024', 'Follow-up #1 June 2025')",
                "study_type": "baseline" or "follow_up" or "unknown",
                [DYNAMIC METRICS - extract whatever sleep metrics you find in the document]
            }
        ]

        METRIC EXTRACTION GUIDELINES:
        - Extract ANY sleep study metric you find (AHI, ODI, O2 nadir, sleep efficiency, etc.)
        - Use consistent field names: snake_case with descriptive names
        - Common metrics to look for (but not limited to):
          * ahi (Apnea-Hypopnea Index)
          * odi (Oxygen Desaturation Index) 
          * o2_nadir_pct (Oxygen nadir percentage)
          * time_below_90_pct (Time with O2 < 90%)
          * sleep_efficiency_pct (Sleep efficiency percentage)
          * total_sleep_time_hours (Total sleep time in hours)
          * rem_ahi (REM AHI)
          * supine_ahi (Supine AHI) 
          * arousal_index (Arousal index)
          * snoring_pct (Snoring percentage)
          * tst_sleep_efficiency (TST/Sleep efficiency)
        - If you find other metrics not listed above, include them with descriptive names
        - Use numbers for all metric values (convert percentages to numbers without % symbol)
        - Use null for metrics not found or not applicable for that time point

        TEMPORAL EXTRACTION INSTRUCTIONS:
        - If you find a comparison table (Baseline vs Follow-up), create separate entries for each column
        - Extract the actual dates or time references (e.g., "June 2025", "Aug 2025", "2024")
        - If only one time point, return array with single object
        - Pay attention to table headers and column titles
        - Look for temporal keywords: baseline, follow-up, before/after treatment, pre/post

        Return ONLY the JSON array above, nothing else."""
        
        user_prompt = f"""Analyze this sleep study document:

{content}

Extract all sleep study metrics and return as JSON."""
        
        # Combine for Bedrock compatibility
        combined_prompt = f"{system_instructions}\n\n{user_prompt}"
        messages = [{"role": "user", "content": combined_prompt}]
        
        current_app.logger.info(f"Direct file-to-LLM analysis for {filename}")
        response = bedrock_query_enhanced(messages, max_tokens=1500, temperature=0.1, top_p=0.9)
        
        if isinstance(response, dict) and response.get("success"):
            raw_response = response.get("response", "")
            
            try:
                import json
                import re
                
                # Extract JSON from the response (could be array or object)
                json_match = re.search(r'(\[.*\]|\{.*\})', raw_response, re.DOTALL)
                if json_match:
                    json_str = json_match.group(0)
                    extracted_data = json.loads(json_str)
                else:
                    # Fallback: try parsing the whole response
                    extracted_data = json.loads(raw_response)
                
                # Handle both array (temporal) and object (single time point) formats
                if isinstance(extracted_data, list):
                    # Array of time points
                    time_points = extracted_data
                    metrics_found = sum(
                        sum(1 for k, v in tp.items() if v is not None and k not in ['date', 'study_type'])
                        for tp in time_points
                    )
                else:
                    # Single object - convert to array format
                    time_points = [extracted_data]
                    metrics_found = sum(1 for v in extracted_data.values() if v is not None)
                
                extracted_data = time_points  # Always return as array
                
                return jsonify({
                    "success": True,
                    "method": "Direct File-to-LLM (Pure Claude-like)",
                    "filename": filename,
                    "metrics_found": metrics_found,
                    "extracted_data": extracted_data,
                    "raw_llm_response": raw_response,
                    "file_size": len(content)
                })
                
            except json.JSONDecodeError as e:
                return jsonify({
                    "success": False,
                    "error": f"LLM response was not valid JSON: {str(e)}",
                    "raw_llm_response": raw_response,
                    "method": "Direct File-to-LLM"
                })
        else:
            return jsonify({
                "success": False,
                "error": "LLM query failed",
                "method": "Direct File-to-LLM"
            })
            
    except Exception as e:
        current_app.logger.error(f"Error in direct_llm_file: {str(e)}")
        return jsonify({"success": False, "error": str(e)})

@o2_extraction.route('/direct-llm-analysis', methods=['POST'])
def direct_llm_analysis():
    """Direct LLM analysis like Claude - no preprocessing, pure document to LLM"""
    try:
        from flask_app.config.bedrock_config import query_bedrock_claude_enhanced as bedrock_query_enhanced
        
        data = request.get_json()
        text = data.get('text', '')
        filename = data.get('filename', 'document')
        
        if not text:
            return jsonify({"success": False, "error": "No text provided"})
        
        # Pure LLM analysis - exactly like Claude direct interface
        system_instructions = """You are a sleep medicine specialist analyzing a sleep study document. 
        Extract ALL sleep study metrics exactly as you would when reviewing a patient's report.
        
        Focus on finding these key metrics:
        - AHI (Apnea-Hypopnea Index)
        - ODI (Oxygen Desaturation Index) 
        - O2 Nadir (lowest oxygen saturation)
        - Time with SpO2 < 90% (percent of sleep time)
        - Sleep efficiency
        - Total sleep time
        - REM AHI, Supine AHI
        - Any other respiratory/sleep metrics
        
        Return a JSON object with all found metrics. Use null for metrics not found.
        Be thorough - look for metrics described in various ways throughout the document."""
        
        user_prompt = f"""Analyze this sleep study document and extract all sleep metrics:

{text}

Return JSON with all extracted sleep study metrics."""
        
        # Combine prompts for Bedrock compatibility
        combined_prompt = f"{system_instructions}\n\n{user_prompt}"
        messages = [{"role": "user", "content": combined_prompt}]
        
        current_app.logger.info(f"Direct LLM analysis for {filename}")
        response = bedrock_query_enhanced(messages, max_tokens=2000, temperature=0.1, top_p=0.9)
        
        if isinstance(response, dict) and response.get("success"):
            raw_response = response.get("response", "")
            
            try:
                import json
                # Try to parse JSON response
                extracted_data = json.loads(raw_response)
                
                # Count non-null metrics
                metrics_found = sum(1 for v in extracted_data.values() if v is not None)
                
                return jsonify({
                    "success": True,
                    "method": "Direct LLM Analysis (Claude-like)",
                    "filename": filename,
                    "metrics_found": metrics_found,
                    "extracted_data": extracted_data,
                    "raw_llm_response": raw_response
                })
                
            except json.JSONDecodeError as e:
                return jsonify({
                    "success": False,
                    "error": f"LLM response was not valid JSON: {str(e)}",
                    "raw_llm_response": raw_response,
                    "method": "Direct LLM Analysis (Claude-like)"
                })
        else:
            error_msg = "LLM query failed"
            if isinstance(response, dict):
                error_msg = response.get("error", error_msg)
            return jsonify({
                "success": False,
                "error": error_msg,
                "method": "Direct LLM Analysis (Claude-like)"
            })
            
    except Exception as e:
        current_app.logger.error(f"Error in direct_llm_analysis: {str(e)}")
        return jsonify({"success": False, "error": str(e)})

@o2_extraction.route('/extract-all-metrics', methods=['POST'])
def extract_all_metrics():
    """Extract all sleep study metrics from the provided text"""
    try:
        from flask_app.config.document_observation_extractor_phase2 import (
            extract_observations_with_llm,
            parse_observations_from_text,
            extract_specific_numerical_fields
        )
        from flask_app.config.extract_o2_lt90_only import normalize_text
        
        data = request.get_json()
        text = data.get('text', '')
        filename = data.get('filename', 'test_document')
        
        if not text:
            return jsonify({"success": False, "error": "No text provided"})
        
        # Normalize the text
        normalized_text = normalize_text(text)
        
        # Extract observations using the main LLM method
        observation_text = extract_observations_with_llm(normalized_text, 'sleep_study', filename)
        
        # Parse observations from the LLM response
        observations = parse_observations_from_text(observation_text)
        
        # Extract specific numerical fields
        numerical_data = extract_specific_numerical_fields([observation_text])
        
        # Test specific O2 patterns as well
        o2_patterns = test_o2_specific_patterns(normalized_text)
        
        return jsonify({
            "success": True,
            "filename": filename,
            "normalized_text": normalized_text[:1000] + "..." if len(normalized_text) > 1000 else normalized_text,
            "llm_observation_text": observation_text,
            "parsed_observations": observations,
            "numerical_data": numerical_data,
            "o2_specific_patterns": o2_patterns,
            "total_observations": len(observations)
        })
        
    except Exception as e:
        current_app.logger.error(f"Error in extract_all_metrics: {str(e)}")
        return jsonify({"success": False, "error": str(e)})

def test_o2_specific_patterns(text: str) -> dict:
    """Test specific O2 saturation patterns"""
    patterns = {
        "time_below_90": [
            (r"Less than 90%\s+O2\s+(\d+(?:\.\d+)?)\s*%?", "Less than 90% O2 X%"),
            (r"time\s*<\s*90%[^\n\r\d]{0,40}(\d+(?:\.\d+)?)\s*%?", "time < 90% ... X%"),
            (r"SpO2\s*<\s*90%[^\n\r\d]{0,40}(\d+(?:\.\d+)?)\s*%?", "SpO2 < 90% ... X%"),
            (r"(\d+(?:\.\d+)?)\s*%\s+time\s+below\s+90%", "X% time below 90%"),
        ],
        "o2_nadir": [
            (r"O2\s+Nadir\s+(\d+)\s*%?", "O2 Nadir X%"),
            (r"oxygen\s+nadir[^\n\r\d]{0,20}(\d+)\s*%?", "oxygen nadir ... X%"),
            (r"lowest\s+oxygen[^\n\r\d]{0,20}(\d+)\s*%?", "lowest oxygen ... X%"),
        ],
        "ahi": [
            (r"AHI\s+(\d+(?:\.\d+)?)", "AHI X"),
            (r"apnea[^h]*hypopnea\s+index[^\n\r\d]{0,20}(\d+(?:\.\d+)?)", "apnea hypopnea index ... X"),
        ],
        "rdi": [
            (r"RDI\s+(\d+(?:\.\d+)?)", "RDI X"),
            (r"respiratory\s+disturbance\s+index[^\n\r\d]{0,20}(\d+(?:\.\d+)?)", "respiratory disturbance index ... X"),
        ],
        "odi": [
            (r"ODI\s+(\d+(?:\.\d+)?)", "ODI X"),
            (r"oxygen\s+desaturation\s+index[^\n\r\d]{0,20}(\d+(?:\.\d+)?)", "oxygen desaturation index ... X"),
        ]
    }
    
    results = {}
    for category, pattern_list in patterns.items():
        results[category] = []
        for pattern, description in pattern_list:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                results[category].append({
                    "pattern": description,
                    "matches": matches,
                    "regex": pattern
                })
    
    return results
