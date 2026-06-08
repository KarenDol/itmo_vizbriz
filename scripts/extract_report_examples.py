#!/usr/bin/env python3
"""
Convert DOCX report examples into plain-text files so they can be fed to the LLM prompt.

Usage:
    python3 scripts/extract_report_examples.py \
        --input-dir "/home/ec2-user/patient_data/Report Examples" \
        --output-dir "/home/ec2-user/patient_data/Report Examples/txt"

If --output-dir is omitted, TXT files are written alongside the DOCX files.
Requires `python-docx` (install via `pip install python-docx`).
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path
from typing import Iterable, List

try:
    from docx import Document
except ImportError as exc:  # pragma: no cover
    raise SystemExit(
        "python-docx is required to run this script. Install via `pip install python-docx`."
    ) from exc


def extract_paragraphs(doc: Document) -> List[str]:
    """Return paragraph text while preserving deliberate blank lines."""
    lines: List[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        lines.append(text if text else "")
    return lines


def extract_tables(doc: Document) -> List[str]:
    """Serialize table rows into pipe-delimited strings."""
    rows: List[str] = []
    for table in doc.tables:
        for row in table.rows:
            cell_text = [cell.text.strip() for cell in row.cells]
            if any(cell_text):
                rows.append(" | ".join(cell_text))
        if table.rows:
            rows.append("")
    return rows


def extract_text_from_docx(docx_path: Path) -> str:
    """Convert a DOCX file into a single plain-text blob."""
    document = Document(docx_path)
    parts: List[str] = []
    parts.extend(extract_paragraphs(document))
    table_rows = extract_tables(document)
    if table_rows:
        parts.append("")
        parts.append("# Tables")
        parts.extend(table_rows)

    text_lines: List[str] = []
    previous_blank = False
    for line in parts:
        if line.strip():
            text_lines.append(line)
            previous_blank = False
        else:
            if not previous_blank:
                text_lines.append("")
            previous_blank = True

    result = "\n".join(text_lines).strip()
    return result + "\n" if result else ""


def write_text(output_path: Path, text: str) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(text, encoding="utf-8")


def iter_docx_files(input_dir: Path) -> Iterable[Path]:
    yield from sorted(input_dir.glob("*.docx"))


def convert_directory(input_dir: Path, output_dir: Path | None) -> None:
    files = list(iter_docx_files(input_dir))
    if not files:
        print(f"No DOCX files found in {input_dir}")
        return

    for docx_path in files:
        target_dir = output_dir if output_dir else docx_path.parent
        txt_path = target_dir / (docx_path.stem + ".txt")
        print(f"Converting {docx_path.name} -> {txt_path}")
        text = extract_text_from_docx(docx_path)
        write_text(txt_path, text)


def parse_args(argv: List[str]) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Convert DOCX reports to TXT.")
    parser.add_argument("--input-dir", type=Path, required=True, help="Directory containing DOCX files.")
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=None,
        help="Directory to write TXT files (defaults to the input directory).",
    )
    return parser.parse_args(argv)


def main(argv: List[str] | None = None) -> None:
    args = parse_args(argv or sys.argv[1:])
    input_dir: Path = args.input_dir.expanduser().resolve()
    if not input_dir.exists() or not input_dir.is_dir():
        raise SystemExit(f"Input directory is invalid: {input_dir}")

    output_dir = args.output_dir.expanduser().resolve() if args.output_dir else None
    if output_dir:
        output_dir.mkdir(parents=True, exist_ok=True)

    convert_directory(input_dir, output_dir)


if __name__ == "__main__":
    main()
