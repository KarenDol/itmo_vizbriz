import os
import json
import logging
from typing import List, Dict, Any, Optional
# Lazy import inside functions to avoid hard dependency at app startup
from datetime import datetime

logger = logging.getLogger(__name__)

# Optional Bedrock integration (mirrors other modules)
try:
    from flask_app.config.bedrock_config import query_bedrock_claude_enhanced as bedrock_query_enhanced
except Exception:
    bedrock_query_enhanced = None


def scan_directory(base_dir: str, limit: int = 0) -> List[str]:
    if not os.path.isdir(base_dir):
        raise FileNotFoundError(f"Directory not found: {base_dir}")
    files = [os.path.join(base_dir, f) for f in os.listdir(base_dir) if f.lower().endswith(('.pdf', '.docx'))]
    files.sort()
    if limit and limit > 0:
        files = files[:limit]
    return files


def extract_text_from_file(file_path: str) -> str:
    text = ''
    try:
        if file_path.lower().endswith('.pdf'):
            import fitz  # PyMuPDF
            with fitz.open(file_path) as doc:
                for page in doc:
                    text += page.get_text() + '\n'
        elif file_path.lower().endswith('.docx'):
            from docx import Document  # python-docx
            doc = Document(file_path)
            for p in doc.paragraphs:
                if p.text:
                    text += p.text + '\n'
            # Include table contents (many clinical Word docs store data in tables)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text += cell.text.strip() + '\n'
            # Include headers and footers (often contain key metadata)
            try:
                for section in doc.sections:
                    header = section.header
                    if header:
                        for p in getattr(header, 'paragraphs', []) or []:
                            if p.text:
                                text += p.text.strip() + '\n'
                        for t in getattr(header, 'tables', []) or []:
                            for r in t.rows:
                                for c in r.cells:
                                    if c.text:
                                        text += c.text.strip() + '\n'
                    footer = section.footer
                    if footer:
                        for p in getattr(footer, 'paragraphs', []) or []:
                            if p.text:
                                text += p.text.strip() + '\n'
                        for t in getattr(footer, 'tables', []) or []:
                            for r in t.rows:
                                for c in r.cells:
                                    if c.text:
                                        text += c.text.strip() + '\n'
            except Exception:
                pass

            # Optional fallbacks if little text was found (text in shapes/textboxes)
            try:
                if len((text or '').strip()) < 200:
                    # Try mammoth (if installed) to extract more robustly
                    try:
                        import mammoth  # type: ignore
                        with open(file_path, 'rb') as f:
                            result = mammoth.convert_to_markdown(f)
                            md = (result.value or '').strip()
                            if md:
                                text += ('\n' + md)
                    except Exception:
                        pass
                    # Try docx2txt as another best-effort fallback
                    try:
                        import docx2txt  # type: ignore
                        extracted = (docx2txt.process(file_path) or '').strip()
                        if extracted:
                            text += ('\n' + extracted)
                    except Exception:
                        pass
            except Exception:
                pass
    except Exception:
        # best-effort; return what we have
        pass
    return (text or '').strip()


def _parse_simple_metrics(text: str) -> Dict[str, Any]:
    # Best-effort regex for AHI, ODI, nadir, severity words
    res: Dict[str, Any] = {}
    try:
        import re
        m = re.search(r'AHI\s*[:=]?\s*(\d+(?:\.\d+)?)', text, re.I)
        if m:
            res['ahi'] = float(m.group(1))
        m = re.search(r'ODI\s*[:=]?\s*(\d+(?:\.\d+)?)', text, re.I)
        if m:
            res['odi'] = float(m.group(1))
        m = re.search(r'(SpO2|O2\s*nadir)\s*[:=]?\s*(\d+(?:\.\d+)?)%?', text, re.I)
        if m:
            res['o2_nadir_pct'] = float(m.group(2))
        if re.search(r'\bsevere\b', text, re.I):
            res['severity'] = 'severe'
        elif re.search(r'\bmoderate\b', text, re.I):
            res['severity'] = 'moderate'
        elif re.search(r'\bmild\b', text, re.I):
            res['severity'] = 'mild'
    except Exception:
        pass
    return res


def normalize_to_patient_case_json_v1(
    file_path: str,
    file_name: str,
    text_content: str,
    patient_id: Optional[str],
    document_type: str = 'canonical',
    version: int = 1,
) -> Dict[str, Any]:
    now_iso = datetime.utcnow().isoformat()
    snapshot: Dict[str, Any] = {
        'schema_version': '1.0',
        'document_type': document_type,
        'patient_id': str(patient_id) if patient_id else None,
        'as_of': now_iso,
        'version': version,
        # Per-report metadata stays for traceability even in canonical
        'report_meta': {
            'report_id': file_name,
            'source_report_type': 'case_report',
            'source_uri': file_path,
            'created_at': now_iso,
            'author_role': 'system',
        },
        'canonical_meta': {
            'version': version,
            'report_refs': [
                {
                    'report_id': file_name,
                    'source_uri': file_path,
                    'ingested_at': now_iso,
                }
            ],
        },
        'sleep_study': {},
        'observations': {'summary': []},
        'treatment_considerations': {},
        'device_design': {},
        'follow_up_plan': {},
        'provenance': [],
        'validation': {
            'errors': [],
            'warnings': [],
        },
        'completeness_flags': {
            'has_sleep_study': False,
            'has_anatomy_imaging': False,
            'has_tmj_info': False,
        },
    }
    metrics = _parse_simple_metrics(text_content or '')
    if metrics:
        snapshot['sleep_study'] = metrics
        snapshot['completeness_flags']['has_sleep_study'] = True
        # provenance example entries for metrics present
        for k in metrics.keys():
            snapshot['provenance'].append({
                'path': f'$.sleep_study.{k}',
                'report_id': file_name,
                'source_uri': file_path,
                'note': 'Parsed from document text'
            })

    # Optionally enrich snapshot with LLM extraction if enabled and available
    if os.getenv('ENABLE_LLM_INGEST', '1') == '1' and bedrock_query_enhanced is not None and text_content:
        try:
            enriched = _extract_structured_fields_with_llm(text_content=text_content, file_name=file_name)
            if isinstance(enriched, dict):
                _merge_snapshot(snapshot, enriched)
        except Exception:
            # Best-effort enrichment; keep baseline snapshot on failure
            logger.info("LLM enrichment failed; continuing with baseline snapshot", exc_info=True)
    return snapshot


def explode_observations_from_snapshot(snapshot: Dict[str, Any]) -> List[Dict[str, Any]]:
    obs: List[Dict[str, Any]] = []
    ss = snapshot.get('sleep_study') or {}
    if 'ahi' in ss:
        obs.append({'path': 'sleep_study.ahi', 'value': str(ss['ahi'])})
    if 'odi' in ss:
        obs.append({'path': 'sleep_study.odi', 'value': str(ss['odi'])})
    if 'o2_nadir_pct' in ss:
        obs.append({'path': 'sleep_study.o2_nadir_pct', 'value': str(ss['o2_nadir_pct'])})
    if 'severity' in ss:
        obs.append({'path': 'sleep_study.severity', 'value': ss['severity']})

    # Observations summary (list of strings)
    obs_summary = (snapshot.get('observations') or {}).get('summary') or []
    for item in obs_summary:
        if isinstance(item, str) and item.strip():
            obs.append({'path': 'observations.summary', 'value': item.strip()})

    # Anatomy/imaging object
    anatomy = (snapshot.get('observations') or {}).get('anatomy_imaging') or {}
    for k, v in anatomy.items():
        if v is not None and v != '':
            obs.append({'path': f'observations.anatomy_imaging.{k}', 'value': str(v)})

    # TMJ flags
    tmj = (snapshot.get('observations') or {}).get('tmj_flags') or {}
    for k, v in tmj.items():
        if v is not None and v is not False:
            obs.append({'path': f'observations.tmj_flags.{k}', 'value': str(v)})

    # Treatment considerations
    tr = snapshot.get('treatment_considerations') or {}
    for k, v in tr.items():
        if v is None:
            continue
        if isinstance(v, list):
            for item in v:
                obs.append({'path': f'treatment_considerations.{k}', 'value': str(item)})
        else:
            obs.append({'path': f'treatment_considerations.{k}', 'value': str(v)})

    # Device design
    dd = snapshot.get('device_design') or {}
    for k, v in dd.items():
        if v is None:
            continue
        if isinstance(v, list):
            for item in v:
                obs.append({'path': f'device_design.{k}', 'value': str(item)})
        else:
            obs.append({'path': f'device_design.{k}', 'value': str(v)})

    # Follow-up plan
    fu = snapshot.get('follow_up_plan') or {}
    for k, v in fu.items():
        if v is None:
            continue
        if isinstance(v, list):
            for item in v:
                obs.append({'path': f'follow_up_plan.{k}', 'value': str(item)})
        else:
            obs.append({'path': f'follow_up_plan.{k}', 'value': str(v)})

    return obs


# ---- Helpers for LLM enrichment ----

def _extract_structured_fields_with_llm(*, text_content: str, file_name: str) -> Dict[str, Any]:
    """
    Use Bedrock (if available) to extract structured fields matching
    the Patient Case JSON v1 partial schema from free text.

    Returns a dict with keys among: observations, sleep_study, treatment_considerations,
    device_design, follow_up_plan, demographics.
    """
    if bedrock_query_enhanced is None:
        return {}

    system_prompt = (
        "You are a clinical extraction assistant. Extract structured data from the report text into the following schema subset.\n"
        "Top-level keys allowed: observations, sleep_study, treatment_considerations, device_design, follow_up_plan, demographics.\n"
        "- observations.summary: array of strings\n"
        "- observations.anatomy_imaging: object { primary_obstruction_site?, soft_palate_uvula?, tongue_base?, bite_jaw?, hyoid?, nose_sinus?, tmj? }\n"
        "- observations.tmj_flags: object { pain?: boolean, clicking?: boolean, side?: 'left'|'right'|'bilateral'|null }\n"
        "- sleep_study: object { study_type?, sleep_duration_h?, sleep_efficiency_pct?, ahi?, odi?, desaturation_events?, o2_nadir_pct?, snoring: { avg_db?, max_db? } }\n"
        "- treatment_considerations: { primary_pathway?: string[], adjuncts?: string[], cautions?: string[], rationale?: string }\n"
        "- device_design: { mandibular_advancement_mm?, advancement_plan?, vertical_opening_mm?, anterior_window?: 'small'|'medium'|'large'|null, retention_features?: string[], material?, coverage?, initial_accessories?: string[] }\n"
        "- follow_up_plan: { evaluations?: array of { type: 'ENT'|'TMJ'|'SleepStudy'|'DISE'|'General', reason?, timeframe? }, lifestyle?: string[], positional_therapy?: boolean|null, retest_after_init_months?: number|null }\n"
        "- demographics (optional): { sex?: 'M'|'F'|'X'|null, age_years?, height_cm?, weight_kg?, bmi? }\n"
        "Rules: Return strictly valid JSON only (no markdown), omit keys you cannot support, use correct numeric/boolean types."
    )

    user_prompt = (
        f"FILE: {file_name}\n"
        "Report Text:\n" + (text_content[:60000] if len(text_content) > 60000 else text_content) + "\n\n"
        "Output JSON only."
    )

    messages = [
        {"role": "assistant", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ]

    result = bedrock_query_enhanced(messages, max_tokens=800, temperature=0.2, top_p=0.9)
    if not isinstance(result, dict) or not result.get("success"):
        return {}
    raw = result.get("response", "").strip()

    # Try to isolate JSON
    parsed: Dict[str, Any] = {}
    try:
        parsed = json.loads(raw)
    except Exception:
        try:
            import re
            match = re.search(r"\{[\s\S]*\}", raw)
            if match:
                parsed = json.loads(match.group(0))
        except Exception:
            parsed = {}
    # Sanitize unexpected keys
    if isinstance(parsed, dict):
        allowed = {
            'observations', 'sleep_study', 'treatment_considerations',
            'device_design', 'follow_up_plan', 'demographics'
        }
        parsed = {k: v for k, v in parsed.items() if k in allowed}
    else:
        parsed = {}
    return parsed


def _merge_snapshot(base: Dict[str, Any], extra: Dict[str, Any]) -> None:
    """Shallow/structured merge for known sections into the snapshot."""
    if not isinstance(extra, dict):
        return
    # observations object
    if isinstance(extra.get('observations'), dict):
        base.setdefault('observations', {})
        # summary
        if isinstance(extra['observations'].get('summary'), list):
            base['observations'].setdefault('summary', [])
            # extend with unique items preserving order
            seen = set(base['observations']['summary'])
            for item in extra['observations']['summary']:
                if isinstance(item, str) and item not in seen:
                    base['observations']['summary'].append(item)
                    seen.add(item)
        # anatomy_imaging
        if isinstance(extra['observations'].get('anatomy_imaging'), dict):
            base['observations'].setdefault('anatomy_imaging', {})
            for k, v in extra['observations']['anatomy_imaging'].items():
                if v is not None and v != '':
                    base['observations']['anatomy_imaging'][k] = v
        # tmj_flags
        if isinstance(extra['observations'].get('tmj_flags'), dict):
            base['observations'].setdefault('tmj_flags', {})
            for k, v in extra['observations']['tmj_flags'].items():
                if v is not None:
                    base['observations']['tmj_flags'][k] = v

    # sleep_study
    if isinstance(extra.get('sleep_study'), dict):
        base.setdefault('sleep_study', {})
        for k, v in extra['sleep_study'].items():
            if v is not None:
                base['sleep_study'][k] = v
        if base.get('sleep_study'):
            base.setdefault('completeness_flags', {}).update({'has_sleep_study': True})

    # treatment_considerations
    if isinstance(extra.get('treatment_considerations'), dict):
        base.setdefault('treatment_considerations', {})
        for k, v in extra['treatment_considerations'].items():
            if v is not None:
                base['treatment_considerations'][k] = v

    # device_design
    if isinstance(extra.get('device_design'), dict):
        base.setdefault('device_design', {})
        for k, v in extra['device_design'].items():
            if v is not None:
                base['device_design'][k] = v

    # follow_up_plan
    if isinstance(extra.get('follow_up_plan'), dict):
        base.setdefault('follow_up_plan', {})
        for k, v in extra['follow_up_plan'].items():
            if v is not None:
                base['follow_up_plan'][k] = v

    # demographics (optional)
    if isinstance(extra.get('demographics'), dict):
        base['demographics'] = extra['demographics']


